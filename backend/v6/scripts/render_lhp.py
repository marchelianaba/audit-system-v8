#!/usr/bin/env python3
"""
render_lhp.py — Renderer LHP DOCX placeholder-driven (v4.0.4).

Konsumsi `_KKP/temuan.json` + `context.md` + `_PKP/sasaran-assignment.json`
+ template `templates/_skeleton-lhp/template-lhp-[jenis].docx`, lalu lakukan
find-replace placeholder `{{...}}` untuk menghasilkan
`_LHP/LHP-SUBSTANSI-[nomor-st-slug].docx`.

Placeholder yang didukung:
  Identitas: NOMOR_ST, TANGGAL_ST, OBYEK, PENERIMA_LHP, PERIODE_PELAKSANAAN,
             BULAN_TAHUN, NOMOR_LHR, NOMOR_NOTA_DINAS, TANGGAL_NOTA_DINAS,
             HAL_LHR, JUDUL_LHR_LINE_1..3, JUDUL_LHR_INLINE, DASAR_PERMINTAAN,
             NAMA_AUDITI, TANGGAL_EXIT_MEETING
  Tim     : NAMA_INSPEKTUR, NIP_INSPEKTUR, TTD_INSPEKTUR, TEMBUSAN_LIST
  Konten  : DASAR_HUKUM_LIST, TUJUAN_REVIU, SASARAN_LIST, RUANG_LINGKUP,
             METODOLOGI_REVIU, GAMBARAN_UMUM, HASIL_REVIU_INTRO,
             HASIL_REVIU_LOOP, SIMPULAN_REVIU

Yang tidak diisi otomatis (harus dilengkapi auditor manual via INTEGRAL/SIMWAS):
  NOMOR_LHR, NOMOR_NOTA_DINAS, TANGGAL_NOTA_DINAS, TTD_INSPEKTUR

Contoh:
  python3 scripts/render_lhp.py --penugasan penugasan/[nama] \
      --rekomendasi-file penugasan/[nama]/_LHP/rekomendasi.json \
      --judul "Pengadaan DC/DRC PSrE Induk Tahun 2026" \
      --auditi "Direktorat Pengawasan Sertifikasi dan Transaksi Elektronik"
"""
from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write("python-docx tidak terpasang. Jalankan: pip install python-docx\n")
    sys.exit(2)


JENIS_LABEL = {
    "reviu-pengadaan": "Reviu Pengadaan",
    "reviu-rka-kl": "Reviu Rencana Kerja dan Anggaran",
    "audit-pengadaan": "Audit Pengadaan",
    "audit-kinerja": "Audit Kinerja",
    "evaluasi-sakip": "Evaluasi SAKIP",
    "evaluasi-spip": "Evaluasi SPIP",
    "evaluasi-reformasi-birokrasi": "Evaluasi Reformasi Birokrasi",
    "evaluasi-manajemen-risiko": "Evaluasi Manajemen Risiko",
    "pemantauan-pengadaan": "Pemantauan Pengadaan",
    "pemantauan-tindak-lanjut": "Pemantauan Tindak Lanjut Hasil Pengawasan",
}


# ---------- Context parser ----------
def parse_context(context_path: Path) -> dict:
    """Ekstrak field dari tabel context.md."""
    out = {}
    if not context_path.exists():
        return out
    text = context_path.read_text(encoding="utf-8")
    field_map = {
        "nomor st": "nomor_st", "tanggal st": "tanggal_st",
        "objek": "obyek", "obyek": "obyek",
        "paket pengadaan": "paket",
        "tahun anggaran": "tahun_anggaran",
        "periode pelaksanaan": "periode",
        "jangka waktu": "jangka_waktu",
        "tingkat risiko": "tingkat_risiko",
        "tingkat keyakinan": "tingkat_keyakinan",
        "penerima lhp": "penerima_lhp",
        "dasar penugasan": "dasar_penugasan",
    }
    for line in text.splitlines():
        # Format tabel: | Key | Value |
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 2:
                k_norm = cells[0].lower()
                for k_pattern, target_field in field_map.items():
                    if k_pattern in k_norm:
                        if target_field not in out or not out[target_field]:
                            out[target_field] = cells[1]
                        break
        # Format list: - Key: Value
        elif line.startswith("-"):
            stripped = line.lstrip("- ").strip()
            if ":" in stripped:
                k_raw, _, v_raw = stripped.partition(":")
                k_norm = k_raw.strip().lower()
                v_val = v_raw.strip()
                for k_pattern, target_field in field_map.items():
                    if k_pattern in k_norm:
                        if target_field not in out or not out[target_field]:
                            out[target_field] = v_val
                        break
    # Tujuan + Ruang Lingkup (inline format Tujuan: ... / Ruang Lingkup: ...)
    m = re.search(r"^Tujuan\s*:\s*(.+?)(?=\n\n|\n##|\nRuang Lingkup)", text, re.MULTILINE | re.DOTALL)
    if m:
        out["tujuan"] = m.group(1).strip()
    m = re.search(r"^Ruang Lingkup\s*:\s*(.+?)(?=\n\n|\n##)", text, re.MULTILINE | re.DOTALL)
    if m:
        out["ruang_lingkup"] = m.group(1).strip()
    # Ringkasan Obyek — heading ## Ringkasan Obyek diikuti paragraf isi
    m = re.search(r"##\s*Ringkasan Obyek\s*\n+([\s\S]+?)(?=\n##|\Z)", text, re.IGNORECASE)
    if m:
        out["ringkasan_obyek"] = m.group(1).strip()
    # Tim — tabel
    tim = []
    in_tim = False
    for line in text.splitlines():
        if line.strip().startswith("## Tim"):
            in_tim = True
            continue
        if in_tim and line.startswith("|"):
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 5 and cells[0].isdigit():
                tim.append({"no": cells[0], "nama": cells[1], "nip": cells[2],
                           "jabatan": cells[3], "jabfung": cells[4]})
        elif in_tim and not line.startswith("|") and line.strip() and not line.startswith("---"):
            in_tim = False
    out["tim"] = tim
    return out


# ---------- Find-replace ----------
def replace_in_paragraph(p, mapping: dict):
    """Replace {{KEY}} placeholders in a paragraph, preserving formatting of
    the first run that contains the placeholder."""
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return False
    new = full
    for k, v in mapping.items():
        new = new.replace("{{" + k + "}}", str(v))
    if new == full:
        return False
    # Replace all runs with single run carrying full new text, preserving
    # formatting of first run.
    if p.runs:
        first = p.runs[0]
        first.text = new
        for r in p.runs[1:]:
            r.text = ""
    return True


def replace_in_doc(doc, mapping: dict):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)


def replace_label_in_doc(doc, pairs: list):
    """Ganti teks plain (bukan {{placeholder}}) di seluruh paragraf.

    `pairs` — list of (old_str, new_str).
    Berguna untuk mengganti label section seperti
    'Dasar Pelaksanaan Reviu' → 'Dasar Pelaksanaan Evaluasi'.
    """
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        new = full
        for old, rep in pairs:
            new = new.replace(old, rep)
        if new != full and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""


# ---------- Loop helpers ----------
def expand_paragraph_placeholder_to_blocks(doc, placeholder: str, blocks: list):
    """Replace SETIAP paragraph yang mengandung `placeholder` dengan list of
    paragraph blocks. Mengembalikan jumlah occurrences yang berhasil di-expand.
    """
    targets = [p for p in doc.paragraphs if placeholder in p.text]
    if not targets:
        return 0

    def make_p(text, fmt=None):
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        fmt = fmt or {}
        new = doc.add_paragraph()
        if fmt.get("align") == "justify":
            new.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif fmt.get("align") == "center":
            new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif fmt.get("align") == "right":
            new.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if fmt.get("indent"):
            new.paragraph_format.left_indent = Cm(fmt["indent"])
        # Spasi 1,5 via XML (paling reliable lintas versi python-docx)
        pPr = new._element.get_or_add_pPr()
        sp_el = _OE("w:spacing")
        sp_el.set(_qn("w:line"), "360")
        sp_el.set(_qn("w:lineRule"), "auto")
        pPr.append(sp_el)
        r = new.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        r.font.size = Pt(fmt.get("size", 12))
        r.font.name = "Arial"
        return new

    expanded = 0
    for target_p in targets:
        target_el = target_p._element
        new_paras = [make_p(t, f) for t, f in blocks]
        for new_p in new_paras:
            target_el.addprevious(new_p._element)
        target_el.getparent().remove(target_el)
        expanded += 1
    return expanded


# ---------- Main composition ----------
def build_mapping(pen_dir: Path, args) -> tuple[dict, dict, dict, dict, dict]:
    """Bangun mapping placeholder dari context + temuan + sasaran."""
    ctx = parse_context(pen_dir / "context.md")
    temuan_path = pen_dir / "_KKP" / "temuan.json"
    sasaran_path = pen_dir / "_PKP" / "sasaran-assignment.json"
    if not temuan_path.exists():
        sys.stderr.write(f"temuan.json tidak ditemukan: {temuan_path}\n"); sys.exit(1)
    if not sasaran_path.exists():
        sys.stderr.write(f"sasaran-assignment.json tidak ditemukan: {sasaran_path}\n"); sys.exit(1)
    kkp = json.loads(temuan_path.read_text(encoding="utf-8"))
    sa = json.loads(sasaran_path.read_text(encoding="utf-8"))

    # Rekomendasi mapping (per id_temuan)
    rek = {}
    if args.rekomendasi_file and Path(args.rekomendasi_file).exists():
        rek = json.loads(Path(args.rekomendasi_file).read_text(encoding="utf-8"))

    jenis = kkp["penugasan"]["jenis_pengawasan"]
    obyek = ctx.get("obyek") or kkp["penugasan"]["obyek"]
    penerima_lhp = ctx.get("penerima_lhp") or args.penerima or "[DIISI AUDITOR]"
    nomor_st = ctx.get("nomor_st") or kkp["penugasan"]["nomor_st"]
    tanggal_st = ctx.get("tanggal_st") or kkp["penugasan"]["tanggal_st"]
    periode = ctx.get("periode") or "[DIISI AUDITOR]"

    # Inspektur (PM)
    pm = next((m for m in ctx.get("tim", []) if "Pengendali Mutu" in m.get("jabatan", "")), None)
    nama_inspektur = pm["nama"] if pm else "[DIISI AUDITOR]"
    nip_inspektur = pm["nip"] if pm else "[DIISI]"

    judul_lhr = args.judul or f"{kkp['penugasan']['obyek']}"
    # Pecah judul jadi 3 baris kalau panjang
    words = judul_lhr.split()
    if len(words) > 12:
        third = len(words) // 3
        line1 = " ".join(words[:third])
        line2 = " ".join(words[third:2*third])
        line3 = " ".join(words[2*third:])
    else:
        line1 = judul_lhr
        line2 = ""
        line3 = ""

    nama_auditi = args.auditi or obyek

    # Bulan/Tahun untuk halaman judul
    bulan_id = ["JANUARI","FEBRUARI","MARET","APRIL","MEI","JUNI","JULI","AGUSTUS","SEPTEMBER","OKTOBER","NOVEMBER","DESEMBER"]
    today = datetime.now()
    bulan_tahun = f"{bulan_id[today.month-1]} {today.year}"

    mapping = {
        "NOMOR_ST": nomor_st,
        "TANGGAL_ST": tanggal_st if tanggal_st and tanggal_st != "None" else "[DIISI AUDITOR]",
        "OBYEK": obyek,
        "PENERIMA_LHP": penerima_lhp,
        "PERIODE_PELAKSANAAN": periode,
        "BULAN_TAHUN": bulan_tahun,
        "NOMOR_LHR": "[DIISI AUDITOR — dari SIMWAS]",
        "NOMOR_NOTA_DINAS": getattr(args, "nomor_nota_dinas", None) or "[DIISI AUDITOR]",
        "TANGGAL_NOTA_DINAS": getattr(args, "tanggal_nota_dinas", None) or "[DIISI AUDITOR]",
        "HAL_LHR": judul_lhr,
        "JUDUL_LHR_LINE_1": line1.upper() if line1 else "",
        "JUDUL_LHR_LINE_2": line2.upper() if line2 else "",
        "JUDUL_LHR_LINE_3": line3.upper() if line3 else "",
        "JUDUL_LHR_INLINE": judul_lhr,
        "DASAR_PERMINTAAN": args.dasar_permintaan or _strip_st_from_dasar(ctx.get("dasar_penugasan") or "") or "[DIISI — Dasar Pengawasan dari Kartu Penugasan]",
        "NAMA_AUDITI": nama_auditi,
        "TANGGAL_EXIT_MEETING": args.tanggal_exit_meeting or "[DIISI AUDITOR]",
        "NAMA_INSPEKTUR": nama_inspektur,
        "NIP_INSPEKTUR": nip_inspektur,
        "TTD_INSPEKTUR": "[DIISI AUDITOR — TTD]",
        "TUJUAN_REVIU": ctx.get("tujuan", "[DIISI dari context.md atau argumen]"),
        "RUANG_LINGKUP": ctx.get("ruang_lingkup", "[DIISI dari context.md]"),
        "METODOLOGI_REVIU": (
            "Reviu dilaksanakan dengan melakukan penelaahan dokumen (desk review) "
            "atas seluruh data dukung yang diterima dari auditi, dipadukan dengan "
            "klarifikasi tertulis kepada PPK dan tim teknis. Tim juga memanfaatkan "
            "pipeline pre-digest dan cross-check otomatis audit-system-v4 untuk "
            "mendeteksi anomali struktural antar dokumen."
        ),
        "GAMBARAN_UMUM": args.gambaran_umum or ctx.get("ringkasan_obyek") or "[DIISI — gambaran obyek pengadaan, nilai HPS, mekanisme pengadaan]",
        "HASIL_REVIU_INTRO": (
            f"Berdasarkan penelaahan atas dokumen perencanaan, tim Inspektorat II "
            f"memperoleh {len(kkp['temuan'])} ({_terbilang(len(kkp['temuan']))}) "
            f"catatan reviu yang dikelompokkan ke dalam {len(sa['sasaran'])} "
            f"({_terbilang(len(sa['sasaran']))}) aspek sesuai sasaran pengawasan. "
            f"Catatan-catatan ini dirumuskan dengan paradigma reviu (Kondisi-Kriteria-"
            f"Akibat-Rekomendasi) dengan tingkat keyakinan terbatas sebagaimana "
            f"diatur dalam SAIPI dan Perlem LKPP 12/2021."
        ),
    }
    return mapping, kkp, sa, rek, ctx


def _terbilang(n: int) -> str:
    digits = ["nol","satu","dua","tiga","empat","lima","enam","tujuh","delapan",
              "sembilan","sepuluh","sebelas"]
    if 0 <= n < len(digits):
        return digits[n]
    return str(n)


def _strip_st_from_dasar(dasar: str) -> str:
    """Hapus referensi Surat Tugas dari field dasar_penugasan.

    KP sering mencantumkan ST di kolom Dasar Pengawasan, tapi di laporan
    ST sudah disebut tersendiri di kalimat 'kami telah menerbitkan ST Nomor ...'.
    """
    cleaned = re.sub(
        r"(?:dan\s+|,\s*)?(?:[—–\-]\s*)?Surat\s+Tugas\s+(?:Nomor\s+)?[\w./,]+(?:\s+tanggal\s+\d+\s+\w+\s+\d{4})?",
        "",
        dasar,
        flags=re.IGNORECASE,
    ).strip().rstrip(";,").strip()
    return cleaned


def build_dasar_hukum_blocks(jenis: str) -> list:
    items = [
        "Peraturan Pemerintah Nomor 60 Tahun 2008 tentang Sistem Pengendalian Intern Pemerintah;",
    ]
    if "pengadaan" in jenis:
        items += [
            "Peraturan Presiden Nomor 16 Tahun 2018 sebagaimana diubah dengan Peraturan Presiden Nomor 12 Tahun 2021 tentang Pengadaan Barang/Jasa Pemerintah;",
            "Peraturan Lembaga Kebijakan Pengadaan Barang/Jasa Pemerintah Nomor 12 Tahun 2021 tentang Pedoman Pelaksanaan Pengadaan Barang/Jasa Pemerintah melalui Penyedia;",
        ]
    items += [
        "Standar Audit Intern Pemerintah Indonesia (SAIPI) AAIPI 2021 (PER-01/AAIPI/DPN/2021);",
        "Program Kerja Pengawasan Tahunan Inspektorat Jenderal Komdigi.",
    ]
    blocks = []
    for i, txt in enumerate(items, start=1):
        blocks.append((f"{i}. {txt}", {"align": "justify"}))
    return blocks


def build_sasaran_blocks(sa: dict) -> list:
    blocks = []
    for i, s in enumerate(sa["sasaran"], start=1):
        blocks.append((f"{i}. {s['deskripsi']}", {"align": "justify", "indent": 0.8}))
    return blocks


def build_hasil_reviu_blocks(kkp: dict, sa: dict, rekomendasi: dict, jenis: str,
                              include_rek: bool = True) -> list:
    blocks = []
    per_sasaran = {}
    for t in kkp["temuan"]:
        per_sasaran.setdefault(t["sasaran_id"], []).append(t)

    aspek_letter_seq = ["F.1", "F.2", "F.3", "F.4", "F.5", "F.6", "F.7", "F.8"]

    for idx, sid in enumerate(sorted(per_sasaran.keys())):
        sasaran_obj = next(x for x in sa["sasaran"] if x["sasaran_id"] == sid)
        section_title = f"{aspek_letter_seq[idx]}. {sasaran_obj['deskripsi']}"
        blocks.append((section_title, {"bold": True, "size": 12}))

        for ti, t in enumerate(per_sasaran[sid], start=1):
            blocks.append((f"{ti}. {t['judul_temuan']}", {"bold": True}))

            blocks.append(("Kondisi:", {"bold": True}))
            blocks.append((t["kondisi"], {"align": "justify", "indent": 0.3}))

            blocks.append(("Kriteria:", {"bold": True}))
            blocks.append((t["kriteria"], {"align": "justify", "indent": 0.3}))

            if jenis.startswith("audit") and t.get("sebab"):
                blocks.append(("Sebab:", {"bold": True}))
                blocks.append((t["sebab"], {"align": "justify", "indent": 0.3}))

            if t.get("akibat"):
                blocks.append(("Akibat:", {"bold": True}))
                blocks.append((t["akibat"], {"align": "justify", "indent": 0.3}))

            if include_rek:
                blocks.append(("Rekomendasi:", {"bold": True}))
                rek = rekomendasi.get(t["id_temuan"], "[Rekomendasi disusun bersama Pengendali Teknis berdasarkan exit meeting]")
                blocks.append((rek, {"align": "justify", "indent": 0.3}))

            sumber = "; ".join(f"{ds['file']} hal. {ds['halaman']}" for ds in t.get("dokumen_sumber", []))
            blocks.append((f"Sumber dokumen: {sumber}", {"italic": True, "size": 9, "indent": 0.3, "align": "justify"}))
            blocks.append(("", {}))
    return blocks


def build_simpulan(kkp: dict, sa: dict) -> str:
    n = len(kkp["temuan"])
    return (
        f"Berdasarkan hasil reviu terbatas yang kami lakukan, dokumen perencanaan "
        f"yang menjadi obyek reviu secara umum telah disusun untuk mendukung "
        f"pelaksanaan kegiatan auditi dan telah memuat persyaratan utama. Namun, "
        f"terdapat {n} ({_terbilang(n)}) catatan reviu yang perlu ditindaklanjuti "
        f"sebelum dimulainya tahap pemilihan penyedia/pelaksanaan kegiatan. "
        f"Dengan keterbatasan keyakinan reviu (limited assurance), kami tidak "
        f"memberikan opini final atas kelayakan harga maupun kelayakan teknis; "
        f"opini tersebut menjadi tanggung jawab Pejabat Pembuat Komitmen, Pokja "
        f"Pemilihan, dan tim teknis pengadaan pada saat tender dilaksanakan."
    )


def build_tembusan_blocks(custom: str = None) -> list:
    """Bangun daftar tembusan.

    `custom` — string dipisah titik-koma, misal:
        "Inspektur Jenderal; Sekretaris Itjen; Arsip."
    Jika tidak diisi, placeholder dihapus (tidak ada tembusan).
    """
    if not custom:
        return []
    items = [s.strip() for s in custom.split(";") if s.strip()]
    return [(f"{i}. {txt}", {"size": 10}) for i, txt in enumerate(items, start=1)]


def insert_survei_paragraph(doc, survei_link: str) -> None:
    """Sisipkan paragraf link survei kepuasan sebelum kalimat 'Demikian Nota Dinas'.

    Dipanggil hanya jika --survei diberikan. Berlaku untuk semua jenis laporan
    karena semua template Nota Dinas mengandung 'Demikian Nota Dinas'.
    """
    from docx.oxml import OxmlElement as _OE
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    XML_NS = "http://www.w3.org/XML/1998/namespace"
    for p in doc.paragraphs:
        if "Demikian Nota Dinas" in p.text:
            new_p = _OE("w:p")
            r = _OE("w:r")
            t = _OE("w:t")
            t.text = (
                f"Kami mohon kesediaan Saudara untuk mengisi survei kepuasan kami "
                f"melalui tautan berikut: {survei_link}"
            )
            t.set(f"{{{XML_NS}}}space", "preserve")
            r.append(t)
            new_p.append(r)
            p._element.addprevious(new_p)
            break


# ============================================================
# AUDIT-UMUM specific builders (A–H + Lampiran 1–2)
# ============================================================

def build_dasar_audit_umum(ctx: dict, kkp: dict, args) -> list:
    """A. DASAR — ST + ND permintaan."""
    dasar = args.dasar_permintaan or _strip_st_from_dasar(ctx.get("dasar_penugasan") or "")
    nomor_st = ctx.get("nomor_st") or kkp["penugasan"]["nomor_st"]
    tanggal_st_raw = ctx.get("tanggal_st") or kkp["penugasan"].get("tanggal_st")
    tanggal_st = tanggal_st_raw if tanggal_st_raw and tanggal_st_raw != "None" else "[DIISI AUDITOR]"
    blocks = []
    no = 1
    if dasar:
        blocks.append((f"{no}. {dasar};", {"align": "justify"}))
        no += 1
    blocks.append((f"{no}. Surat Tugas Nomor {nomor_st} tanggal {tanggal_st}.", {"align": "justify"}))
    return blocks


def build_tujuan_audit_umum(ctx: dict, sa: dict) -> list:
    """B. TUJUAN — dari context.md atau deskripsi sasaran."""
    tujuan = ctx.get("tujuan") or sa.get("tujuan_penugasan") or ""
    if tujuan:
        return [(tujuan, {"align": "justify"})]
    blocks = []
    for i, s in enumerate(sa.get("sasaran", []), start=1):
        desc = s.get("deskripsi", "")
        if desc:
            blocks.append((f"{i}. {desc}", {"align": "justify"}))
    if not blocks:
        blocks.append(("[TUJUAN AUDIT — diisi dari Kartu Penugasan]", {"align": "justify"}))
    return blocks


def build_tujuan_audit_umum_audit(ctx: dict, sa: dict) -> list:
    """B. TUJUAN khusus skill audit — format baku: 1. Tujuan... 2. Sasaran (a,b,c)."""
    tujuan = ctx.get("tujuan") or sa.get("tujuan_penugasan") or "[TUJUAN AUDIT — diisi dari Kartu Penugasan]"
    blocks = [(f"1. Tujuan dari dilaksanakannya audit adalah {tujuan}", {"align": "justify"})]
    sasaran = sa.get("sasaran", [])
    if not sasaran:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah [SASARAN — diisi dari Kartu Penugasan]",
                       {"align": "justify"}))
    elif len(sasaran) == 1:
        desc = sasaran[0].get("deskripsi", "[SASARAN — diisi dari Kartu Penugasan]")
        blocks.append((f"2. Sasaran dari dilaksanakannya audit adalah {desc}", {"align": "justify"}))
    else:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah:", {"align": "justify"}))
        for i, s in enumerate(sasaran):
            desc = s.get("deskripsi", "")
            if desc:
                blocks.append((f"    {chr(ord('a') + i)}. {desc}", {"align": "justify", "indent": 0.5}))
    return blocks


def build_metodologi_audit_umum(ctx: dict, sa: dict) -> list:
    """D. METODOLOGI — keyakinan memadai (reasonable assurance)."""
    custom = ctx.get("metodologi") or sa.get("metodologi") or ""
    if custom:
        return [(custom, {"align": "justify"})]
    return [(
        "Audit dilaksanakan melalui serangkaian prosedur: (1) penelaahan dokumen "
        "(desk review) terhadap seluruh kriteria yang diunggah auditor dan dokumen "
        "objek audit yang tersedia; (2) pengujian kesesuaian per kriteria menggunakan "
        "paradigma Kondisi-Kriteria-Sebab-Akibat (KKSA); (3) analytical review atas "
        "data kuantitatif yang relevan; dan (4) konfirmasi/klarifikasi atas hal-hal "
        "yang memerlukan penjelasan lebih lanjut. Audit dilaksanakan sesuai dengan "
        "Standar Audit Intern Pemerintah Indonesia (SAIPI) yang ditetapkan oleh "
        "Asosiasi Auditor Intern Pemerintah Indonesia (PER-01/AAIPI/DPN/2021), "
        "dengan tingkat keyakinan memadai (reasonable assurance).",
        {"align": "justify"}
    )]


def build_rekomendasi_audit_umum(kkp: dict, rek: dict) -> list:
    """G. REKOMENDASI — daftar rekomendasi per temuan."""
    blocks = []
    for t in kkp.get("temuan", []):
        r_text = (
            rek.get(t.get("id_temuan", ""))
            or t.get("rekomendasi")
            or "[Rekomendasi disusun Ketua Tim berdasarkan exit meeting]"
        )
        blocks.append((f"{t.get('id_temuan', '')}. {t['judul_temuan']}", {"bold": True}))
        blocks.append((r_text, {"align": "justify", "indent": 0.5}))
        blocks.append(("", {}))
    if not blocks:
        blocks.append(("[Rekomendasi akan disusun Ketua Tim berdasarkan temuan di atas]",
                       {"align": "justify", "italic": True}))
    return blocks


def build_apresiasi_audit_umum(penerima_lhp: str) -> list:
    """H. APRESIASI DAN PENUTUP."""
    return [(
        f"Berdasarkan kegiatan audit yang dilakukan, Tim menyatakan apresiasi kepada "
        f"{penerima_lhp} dalam membantu proses audit sehingga dapat diselesaikan tepat waktu.",
        {"align": "justify"}
    )]


def build_matriks_temuan_audit_umum(kkp: dict) -> list:
    """LAMPIRAN 1: MATRIKS TEMUAN (CCSAA)."""
    blocks = []
    for t in kkp.get("temuan", []):
        risiko = t.get("level_risiko", "").upper()
        nilai = t.get("nilai_rp", 0) or 0
        label = f"[{t.get('id_temuan', '')}] {t['judul_temuan']}"
        if risiko:
            label += f" — {risiko}"
        if nilai > 0:
            label += f" (Rp {nilai:,.0f})"
        blocks.append((label, {"bold": True}))
        blocks.append((f"Kondisi : {t['kondisi']}", {"align": "justify", "indent": 0.5}))
        blocks.append((f"Kriteria: {t['kriteria']}", {"align": "justify", "indent": 0.5}))
        if t.get("sebab"):
            blocks.append((f"Sebab   : {t['sebab']}", {"align": "justify", "indent": 0.5}))
        blocks.append((f"Akibat  : {t.get('akibat', '—')}", {"align": "justify", "indent": 0.5}))
        sumber = "; ".join(
            f"{ds['file']} hal. {ds.get('halaman', '?')}"
            for ds in t.get("dokumen_sumber", [])
        )
        if sumber:
            blocks.append((f"Sumber  : {sumber}", {"italic": True, "size": 9, "indent": 0.5}))
        blocks.append(("", {}))
    if not blocks:
        blocks.append(("(Tidak ada temuan yang terdaftar)", {"italic": True}))
    return blocks


def build_daftar_dokumen_audit_umum(kkp: dict) -> list:
    """LAMPIRAN 2: DAFTAR DOKUMEN SUMBER."""
    seen: dict = {}
    for t in kkp.get("temuan", []):
        for d in t.get("dokumen_sumber", []):
            fname = d.get("file", "")
            if fname and fname not in seen:
                seen[fname] = d.get("ringkasan", "")
    blocks = []
    for i, (fname, ringkasan) in enumerate(seen.items(), start=1):
        entry = f"{i}. {fname}"
        if ringkasan:
            entry += f" — {ringkasan}"
        blocks.append((entry, {"align": "justify"}))
    if not blocks:
        blocks.append(("(Tidak ada dokumen sumber terdaftar)", {"italic": True}))
    return blocks


def expand_placeholder_to_matriks_table(doc, placeholder: str, kkp: dict, rek: dict) -> int:
    """Replace {{placeholder}} dengan tabel CCSAA dua kolom (Label | Uraian).

    Setiap temuan mendapat baris header ter-merge, diikuti baris
    Kondisi / Kriteria / Sebab / Akibat / Rekomendasi.
    """
    targets = [p for p in doc.paragraphs if placeholder in p.text]
    if not targets:
        return 0

    for target_p in targets:
        target_el = target_p._element
        temuan_list = kkp.get("temuan", [])

        if not temuan_list:
            if target_p.runs:
                target_p.runs[0].text = "(Tidak ada temuan yang terdaftar)"
            else:
                target_p.add_run("(Tidak ada temuan yang terdaftar)").italic = True
            continue

        # Buat tabel 2 kolom di akhir dokumen, lalu pindahkan
        table = doc.add_table(rows=0, cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for t in temuan_list:
            # ── Baris header (merged) ──────────────────────────────────
            hrow = table.add_row()
            hrow.cells[0].merge(hrow.cells[1])
            hpara = hrow.cells[0].paragraphs[0]
            head_txt = f"{t.get('id_temuan', '')}. {t.get('judul_temuan', '')}"
            if t.get("level_risiko"):
                head_txt += f"  [Risiko: {t['level_risiko'].upper()}]"
            if t.get("nilai_rp"):
                head_txt += f"  (Rp {t['nilai_rp']:,.0f})"
            hpara.add_run(head_txt).bold = True
            hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ── Baris CCSAA ───────────────────────────────────────────
            rek_text = (
                rek.get(t.get("id_temuan", ""))
                or t.get("rekomendasi")
                or "[Rekomendasi disusun berdasarkan exit meeting]"
            )
            rows_data = [
                ("Kondisi",      t.get("kondisi") or "—"),
                ("Kriteria",     t.get("kriteria") or "—"),
                ("Sebab",        t.get("sebab") or ""),
                ("Akibat",       t.get("akibat") or "—"),
                ("Rekomendasi",  rek_text),
            ]
            for label, content in rows_data:
                if not content:
                    continue
                row = table.add_row()
                # Kolom label — lebar tetap via preferred width
                label_cell = row.cells[0]
                label_cell.paragraphs[0].add_run(label + ":").bold = True
                # Kolom isi
                content_para = row.cells[1].paragraphs[0]
                run = content_para.add_run(content)
                run.font.size = Pt(10)
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Pindahkan tabel ke posisi placeholder
        tbl_elem = table._element
        tbl_elem.getparent().remove(tbl_elem)
        target_el.addprevious(tbl_elem)
        target_el.getparent().remove(target_el)

    return len(targets)


# ============================================================
# AUDIT-PENGADAAN specific builders
# ============================================================

def build_ringkasan_eksekutif_pengadaan(kkp: dict, sa: dict, rek: dict) -> list:
    """RINGKASAN EKSEKUTIF — ikhtisar temuan dan rekomendasi."""
    n_temuan = len(kkp.get("temuan", []))
    n_sasaran = len(sa.get("sasaran", []))
    nilai_total = sum(t.get("nilai_rp", 0) or 0 for t in kkp.get("temuan", []))
    obyek = kkp["penugasan"]["obyek"]
    blocks = [(
        f"Berdasarkan audit yang dilaksanakan terhadap {obyek}, tim Inspektorat II "
        f"menemukan {n_temuan} ({_terbilang(n_temuan)}) temuan yang tersebar pada "
        f"{n_sasaran} ({_terbilang(n_sasaran)}) aspek sasaran audit.",
        {"align": "justify"},
    )]
    if nilai_total > 0:
        blocks.append((
            f"Total nilai temuan yang teridentifikasi sebesar Rp {nilai_total:,.0f}.",
            {"align": "justify"},
        ))
    blocks.append((
        "Atas temuan-temuan tersebut, Inspektorat II menyampaikan rekomendasi "
        "kepada pimpinan auditi untuk segera ditindaklanjuti sesuai ketentuan "
        "perundang-undangan yang berlaku.",
        {"align": "justify"},
    ))
    if kkp.get("temuan"):
        blocks.append(("Pokok-pokok temuan:", {"bold": True}))
        for i, t in enumerate(kkp["temuan"], start=1):
            blocks.append((f"{i}. {t['judul_temuan']}", {"align": "justify", "indent": 0.5}))
    return blocks


def build_tujuan_sasaran_pengadaan(ctx: dict, sa: dict) -> list:
    """BAB I — Tujuan dan Sasaran Audit (format baku: 1. Tujuan... 2. Sasaran a,b,c)."""
    tujuan = ctx.get("tujuan") or sa.get("tujuan_penugasan") or "[TUJUAN AUDIT — diisi dari Kartu Penugasan]"
    blocks = [(f"1. Tujuan dari dilaksanakannya audit adalah {tujuan}", {"align": "justify"})]
    sasaran = sa.get("sasaran", [])
    if not sasaran:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah [SASARAN — diisi dari Kartu Penugasan]",
                       {"align": "justify"}))
    elif len(sasaran) == 1:
        desc = sasaran[0].get("deskripsi", "[SASARAN — diisi dari Kartu Penugasan]")
        blocks.append((f"2. Sasaran dari dilaksanakannya audit adalah {desc}", {"align": "justify"}))
    else:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah:", {"align": "justify"}))
        for i, s in enumerate(sasaran):
            desc = s.get("deskripsi", "")
            if desc:
                blocks.append((f"    {chr(ord('a') + i)}. {desc}", {"align": "justify", "indent": 0.5}))
    return blocks


def build_ruang_lingkup_pengadaan(ctx: dict, sa: dict) -> list:
    """BAB I — Ruang Lingkup Audit."""
    ruang_lingkup = ctx.get("ruang_lingkup") or ""
    if ruang_lingkup:
        return [(ruang_lingkup, {"align": "justify"})]
    blocks = [("Ruang lingkup audit mencakup pengujian atas:", {"align": "justify"})]
    for i, s in enumerate(sa.get("sasaran", []), start=1):
        blocks.append((f"{i}. {s['deskripsi']}", {"align": "justify", "indent": 0.5}))
    return blocks


def build_metodologi_pengadaan(ctx: dict, sa: dict) -> list:
    """BAB I — Metodologi Audit Pengadaan (reasonable assurance)."""
    custom = ctx.get("metodologi") or sa.get("metodologi") or ""
    if custom:
        return [(custom, {"align": "justify"})]
    return [(
        "Audit pengadaan dilaksanakan melalui serangkaian prosedur: "
        "(1) penelaahan dokumen (desk review) terhadap kontrak, dokumen pengadaan, "
        "dan dokumen penunjang; (2) pengujian substantif menggunakan paradigma "
        "Kondisi-Kriteria-Sebab-Akibat (KKSA); (3) analytical review atas data "
        "keuangan dan fisik yang relevan; serta (4) konfirmasi/klarifikasi kepada "
        "pihak-pihak terkait. Audit dilaksanakan sesuai Standar Audit Intern "
        "Pemerintah Indonesia (SAIPI) yang ditetapkan Asosiasi Auditor Intern "
        "Pemerintah Indonesia (PER-01/AAIPI/DPN/2021), dengan tingkat keyakinan "
        "memadai (reasonable assurance).",
        {"align": "justify"},
    )]


def build_penutup_pengadaan(obyek: str) -> list:
    """BAB V — Penutup (legacy, dipertahankan untuk kompatibilitas)."""
    return [(
        f"Atas kerja sama dan keterbukaan {obyek} dalam mendukung pelaksanaan audit "
        f"ini, Inspektorat II menyampaikan apresiasi dan terima kasih. "
        f"Laporan Hasil Audit ini kami sampaikan untuk dapat digunakan sebagaimana "
        f"mestinya. Apabila terdapat hal-hal yang perlu diklarifikasi lebih lanjut, "
        f"kami terbuka untuk berdiskusi dengan pimpinan auditi.",
        {"align": "justify"},
    )]


def build_temuan_rekomendasi_pengadaan(kkp: dict, rek: dict) -> list:
    """Bab 5: Temuan dan Rekomendasi — detail CCSAA per temuan."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("Tidak terdapat temuan yang signifikan berdasarkan audit yang dilaksanakan.",
                 {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul    = t.get("judul_temuan") or f"Temuan {i}"
        kondisi  = t.get("kondisi") or ""
        kriteria = t.get("kriteria") or ""
        sebab    = t.get("sebab") or ""
        akibat   = t.get("akibat") or t.get("dampak") or ""
        nomor    = str(t.get("nomor") or i)
        r        = rek.get(nomor) or rek.get(str(i)) or {}
        rekomendasi = r.get("rekomendasi") if isinstance(r, dict) else r
        if not rekomendasi:
            rekomendasi = t.get("rekomendasi") or ""
        blocks.append((f"Temuan {i}: {judul}", {"bold": True}))
        if kondisi:
            blocks.append(("Kondisi:", {"bold": True}))
            blocks.append((kondisi, {"size": 10}))
        if kriteria:
            blocks.append(("Kriteria:", {"bold": True}))
            blocks.append((kriteria, {"size": 10}))
        if sebab:
            blocks.append(("Sebab:", {"bold": True}))
            blocks.append((sebab, {"size": 10}))
        if akibat:
            blocks.append(("Akibat:", {"bold": True}))
            blocks.append((akibat, {"size": 10}))
        if rekomendasi:
            blocks.append(("Rekomendasi:", {"bold": True}))
            blocks.append((rekomendasi, {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_kesimpulan_audit_pengadaan(kkp: dict, ctx: dict) -> list:
    """Bab 6: Kesimpulan audit pengadaan — keyakinan memadai."""
    n    = len(kkp.get("temuan", []))
    obyek = ctx.get("obyek") or "[objek audit]"
    if n == 0:
        return [(
            f"Berdasarkan audit yang telah dilaksanakan terhadap {obyek} "
            f"dengan tingkat keyakinan memadai, Inspektorat II menyimpulkan bahwa "
            f"proses pengadaan barang/jasa telah dilaksanakan sesuai dengan ketentuan "
            f"peraturan perundang-undangan yang berlaku. Tidak terdapat temuan yang "
            f"bersifat material.",
            {"align": "justify"},
        )]
    kritis = len([t for t in kkp.get("temuan", [])
                  if t.get("level_risiko", "").lower() in ("kritis", "tinggi", "merah")])
    kalimat_kritis = (
        f" Terdapat {kritis} ({_terbilang(kritis)}) temuan bersifat kritis/tinggi "
        f"yang memerlukan eskalasi dan tindak lanjut segera."
        if kritis else ""
    )
    return [
        (f"Berdasarkan audit yang telah dilaksanakan terhadap {obyek} "
         f"dengan tingkat keyakinan memadai, Inspektorat II menyimpulkan bahwa "
         f"terdapat {n} ({_terbilang(n)}) temuan yang perlu mendapat perhatian "
         f"dan tindak lanjut segera.{kalimat_kritis}",
         {"align": "justify"}),
        ("Seluruh rekomendasi wajib ditindaklanjuti oleh {{NAMA_AUDITI}} sesuai dengan "
         "tenggat waktu yang ditetapkan. Hasil tindak lanjut dilaporkan kepada "
         "Inspektorat II melalui mekanisme pemantauan tindak lanjut rekomendasi.",
         {"align": "justify"}),
    ]


# ============================================================
# PER-SKILL-GROUP builders: metodologi, simpulan, label pairs
# ============================================================

# ── Label-pair helpers ───────────────────────────────────────────────────────

def _label_pairs(verb: str) -> list:
    """Pasangan (lama, baru) untuk mengganti label section.

    Menangani template yang menggunakan kata 'Reviu' MAUPUN 'Pengawasan'
    sebagai kata generik, agar output konsisten per skill.
    """
    pairs = []
    for src in ("Reviu", "Pengawasan"):
        if src == verb:
            continue   # tidak perlu diganti kalau sama
        pairs += [
            (f"Dasar Pelaksanaan {src}",  f"Dasar Pelaksanaan {verb}"),
            (f"Tujuan dan Sasaran {src}", f"Tujuan dan Sasaran {verb}"),
            (f"Ruang Lingkup {src}",      f"Ruang Lingkup {verb}"),
            (f"Metodologi {src}",         f"Metodologi {verb}"),
            (f"Hasil {src}",              f"Hasil {verb}"),
            (f"Simpulan {src}",           f"Simpulan {verb}"),
            (f"Pelaksanaan {src}",        f"Pelaksanaan {verb}"),
        ]
    return pairs


# ── Metodologi ───────────────────────────────────────────────────────────────

def _build_metodologi_reviu(ctx: dict) -> str:
    """Metodologi reviu — keyakinan terbatas."""
    custom = ctx.get("metodologi") or ""
    if custom:
        return custom
    return (
        "Reviu dilaksanakan melalui penelaahan dokumen (desk review) atas dokumen "
        "yang disampaikan auditi, dipadukan dengan konfirmasi/klarifikasi terhadap "
        "hal-hal yang memerlukan penjelasan lebih lanjut."
    )


def _build_metodologi_evaluasi(ctx: dict) -> str:
    """Metodologi evaluasi — keyakinan terbatas."""
    custom = ctx.get("metodologi") or ""
    if custom:
        return custom
    return (
        "Evaluasi dilaksanakan melalui penelaahan dokumen, analisis data, dan "
        "klarifikasi kepada pihak terkait. Evaluasi mengacu pada kriteria yang "
        "ditetapkan dalam peraturan perundang-undangan dan pedoman yang berlaku."
    )


def _build_metodologi_audit_kinerja(ctx: dict) -> str:
    """Metodologi audit kinerja — reasonable assurance."""
    custom = ctx.get("metodologi") or ""
    if custom:
        return custom
    return (
        "Audit kinerja dilaksanakan melalui: (1) penelaahan dokumen perencanaan dan "
        "pelaporan kinerja; (2) pengujian substantif atas efektivitas dan efisiensi "
        "pelaksanaan program/kegiatan menggunakan paradigma Kondisi-Kriteria-Sebab-Akibat "
        "(KKSA); (3) wawancara dengan pemangku program dan pelaksana teknis; serta "
        "(4) analytical review atas data kinerja dan keuangan."
    )


def _build_metodologi_pemantauan(ctx: dict) -> str:
    """Metodologi pemantauan — tanpa level keyakinan."""
    custom = ctx.get("metodologi") or ""
    if custom:
        return custom
    return (
        "Pemantauan dilaksanakan melalui penelaahan dokumen kemajuan pelaksanaan, "
        "konfirmasi kepada pihak terkait, dan analisis realisasi dibandingkan dengan "
        "target yang ditetapkan. Pemantauan tidak memberikan opini keyakinan dan "
        "hanya menyampaikan informasi status pelaksanaan berdasarkan data yang tersedia "
        "pada saat pemantauan dilakukan."
    )


# ── Simpulan ─────────────────────────────────────────────────────────────────

def build_simpulan_reviu(kkp: dict, sa: dict) -> str:
    """Simpulan untuk semua jenis reviu (keyakinan terbatas)."""
    n = len(kkp.get("temuan", []))
    return (
        f"Berdasarkan hasil reviu terbatas yang kami lakukan, dokumen yang menjadi "
        f"obyek reviu secara umum telah memenuhi persyaratan utama. Namun, terdapat "
        f"{n} ({_terbilang(n)}) catatan reviu yang perlu ditindaklanjuti. "
        f"Dengan keterbatasan keyakinan reviu (limited assurance), kami tidak "
        f"memberikan opini final atas kelayakan substansi; hal tersebut menjadi "
        f"tanggung jawab pejabat yang berwenang sesuai peraturan perundang-undangan "
        f"yang berlaku."
    )


def build_simpulan_evaluasi(kkp: dict, sa: dict) -> str:
    """Simpulan untuk semua jenis evaluasi (keyakinan terbatas)."""
    n = len(kkp.get("temuan", []))
    return (
        f"Berdasarkan hasil evaluasi yang kami lakukan, terdapat {n} ({_terbilang(n)}) "
        f"catatan evaluasi yang perlu mendapat perhatian dan tindak lanjut dari pimpinan "
        f"auditi. Evaluasi dilaksanakan dengan keyakinan terbatas berdasarkan dokumen "
        f"dan data yang tersedia. Rekomendasi yang disampaikan bersifat korektif dan "
        f"konstruktif guna peningkatan kualitas tata kelola, manajemen risiko, dan "
        f"pengendalian intern organisasi."
    )


def build_simpulan_audit_kinerja(kkp: dict, sa: dict) -> str:
    """Simpulan untuk audit kinerja (reasonable assurance)."""
    n = len(kkp.get("temuan", []))
    return (
        f"Berdasarkan audit kinerja yang kami lakukan dengan tingkat keyakinan memadai "
        f"(reasonable assurance), program/kegiatan yang diaudit secara umum telah "
        f"dilaksanakan, namun terdapat {n} ({_terbilang(n)}) temuan yang memerlukan "
        f"perhatian dan tindak lanjut dari pimpinan auditi. Rekomendasi yang disampaikan "
        f"ditujukan untuk meningkatkan efektivitas dan efisiensi pencapaian tujuan "
        f"program/kegiatan sesuai dengan perencanaan yang telah ditetapkan."
    )


def build_gambaran_umum_kinerja(ctx: dict, kkp: dict) -> list:
    """Bab 2 LHA Kinerja — 4 sub-seksi sesuai SKILL.md (2.1–2.4)."""
    obyek = ctx.get("obyek") or "[Program]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    pen   = kkp.get("penugasan", {})
    return [
        (ctx.get("tujuan_program") or
         f"[DIISI — Tujuan dan desain program {obyek} TA {ta}: tujuan strategis, "
         f"landasan kebijakan, dan keluaran yang diharapkan.]",
         {"align": "justify"}),
        ("[DIISI — Logika intervensi: Input (anggaran, SDM) → Proses (pelaksanaan "
         "kegiatan utama) → Output (keluaran langsung) → Outcome (dampak/manfaat bagi "
         "penerima manfaat).]",
         {"align": "justify"}),
        (ctx.get("anggaran_gambaran") or
         f"[DIISI — Alokasi anggaran {obyek} TA {ta}: pagu, realisasi, dan proporsi "
         f"per komponen belanja. Sumber daya SDM yang terlibat.]",
         {"align": "justify"}),
        (ctx.get("mekanisme_pelaksanaan") or
         "[DIISI — Satuan kerja pelaksana, mekanisme pengelolaan, jadwal pelaksanaan "
         "kegiatan, dan mitra/pihak terkait yang terlibat.]",
         {"align": "justify"}),
    ]


def build_simpulan_audit_kinerja_blocks(kkp: dict, sa: dict) -> list:
    """Simpulan audit kinerja sebagai list of blocks (untuk expand_paragraph)."""
    return [(build_simpulan_audit_kinerja(kkp, sa), {"align": "justify"})]


# ── PEMANTAUAN builders shared ───────────────────────────────────────────────

def build_latar_belakang_umum(ctx: dict, kkp: dict, jenis_label: str) -> list:
    """A1. Latar Belakang generik untuk pemantauan."""
    obyek = ctx.get("obyek") or "[Objek]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Dalam rangka pengawasan intern berbasis risiko, Inspektorat Jenderal "
         f"Kementerian Komunikasi dan Digital melaksanakan {jenis_label} terhadap "
         f"{obyek} Tahun {ta}.", {"align": "justify"}),
    ]


def build_profil_pekerjaan_pengadaan(ctx: dict, kkp: dict) -> list:
    """B. Profil paket pekerjaan pengadaan."""
    return [
        ("[DIISI — Profil paket pengadaan: nama paket, nomor kontrak, nilai kontrak, "
         "nama penyedia, jangka waktu, dan ruang lingkup pekerjaan.]",
         {"italic": True}),
    ]


def build_status_pelaksanaan_pengadaan(ctx: dict, kkp: dict) -> list:
    """C. Status pelaksanaan — progres fisik vs pembayaran + warna."""
    return [
        ("[DIISI — Status pelaksanaan per tanggal cut-off: % progres fisik, "
         "% realisasi pembayaran, sisa waktu pelaksanaan, status (HIJAU/KUNING/MERAH).]",
         {"italic": True}),
    ]


def build_isu_pengadaan(kkp: dict) -> list:
    """D. Isu dan Permasalahan — dari temuan."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("Tidak terdapat isu atau permasalahan signifikan pada saat pemantauan "
                 "dilaksanakan.", {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or f"Isu {i}"
        cond  = t.get("kondisi") or ""
        krit  = t.get("kriteria") or ""
        akib  = t.get("akibat") or t.get("dampak") or ""
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if cond:
            blocks.append((f"Kondisi          : {cond}", {"size": 10}))
        if krit:
            blocks.append((f"Kriteria         : {krit}", {"size": 10}))
        if akib:
            blocks.append((f"Potensi Risiko   : {akib}", {"size": 10}))
        if rek:
            blocks.append((f"Rekomendasi      : {rek}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_simpulan_rekomendasi_pengadaan(kkp: dict) -> list:
    """G. Simpulan dan Rekomendasi pemantauan pengadaan."""
    n_isu = len(kkp.get("temuan", []))
    blocks = [
        (f"Berdasarkan pemantauan yang dilaksanakan, terdapat {n_isu} "
         f"({_terbilang(n_isu)}) isu yang perlu mendapat perhatian.",
         {"align": "justify"}),
    ]
    reks = [t.get("rekomendasi", "").strip() for t in kkp.get("temuan", [])
            if t.get("rekomendasi")]
    if reks:
        blocks.append(("Rekomendasi:", {"bold": True}))
        for i, r in enumerate(reks, 1):
            blocks.append((f"{i}. {r}", {"size": 10}))
    return blocks


def build_ringkasan_tlhp(ctx: dict, kkp: dict) -> list:
    """Seksi 1. Ringkasan Eksekutif TLHP."""
    obyek = ctx.get("obyek") or "[Objek]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    temuan = kkp.get("temuan", [])
    total  = len(temuan)
    selesai = len([t for t in temuan if t.get("status_tl", "").lower() == "selesai"])
    proses  = len([t for t in temuan if t.get("status_tl", "").lower() == "dalam proses"])
    belum   = total - selesai - proses
    return [
        (f"Pemantauan Tindak Lanjut Hasil Pengawasan (TLHP) terhadap {obyek} "
         f"dilaksanakan untuk memastikan rekomendasi hasil pengawasan telah "
         f"ditindaklanjuti secara memadai.", {"align": "justify"}),
        (f"Per tanggal cut-off pemantauan: dari {total} rekomendasi yang dipantau, "
         f"{selesai} telah selesai, {proses} dalam proses, dan {belum} belum "
         f"ditindaklanjuti.", {"align": "justify"}),
    ]


def build_statistik_tlhp(kkp: dict) -> list:
    """Seksi 2. Statistik umum TLHP."""
    temuan = kkp.get("temuan", [])
    total  = len(temuan)
    if not total:
        return [("[DIISI — Tabel statistik: total rekomendasi, status (Selesai / "
                 "Dalam Proses / Belum / Tidak Dapat Ditindaklanjuti), "
                 "dan persentase per status.]", {"italic": True})]
    from collections import Counter
    status_count = Counter(t.get("status_tl", "Belum Ditindaklanjuti") for t in temuan)
    blocks = [(f"Total rekomendasi yang dipantau: {total}", {"bold": True})]
    for status, count in status_count.most_common():
        pct = f"{100*count//total}%" if total else "0%"
        blocks.append((f"  {status}: {count} ({pct})", {"size": 10}))
    return blocks


def build_aging_per_pic(kkp: dict) -> list:
    """Seksi 3. Aging rekomendasi per PIC."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("[DIISI — Tabel aging rekomendasi per PIC: nama PIC, jumlah "
                 "rekomendasi, umur rata-rata (hari), dan status.]", {"italic": True})]
    from collections import defaultdict
    pic_map = defaultdict(list)
    for t in temuan:
        pic = t.get("pic") or t.get("penanggung_jawab") or "Tidak Diketahui"
        pic_map[pic].append(t)
    blocks = []
    for pic, items in sorted(pic_map.items()):
        n_selesai = len([i for i in items if "selesai" in i.get("status_tl","").lower()])
        blocks.append((f"{pic}: {len(items)} rekomendasi ({n_selesai} selesai)",
                       {"size": 10}))
    return blocks


def build_rekomendasi_kritis_tlhp(kkp: dict) -> list:
    """Seksi 4. Rekomendasi kritis > 365 hari belum ditindaklanjuti."""
    temuan = kkp.get("temuan", [])
    kritis = [t for t in temuan if int(t.get("umur_hari", 0) or 0) > 365
              and "selesai" not in t.get("status_tl", "").lower()]
    if not kritis:
        return [("Tidak terdapat rekomendasi yang telah melampaui 365 hari tanpa "
                 "tindak lanjut.", {"align": "justify"})]
    blocks = []
    for i, t in enumerate(kritis, 1):
        judul = t.get("judul_temuan") or t.get("kondisi") or f"Rekomendasi {i}"
        umur  = t.get("umur_hari", "> 365")
        rek   = t.get("rekomendasi") or "[rekomendasi]"
        blocks.append((f"{i}. {judul} ({umur} hari)", {"bold": True}))
        blocks.append((f"   Rekomendasi: {rek}", {"size": 10}))
    return blocks


def build_percepatan_tlhp(kkp: dict) -> list:
    """Seksi 5. Rekomendasi percepatan tindak lanjut."""
    reks = [t.get("rekomendasi", "").strip() for t in kkp.get("temuan", [])
            if t.get("rekomendasi") and
            "selesai" not in t.get("status_tl", "").lower()]
    if not reks:
        return [("Seluruh rekomendasi yang dipantau telah selesai ditindaklanjuti. "
                 "Tidak diperlukan rekomendasi percepatan.", {"align": "justify"})]
    blocks = [("Agar percepatan tindak lanjut segera dilaksanakan atas rekomendasi "
               "berikut:", {"align": "justify"})]
    for i, r in enumerate(reks, 1):
        blocks.append((f"{i}. {r}", {"size": 10}))
    return blocks


def build_ringkasan_status_umum(ctx: dict, kkp: dict) -> list:
    """E. Ringkasan Status pemantauan umum (% capaian + warna)."""
    temuan = kkp.get("temuan", [])
    n = len(temuan)
    if not temuan:
        return [("[DIISI — Tabel agregat status: item yang dipantau, target, "
                 "realisasi, % capaian, dan status warna (HIJAU/KUNING/MERAH).]",
                 {"italic": True})]
    hijau = len([t for t in temuan if t.get("status", "").upper() == "HIJAU"])
    kuning = len([t for t in temuan if t.get("status", "").upper() == "KUNING"])
    merah  = len([t for t in temuan if t.get("status", "").upper() == "MERAH"])
    return [
        (f"Dari {n} item yang dipantau: {hijau} HIJAU (on-track), "
         f"{kuning} KUNING (perlu perhatian), {merah} MERAH (kritis).",
         {"align": "justify"}),
    ]


def build_hasil_per_item_umum(kkp: dict) -> list:
    """F. Hasil pemantauan per item — fokus KUNING/MERAH."""
    temuan = kkp.get("temuan", [])
    items  = [t for t in temuan if t.get("status", "").upper() in ("KUNING", "MERAH")]
    if not items:
        items = temuan  # fallback: tampilkan semua
    if not items:
        return [("Tidak terdapat item yang berstatus KUNING atau MERAH.", {})]
    blocks = []
    for i, t in enumerate(items, 1):
        judul  = t.get("judul_temuan") or t.get("kondisi") or f"Item {i}"
        status = t.get("status") or "KUNING"
        target = t.get("target") or "[target]"
        real   = t.get("realisasi") or "[realisasi]"
        dev    = t.get("kondisi") or "[penyebab deviasi]"
        rek    = t.get("rekomendasi") or "[rekomendasi percepatan]"
        blocks.append((f"{i}. [{status}] {judul}", {"bold": True}))
        blocks.append((f"   Target: {target} | Realisasi: {real}", {"size": 10}))
        blocks.append((f"   Penyebab Deviasi: {dev}", {"size": 10}))
        blocks.append((f"   Rekomendasi     : {rek}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_rekomendasi_pemantauan_umum(kkp: dict) -> list:
    """G. Rekomendasi dan Tindakan Percepatan pemantauan umum."""
    reks = [t.get("rekomendasi", "").strip() for t in kkp.get("temuan", [])
            if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi tindakan percepatan untuk item KUNING/MERAH]",
                 {"italic": True})]
    return [(f"{i}. {r}", {"size": 10}) for i, r in enumerate(reks, 1)]


# ── REVIU builders ──────────────────────────────────────────────────────────

def build_gambaran_umum_reviu_pengadaan(ctx: dict, kkp: dict) -> list:
    """B. Gambaran umum paket pengadaan yang direviu."""
    return [
        ("[DIISI — Gambaran umum paket: nama paket, nilai pagu, HPS, metode pemilihan, "
         "sumber dana, dan ruang lingkup pekerjaan yang direviu.]",
         {"italic": True}),
    ]


def build_hasil_perencanaan_reviu_pengadaan(ctx: dict, kkp: dict) -> list:
    """C.1 Hasil reviu perencanaan pengadaan."""
    temuan = [t for t in kkp.get("temuan", [])
              if t.get("area", "").lower() in ("perencanaan", "c1", "")]
    if not temuan:
        return [("Hasil reviu perencanaan pengadaan tidak ditemukan catatan. "
                 "Perencanaan pengadaan telah memadai.", {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or f"Catatan {i}"
        cond  = t.get("kondisi") or ""
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if cond:
            blocks.append((f"Kondisi: {cond}", {"size": 10}))
        if rek:
            blocks.append((f"Rekomendasi: {rek}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_hasil_pemilihan_reviu_pengadaan(ctx: dict, kkp: dict) -> list:
    """C.2 Hasil reviu pemilihan penyedia."""
    temuan = [t for t in kkp.get("temuan", [])
              if t.get("area", "").lower() in ("pemilihan", "c2")]
    if not temuan:
        return [("Hasil reviu pemilihan penyedia tidak ditemukan catatan. "
                 "Proses pemilihan penyedia telah sesuai ketentuan.", {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or f"Catatan {i}"
        cond  = t.get("kondisi") or ""
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if cond:
            blocks.append((f"Kondisi: {cond}", {"size": 10}))
        if rek:
            blocks.append((f"Rekomendasi: {rek}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_simpulan_reviu(ctx: dict, kkp: dict, sa: dict) -> list:
    """D. Simpulan reviu (pengadaan / rka-kl / umum)."""
    n = len(kkp.get("temuan", []))
    obyek = ctx.get("obyek") or "[Objek reviu]"
    return [
        (f"Berdasarkan reviu yang telah dilaksanakan terhadap {obyek}, "
         f"terdapat {n} ({_terbilang(n)}) catatan yang perlu ditindaklanjuti.",
         {"align": "justify"}),
    ]


def build_rekomendasi_reviu(kkp: dict) -> list:
    """E. Rekomendasi reviu."""
    reks = [t.get("rekomendasi", "").strip() for t in kkp.get("temuan", [])
            if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi berdasarkan catatan hasil reviu]",
                 {"italic": True})]
    return [(f"{i}. {r}", {"size": 10}) for i, r in enumerate(reks, 1)]


def _build_aspek_reviu_rkakl(kkp: dict, kode_area: str, default_ok: str) -> list:
    """Helper: ambil temuan per aspek RKA-K/L berdasarkan kode area."""
    kode = kode_area.lower()
    temuan = [t for t in kkp.get("temuan", [])
              if t.get("area", "").lower().startswith(kode)]
    if not temuan:
        return [(default_ok, {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or f"Catatan {i}"
        cond  = t.get("kondisi") or ""
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if cond:
            blocks.append((f"Kondisi: {cond}", {"size": 10}))
        if rek:
            blocks.append((f"Rekomendasi: {rek}", {"size": 10}))
    return blocks


def build_hasil_reviu_umum_aspek(kkp: dict) -> list:
    """D. Hasil reviu umum — status per aspek (TERPENUHI / CATATAN / TIDAK TERPENUHI)."""
    aspek = kkp.get("aspek_reviu", [])
    if not aspek:
        return [("[DIISI — Status per aspek yang direviu: "
                 "TERPENUHI / TERPENUHI DENGAN CATATAN / TIDAK TERPENUHI]",
                 {"italic": True})]
    blocks = []
    SIMBOL = {"terpenuhi": "[OK]", "catatan": "[!]", "tidak terpenuhi": "[X]"}
    for a in aspek:
        nama   = a.get("nama", "[Aspek]")
        status = a.get("status", "catatan").lower()
        simbol = SIMBOL.get(status, "[!]")
        ket    = a.get("keterangan", "")
        blocks.append((f"{simbol} {nama}: {status.upper()}", {"bold": True}))
        if ket:
            blocks.append((f"   {ket}", {"size": 10}))
    return blocks


def build_catatan_rekomendasi_reviu_umum(kkp: dict) -> list:
    """E. Catatan dan rekomendasi reviu umum."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("Tidak terdapat catatan atau rekomendasi berdasarkan hasil reviu.",
                 {"align": "justify"})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or t.get("kondisi") or f"Catatan {i}"
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if rek:
            blocks.append((f"   Rekomendasi: {rek}", {"size": 10}))
    return blocks


def build_simpulan_reviu_umum(ctx: dict, kkp: dict) -> list:
    """F. Simpulan reviu umum."""
    return build_simpulan_reviu(ctx, kkp, {})


# ── akhir REVIU builders ─────────────────────────────────────────────────────


def build_simpulan_pemantauan(kkp: dict, sa: dict) -> str:
    """Simpulan untuk semua jenis pemantauan (tanpa level keyakinan)."""
    n = len(kkp.get("temuan", []))
    return (
        f"Berdasarkan pemantauan yang kami lakukan, terdapat {n} ({_terbilang(n)}) isu "
        f"atau permasalahan yang memerlukan perhatian dan tindak lanjut segera. "
        f"Laporan ini menyampaikan informasi status pelaksanaan berdasarkan data yang "
        f"tersedia pada saat pemantauan dilakukan dan tidak memberikan opini keyakinan. "
        f"Rekomendasi disampaikan sebagai masukan bagi pimpinan untuk pengambilan "
        f"keputusan dan percepatan pelaksanaan."
    )


# ============================================================
# EVALUASI-SAKIP builders  (LHE AKIP — I–V)
# ============================================================

_PREDIKAT_INTERPRETASI = {
    "AA": "Sangat Memuaskan",
    "A":  "Memuaskan",
    "BB": "Sangat Baik",
    "B":  "Baik",
    "CC": "Cukup",
    "C":  "Kurang",
    "D":  "Sangat Kurang",
}

_KOMPONEN_KEYS = [
    ("perencanaan",       "Perencanaan Kinerja",                        "30,00"),
    ("pengukuran",        "Pengukuran Kinerja",                         "30,00"),
    ("pelaporan",         "Pelaporan Kinerja",                          "15,00"),
    ("evaluasi_internal", "Evaluasi Akuntabilitas Kinerja Internal",    "25,00"),
]

# Kata kunci dalam judul/deskripsi untuk maping otomatis komponen
_KOMPONEN_KEYWORDS = {
    "perencanaan":       ["perencanaan", "renstra", "renja", "pk ", "perjanjian kinerja", "iku", "indikator"],
    "pengukuran":        ["pengukuran", "monev", "monitoring", "capaian"],
    "pelaporan":         ["pelaporan", "lkj", "laporan kinerja"],
    "evaluasi_internal": ["evaluasi", "akip", "akuntabilitas", "internal"],
}


def _read_penilaian_lke(pen_dir: Path) -> dict:
    """Coba baca _KKP/penilaian_lke.json. Return {} jika tidak ada."""
    path = pen_dir / "_KKP" / "penilaian_lke.json"
    if path.exists():
        try:
            import json as _json
            return _json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _guess_komponen(t: dict) -> str:
    """Tebak komponen dari field `komponen`, `sasaran_id`, atau teks judul/kondisi."""
    if t.get("komponen"):
        raw = t["komponen"].lower()
        for key in _KOMPONEN_KEYWORDS:
            if key in raw or any(kw in raw for kw in _KOMPONEN_KEYWORDS[key]):
                return key
        return t["komponen"].lower()
    # Coba sasaran_id: S01→perencanaan, S02→pengukuran, S03→pelaporan, S04→evaluasi
    sid = t.get("sasaran_id", "")
    mapping = {"S01": "perencanaan", "S02": "pengukuran",
               "S03": "pelaporan",   "S04": "evaluasi_internal"}
    if sid in mapping:
        return mapping[sid]
    # Fallback: cari kata kunci di judul
    haystack = (t.get("judul_temuan", "") + " " + t.get("kondisi", "")).lower()
    for key, keywords in _KOMPONEN_KEYWORDS.items():
        if any(kw in haystack for kw in keywords):
            return key
    return "perencanaan"  # default


def build_ringkasan_eksekutif_sakip(kkp: dict, ctx: dict, penilaian: dict) -> list:
    """Ringkasan Eksekutif 2–3 paragraf."""
    obyek  = ctx.get("obyek") or ctx.get("nama_auditi") or "[Objek Evaluasi]"
    ta     = ctx.get("tahun_anggaran") or "[Tahun]"
    total  = penilaian.get("total", {})
    nilai  = total.get("nilai")
    pred   = total.get("predikat", "[DIISI]")
    interp = total.get("interpretasi") or _PREDIKAT_INTERPRETASI.get(pred, "[DIISI interpretasi]")
    n_temuan = len(kkp.get("temuan", []))

    if nilai is not None:
        nilai_str = f"{nilai:.2f}".replace(".", ",")
        p1 = (f"Berdasarkan evaluasi Akuntabilitas Kinerja Instansi Pemerintah (AKIP) "
              f"yang dilaksanakan terhadap {obyek} Tahun {ta}, diperoleh nilai sebesar "
              f"{nilai_str} ({pred}) atau {interp}.", {"align": "justify"})
    else:
        p1 = (f"Berdasarkan evaluasi Akuntabilitas Kinerja Instansi Pemerintah (AKIP) "
              f"yang dilaksanakan terhadap {obyek} Tahun {ta}, hasil penilaian akan "
              f"disampaikan sesuai nilai Lembar Kerja Evaluasi (LKE) yang telah terisi.",
              {"align": "justify"})

    p2 = (f"Secara umum, {obyek} telah memiliki dokumen perencanaan, pengukuran, dan "
          f"pelaporan kinerja yang memadai. Namun demikian, masih terdapat "
          f"{n_temuan} catatan/Area of Improvement (AoI) yang perlu mendapat perhatian "
          f"untuk peningkatan akuntabilitas kinerja ke depan.",
          {"align": "justify"})
    return [p1, p2]


def build_dasar_hukum_sakip(ctx: dict, kkp: dict) -> list:
    """I.A Dasar Hukum — regulasi + ST."""
    nomor_st = kkp.get("penugasan", {}).get("nomor_st") or ctx.get("nomor_st") or "[Nomor ST]"
    tanggal_st = kkp.get("penugasan", {}).get("tanggal_st") or ctx.get("tanggal_st") or "[Tanggal ST]"
    lines = [
        ("1. Peraturan Presiden Nomor 29 Tahun 2014 tentang Sistem Akuntabilitas "
         "Kinerja Instansi Pemerintah;", {"size": 10}),
        ("2. Peraturan Menteri Pendayagunaan Aparatur Negara dan Reformasi Birokrasi "
         "Nomor 88 Tahun 2021 tentang Evaluasi Akuntabilitas Kinerja Instansi Pemerintah;",
         {"size": 10}),
        ("3. Peraturan Pemerintah Nomor 8 Tahun 2006 tentang Pelaporan Keuangan dan "
         "Kinerja Instansi Pemerintah;", {"size": 10}),
        (f"4. Surat Tugas Inspektur Jenderal Nomor {nomor_st} tanggal {tanggal_st}.",
         {"size": 10}),
    ]
    return lines


def build_tujuan_sakip(ctx: dict, sa: dict) -> list:
    """I.B Tujuan Evaluasi."""
    obyek = ctx.get("obyek") or "[Objek Evaluasi]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Evaluasi AKIP bertujuan untuk:", {"align": "justify"}),
        (f"1. Menilai tingkat akuntabilitas kinerja {obyek} dalam penyelenggaraan "
         f"pemerintahan Tahun {ta};", {"size": 10}),
        (f"2. Memberikan saran perbaikan untuk peningkatan kinerja dan penguatan "
         f"akuntabilitas kinerja {obyek}.", {"size": 10}),
    ]


def build_metodologi_sakip(ctx: dict) -> list:
    """I.D Metodologi Evaluasi AKIP."""
    return [
        ("Evaluasi dilakukan melalui:", {"align": "justify"}),
        ("1. Reviu dokumen perencanaan, pengukuran, dan pelaporan kinerja berdasarkan "
         "bukti dukung yang dikumpulkan auditi;", {"size": 10}),
        ("2. Wawancara/klarifikasi dengan pejabat dan pegawai terkait;", {"size": 10}),
        ("3. Pengisian Lembar Kerja Evaluasi (LKE) sesuai PermenPAN-RB Nomor 88 "
         "Tahun 2021 dengan penilaian predikat APIP per kriteria berdasarkan bukti dukung.",
         {"size": 10}),
    ]


def build_tindak_lanjut_sakip(ctx: dict, kkp: dict) -> list:
    """II. Tindak Lanjut — placeholder jika tidak ada data historis."""
    obyek = ctx.get("obyek") or "[Objek Evaluasi]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    try:
        ta_int = int(ta)
        ta_prev = str(ta_int - 1)
    except Exception:
        ta_prev = "[Tahun Sebelumnya]"

    return [
        ("A.  Tindak Lanjut atas Rekomendasi Kementerian PAN-RB", {"bold": True}),
        (f"[DIISI — Tindak lanjut atas rekomendasi KemenPAN-RB atas evaluasi AKIP "
         f"{ta_prev}, atau nyatakan tidak ada jika evaluasi baru pertama kali dilaksanakan]",
         {"italic": True}),
        ("B.  Tindak Lanjut atas Rekomendasi Inspektorat Jenderal", {"bold": True}),
        (f"[DIISI — Tindak lanjut atas rekomendasi Inspektorat Jenderal atas evaluasi "
         f"AKIP {ta_prev}, atau nyatakan tidak ada jika evaluasi baru pertama kali]",
         {"italic": True}),
    ]


def _build_nilai_str(penilaian: dict, key: str) -> str:
    komponen = penilaian.get("komponen", {})
    data = komponen.get(key, {})
    nilai = data.get("nilai")
    pred  = data.get("predikat", "")
    if nilai is not None:
        return f"{nilai:.2f}".replace(".", ",") + (f" ({pred})" if pred else "")
    return "[DIISI]"


def build_hasil_komponen_sakip(temuan_list: list, komponen_nama: str,
                                bobot: str, nilai_str: str) -> list:
    """Narasi per komponen: hal positif lalu catatan kekurangan."""
    blocks = [
        (f"(Nilai: {nilai_str} dari {bobot})", {"italic": True}),
        (f"{komponen_nama.split()[0]} telah terlaksana dengan memadai sesuai dengan "
         f"ketentuan yang berlaku. Dokumen dan bukti dukung yang diperlukan telah tersedia "
         f"dan dapat diverifikasi.", {"align": "justify"}),
    ]
    kekurangan = []
    for t in temuan_list:
        kondisi = t.get("kondisi") or t.get("judul_temuan") or ""
        if kondisi:
            kekurangan.append(kondisi.strip())

    if kekurangan:
        blocks.append(("Namun demikian, masih terdapat hal yang perlu diperhatikan, "
                        "antara lain:", {"align": "justify"}))
        for i, k in enumerate(kekurangan, 1):
            blocks.append((f"{i}. {k}", {"size": 10}))
    else:
        blocks.append(("Tidak terdapat catatan kekurangan yang signifikan pada komponen "
                        "ini berdasarkan bukti dukung yang diperoleh.", {"align": "justify"}))
    return blocks


def build_tabel_nilai_lke_blocks(penilaian: dict) -> list:
    """Tabel rekapitulasi sebagai blok teks (tabel Word diproses tersendiri)."""
    komponen_data = penilaian.get("komponen", {})
    if not komponen_data:
        return [("[DIISI — Tabel rekapitulasi nilai per komponen LKE]", {"italic": True})]
    lines = [("Rekapitulasi Nilai AKIP:", {"bold": True})]
    total_nilai = 0.0
    total_pred  = penilaian.get("total", {}).get("predikat", "")
    for key, nama, bobot in _KOMPONEN_KEYS:
        data  = komponen_data.get(key, {})
        nilai = data.get("nilai", 0.0)
        pred  = data.get("predikat", "")
        total_nilai += nilai if isinstance(nilai, (int, float)) else 0.0
        nilai_str = f"{nilai:.2f}".replace(".", ",") if isinstance(nilai, (int, float)) else "[DIISI]"
        lines.append((f"  {nama}: {nilai_str} / {bobot}  ({pred})", {"size": 10}))
    total_str = f"{total_nilai:.2f}".replace(".", ",")
    interp = penilaian.get("total", {}).get("interpretasi", "")
    interp_txt = f"  ({interp})" if interp else ""
    lines.append((f"  NILAI AKIP: {total_str} / 100,00  ({total_pred}){interp_txt}",
                  {"bold": True}))
    return lines


def build_rekomendasi_sakip(kkp: dict) -> list:
    """IV. Rekomendasi — diawali 'Agar'.

    Jika rekomendasi mengandung pola '... agar ...' (mis. 'Pimpinan X agar menyusun'),
    ubah menjadi 'Agar X menyusun' dengan menghapus 'agar' yang ada di tengah.
    """
    import re as _re
    temuan = kkp.get("temuan", [])
    reks = [t.get("rekomendasi", "").strip() for t in temuan if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi evaluasi AKIP, diawali kata 'Agar']",
                 {"italic": True})]
    blocks = []
    for i, rek in enumerate(reks, 1):
        if not rek.lower().startswith("agar"):
            # Hapus "agar" internal jika ada, lalu prefix "Agar"
            rek_clean = _re.sub(r'\s+agar\s+', ' ', rek, flags=_re.IGNORECASE, count=1).strip()
            rek = "Agar " + rek_clean[0].lower() + rek_clean[1:]
        blocks.append((f"{i}. {rek}", {"size": 10}))
    return blocks


# ============================================================
# EVALUASI-SPIP builders  (LHE PK SPIP Terintegrasi — I–VI)
# ============================================================

_KOMPONEN_KEYS_SPIP = [
    ("penetapan_tujuan", "Penetapan Tujuan",   "40%"),
    ("struktur_proses",  "Struktur dan Proses", "30%"),
    ("pencapaian_tujuan","Pencapaian Tujuan",   "30%"),
]

_KOMPONEN_KEYWORDS_SPIP = {
    "penetapan_tujuan":  ["penetapan tujuan", "sasaran strategis", "kualitas sasaran",
                          "sasaran oPD", "rencana strategis", "kinerja sasaran"],
    "struktur_proses":   ["lingkungan pengendalian", "penilaian risiko", "kegiatan pengendalian",
                          "informasi", "komunikasi", "pemantauan", "struktur", "proses",
                          "unsur i", "unsur ii", "unsur iii", "unsur iv", "unsur v",
                          "integritas", "kompetensi", "kepemimpinan", "identifikasi risiko",
                          "analisis risiko"],
    "pencapaian_tujuan": ["pencapaian tujuan", "opini", "capaian outcome", "capaian output",
                          "pelaporan keuangan", "aset", "ketaatan", "efektivitas pencapaian"],
}

_MATURITAS_SPIP = {
    1: "Rintisan",
    2: "Berkembang",
    3: "Terdefinisi",
    4: "Terkelola dan Terukur",
    5: "Optimum",
}


def _guess_komponen_spip(t: dict) -> str:
    """Tebak komponen SPIP dari field komponen, sasaran_id, atau teks."""
    raw_k = (t.get("komponen") or "").lower()
    if raw_k:
        for key, kws in _KOMPONEN_KEYWORDS_SPIP.items():
            if any(kw in raw_k for kw in [key] + kws):
                return key
    sid = t.get("sasaran_id", "")
    sid_map = {"S01": "penetapan_tujuan", "S02": "struktur_proses", "S03": "pencapaian_tujuan"}
    if sid in sid_map:
        return sid_map[sid]
    haystack = (t.get("judul_temuan", "") + " " + t.get("kondisi", "")).lower()
    for key, kws in _KOMPONEN_KEYWORDS_SPIP.items():
        if any(kw in haystack for kw in kws):
            return key
    return "penetapan_tujuan"


def _tingkat_maturitas_spip(nilai) -> str:
    """Konversi nilai skor ke nama tingkat maturitas SPIP."""
    try:
        v = float(nilai)
    except (TypeError, ValueError):
        return "[DIISI]"
    if v < 2.0:
        return "Rintisan"
    if v < 3.0:
        return "Berkembang"
    if v < 4.0:
        return "Terdefinisi"
    if v < 4.5:
        return "Terkelola dan Terukur"
    return "Optimum"


def build_ringkasan_eksekutif_spip(kkp: dict, ctx: dict, penilaian: dict) -> list:
    """Ringkasan Eksekutif LHE PK SPIP (nilai PK vs PM + tingkat maturitas)."""
    obyek = ctx.get("obyek") or ctx.get("nama_auditi") or "[Objek PK]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    total = penilaian.get("total", {})
    nilai_pk  = total.get("nilai_pk")
    nilai_pm  = total.get("nilai_pm")
    tingkat   = total.get("tingkat_maturitas") or (
        _tingkat_maturitas_spip(nilai_pk) if nilai_pk is not None else "[DIISI]")
    n_aoi = len([t for t in kkp.get("temuan", []) if t])

    if nilai_pk is not None:
        pk_str = f"{float(nilai_pk):.2f}".replace(".", ",")
        pm_str = (f"{float(nilai_pm):.2f}".replace(".", ",")
                  if nilai_pm is not None else "[DIISI PM]")
        p1 = (f"Berdasarkan Penjaminan Kualitas (PK) atas Penilaian Mandiri (PM) "
              f"Maturitas Penyelenggaraan SPIP Terintegrasi terhadap {obyek} Tahun "
              f"{ta}, diperoleh nilai maturitas SPIP hasil PK sebesar {pk_str} "
              f"(Level — {tingkat}), dibandingkan dengan nilai PM sebesar {pm_str}.",
              {"align": "justify"})
    else:
        p1 = (f"Berdasarkan Penjaminan Kualitas (PK) atas Penilaian Mandiri (PM) "
              f"Maturitas Penyelenggaraan SPIP Terintegrasi terhadap {obyek} Tahun "
              f"{ta}, nilai maturitas SPIP akan diperoleh setelah LKE SPIP terisi.",
              {"align": "justify"})

    p2 = (f"Tim PK mengidentifikasi {n_aoi} Area of Improvement (AoI) yang perlu "
          f"ditindaklanjuti untuk peningkatan maturitas SPIP, Manajemen Risiko Indeks "
          f"(MRI), dan Indeks Efektivitas Pengendalian Korupsi (IEPK) {obyek}.",
          {"align": "justify"})
    return [p1, p2]


def build_dasar_hukum_spip(ctx: dict, kkp: dict) -> list:
    """I.A Dasar Hukum PK SPIP — PerBPKP 5/2021 + ST."""
    nomor_st   = kkp.get("penugasan", {}).get("nomor_st") or ctx.get("nomor_st") or "[Nomor ST]"
    tanggal_st = kkp.get("penugasan", {}).get("tanggal_st") or ctx.get("tanggal_st") or "[Tanggal ST]"
    return [
        ("1. Peraturan Pemerintah Nomor 60 Tahun 2008 tentang Sistem Pengendalian "
         "Intern Pemerintah;", {"size": 10}),
        ("2. Peraturan Kepala BPKP Nomor 5 Tahun 2021 tentang Penilaian Maturitas "
         "Penyelenggaraan Sistem Pengendalian Intern Pemerintah Terintegrasi pada "
         "Kementerian/Lembaga/Pemerintah Daerah;", {"size": 10}),
        ("3. Peraturan Kepala BPKP terkait Manajemen Risiko Indeks (MRI) dan Indeks "
         "Efektivitas Pengendalian Korupsi (IEPK);", {"size": 10}),
        (f"4. Surat Tugas Inspektur Jenderal Nomor {nomor_st} tanggal {tanggal_st}.",
         {"size": 10}),
    ]


def build_tujuan_spip(ctx: dict, sa: dict) -> list:
    """I.B Tujuan PK SPIP."""
    obyek = ctx.get("obyek") or "[Objek PK]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        ("Penjaminan Kualitas (PK) Maturitas SPIP Terintegrasi bertujuan untuk:",
         {"align": "justify"}),
        (f"1. Memastikan proses Penilaian Mandiri (PM) oleh {obyek} telah dilaksanakan "
         f"sesuai ketentuan PerBPKP 5/2021;", {"size": 10}),
        (f"2. Menilai kewajaran Nilai PK secara independen berdasarkan analisis bukti "
         f"dokumen, dan membandingkannya dengan Nilai PM;", {"size": 10}),
        (f"3. Mengidentifikasi Area of Improvement (AoI) untuk peningkatan maturitas "
         f"SPIP Terintegrasi {obyek} Tahun {ta}.", {"size": 10}),
    ]


def build_ruang_lingkup_spip(ctx: dict, sa: dict) -> list:
    """I.C Ruang Lingkup PK SPIP."""
    obyek = ctx.get("obyek") or "[Objek PK]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    satker = ctx.get("satker_pk") or "satker wajib sesuai PerBPKP 5/2021"
    return [
        (f"PK dilaksanakan terhadap PM Maturitas SPIP Terintegrasi {obyek} Tahun {ta}, "
         f"mencakup penilaian atas tiga komponen:", {"align": "justify"}),
        ("1. Penetapan Tujuan (bobot 40%) — kualitas sasaran strategis dan strategi "
         "pencapaiannya;", {"size": 10}),
        ("2. Struktur dan Proses (bobot 30%) — 25 subunsur dalam 5 unsur SPIP "
         "(Lingkungan Pengendalian, Penilaian Risiko, Kegiatan Pengendalian, Informasi "
         "dan Komunikasi, Pemantauan);", {"size": 10}),
        ("3. Pencapaian Tujuan (bobot 30%) — efektivitas, keandalan pelaporan keuangan, "
         "pengamanan aset, dan ketaatan perundang-undangan.", {"size": 10}),
        (f"Satker yang menjadi objek PK adalah {satker}.", {"align": "justify"}),
    ]


def build_metodologi_spip(ctx: dict) -> list:
    """I.D Metodologi PK SPIP."""
    return [
        ("PK dilaksanakan melalui:", {"align": "justify"}),
        ("1. Reviu draf Laporan Hasil Penilaian Mandiri (LHPM) beserta Kertas Kerja "
         "Asesor (KKE/KK) yang digunakan asesor manajemen;", {"size": 10}),
        ("2. Analisis dokumen pendukung per unsur/subunsur SPIP (SOP, SK, laporan, "
         "notulen, data kinerja) yang tersedia di folder dokumen;", {"size": 10}),
        ("3. Pengisian kolom Nilai PK pada Lembar Kerja Evaluasi (LKE) SPIP secara "
         "independen berdasarkan bukti dokumen;", {"size": 10}),
        ("4. Validasi proses PM dan penilaian kewajaran skor menggunakan KK LEAD I "
         "sesuai PerBPKP Nomor 5 Tahun 2021.", {"size": 10}),
    ]


def build_tindak_lanjut_spip(ctx: dict, kkp: dict) -> list:
    """II. Tindak Lanjut atas hasil PK tahun sebelumnya."""
    obyek = ctx.get("obyek") or "[Objek PK]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    try:
        ta_prev = str(int(ta) - 1)
    except Exception:
        ta_prev = "[Tahun Sebelumnya]"
    return [
        ("A.  Tindak Lanjut atas Rekomendasi PK SPIP Tahun Sebelumnya",
         {"bold": True}),
        (f"[DIISI — Tindak lanjut {obyek} atas AoI dan rekomendasi dari LHE PK SPIP "
         f"Tahun {ta_prev}, atau nyatakan 'Tidak ada — PK pertama kali dilaksanakan'.]",
         {"italic": True}),
        ("B.  Tindak Lanjut atas Temuan BPKP dan Aparat Pengawas Lainnya",
         {"bold": True}),
        (f"[DIISI — Temuan/rekomendasi BPKP atau inspektorat lain terkait SPIP yang "
         f"belum ditindaklanjuti {obyek}.]", {"italic": True}),
    ]


def build_tabel_nilai_spip_blocks(penilaian: dict) -> list:
    """Rekapitulasi nilai maturitas SPIP PK vs PM sebagai blok teks."""
    komponen_data = penilaian.get("komponen", {})
    if not komponen_data:
        return [("[DIISI — Tabel rekapitulasi nilai maturitas SPIP (Nilai PK vs PM "
                 "per komponen dan total)]", {"italic": True})]
    lines = [("Rekapitulasi Nilai Maturitas SPIP Terintegrasi:", {"bold": True})]
    total_pk = 0.0
    for key, nama, bobot in _KOMPONEN_KEYS_SPIP:
        data     = komponen_data.get(key, {})
        nilai_pk = data.get("nilai_pk", 0.0)
        nilai_pm = data.get("nilai_pm", 0.0)
        total_pk += float(nilai_pk) if isinstance(nilai_pk, (int, float)) else 0.0
        pk_str   = f"{float(nilai_pk):.2f}".replace(".", ",") if isinstance(nilai_pk, (int, float)) else "[DIISI]"
        pm_str   = f"{float(nilai_pm):.2f}".replace(".", ",") if isinstance(nilai_pm, (int, float)) else "[DIISI]"
        lines.append((f"  {nama} ({bobot}): Nilai PK = {pk_str} | Nilai PM = {pm_str}",
                      {"size": 10}))
    total_str  = f"{total_pk:.2f}".replace(".", ",")
    tingkat    = _tingkat_maturitas_spip(total_pk) if total_pk > 0 else "[DIISI]"
    lines.append((f"  NILAI MATURITAS SPIP (PK): {total_str} — Level {tingkat}",
                  {"bold": True}))
    return lines


def build_hasil_komponen_spip(temuan_list: list, komponen_nama: str,
                               nilai_str: str) -> list:
    """Narasi hasil PK per komponen SPIP."""
    blocks = [(f"(Nilai PK: {nilai_str})", {"italic": True})]
    if not temuan_list:
        blocks.append((f"Berdasarkan hasil PK atas komponen {komponen_nama}, "
                       "tidak terdapat Area of Improvement yang signifikan. "
                       "Nilai PK dikonfirmasi sesuai dengan bukti dokumen yang tersedia.",
                       {"align": "justify"}))
        return blocks
    blocks.append((f"Berdasarkan hasil PK atas komponen {komponen_nama}, "
                   "terdapat hal-hal yang perlu mendapat perhatian:",
                   {"align": "justify"}))
    for i, t in enumerate(temuan_list, 1):
        kondisi = (t.get("kondisi") or t.get("judul_temuan") or "").strip()
        if kondisi:
            blocks.append((f"{i}. {kondisi}", {"size": 10}))
    return blocks


def build_aoi_spip(kkp: dict) -> list:
    """IV. Area of Improvement — dari temuan dengan AoI."""
    temuan = kkp.get("temuan", [])
    aoi_items = [t for t in temuan if t]
    if not aoi_items:
        return [("[DIISI — Area of Improvement berdasarkan subunsur/parameter dengan "
                 "Nilai PK ≤ 3 atau yang direvisi turun dari PM]", {"italic": True})]
    blocks = []
    for i, t in enumerate(aoi_items, 1):
        judul   = t.get("judul_temuan") or t.get("kondisi") or f"AoI {i}"
        kondisi = t.get("kondisi") or ""
        rek     = t.get("rekomendasi") or ""
        blocks.append((f"AoI {i} — {judul}", {"bold": True}))
        if kondisi:
            blocks.append((f"Kondisi  : {kondisi}", {"size": 10}))
        if rek:
            blocks.append((f"Perbaikan: {rek}", {"size": 10}))
    return blocks


def build_rekomendasi_spip(kkp: dict) -> list:
    """V. Rekomendasi PK SPIP — diawali 'Agar'."""
    import re as _re
    temuan = kkp.get("temuan", [])
    reks   = [t.get("rekomendasi", "").strip() for t in temuan if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi perbaikan SPIP, diawali kata 'Agar']",
                 {"italic": True})]
    blocks = []
    for i, rek in enumerate(reks, 1):
        if not rek.lower().startswith("agar"):
            rek_clean = _re.sub(r'\s+agar\s+', ' ', rek, flags=_re.IGNORECASE, count=1).strip()
            rek = "Agar " + rek_clean[0].lower() + rek_clean[1:]
        blocks.append((f"{i}. {rek}", {"size": 10}))
    return blocks



# ============================================================
# EVALUASI-MANAJEMEN-RISIKO builders  (LHE MR — A–H)
# ============================================================

def build_dasar_mr(ctx: dict, kkp: dict) -> list:
    """A. Dasar Pelaksanaan Evaluasi MR."""
    pen  = kkp.get("penugasan", {})
    n_st = pen.get("nomor_st") or ctx.get("nomor_st") or "[Nomor ST]"
    t_st = pen.get("tanggal_st") or ctx.get("tanggal_st") or "[Tanggal ST]"
    return [
        ("1. Program Kerja Pengawasan Tahunan (PKPT) Inspektorat Jenderal "
         f"Tahun {ctx.get('tahun_anggaran') or '[Tahun]'};", {"size": 10}),
        (f"2. Surat Tugas Inspektur Jenderal Nomor {n_st} tanggal {t_st} "
         "tentang Evaluasi Efektivitas Manajemen Risiko.", {"size": 10}),
    ]


def build_tujuan_mr(ctx: dict) -> list:
    """B. Tujuan Evaluasi MR."""
    obyek = ctx.get("obyek") or "[Objek]"
    return [
        ("a. Memberikan keyakinan terbatas atas pelaksanaan manajemen risiko di "
         f"lingkungan {obyek};", {"size": 10}),
        ("b. Menilai kesesuaian penerapan manajemen risiko dengan Pedoman Menteri "
         "Komunikasi dan Informatika Nomor 6 Tahun 2017 dan ISO 31000:2018;",
         {"size": 10}),
        ("c. Mengidentifikasi kelemahan dan area perbaikan dalam penerapan "
         "manajemen risiko guna mendukung pencapaian tujuan organisasi.",
         {"size": 10}),
    ]


def build_ruang_lingkup_mr(ctx: dict, sa: dict) -> list:
    """C. Ruang Lingkup Evaluasi MR."""
    obyek = ctx.get("obyek") or "[Objek]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Evaluasi mencakup pelaksanaan manajemen risiko di lingkungan {obyek} "
         f"Tahun {ta}, meliputi aspek:", {"align": "justify"}),
        ("1. Struktur organisasi manajemen risiko (KMR dan UPR);", {"size": 10}),
        ("2. Penetapan konteks manajemen risiko (Formulir 1);", {"size": 10}),
        ("3. Kualitas profil dan peta risiko (Formulir 2);", {"size": 10}),
        ("4. Rencana penanganan risiko (Formulir 3);", {"size": 10}),
        ("5. Pemantauan dan pelaporan manajemen risiko secara triwulanan "
         "(Formulir 4 dan 5);", {"size": 10}),
        ("6. Tingkat kematangan penerapan manajemen risiko (TKPMR).", {"size": 10}),
    ]


def build_gambaran_umum_mr(ctx: dict, kkp: dict) -> list:
    """E. Gambaran Umum kondisi MR saat ini."""
    obyek = ctx.get("obyek") or "[Objek]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Pelaksanaan manajemen risiko di lingkungan {obyek} Tahun {ta} "
         f"berpedoman pada Peraturan Menteri Komunikasi dan Informatika Nomor 6 "
         f"Tahun 2017 tentang Manajemen Risiko di Lingkungan Kementerian "
         f"Komunikasi dan Informatika.", {"align": "justify"}),
        ("[DIISI — Kondisi umum implementasi MR: struktur KMR/UPR yang sudah "
         "ditetapkan, sistem informasi yang digunakan, periode dokumen MR yang "
         "menjadi objek evaluasi, dan kondisi umum kepatuhan terhadap pedoman.]",
         {"italic": True}),
    ]


def build_hasil_evaluasi_mr(kkp: dict) -> list:
    """F. Hasil Evaluasi — catatan naratif KKSA bernomor."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("[DIISI — Catatan evaluasi manajemen risiko sesuai format KKSA]",
                 {"italic": True})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul    = t.get("judul_temuan") or t.get("kondisi") or f"Catatan {i}"
        kondisi  = t.get("kondisi") or ""
        kriteria = t.get("kriteria") or ""
        sebab    = t.get("sebab") or "Tidak ditemukan penyebab berdasarkan bukti yang tersedia."
        akibat   = t.get("akibat") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if kondisi:
            blocks.append(("Kondisi:", {"bold": True, "size": 10}))
            blocks.append((kondisi, {"size": 10}))
        if kriteria:
            blocks.append(("Kriteria:", {"bold": True, "size": 10}))
            blocks.append((kriteria, {"size": 10}))
        if sebab:
            blocks.append(("Sebab:", {"bold": True, "size": 10}))
            blocks.append((sebab, {"size": 10}))
        if akibat:
            blocks.append(("Akibat:", {"bold": True, "size": 10}))
            blocks.append((akibat, {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_rekomendasi_mr(kkp: dict) -> list:
    """G. Rekomendasi — dikompilasi terpisah dari F, diawali kata kerja aktif."""
    temuan = kkp.get("temuan", [])
    reks   = [t.get("rekomendasi", "").strip() for t in temuan if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi perbaikan manajemen risiko, diawali kata "
                 "kerja aktif (Menyusun, Melengkapi, Melaksanakan, dst.)]",
                 {"italic": True})]
    blocks = []
    for i, rek in enumerate(reks, 1):
        blocks.append((f"{i}. {rek}", {"size": 10}))
    return blocks


# ============================================================
# EVALUASI-REFORMASI-BIROKRASI builders  (LHEI RB — 1–4)
# ============================================================

def build_tujuan_evaluasi_rb(ctx: dict) -> list:
    """Tujuan evaluasi RB — sebelum seksi 1."""
    obyek = ctx.get("obyek") or "[Unit Kerja]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Evaluasi pelaksanaan Reformasi Birokrasi di lingkungan {obyek} Tahun "
         f"{ta} dilaksanakan dengan tujuan untuk:",  {"align": "justify"}),
        ("a. Memastikan Rencana Aksi Reformasi Birokrasi berisi solusi pemecahan "
         "masalah tata kelola yang nyata, berkualitas baik, dan layak sebagai "
         "pedoman pelaksanaan RB;", {"size": 10}),
        ("b. Menilai kesesuaian pelaksanaan Rencana Aksi RB berdasarkan 4 dimensi "
         "(Ketepatan Pelaksanaan, Ketercapaian Output, Kualitas Pelaksanaan, "
         "Kesesuaian Waktu); dan", {"size": 10}),
        ("c. Memberikan saran perbaikan untuk peningkatan kualitas pelaksanaan "
         "Reformasi Birokrasi.", {"size": 10}),
    ]


def build_metode_evaluasi_rb(ctx: dict) -> list:
    """Metode evaluasi RB."""
    return [
        ("Evaluasi dilaksanakan dengan cara mempelajari dan menelaah dokumen "
         "Rencana Aksi Reformasi Birokrasi dan Road Map RB untuk mendapatkan "
         "informasi tentang 4 dimensi evaluasi, serta melakukan koordinasi "
         "dengan PIC kegiatan terkait.", {"align": "justify"}),
    ]


def build_gambaran_umum_rb(ctx: dict, kkp: dict) -> list:
    """Seksi 1 — Gambaran Umum Pelaksanaan RB (General + Tematik)."""
    obyek = ctx.get("obyek") or "[Unit Kerja]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        ("a.  Reformasi Birokrasi General", {"bold": True}),
        (f"[DIISI — Apakah terdapat RB General di lingkungan {obyek}? Jika tidak: "
         f"'Tidak terdapat RB General di Lingkungan {obyek} Tahun {ta}.']",
         {"italic": True}),
        ("b.  Reformasi Birokrasi Tematik", {"bold": True}),
        ("[DIISI — Tabel tema RB Tematik yang dijalankan, meliputi nama tema, "
         "indikator, dan target output per tema sesuai Renaksi.]",
         {"italic": True}),
    ]


def build_dampak_rb(ctx: dict, kkp: dict) -> list:
    """Seksi 2 — Analisis Dampak RB Tematik per tema."""
    temuan = kkp.get("temuan", [])
    tema_items = [t for t in temuan if (t.get("komponen") or "").lower().startswith("tematik")]
    if not tema_items:
        return [("[DIISI — Narasi analisis dampak per tema RB Tematik. "
                 "Sertakan data kuantitatif (angka, persentase, target vs realisasi) "
                 "yang dikutip dari laporan capaian, PPATK, Bareskrim, atau data internal.]",
                 {"italic": True})]
    blocks = []
    for i, t in enumerate(tema_items, ord("a")):
        tema  = t.get("komponen") or f"Tema {chr(i)}"
        naras = t.get("kondisi") or t.get("judul_temuan") or "[DIISI narasi dampak tema]"
        blocks.append((f"{chr(i)}.  {tema}", {"bold": True}))
        blocks.append((naras, {"align": "justify"}))
    return blocks


def build_hasil_renaksi_rb(ctx: dict, kkp: dict) -> list:
    """Seksi 3 — Hasil evaluasi pelaksanaan Rencana Aksi (4 dimensi)."""
    obyek = ctx.get("obyek") or "[Unit Kerja]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    # Simpulan per dimensi berdasarkan temuan
    temuan  = kkp.get("temuan", [])
    n_total = len(temuan)
    n_sesuai = len([t for t in temuan if t.get("nilai", "") == "Sesuai"])
    n_tidak  = n_total - n_sesuai
    blocks = [
        ("[DIISI — Tabel rekapitulasi hasil evaluasi per Rencana Aksi × 4 dimensi "
         "(Ketepatan Pelaksanaan / Ketercapaian Output / Kualitas Pelaksanaan / "
         "Kesesuaian Waktu) dengan persentase Sesuai dan Tidak Sesuai per dimensi.]",
         {"italic": True}),
        (f"Berdasarkan hasil evaluasi tersebut, dapat disimpulkan pelaksanaan "
         f"Reformasi Birokrasi di lingkungan {obyek} Tahun {ta}:",
         {"align": "justify"}),
        ("1. Pelaksanaan komponen kegiatan [sesuai/tidak sesuai] dengan maksud "
         "kegiatan saat penyusunan Rencana Aksi;", {"size": 10}),
        ("2. Output aksi [telah/belum] tercapai sesuai target triwulan;",
         {"size": 10}),
        ("3. Pelaksanaan aksi [telah/belum] direncanakan, dilaksanakan, dan "
         "dilaporkan dengan baik;", {"size": 10}),
        ("4. Realisasi waktu [sesuai/tidak sesuai] target waktu Rencana Aksi.",
         {"size": 10}),
    ]
    return blocks


# ============================================================
# EVALUASI-UMUM builders  (LHE A–I: dimensi/skor/predikat)
# ============================================================

def build_gambaran_umum_evaluasi_umum(ctx: dict, kkp: dict) -> list:
    """D. Gambaran Umum Objek Evaluasi."""
    obyek = ctx.get("obyek") or "[Objek Evaluasi]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (ctx.get("ringkasan_obyek") or
         f"[DIISI — Gambaran umum {obyek} yang menjadi objek evaluasi Tahun {ta}: "
         f"tugas/fungsi, ruang lingkup kegiatan, dan kondisi umum pelaksanaan yang "
         f"relevan dengan dimensi yang dievaluasi.]",
         {"align": "justify"}),
    ]


def build_hasil_evaluasi_umum_skor(ctx: dict, kkp: dict) -> list:
    """E.1 Skor per Dimensi — dari penilaian_lke atau placeholder."""
    return [
        ("[DIISI — Tabel skor per dimensi evaluasi: nama dimensi, bobot (%), "
         "skor yang diperoleh, dan interpretasinya (Sangat Baik/Baik/Cukup/Kurang)]",
         {"italic": True}),
    ]


def build_predikat_evaluasi_umum(ctx: dict, kkp: dict) -> list:
    """E.2 Predikat dan Posisi keseluruhan."""
    return [
        ("[DIISI — Predikat keseluruhan hasil evaluasi (Sangat Baik/Baik/Cukup/Kurang) "
         "berdasarkan total skor agregat, dan posisi relatif dibandingkan target atau "
         "instansi sejenis.]",
         {"italic": True}),
    ]


def build_analisis_per_dimensi(ctx: dict, kkp: dict) -> list:
    """E.3 Analisis per Dimensi — dari temuan yang dikelompokkan."""
    temuan  = kkp.get("temuan", [])
    if not temuan:
        return [("[DIISI — Narasi analisis per dimensi evaluasi: kondisi yang ditemukan, "
                 "gap terhadap kriteria, dan hal-hal yang telah baik.]",
                 {"italic": True})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or t.get("kondisi") or f"Catatan {i}"
        cond  = t.get("kondisi") or ""
        krit  = t.get("kriteria") or ""
        akib  = t.get("akibat") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        if cond:
            blocks.append((f"Kondisi  : {cond}", {"size": 10}))
        if krit:
            blocks.append((f"Kriteria : {krit}", {"size": 10}))
        if akib:
            blocks.append((f"Akibat   : {akib}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_temuan_evaluasi_umum(ctx: dict, kkp: dict) -> list:
    """F. Temuan & Catatan — catatan bernomor KKSA."""
    temuan = kkp.get("temuan", [])
    if not temuan:
        return [("[DIISI — Temuan dan catatan evaluasi yang memerlukan perhatian "
                 "dan tindak lanjut]", {"italic": True})]
    blocks = []
    for i, t in enumerate(temuan, 1):
        judul = t.get("judul_temuan") or f"Catatan {i}"
        sebab = t.get("sebab") or "Tidak ditemukan penyebab berdasarkan bukti yang tersedia."
        rek   = t.get("rekomendasi") or ""
        blocks.append((f"{i}. {judul}", {"bold": True}))
        for lbl, val in [("Kondisi", t.get("kondisi")),
                         ("Kriteria", t.get("kriteria")),
                         ("Sebab", sebab),
                         ("Akibat", t.get("akibat"))]:
            if val:
                blocks.append((f"{lbl}: {val}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_rekomendasi_evaluasi_umum(kkp: dict) -> list:
    """G. Rekomendasi — dikompilasi sistem-level terpisah dari F."""
    reks = [t.get("rekomendasi", "").strip() for t in kkp.get("temuan", [])
            if t.get("rekomendasi")]
    if not reks:
        return [("[DIISI — Rekomendasi perbaikan sistem-level agar:]",
                 {"italic": True})]
    return [(f"{i}. {r}", {"size": 10}) for i, r in enumerate(reks, 1)]


def build_simpulan_evaluasi_umum(ctx: dict, kkp: dict) -> list:
    """H. Simpulan evaluasi umum (keyakinan terbatas)."""
    n = len(kkp.get("temuan", []))
    obyek = ctx.get("obyek") or "[Objek]"
    ta    = ctx.get("tahun_anggaran") or "[Tahun]"
    return [
        (f"Berdasarkan evaluasi yang dilaksanakan dengan keyakinan terbatas, "
         f"{obyek} Tahun {ta} secara umum telah menjalankan fungsinya, namun "
         f"terdapat {n} ({_terbilang(n)}) catatan yang memerlukan perhatian dan "
         f"tindak lanjut guna peningkatan kualitas tata kelola, manajemen risiko, "
         f"dan pencapaian tujuan organisasi.",
         {"align": "justify"}),
    ]


# ============================================================
# AUDIT-KINERJA builders (Bab 1–6 LHA Kinerja)
# ============================================================

def build_dasar_kinerja(ctx: dict, kkp: dict, args) -> list:
    """1.2 Dasar Penugasan — sama dengan audit umum."""
    return build_dasar_audit_umum(ctx, kkp, args)


def build_tujuan_kinerja(ctx: dict, sa: dict) -> list:
    """1.3 Tujuan dan Sasaran Audit Kinerja (format baku: 1. Tujuan... 2. Sasaran a,b,c)."""
    tujuan = sa.get("tujuan_penugasan") or ctx.get("tujuan") or ""
    obyek = kkp_obyek(ctx)
    ta = ctx.get("tahun_anggaran") or ""
    tujuan_text = tujuan or (
        f"untuk menilai efektivitas dan efisiensi pelaksanaan program/kegiatan "
        f"{obyek} Tahun Anggaran {ta}"
    )
    blocks = [(f"1. Tujuan dari dilaksanakannya audit adalah {tujuan_text}.", {"align": "justify"})]
    sasaran = sa.get("sasaran", [])
    if not sasaran:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah [SASARAN — diisi dari Kartu Penugasan]",
                       {"align": "justify"}))
    elif len(sasaran) == 1:
        desc = sasaran[0].get("deskripsi", "[SASARAN — diisi dari Kartu Penugasan]")
        blocks.append((f"2. Sasaran dari dilaksanakannya audit adalah {desc}", {"align": "justify"}))
    else:
        blocks.append(("2. Sasaran dari dilaksanakannya audit adalah:", {"align": "justify"}))
        for i, s in enumerate(sasaran):
            desc = s.get("deskripsi", "")
            if desc:
                blocks.append((f"    {chr(ord('a') + i)}. {desc}", {"align": "justify", "indent": 0.5}))
    return blocks


def build_pertanyaan_audit_kinerja(ctx: dict, sa: dict) -> list:
    """1.4 Pertanyaan Audit — diturunkan dari sasaran."""
    sasaran = sa.get("sasaran", [])
    if not sasaran:
        return [("Pertanyaan audit kinerja disusun oleh Ketua Tim berdasarkan "
                 "Kerangka Penugasan (KP).", {"align": "justify"})]
    blocks = [("Pertanyaan audit kinerja yang harus dijawab adalah:", {"align": "justify"})]
    for i, s in enumerate(sasaran, 1):
        desc = s.get("deskripsi", "")
        if desc:
            # Jadikan pertanyaan
            q = desc.rstrip(".")
            blocks.append((f"{i}. Apakah {q[0].lower()}{q[1:]} telah terlaksana secara "
                            f"efektif dan efisien?", {"size": 10}))
    return blocks


def build_komposisi_tim_blocks(ctx: dict) -> list:
    """1.7 Komposisi Tim dan Jangka Waktu."""
    tim = ctx.get("tim", [])
    periode = ctx.get("periode_pelaksanaan") or "[DIISI]"
    blocks = [(f"Jangka waktu pelaksanaan: {periode}.", {"align": "justify"}),
              ("Susunan tim audit kinerja:", {"align": "justify"})]
    if not tim:
        blocks.append(("[DIISI — Komposisi tim audit kinerja]", {"size": 10}))
        return blocks
    for m in tim:
        nama = m.get("nama", "")
        jabatan = m.get("jabatan", "")
        jabfung = m.get("jabfung", "")
        nip = m.get("nip", "")
        blocks.append((f"- {nama} (NIP {nip}) — {jabatan} / {jabfung}", {"size": 10}))
    return blocks


def build_ruang_lingkup_kinerja(ctx: dict, sa: dict) -> list:
    """1.5 Ruang Lingkup."""
    rl = ctx.get("ruang_lingkup") or ""
    sasaran = sa.get("sasaran", [])
    if rl:
        return [(rl, {"align": "justify"})]
    descs = [s.get("deskripsi", "") for s in sasaran if s.get("deskripsi")]
    if descs:
        blk = [("Ruang lingkup audit kinerja meliputi:", {"align": "justify"})]
        for i, d in enumerate(descs, 1):
            blk.append((f"{i}. {d}", {"size": 10}))
        return blk
    obyek = kkp_obyek(ctx)
    ta = ctx.get("tahun_anggaran") or ""
    return [(f"Ruang lingkup audit kinerja mencakup pelaksanaan program/kegiatan {obyek} "
             f"Tahun Anggaran {ta} sesuai sasaran dalam Kerangka Penugasan.",
             {"align": "justify"})]


def _filter_temuan_by_dimensi(kkp: dict, dimensi_target: str) -> dict:
    """Return salinan kkp hanya dengan temuan yang cocok dimensi."""
    import copy
    kkp2 = copy.deepcopy(kkp)
    all_t = kkp2.get("temuan", [])
    if dimensi_target == "efektivitas":
        # efektivitas = semua temuan yang BUKAN efisiensi (atau tidak ada dimensi)
        kkp2["temuan"] = [t for t in all_t
                          if t.get("dimensi", "efektivitas") != "efisiensi"]
    else:
        kkp2["temuan"] = [t for t in all_t
                          if t.get("dimensi", "") == "efisiensi"]
    return kkp2


def build_hasil_efektivitas(kkp: dict, sa: dict, rek: dict) -> list:
    """4.1 Efektivitas."""
    kkp2 = _filter_temuan_by_dimensi(kkp, "efektivitas")
    if not kkp2.get("temuan"):
        return [("Tidak terdapat catatan terkait efektivitas pencapaian target "
                 "pada periode audit ini.", {"align": "justify"})]
    return build_hasil_reviu_blocks(kkp2, sa, rek, "audit-kinerja")


def build_hasil_efisiensi(kkp: dict, sa: dict, rek: dict) -> list:
    """4.2 Efisiensi."""
    kkp2 = _filter_temuan_by_dimensi(kkp, "efisiensi")
    if not kkp2.get("temuan"):
        return [("Tidak terdapat catatan terkait efisiensi penggunaan sumber daya "
                 "pada periode audit ini.", {"align": "justify"})]
    return build_hasil_reviu_blocks(kkp2, sa, rek, "audit-kinerja")


def build_rekomendasi_kinerja(kkp: dict, rek: dict) -> list:
    """Bab 6 Rekomendasi — reuse builder audit umum."""
    return build_rekomendasi_audit_umum(kkp, rek)


def kkp_obyek(ctx: dict) -> str:
    return ctx.get("obyek") or ctx.get("nama_auditi") or "[Objek Audit]"


# ============================================================
# KONSULTASI builders (konsultansi-umum & konsultasi-pengadaan)
# ============================================================

# ── konsultansi-umum A-G builders ───────────────────────────────────────────

def build_dasar_konsultansi_umum(ctx: dict, kkp: dict, args) -> list:
    """A. Dasar — ND permintaan + ST."""
    nd    = ctx.get("nd_permintaan") or getattr(args, "dasar_permintaan", None) or "{{DASAR_PERMINTAAN}}"
    nomor = ctx.get("nomor_st") or "{{NOMOR_ST}}"
    tgl   = ctx.get("tanggal_st") or "{{TANGGAL_ST}}"
    return [
        (f"Surat Tugas Inspektur Jenderal Nomor {nomor} tanggal {tgl} tentang "
         f"Konsultansi terhadap {{NAMA_AUDITI}}.", {"align": "justify"}),
        (f"Dasar permintaan: {nd}.", {"align": "justify"}),
    ]


def build_pertanyaan_konsultansi_umum(kkp: dict, sa: dict) -> list:
    """B. Pertanyaan — daftar pertanyaan yang dijawab."""
    # Coba dari pertanyaan_terformulasi (JSON schema konsultansi)
    pertanyaan = kkp.get("pertanyaan_terformulasi", [])
    if pertanyaan:
        return [(f"{p.get('id','')}.  {p.get('teks','')}", {"align": "justify"})
                for p in pertanyaan]
    # Fallback ke sasaran atau tujuan
    tujuan = sa.get("tujuan_penugasan") or ""
    if tujuan:
        return [(tujuan, {"align": "justify"})]
    blocks = []
    for i, s in enumerate(sa.get("sasaran", []), 1):
        blocks.append((f"{i}. {s['deskripsi']}", {"align": "justify"}))
    if not blocks:
        blocks.append(("[DIISI — Daftar pertanyaan yang dikonsultasikan]",
                       {"italic": True}))
    return blocks


def build_dasar_hukum_konsultansi_umum(kkp: dict) -> list:
    """C. Dasar Hukum — kompilasi referensi dari kkp.dasar_hukum."""
    dh = kkp.get("dasar_hukum", [])
    if dh:
        return [(f"{d.get('id','')}.  {d.get('sumber','')} — {d.get('pasal','')}",
                 {"size": 10})
                for d in dh]
    return [("[DIISI — Peraturan dan dasar hukum yang relevan dengan pertanyaan]",
             {"italic": True})]


def build_telaah_konsultansi_umum(kkp: dict) -> list:
    """D. Telaah / Analisis — narasi analisis per pertanyaan."""
    pendapat_list = kkp.get("pendapat", [])
    if pendapat_list:
        blocks = []
        for p in pendapat_list:
            qid      = p.get("pertanyaan_id", "")
            analisis = p.get("analisis", "")
            if qid:
                blocks.append((f"Pertanyaan {qid}:", {"bold": True}))
            if analisis:
                blocks.append((analisis, {"align": "justify"}))
            blocks.append(("", {}))
        return blocks
    # Fallback: pakai temuan sebagai pertanyaan/analisis
    temuan = kkp.get("temuan", [])
    if temuan:
        blocks = []
        for i, t in enumerate(temuan, 1):
            judul    = t.get("judul_temuan") or f"Pertanyaan {i}"
            kondisi  = t.get("kondisi") or ""
            kriteria = t.get("kriteria") or ""
            blocks.append((f"{i}. {judul}", {"bold": True}))
            if kondisi:
                blocks.append((kondisi, {"align": "justify"}))
            if kriteria:
                blocks.append((f"Ketentuan: {kriteria}", {"size": 10}))
            blocks.append(("", {}))
        return blocks
    return [("[DIISI — Telaah dan analisis per pertanyaan berdasarkan dasar hukum]",
             {"italic": True})]


def build_pendapat_konsultansi_umum(kkp: dict) -> list:
    """E. Pendapat / Saran — jawaban ringkas per pertanyaan."""
    pendapat_list = kkp.get("pendapat", [])
    if pendapat_list:
        blocks = []
        for p in pendapat_list:
            qid      = p.get("pertanyaan_id", "")
            pendapat = p.get("pendapat", "")
            if qid:
                blocks.append((f"Pertanyaan {qid}:", {"bold": True}))
            if pendapat:
                blocks.append((f"Kami berpendapat bahwa {pendapat}", {"align": "justify"}))
            blocks.append(("", {}))
        return blocks
    # Fallback: ambil rekomendasi dari temuan sebagai pendapat
    temuan = kkp.get("temuan", [])
    if temuan:
        blocks = []
        for i, t in enumerate(temuan, 1):
            judul    = t.get("judul_temuan") or f"Pertanyaan {i}"
            rekomend = t.get("rekomendasi") or "[DIISI — pendapat/saran]"
            blocks.append((f"{i}. {judul}:", {"bold": True}))
            blocks.append((f"Kami menyarankan agar {rekomend}", {"align": "justify"}))
            blocks.append(("", {}))
        return blocks
    return [("[DIISI — Pendapat/saran per pertanyaan dalam bahasa advisory (berpendapat bahwa... / menyarankan agar...)]",
             {"italic": True})]


def build_asumsi_batasan_konsultansi_umum(kkp: dict, ctx: dict) -> list:
    """F. Asumsi & Batasan — eksplisit menyebutkan apa yang tidak dijawab."""
    asumsi_custom = ctx.get("asumsi_batasan") or kkp.get("asumsi_batasan") or ""
    if asumsi_custom:
        return [(asumsi_custom, {"align": "justify"})]
    return [
        ("Pendapat ini diberikan berdasarkan informasi yang disampaikan dalam "
         "permintaan konsultasi dan dokumen pendukung yang diterima. Apabila "
         "terdapat informasi tambahan atau perubahan fakta yang belum kami "
         "pertimbangkan, pendapat ini dapat berubah.",
         {"align": "justify"}),
        ("Pendapat ini bersifat tidak mengikat dan tidak menggantikan kewenangan "
         "pejabat yang berwenang dalam pengambilan keputusan operasional.",
         {"align": "justify"}),
        ("Hal-hal di luar pertanyaan yang dikonsultasikan tidak termasuk dalam "
         "ruang lingkup memo ini.",
         {"align": "justify"}),
    ]


# ── legacy builders (tetap dipakai konsultasi-pengadaan) ─────────────────────

def build_pertanyaan_konsultasi(kkp: dict, sa: dict) -> list:
    """I. PERTANYAAN / PERMASALAHAN."""
    tujuan = sa.get("tujuan_penugasan") or ""
    if tujuan:
        return [(tujuan, {"align": "justify"})]
    blocks = []
    for i, s in enumerate(sa.get("sasaran", []), start=1):
        blocks.append((f"{i}. {s['deskripsi']}", {"align": "justify"}))
    if not blocks:
        blocks.append(("[PERTANYAAN KONSULTASI — diisi dari Kartu Penugasan]",
                       {"align": "justify", "italic": True}))
    return blocks


def build_analisis_konsultasi(kkp: dict, sa: dict, rek: dict) -> list:
    """III. ANALISIS / PENDAPAT — narasi per butir konsultasi."""
    blocks = []
    for t in kkp.get("temuan", []):
        tid = t.get("id_temuan", "")
        blocks.append((f"{tid}. {t['judul_temuan']}", {"bold": True}))
        blocks.append(("Kondisi:", {"bold": True}))
        blocks.append((t.get("kondisi") or "—", {"align": "justify", "indent": 0.3}))
        blocks.append(("Kriteria / Ketentuan:", {"bold": True}))
        blocks.append((t.get("kriteria") or "—", {"align": "justify", "indent": 0.3}))
        blocks.append(("Pendapat / Saran:", {"bold": True}))
        saran = (rek.get(tid) or t.get("rekomendasi")
                 or "[DIISI — pendapat/saran konsultasi]")
        blocks.append((saran, {"align": "justify", "indent": 0.3}))
        blocks.append(("", {}))
    if not blocks:
        blocks.append(("[ANALISIS — diisi berdasarkan dokumen dan pertanyaan konsultasi]",
                       {"align": "justify", "italic": True}))
    return blocks


def build_simpulan_konsultasi(kkp: dict, sa: dict, rek: dict) -> str:
    """IV. KESIMPULAN DAN SARAN — advisory, tidak mengikat (legacy)."""
    n = len(kkp.get("temuan", []))
    return (
        f"Berdasarkan penelaahan kami terhadap permasalahan yang dikonsultasikan, "
        f"kami berpendapat bahwa terdapat {n} ({_terbilang(n)}) hal yang perlu "
        f"diperhatikan sebagaimana diuraikan di atas. Kami menyarankan agar pimpinan "
        f"mempertimbangkan saran tersebut sesuai kewenangan masing-masing. "
        f"Memo ini bersifat advisory dan tidak mengikat secara hukum; pelaksanaan "
        f"keputusan akhir tetap menjadi tanggung jawab pejabat yang berwenang."
    )


# ── konsultasi-pengadaan (v3.0) builders ────────────────────────────────────

def build_kegiatan_pendampingan(kkp: dict) -> list:
    """Bab I — Log kegiatan pendampingan yang telah diselesaikan."""
    kegiatan = kkp.get("kegiatan_pendampingan", [])
    if not kegiatan:
        # Fallback: coba temuan sebagai kegiatan
        kegiatan = kkp.get("temuan", [])
    if not kegiatan:
        return [("[DIISI — Daftar kegiatan pendampingan yang telah diselesaikan]",
                 {"italic": True})]
    blocks = []
    for i, k in enumerate(kegiatan, 1):
        tgl   = k.get("tanggal") or k.get("tanggal_temuan") or "[tanggal]"
        jenis = k.get("jenis_kegiatan") or k.get("judul_temuan") or "[jenis kegiatan]"
        pihak = k.get("pihak_didampingi") or "[pihak]"
        desk  = k.get("deskripsi") or k.get("kondisi") or "[deskripsi]"
        hasil = k.get("hasil") or k.get("rekomendasi") or "[hasil]"
        blocks.append((f"Kegiatan {i}: {jenis}", {"bold": True}))
        blocks.append((f"  Tanggal         : {tgl}", {"size": 10}))
        blocks.append((f"  Pihak Didampingi: {pihak}", {"size": 10}))
        blocks.append((f"  Deskripsi       : {desk}", {"size": 10}))
        blocks.append((f"  Hasil           : {hasil}", {"size": 10}))
        blocks.append(("", {}))
    return blocks


def build_dokumen_pendukung_pendampingan(kkp: dict) -> list:
    """Bab I (lanjutan) — Dokumen pendukung per kegiatan."""
    kegiatan = kkp.get("kegiatan_pendampingan", []) or kkp.get("temuan", [])
    if not kegiatan:
        return [("[DIISI — Dokumen pendukung per kegiatan (notulen, draft dokumen, dll)]",
                 {"italic": True})]
    blocks = []
    for i, k in enumerate(kegiatan, 1):
        tgl  = k.get("tanggal") or k.get("tanggal_temuan") or "[tanggal]"
        docs = k.get("dokumen_pendukung") or k.get("dokumen_sumber") or []
        if isinstance(docs, list) and docs:
            blocks.append((f"Kegiatan {i} ({tgl}):", {"bold": True}))
            for d in docs:
                ref = d.get("nama_file") or d.get("nama") or str(d)
                blocks.append((f"  • {ref}", {"size": 10}))
        elif isinstance(docs, str) and docs:
            blocks.append((f"Kegiatan {i} ({tgl}): {docs}", {"size": 10}))
    if not blocks:
        blocks.append(("[Dokumen pendukung belum tercatat]", {"italic": True}))
    return blocks


def build_tindak_lanjut_pendampingan(kkp: dict) -> list:
    """Bab II — Hal yang masih memerlukan tindak lanjut auditi."""
    kegiatan = kkp.get("kegiatan_pendampingan", []) or kkp.get("temuan", [])
    items = [(i, k) for i, k in enumerate(kegiatan, 1)
             if k.get("tindak_lanjut") or k.get("akibat")]
    if not items:
        return [("Seluruh kegiatan pendampingan telah diselesaikan. Tidak terdapat "
                 "hal yang masih memerlukan tindak lanjut dari auditi.",
                 {"align": "justify"})]
    blocks = []
    for i, k in items:
        tgl   = k.get("tanggal") or k.get("tanggal_temuan") or "[tanggal]"
        jenis = k.get("jenis_kegiatan") or k.get("judul_temuan") or f"Kegiatan {i}"
        tl    = k.get("tindak_lanjut") or k.get("akibat") or "[tindak lanjut]"
        blocks.append((f"{i}. {jenis} ({tgl}):", {"bold": True}))
        blocks.append((f"   {tl}", {"size": 10}))
    return blocks


def build_kesimpulan_pendampingan(kkp: dict, ctx: dict) -> list:
    """Bab III — Kesimpulan pendampingan."""
    kegiatan = kkp.get("kegiatan_pendampingan", []) or kkp.get("temuan", [])
    n = len(kegiatan)
    obyek = ctx.get("obyek") or "[unit kerja]"
    n_tl  = len([k for k in kegiatan if k.get("tindak_lanjut") or k.get("akibat")])
    return [
        (f"Inspektorat II telah melaksanakan {n} ({_terbilang(n)}) kegiatan "
         f"pendampingan pengadaan barang/jasa kepada {obyek}.",
         {"align": "justify"}),
        (f"Dari kegiatan tersebut, terdapat {n_tl} ({_terbilang(n_tl)}) hal yang "
         f"masih memerlukan tindak lanjut dari auditi sebagaimana diuraikan pada Bab II."
         if n_tl else
         "Seluruh agenda pendampingan telah diselesaikan dan tidak terdapat hal yang "
         "masih memerlukan tindak lanjut.",
         {"align": "justify"}),
        ("Pendampingan ini bersifat advisory dan preventif. Tidak memberikan keyakinan "
         "dan tidak mengikat keputusan pejabat yang berwenang. Rekomendasi yang "
         "disampaikan selama pendampingan semata-mata merupakan masukan untuk "
         "dipertimbangkan oleh pimpinan {{NAMA_AUDITI}}.",
         {"align": "justify"}),
    ]


def main() -> int:
    ap = argparse.ArgumentParser(description="Render LHP DOCX dari template + temuan.json")
    ap.add_argument("--penugasan", required=True)
    ap.add_argument("--rekomendasi-file", default=None,
                    help="Path JSON {id_temuan: 'rekomendasi text'}")
    ap.add_argument("--judul", default=None, help="Judul LHR (mis. 'Pengadaan DC/DRC PSrE Induk Tahun 2026')")
    ap.add_argument("--auditi", default=None, help="Nama auditi (default: dari obyek)")
    ap.add_argument("--penerima", default=None,
                    help="Penerima LHP — Yth. siapa laporan dikirim (tanyakan ke ketua tim)")
    ap.add_argument("--dasar-permintaan", default=None,
                    help="Dasar permintaan reviu (override default ST-only)")
    ap.add_argument("--gambaran-umum", default=None, help="Paragraf gambaran umum pengadaan")
    ap.add_argument("--tanggal-exit-meeting", default=None)
    # ── Item opsional yang ditanyakan ke ketua tim saat render ────────────────
    ap.add_argument("--tanggal-nota-dinas", default=None,
                    help="Tanggal Nota Dinas pengantar LHP, misal: '20 Februari 2026'")
    ap.add_argument("--nomor-nota-dinas", default=None,
                    help="Nomor Nota Dinas pengantar LHP")
    ap.add_argument("--survei", default=None,
                    help="URL link survei kepuasan (opsional — disisipkan di Nota Dinas jika diisi)")
    ap.add_argument("--tembusan", default=None,
                    help="Daftar tembusan dipisah titik-koma (opsional — gunakan default jika kosong), "
                         "misal: 'Inspektur Jenderal; Sekretaris Itjen; Arsip.'")
    # ─────────────────────────────────────────────────────────────────────────
    ap.add_argument("--template", default=None,
                    help="Override template path (default: templates/_skeleton-lhp/template-lhp-[jenis].docx)")
    ap.add_argument("--out", default=None,
                    help="Output path (default: _LHP/LHP-SUBSTANSI-[nomor-st-slug].docx)")
    args = ap.parse_args()

    pen_dir = Path(args.penugasan)
    if not pen_dir.exists():
        sys.stderr.write(f"Folder tidak ada: {pen_dir}\n"); return 1

    mapping, kkp, sa, rek, ctx = build_mapping(pen_dir, args)
    jenis = kkp["penugasan"]["jenis_pengawasan"]

    # Resolve template
    if args.template:
        tpl_path = Path(args.template)
    else:
        # Cari skeleton template di dua kemungkinan path (relative ke cwd dan absolute)
        candidates = [
            Path("templates/_skeleton-lhp") / f"template-lhp-{jenis}.docx",
            Path(__file__).resolve().parent.parent / "templates/_skeleton-lhp" / f"template-lhp-{jenis}.docx",
        ]
        tpl_path = next((p for p in candidates if p.exists()), None)
    if tpl_path is None or not tpl_path.exists():
        sys.stderr.write(f"Template skeleton tidak ditemukan untuk jenis '{jenis}'\n")
        return 1

    # Resolve output
    if args.out:
        out_path = Path(args.out)
    else:
        slug = re.sub(r"[^A-Za-z0-9]+", "-", mapping["NOMOR_ST"]).strip("-")
        out_path = pen_dir / "_LHP" / f"LHP-SUBSTANSI-{slug}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Load template
    doc = Document(str(tpl_path))

    # 1a. Override METODOLOGI_REVIU sebelum replace_in_doc (skill-specific)
    if jenis == "audit-kinerja":
        mapping["METODOLOGI_REVIU"] = _build_metodologi_audit_kinerja(ctx)
    elif jenis.startswith("evaluasi"):
        mapping["METODOLOGI_REVIU"] = _build_metodologi_evaluasi(ctx)
    elif jenis.startswith("pemantauan"):
        mapping["METODOLOGI_REVIU"] = _build_metodologi_pemantauan(ctx)
    elif jenis.startswith("reviu"):
        mapping["METODOLOGI_REVIU"] = _build_metodologi_reviu(ctx)
    # audit-umum, audit-pengadaan, konsultasi: punya ekspansi sendiri

    # 1b. Replace simple placeholders (berlaku untuk semua jenis)
    replace_in_doc(doc, mapping)

    # 2. Expand list/loop placeholders — cabang per jenis
    if jenis == "audit-umum":
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_audit_umum_audit(ctx, sa))
        replace_in_doc(doc, {"C_RUANG_LINGKUP": ctx.get("ruang_lingkup") or "[DIISI dari context.md]"})
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_METODOLOGI}}", build_metodologi_audit_umum(ctx, sa))
        replace_in_doc(doc, {"E_GAMBARAN_UMUM": (
            args.gambaran_umum or ctx.get("ringkasan_obyek") or "[DIISI — gambaran umum objek audit]")})
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_HASIL_AUDIT}}", build_hasil_reviu_blocks(kkp, sa, rek, jenis, include_rek=False))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{G_REKOMENDASI}}", build_rekomendasi_audit_umum(kkp, rek))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{H_APRESIASI}}", build_apresiasi_audit_umum(mapping["PENERIMA_LHP"]))
        expand_placeholder_to_matriks_table(doc, "{{LAMPIRAN_1_MATRIKS_TEMUAN}}", kkp, rek)
        expand_paragraph_placeholder_to_blocks(
            doc, "{{LAMPIRAN_2_DOKUMEN_SUMBER}}", build_daftar_dokumen_audit_umum(kkp))
    elif jenis == "audit-pengadaan":
        # ── LHA Pengadaan (Bab 1–6 sesuai SKILL.md) ─────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB1_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB1_TUJUAN}}", build_tujuan_sasaran_pengadaan(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB1_RUANG_LINGKUP}}", build_ruang_lingkup_pengadaan(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB2_GAMBARAN_UMUM}}", [(
                getattr(args, "gambaran_umum", None) or
                ctx.get("ringkasan_obyek") or
                "[DIISI — gambaran umum objek audit]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB3_METODOLOGI}}", build_metodologi_pengadaan(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB4_HASIL_AUDIT}}", build_hasil_reviu_blocks(kkp, sa, rek, jenis, include_rek=False))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB5_TEMUAN_REKOMENDASI}}", build_temuan_rekomendasi_pengadaan(kkp, rek))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB6_KESIMPULAN}}", build_kesimpulan_audit_pengadaan(kkp, ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{LAMPIRAN_1_DAFTAR_DOKUMEN}}", build_daftar_dokumen_audit_umum(kkp))
        expand_placeholder_to_matriks_table(doc, "{{LAMPIRAN_2_MATRIKS_TEMUAN}}", kkp, rek)
    elif jenis == "evaluasi-sakip":
        # ── LHE AKIP  (Ringkasan Eksekutif + I–V) ─────────────────────────────
        penilaian = _read_penilaian_lke(pen_dir)

        expand_paragraph_placeholder_to_blocks(
            doc, "{{RINGKASAN_EKSEKUTIF}}", build_ringkasan_eksekutif_sakip(kkp, ctx, penilaian))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_A_DASAR_HUKUM}}", build_dasar_hukum_sakip(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_B_TUJUAN}}", build_tujuan_sakip(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_D_METODOLOGI}}", build_metodologi_sakip(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{II_TINDAK_LANJUT}}", build_tindak_lanjut_sakip(ctx, kkp))

        # Tabel nilai — teks dulu, bisa diganti tabel Word kalau penilaian_lke.json ada
        expand_paragraph_placeholder_to_blocks(
            doc, "{{TABEL_NILAI_LKE}}", build_tabel_nilai_lke_blocks(penilaian))

        # Hasil per komponen
        all_temuan = kkp.get("temuan", [])
        komponen_temuan = {key: [] for key, _, _ in _KOMPONEN_KEYS}
        for t in all_temuan:
            k = _guess_komponen(t)
            if k not in komponen_temuan:
                k = "perencanaan"
            komponen_temuan[k].append(t)

        for key, nama, bobot in _KOMPONEN_KEYS:
            nilai_str = _build_nilai_str(penilaian, key)
            placeholder_key = {
                "perencanaan":       "III_A_PERENCANAAN",
                "pengukuran":        "III_B_PENGUKURAN",
                "pelaporan":         "III_C_PELAPORAN",
                "evaluasi_internal": "III_D_EVALUASI_INTERNAL",
            }[key]
            expand_paragraph_placeholder_to_blocks(
                doc, f"{{{{{placeholder_key}}}}}",
                build_hasil_komponen_sakip(komponen_temuan[key], nama, bobot, nilai_str))

        expand_paragraph_placeholder_to_blocks(
            doc, "{{IV_REKOMENDASI}}", build_rekomendasi_sakip(kkp))

    elif jenis == "evaluasi-spip":
        # ── LHE PK SPIP Terintegrasi (Ringkasan Eksekutif + I–VI) ───────────
        penilaian = _read_penilaian_lke(pen_dir)

        expand_paragraph_placeholder_to_blocks(
            doc, "{{RINGKASAN_EKSEKUTIF}}", build_ringkasan_eksekutif_spip(kkp, ctx, penilaian))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_A_DASAR_HUKUM}}", build_dasar_hukum_spip(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_B_TUJUAN}}", build_tujuan_spip(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_C_RUANG_LINGKUP}}", build_ruang_lingkup_spip(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{I_D_METODOLOGI}}", build_metodologi_spip(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{II_TINDAK_LANJUT}}", build_tindak_lanjut_spip(ctx, kkp))

        expand_paragraph_placeholder_to_blocks(
            doc, "{{TABEL_NILAI_SPIP}}", build_tabel_nilai_spip_blocks(penilaian))

        # Hasil per komponen (III.A, III.B, III.C)
        all_temuan = kkp.get("temuan", [])
        komponen_temuan_spip = {key: [] for key, _, _ in _KOMPONEN_KEYS_SPIP}
        for t in all_temuan:
            k = _guess_komponen_spip(t)
            if k not in komponen_temuan_spip:
                k = "penetapan_tujuan"
            komponen_temuan_spip[k].append(t)

        ph_map_spip = {
            "penetapan_tujuan":  "III_A_PENETAPAN_TUJUAN",
            "struktur_proses":   "III_B_STRUKTUR_PROSES",
            "pencapaian_tujuan": "III_C_PENCAPAIAN_TUJUAN",
        }
        for key, nama, _ in _KOMPONEN_KEYS_SPIP:
            data    = penilaian.get("komponen", {}).get(key, {})
            nilai   = data.get("nilai_pk") or data.get("nilai")
            ns      = (f"{float(nilai):.2f}".replace(".", ",")
                       if isinstance(nilai, (int, float)) else "[DIISI]")
            expand_paragraph_placeholder_to_blocks(
                doc, f"{{{{{ph_map_spip[key]}}}}}",
                build_hasil_komponen_spip(komponen_temuan_spip[key], nama, ns))

        expand_paragraph_placeholder_to_blocks(
            doc, "{{IV_AOI}}", build_aoi_spip(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{V_REKOMENDASI}}", build_rekomendasi_spip(kkp))

    elif jenis == "evaluasi-manajemen-risiko":
        # ── LHE Manajemen Risiko (A–H, format catatan KKSA) ──────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_mr(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_mr(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_RUANG_LINGKUP}}", build_ruang_lingkup_mr(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_METODOLOGI}}", [(_build_metodologi_evaluasi(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_GAMBARAN_UMUM}}", build_gambaran_umum_mr(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_HASIL_EVALUASI}}", build_hasil_evaluasi_mr(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{G_REKOMENDASI}}", build_rekomendasi_mr(kkp))

    elif jenis == "evaluasi-reformasi-birokrasi":
        # ── LHEI RB (1–4: Gambaran Umum, Dampak, Renaksi, Penutup) ──────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{TUJUAN_EVALUASI_RB}}", build_tujuan_evaluasi_rb(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{METODE_EVALUASI_RB}}", build_metode_evaluasi_rb(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{SEKSI_1_GAMBARAN_UMUM}}", build_gambaran_umum_rb(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{SEKSI_2_DAMPAK}}", build_dampak_rb(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{SEKSI_3_RENAKSI}}", build_hasil_renaksi_rb(ctx, kkp))

    elif jenis == "evaluasi-umum":
        # ── LHE Umum (A–I: dimensi/skor/predikat) ────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_METODOLOGI}}", [(_build_metodologi_evaluasi(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_GAMBARAN_UMUM}}", build_gambaran_umum_evaluasi_umum(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E1_SKOR_DIMENSI}}", build_hasil_evaluasi_umum_skor(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E2_PREDIKAT}}", build_predikat_evaluasi_umum(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E3_ANALISIS}}", build_analisis_per_dimensi(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_TEMUAN}}", build_temuan_evaluasi_umum(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{G_REKOMENDASI}}", build_rekomendasi_evaluasi_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{H_SIMPULAN}}", build_simpulan_evaluasi_umum(ctx, kkp))

    elif jenis == "pemantauan-pengadaan":
        # ── LHP Pemantauan Pengadaan (A–H) ───────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A1_LATAR_BELAKANG}}", build_latar_belakang_umum(ctx, kkp, "Pemantauan Pengadaan"))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A2_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A3_TUJUAN_RUANG_LINGKUP}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A4_METODOLOGI}}", [(_build_metodologi_pemantauan(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A5_PERIODE}}", [(ctx.get("periode_pemantauan") or
                                     ctx.get("periode_pelaksanaan") or "[DIISI — Periode pemantauan]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A6_KOMPOSISI_TIM}}", build_komposisi_tim_blocks(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_PROFIL_PEKERJAAN}}", build_profil_pekerjaan_pengadaan(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_STATUS_PELAKSANAAN}}", build_status_pelaksanaan_pengadaan(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_ISU_PERMASALAHAN}}", build_isu_pengadaan(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_PERUBAHAN_KONTRAK}}", [(
                ctx.get("perubahan_kontrak") or "[DIISI — Perubahan kontrak jika ada, atau: Tidak terdapat perubahan kontrak.]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_TINDAK_LANJUT_SEBELUMNYA}}", [(
                ctx.get("tl_sebelumnya") or "[DIISI — Tindak lanjut atas isu pemantauan sebelumnya.]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{G_SIMPULAN_REKOMENDASI}}", build_simpulan_rekomendasi_pengadaan(kkp))

    elif jenis == "pemantauan-tindak-lanjut":
        # ── LHP Pemantauan TLHP (1–5) ────────────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{RINGKASAN_EKSEKUTIF_TLHP}}", build_ringkasan_tlhp(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{STATISTIK_TLHP}}", build_statistik_tlhp(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{AGING_PER_PIC}}", build_aging_per_pic(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{REKOMENDASI_KRITIS}}", build_rekomendasi_kritis_tlhp(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{REKOMENDASI_PERCEPATAN}}", build_percepatan_tlhp(kkp))

    elif jenis == "pemantauan-umum":
        # ── LHP Pemantauan Umum (A–H) ─────────────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_PERIODE}}", [(
                ctx.get("periode_pemantauan") or ctx.get("periode_pelaksanaan") or
                "[DIISI — Periode pemantauan dan cut-off date]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_METODOLOGI}}", [(_build_metodologi_pemantauan(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_RINGKASAN_STATUS}}", build_ringkasan_status_umum(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_HASIL_PER_ITEM}}", build_hasil_per_item_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{G_REKOMENDASI}}", build_rekomendasi_pemantauan_umum(kkp))

    elif jenis == "reviu-pengadaan":
        # ── LHR Reviu Pengadaan (A–F) ─────────────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A1_LATAR_BELAKANG}}", build_latar_belakang_umum(ctx, kkp, "Reviu Pengadaan Barang/Jasa"))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A2_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A3_TUJUAN}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A4_RUANG_LINGKUP}}", [(ctx.get("ruang_lingkup") or "[DIISI — Ruang lingkup reviu pengadaan]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A5_METODOLOGI}}", [(_build_metodologi_pemantauan(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A6_JANGKA_WAKTU}}", [(ctx.get("periode_pelaksanaan") or "[DIISI — Jangka waktu reviu]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A7_KOMPOSISI_TIM}}", build_komposisi_tim_blocks(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_GAMBARAN_UMUM}}", build_gambaran_umum_reviu_pengadaan(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C1_PERENCANAAN}}", build_hasil_perencanaan_reviu_pengadaan(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C2_PEMILIHAN}}", build_hasil_pemilihan_reviu_pengadaan(ctx, kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_SIMPULAN}}", build_simpulan_reviu(ctx, kkp, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_REKOMENDASI}}", build_rekomendasi_reviu(kkp))

    elif jenis == "reviu-rka-kl":
        # ── LHR Reviu RKA-K/L (A–F, 6 aspek) ────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A1_LATAR_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A2_TUJUAN}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A3_RUANG_LINGKUP}}", [(ctx.get("ruang_lingkup") or "[DIISI — Ruang lingkup reviu RKA-K/L]", {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A4_METODOLOGI}}", [(_build_metodologi_pemantauan(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A5_JANGKA_TIM}}", build_komposisi_tim_blocks(ctx))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_GAMBARAN_UMUM}}", [(
                ctx.get("gambaran_umum_rkakl") or
                "[DIISI — Gambaran umum RKA-K/L: unit kerja, total pagu, komposisi per sumber dana, dan program prioritas]",
                {"italic": True})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C1_SBM_SBK}}", _build_aspek_reviu_rkakl(
                kkp, "c1", "Kelayakan SBM/SBK telah sesuai dengan ketentuan."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C2_KAIDAH}}", _build_aspek_reviu_rkakl(
                kkp, "c2", "Kaidah penganggaran telah dipatuhi."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C3_PENANDAAN}}", _build_aspek_reviu_rkakl(
                kkp, "c3", "Penandaan tematik telah dilakukan dengan tepat."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C4_KELENGKAPAN}}", _build_aspek_reviu_rkakl(
                kkp, "c4", "Dokumen pendukung telah lengkap."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C5_RINCIAN_BARU}}", _build_aspek_reviu_rkakl(
                kkp, "c5", "Kelayakan rincian baru telah memadai."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C6_ALOKASI_TEMATIK}}", _build_aspek_reviu_rkakl(
                kkp, "c6", "Pengalokasian tematik telah sesuai arahan."))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_SIMPULAN}}", build_simpulan_reviu(ctx, kkp, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_REKOMENDASI}}", build_rekomendasi_reviu(kkp))

    elif jenis == "reviu-umum":
        # ── LHR Reviu Umum (A–G) ──────────────────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_audit_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_audit_umum(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_METODOLOGI}}", [(_build_metodologi_pemantauan(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_HASIL_REVIU}}", build_hasil_reviu_umum_aspek(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_CATATAN_REKOMENDASI}}", build_catatan_rekomendasi_reviu_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_SIMPULAN}}", build_simpulan_reviu_umum(ctx, kkp))

    elif jenis == "audit-kinerja":
        # ── LHA Kinerja (Bab 1–6) ────────────────────────────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_kinerja(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_TUJUAN}}", build_tujuan_kinerja(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_PERTANYAAN_AUDIT}}", build_pertanyaan_audit_kinerja(ctx, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_RUANG_LINGKUP}}", build_ruang_lingkup_kinerja(ctx, sa))
        # D_METODOLOGI sudah diisi via mapping (METODOLOGI_REVIU) → placeholder D_METODOLOGI
        # Jika template pakai {{D_METODOLOGI}} (bukan {{METODOLOGI_REVIU}}),
        # perlu expand lagi:
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_METODOLOGI}}", [(_build_metodologi_audit_kinerja(ctx), {})])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_KOMPOSISI_TIM}}", build_komposisi_tim_blocks(ctx))
        # Bab 2 — 4 sub-seksi (2.1 Tujuan Desain, 2.2 Logika, 2.3 Anggaran, 2.4 Pelaksana)
        bab2 = build_gambaran_umum_kinerja(ctx, kkp)
        expand_paragraph_placeholder_to_blocks(doc, "{{BAB2_1_TUJUAN_DESAIN}}", [bab2[0]])
        expand_paragraph_placeholder_to_blocks(doc, "{{BAB2_2_LOGIKA}}",        [bab2[1]])
        expand_paragraph_placeholder_to_blocks(doc, "{{BAB2_3_ANGGARAN}}",      [bab2[2]])
        expand_paragraph_placeholder_to_blocks(doc, "{{BAB2_4_PELAKSANA}}",     [bab2[3]])
        expand_paragraph_placeholder_to_blocks(
            doc, "{{HASIL_EFEKTIVITAS}}", build_hasil_efektivitas(kkp, sa, rek))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{HASIL_EFISIENSI}}", build_hasil_efisiensi(kkp, sa, rek))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{SIMPULAN_REVIU}}", build_simpulan_audit_kinerja_blocks(kkp, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{REKOMENDASI_LOOP}}", build_rekomendasi_kinerja(kkp, rek))

    elif jenis == "konsultansi-umum":
        # ── Memo Konsultasi Umum (A–G sesuai SKILL.md) ───────────────────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{A_DASAR}}", build_dasar_konsultansi_umum(ctx, kkp, args))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{B_PERTANYAAN}}", build_pertanyaan_konsultansi_umum(kkp, sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{C_DASAR_HUKUM}}", build_dasar_hukum_konsultansi_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{D_TELAAH}}", build_telaah_konsultansi_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{E_PENDAPAT}}", build_pendapat_konsultansi_umum(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{F_ASUMSI_BATASAN}}", build_asumsi_batasan_konsultansi_umum(kkp, ctx))

    elif jenis == "konsultasi-pengadaan":
        # ── Laporan Hasil Pendampingan (I–III sesuai SKILL.md v3.0) ──────────
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB1_KEGIATAN}}", build_kegiatan_pendampingan(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB1_DOKUMEN_PENDUKUNG}}", build_dokumen_pendukung_pendampingan(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB2_TINDAK_LANJUT}}", build_tindak_lanjut_pendampingan(kkp))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{BAB3_KESIMPULAN}}", build_kesimpulan_pendampingan(kkp, ctx))

    else:
        # ── Kelompok skill: evaluasi-*, pemantauan-*, reviu-* ─────────────────
        if jenis.startswith("evaluasi"):
            verb = "Evaluasi"
            mapping["METODOLOGI_REVIU"] = _build_metodologi_evaluasi(ctx)
            simpulan_fn = build_simpulan_evaluasi

        elif jenis.startswith("pemantauan"):
            verb = "Pemantauan"
            mapping["METODOLOGI_REVIU"] = _build_metodologi_pemantauan(ctx)
            simpulan_fn = build_simpulan_pemantauan

        else:  # reviu-*
            verb = "Reviu"
            mapping["METODOLOGI_REVIU"] = _build_metodologi_reviu(ctx)
            simpulan_fn = build_simpulan_reviu

        # Perbaiki label section ("Dasar Pelaksanaan Reviu" → "...Audit" dst.)
        replace_label_in_doc(doc, _label_pairs(verb))

        expand_paragraph_placeholder_to_blocks(
            doc, "{{DASAR_HUKUM_LIST}}", build_dasar_hukum_blocks(jenis))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{SASARAN_LIST}}", build_sasaran_blocks(sa))
        expand_paragraph_placeholder_to_blocks(
            doc, "{{HASIL_REVIU_LOOP}}", build_hasil_reviu_blocks(kkp, sa, rek, jenis))
        replace_in_doc(doc, {"SIMPULAN_REVIU": simpulan_fn(kkp, sa)})

    # ── Common: tembusan + survei — berlaku untuk semua skill ─────────────────
    expand_paragraph_placeholder_to_blocks(
        doc, "{{TEMBUSAN_LIST}}", build_tembusan_blocks(args.tembusan))
    if args.survei:
        insert_survei_paragraph(doc, args.survei)

    doc.save(out_path)
    print(f"OK: {out_path}")
    print(f"  paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
