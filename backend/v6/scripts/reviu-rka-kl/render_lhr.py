"""
render_lhr.py — Render anomalies(-master).json + context.md → draft LHR RKA-K/L (docx).

Mode otomatis:
  - **Multi-RO**: input adalah anomalies-master.json (output cross_check.py batch).
    Format: {ringkasan: {...}, anomalies: [{ro_id, ro_nama, ...}, ...]}.
    Laporan disusun mengikuti audit-system-v4/skills/panduan-format-umum/PANDUAN.md
    (Bab A-H + Lampiran 1 Matriks Catatan Reviu).
  - **Single-RO** (legacy): input adalah anomalies.json single (tanpa ro_id).
    Mempertahankan struktur lama (Nota Dinas + Cover + 9 seksi).

CLI:
    python render_lhr.py <ANOMALIES_JSON> [<CONTEXT_MD>]
        [--tor-dir DIR | --tor PATH]
        [--rab-dir DIR | --rab PATH]
        [--judul "..."] [--nomor "..."] [--tanggal "..."] [--penerima "..."]
        -o OUTPUT.docx
"""

from __future__ import annotations
import argparse
import json
import re
import sys
from collections import defaultdict, Counter
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")


ASPEK_TITLES = {
    "A": "Kelayakan Anggaran terhadap Standar Biaya (SBM/SBK/SSB)",
    "B": "Kepatuhan Kaidah Penganggaran (Klasifikasi Akun)",
    "C": "Penandaan Anggaran (Budget Tagging)",
    "D": "Kelengkapan Dokumen Pendukung (Kriteria IR2)",
    "E": "Kelayakan Rincian Anggaran Baru & Konsistensi Internal TOR-RAB-Timeline",
    "F": "Pengalokasian Tematik",
}


# ============================================================
# Helper umum
# ============================================================
def _parse_context_md(path):
    """Parse context.md / KP-Reviu.md ke dict (best-effort)."""
    if not path or not path.exists():
        return {}
    text = path.read_text(encoding="utf-8")
    ctx = {}
    for m in re.finditer(r"^([A-Z][A-Za-z\s/]+?)\s*:\s*(.+?)$", text, re.M):
        key = m.group(1).strip()
        val = m.group(2).strip()
        ctx[key] = val
    for label in ("Surat Tugas", "ND Permintaan", "Periode Reviu", "Dasar Penugasan"):
        m = re.search(rf"\*\*{label}:\*\*\s*(.+?)(?:\n|$)", text)
        if m:
            ctx[label] = m.group(1).strip()
    return ctx


def _add_heading(doc, text, size=12, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _add_para(doc, text, bold=False, italic=False, size=11, align=None,
              indent_cm=None, justify=False):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def _set_cell_border(cell, sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _setup_doc():
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def _detect_mode(ano_data):
    """Return 'multi' jika anomalies-master, 'single' jika legacy single-RO."""
    if isinstance(ano_data, list):
        return "single"
    if not isinstance(ano_data, dict):
        return "single"
    items = ano_data.get("anomalies", [])
    if not items:
        if "ringkasan" in ano_data and "per_ro" in (ano_data.get("ringkasan") or {}):
            return "multi"
        return "single"
    sample = items[0] if isinstance(items[0], dict) else {}
    if "ro_id" in sample and "ro_nama" in sample:
        return "multi"
    return "single"


# ============================================================
# MULTI-RO RENDER
# ============================================================
def _aggregate_multi(anomalies, ringkasan):
    """Group + dedup anomalies untuk render multi-RO."""
    anomalies = [a for a in anomalies if a.get("severity") != "ERROR"]

    per_aspek = defaultdict(list)
    by_ro = {}
    for a in anomalies:
        per_aspek[a.get("aspek", "?")].append(a)
        ro_id = a.get("ro_id")
        ro_nama = a.get("ro_nama", "")
        if ro_id is not None:
            entry = by_ro.setdefault(ro_id, {"ro_nama": ro_nama, "anomalies": []})
            entry["anomalies"].append(a)

    a3_items = [a for a in anomalies if a.get("rule_id") == "A.3"]
    a3_dedup = None
    if a3_items:
        ro_ids = sorted({a.get("ro_id") for a in a3_items if a.get("ro_id") is not None})
        a3_dedup = {
            "rule_id": "A.3",
            "severity": "INFO",
            "aspek": "A",
            "judul": a3_items[0].get("judul"),
            "deskripsi": a3_items[0].get("deskripsi"),
            "draft_catatan": a3_items[0].get("draft_catatan", {}),
            "ro_ids": ro_ids,
            "berlaku_seluruh_ro": True,
        }

    peringatan = [a for a in anomalies if a.get("severity") == "PERINGATAN"]
    rid_counter = Counter(a.get("rule_id") for a in peringatan)
    top_rules = [rid for rid, _ in rid_counter.most_common(5)]

    top_recs = []
    for rid in top_rules:
        sample = next((a for a in peringatan if a.get("rule_id") == rid), None)
        if sample is None:
            continue
        ro_ids_hit = sorted({a.get("ro_id") for a in peringatan
                             if a.get("rule_id") == rid and a.get("ro_id") is not None})
        top_recs.append({
            "rule_id": rid,
            "judul": sample.get("judul", ""),
            "rekomendasi": (sample.get("draft_catatan") or {}).get("rekomendasi", ""),
            "ro_ids": ro_ids_hit,
            "frekuensi": rid_counter[rid],
        })

    return {
        "per_aspek": dict(per_aspek),
        "by_ro": by_ro,
        "a3_dedup": a3_dedup,
        "anomalies_clean": anomalies,
        "peringatan": peringatan,
        "top_recs": top_recs,
        "total_peringatan": len(peringatan),
        "total_info": len([a for a in anomalies if a.get("severity") == "INFO"]),
    }


def _render_aspek_section_multi(doc, aspek_code, items, a3_dedup):
    """Render satu aspek di Bab E (multi-RO)."""
    title = ASPEK_TITLES.get(aspek_code, "")
    _add_heading(doc, f"E.{aspek_code}. {title}", size=11)

    items_to_render = items
    if aspek_code == "A" and a3_dedup is not None:
        non_a3 = [a for a in items if a.get("rule_id") != "A.3"]
        a = a3_dedup
        ro_ids_str = ", ".join(f"#{r}" for r in (a.get("ro_ids") or []))
        p = doc.add_paragraph()
        r = p.add_run(f"{a['rule_id']}. {a.get('judul', '')}")
        r.bold = True
        r.font.size = Pt(11)
        _add_para(doc,
                  f"Catatan ini berlaku untuk seluruh RO yang direviu (RO {ro_ids_str}).",
                  italic=True, indent_cm=0.5, size=10, justify=True)
        draft = a.get("draft_catatan", {}) or {}
        for label, key in [("Kondisi", "kondisi"), ("Kriteria", "kriteria"),
                           ("Akibat", "akibat"), ("Rekomendasi", "rekomendasi")]:
            val = draft.get(key)
            if not val:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(10)
            r2 = p.add_run(val)
            r2.font.size = Pt(10)
        items_to_render = non_a3

    if not items_to_render:
        if aspek_code == "A" and a3_dedup is not None:
            return
        _add_para(doc, "Tidak terdapat catatan pada aspek ini.",
                  italic=True, indent_cm=0.5, size=10)
        return

    items_sorted = sorted(
        items_to_render,
        key=lambda a: (0 if a.get("severity") == "PERINGATAN" else 1,
                       str(a.get("rule_id")), a.get("ro_id") or 0)
    )

    n_peringatan = sum(1 for a in items_sorted if a.get("severity") == "PERINGATAN")
    n_info = sum(1 for a in items_sorted if a.get("severity") == "INFO")
    ro_terdampak = sorted({a.get("ro_id") for a in items_sorted if a.get("ro_id") is not None})
    summary = (f"Pada aspek ini ditemukan {len(items_sorted)} catatan "
               f"({n_peringatan} PERINGATAN, {n_info} INFO) "
               f"pada {len(ro_terdampak)} RO: "
               f"{', '.join('#' + str(r) for r in ro_terdampak)}.")
    _add_para(doc, summary, justify=True, indent_cm=0.5, size=10, italic=True)

    for a in items_sorted:
        rid = a.get("rule_id", "?")
        sev = a.get("severity", "?")
        judul = a.get("judul", "?")
        ro_id = a.get("ro_id")
        ro_nama = a.get("ro_nama", "")
        prefix = f"RO #{ro_id} ({ro_nama}) - " if ro_id is not None else ""

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f"{rid} [{sev}]. {prefix}{judul}")
        r.bold = True
        r.font.size = Pt(10)

        draft = a.get("draft_catatan") or {}
        for label, key in [("Kondisi", "kondisi"), ("Kriteria", "kriteria"),
                           ("Akibat", "akibat"), ("Rekomendasi", "rekomendasi")]:
            val = draft.get(key)
            if not val:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(10)
            r2 = p.add_run(val)
            r2.font.size = Pt(10)


def _render_lampiran_matriks(doc, peringatan, max_rows=25):
    """Lampiran 1 - Matriks Catatan Reviu (PERINGATAN saja)."""
    doc.add_page_break()
    _add_heading(doc, "LAMPIRAN 1 - MATRIKS CATATAN REVIU",
                 size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc,
              f"Daftar catatan dengan severity PERINGATAN ({len(peringatan)} catatan, "
              f"ditampilkan top {min(len(peringatan), max_rows)}).",
              italic=True, justify=True, size=10)
    if not peringatan:
        _add_para(doc, "Tidak terdapat catatan PERINGATAN.", italic=True)
        return

    rows = sorted(peringatan,
                  key=lambda a: (a.get("ro_id") or 0, a.get("aspek") or "?",
                                 str(a.get("rule_id"))))
    rows = rows[:max_rows]

    headers = ["No", "RO", "Aspek", "Rule", "Severity", "Kondisi (judul)", "Rekomendasi"]
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        _set_cell_border(c)
    for i, a in enumerate(rows, start=1):
        ro_id = a.get("ro_id")
        ro_label = f"#{ro_id} {a.get('ro_nama', '')}" if ro_id else "-"
        rec = (a.get("draft_catatan") or {}).get("rekomendasi", "")
        cells_data = [
            str(i),
            ro_label,
            a.get("aspek", "?"),
            a.get("rule_id", "?"),
            a.get("severity", "?"),
            a.get("judul", "")[:120],
            rec[:200] + ("..." if len(rec) > 200 else ""),
        ]
        for j, val in enumerate(cells_data):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            _set_cell_border(c)


def _render_multi(ano_data, ctx, output_path,
                  tor_dir, rab_dir,
                  judul, nomor, tanggal, penerima):
    """Render multi-RO LHR mengikuti Bab A-H + Lampiran 1."""
    anomalies = ano_data.get("anomalies", [])
    ringkasan = ano_data.get("ringkasan", {}) or {}
    agg = _aggregate_multi(anomalies, ringkasan)

    judul_default = "Laporan Hasil Reviu RKA-K/L"
    judul = judul or ctx.get("Hal") or judul_default
    nomor = nomor or "B-[XX]/IJ.3/PW.04.04/[MM]/[YYYY]  >> DIISI AUDITOR"
    tanggal = tanggal or "[DD Bulan YYYY]  >> DIISI AUDITOR"
    penerima = penerima or ctx.get("Penerima Laporan") or "[Jabatan Penerima]  >> DIISI AUDITOR"

    tahun = "[TA]"
    m = re.search(r"\bTA\s*(\d{4})\b|Tahun\s+Anggaran\s+(\d{4})", judul + " " + str(ctx))
    if m:
        tahun = m.group(1) or m.group(2)

    per_ro_summary = ringkasan.get("per_ro") or []
    if not per_ro_summary:
        seen = {}
        for a in anomalies:
            rid = a.get("ro_id")
            if rid is None:
                continue
            if rid not in seen:
                seen[rid] = {"ro_id": rid, "ro_nama": a.get("ro_nama", ""),
                             "jumlah_anomali": 0}
            seen[rid]["jumlah_anomali"] += 1
        per_ro_summary = sorted(seen.values(), key=lambda x: x["ro_id"])

    nomor_st = ctx.get("Nomor ST") or ctx.get("Surat Tugas") or "[Nomor ST]"
    tgl_st = ctx.get("Tanggal ST") or "[tanggal ST]"
    nd_minta = ctx.get("ND Permintaan") or ctx.get("Dasar Penugasan") or "ND Permintaan dari Unit Penyusun"

    doc = _setup_doc()

    # ====== KOP SURAT ======
    for line in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI",
                 "INSPEKTORAT JENDERAL", "INSPEKTORAT II"]:
        _add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _add_para(doc,
              "Jl. Medan Merdeka Barat No. 9, Jakarta 10110 Telp./Fax. (021) 3859627 www.komdigi.go.id",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True)
    doc.add_paragraph()

    # ====== HEADER ======
    t = doc.add_table(rows=3, cols=3)
    t.autofit = False
    rows_data = [
        ("Nomor", nomor, f"Jakarta, {tanggal}"),
        ("Lampiran", "Satu berkas", ""),
        ("Hal", judul, ""),
    ]
    for i, (k, v, right) in enumerate(rows_data):
        for j, val in enumerate([k, ":", v]):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(11)
            if j == 0:
                r.bold = True
    for row in t.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(13.5)
    cell_tgl = t.rows[0].cells[2]
    p2 = cell_tgl.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"Jakarta, {tanggal}")
    r2.font.size = Pt(11)

    doc.add_paragraph()
    _add_para(doc, "Kepada Yth.")
    _add_para(doc, penerima, bold=True)
    _add_para(doc, "di Jakarta")
    doc.add_paragraph()

    # ====== PEMBUKA ======
    pembuka = (
        f"Menindaklanjuti {nd_minta}, kami telah menerbitkan Surat Tugas Nomor "
        f"{nomor_st} tanggal {tgl_st} untuk melakukan reviu atas Rencana Kerja "
        f"dan Anggaran Kementerian/Lembaga (RKA-K/L) Tahun Anggaran {tahun}. "
        f"Bersama ini kami sampaikan {judul} sebagaimana diuraikan di bawah ini."
    )
    _add_para(doc, pembuka, justify=True)

    # ====== A. DASAR ======
    doc.add_paragraph()
    _add_heading(doc, "A. Dasar", size=12)
    dasar_items = [
        f"1. {nd_minta};",
        f"2. Surat Tugas Inspektur II Nomor {nomor_st} tanggal {tgl_st};",
        "3. Peraturan Menteri Keuangan Nomor 107 Tahun 2024 tentang Perubahan atas PMK 62/2023, khususnya Pasal 61.",
    ]
    for item in dasar_items:
        _add_para(doc, item, indent_cm=0.5, justify=True)

    # ====== B. TUJUAN ======
    doc.add_paragraph()
    _add_heading(doc, "B. Tujuan", size=12)
    _add_para(doc,
              f"Memberikan keyakinan terbatas bahwa dokumen RKA-K/L TA {tahun} telah disusun "
              f"sesuai kaidah penganggaran terhadap 6 (enam) aspek reviu sebagaimana diatur "
              f"dalam Pasal 61 ayat (2) PMK 107/2024, yaitu:",
              justify=True)
    for code in ["A", "B", "C", "D", "E", "F"]:
        _add_para(doc, f"{code}. {ASPEK_TITLES[code]};", indent_cm=0.8, justify=True)

    # ====== C. RUANG LINGKUP ======
    doc.add_paragraph()
    _add_heading(doc, "C. Ruang Lingkup", size=12)
    _add_para(doc,
              f"Reviu dilaksanakan atas {len(per_ro_summary)} Rincian Output (RO) berikut:",
              justify=True)
    for ro in per_ro_summary:
        _add_para(doc,
                  f"{ro['ro_id']}. {ro.get('ro_nama', '-')} ({ro.get('jumlah_anomali', 0)} catatan).",
                  indent_cm=0.8, justify=True)
    _add_para(doc,
              "Reviu tidak mencakup pengujian sistem pengendalian intern, pembuktian kerugian "
              "negara, dan analisis substansi kebijakan.",
              justify=True, italic=True)

    # ====== D. METODOLOGI ======
    doc.add_paragraph()
    _add_heading(doc, "D. Metodologi", size=12)
    _add_para(doc, "Reviu dilaksanakan dengan tahapan berikut:", justify=True)
    metod = [
        "1. Desk review dokumen TOR dan RAB seluruh RO yang direviu;",
        "2. Pipeline pre-digest dan cross-check deterministik (21 rules utama + 18 alt-rules) "
        "yang menguji konsistensi internal TOR, kesesuaian TOR-RAB, dan kepatuhan kaidah "
        "penganggaran;",
        "3. Verifikasi manual substantif atas aspek yang memerlukan judgment auditor "
        "(klasifikasi akun, tematik, kelayakan rincian baru);",
        "4. Pembandingan satuan biaya terhadap PMK SBM (TA 2026 sebagai proksi bila "
        "PMK SBM tahun reviu belum terbit).",
    ]
    for item in metod:
        _add_para(doc, item, indent_cm=0.5, justify=True)

    # ====== E. HASIL REVIU ======
    doc.add_paragraph()
    _add_heading(doc, "E. Hasil Reviu", size=12)
    _add_para(doc,
              f"Berdasarkan pelaksanaan reviu, diidentifikasi {len(agg['anomalies_clean'])} "
              f"catatan ({agg['total_peringatan']} PERINGATAN, {agg['total_info']} INFO) "
              f"yang dikelompokkan per aspek sebagai berikut.",
              justify=True)

    for code in ["A", "B", "C", "D", "E", "F"]:
        items = agg["per_aspek"].get(code, [])
        if not items and code != "A":
            continue
        if code == "A" or items:
            _render_aspek_section_multi(doc, code, items, agg.get("a3_dedup"))

    # ====== F. CATATAN & REKOMENDASI ======
    doc.add_paragraph()
    _add_heading(doc, "F. Catatan dan Rekomendasi Prioritas", size=12)
    if not agg["top_recs"]:
        _add_para(doc,
                  "Tidak terdapat catatan PERINGATAN yang memerlukan rekomendasi prioritas.",
                  italic=True, justify=True)
    else:
        _add_para(doc,
                  f"Berikut {len(agg['top_recs'])} rekomendasi prioritas yang dipilih dari "
                  f"PERINGATAN dengan frekuensi tertinggi:",
                  justify=True)
        for i, rec in enumerate(agg["top_recs"], 1):
            ro_str = ", ".join(f"#{r}" for r in rec["ro_ids"][:8])
            if len(rec["ro_ids"]) > 8:
                ro_str += f", ... (+{len(rec['ro_ids']) - 8} RO lain)"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"{i}. [{rec['rule_id']} - {rec['frekuensi']}x] {rec['judul']}")
            r.bold = True
            r.font.size = Pt(11)
            _add_para(doc, f"RO terdampak: {ro_str}.",
                      indent_cm=1.0, italic=True, size=10, justify=True)
            if rec["rekomendasi"]:
                _add_para(doc, f"Rekomendasi: {rec['rekomendasi']}",
                          indent_cm=1.0, size=10, justify=True)

    # ====== G. SIMPULAN ======
    doc.add_paragraph()
    _add_heading(doc, "G. Simpulan", size=12)
    if agg["total_peringatan"] == 0:
        _add_para(doc,
                  f"Berdasarkan hasil reviu secara terbatas, tidak terdapat hal-hal yang "
                  f"membuat kami yakin bahwa RKA-K/L TA {tahun} disusun tidak sesuai dengan "
                  f"PMK 107/2024.",
                  justify=True)
    else:
        aspek_with_p = sorted({a.get("aspek") for a in agg["peringatan"] if a.get("aspek")})
        aspek_str = ", ".join(f"aspek {x} ({ASPEK_TITLES.get(x, '')})"
                              for x in aspek_with_p)
        prioritas_list = "; ".join(f"({i + 1}) {r['judul']}"
                                   for i, r in enumerate(agg["top_recs"]))
        _add_para(doc,
                  f"Berdasarkan hasil reviu, masih ditemukan {agg['total_peringatan']} catatan "
                  f"PERINGATAN dalam {aspek_str}, di antaranya: {prioritas_list}. "
                  f"Untuk perbaikan, kami merekomendasikan agar Saudara melakukan tindak "
                  f"lanjut sebelum DIPA TA {tahun} ditetapkan.",
                  justify=True)

    # ====== H. APRESIASI & PENUTUP ======
    doc.add_paragraph()
    _add_heading(doc, "H. Apresiasi dan Penutup", size=12)
    _add_para(doc,
              "Inspektorat II menyampaikan apresiasi atas bantuan dan kerja sama seluruh "
              "pejabat/pegawai yang mendukung pelaksanaan reviu ini. "
              "Terima kasih telah membantu kami dalam menjaga integritas. "
              "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama "
              "Bapak/Ibu kami ucapkan terima kasih.",
              justify=True)

    # ====== TTD ======
    doc.add_paragraph()
    _add_para(doc, "Inspektur II,", align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    doc.add_paragraph()
    _add_para(doc, "[TTD - DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT,
              size=10, italic=True)
    _add_para(doc, "Muhammad Arief", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    _add_para(doc, "NIP. 19xxxxxxxxxxxxx  >> DIISI AUDITOR",
              align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    # ====== LAMPIRAN 1 ======
    _render_lampiran_matriks(doc, agg["peringatan"], max_rows=25)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path, {
        "mode": "multi",
        "n_ro": len(per_ro_summary),
        "n_anomalies": len(agg["anomalies_clean"]),
        "n_peringatan": agg["total_peringatan"],
        "n_info": agg["total_info"],
        "n_top_recs": len(agg["top_recs"]),
        "a3_dedup_ro_count": len(agg["a3_dedup"]["ro_ids"]) if agg.get("a3_dedup") else 0,
    }


# ============================================================
# SINGLE-RO RENDER (legacy)
# ============================================================
def _render_single(ano_data, ctx, output_path, tor_json_path, rab_json_path):
    """Render single-RO legacy (backward compat dengan signature lama)."""
    if isinstance(ano_data, list):
        ano_data = {"anomalies": ano_data, "summary_by_aspek": {}}

    tor_meta, rab_meta = {}, {}
    if tor_json_path and Path(tor_json_path).exists():
        tor_full = json.loads(Path(tor_json_path).read_text(encoding="utf-8"))
        ident = tor_full.get("identitas_ro") or {}
        biaya = tor_full.get("biaya") or {}
        lokasi = tor_full.get("lokasi") or {}
        tor_meta = {
            "ro": ident.get("ro"),
            "kro": ident.get("kro"),
            "program_kode": ident.get("program_kode"),
            "kegiatan_kode": ident.get("kegiatan_kode"),
            "volume": ident.get("volume_int"),
            "satuan": ident.get("satuan"),
            "lokasi": lokasi.get("lokasi_target", []),
            "biaya_total": biaya.get("total"),
        }
    if rab_json_path and Path(rab_json_path).exists():
        rab_full = json.loads(Path(rab_json_path).read_text(encoding="utf-8"))
        rab_meta = {
            "total_pagu": rab_full.get("total_pagu"),
            "jumlah_komponen": len(rab_full.get("komponen", [])),
        }

    doc = _setup_doc()

    for line in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI", "INSPEKTORAT JENDERAL", "INSPEKTORAT II"]:
        _add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _add_para(doc,
              "Jl. Medan Merdeka Barat No. 9, Jakarta 10110 Telp./Fax. (021) 3859627 www.komdigi.go.id",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True)
    doc.add_paragraph()

    _add_heading(doc, "NOTA DINAS", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc,
              f"Nomor: [DIISI AUDITOR - dari SIMWAS]/IJ.3/{ctx.get('Kode Nomor Surat', 'PW.04.04')}/[bulan]/[tahun]",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    obyek = ctx.get("Obyek", "[Obyek]")
    tahap = ctx.get("Tahap Reviu", "Tahap 1 - Pagu Indikatif")
    unit_diawasi = ctx.get("Unit yang Diawasi", "[Unit yang Diawasi]")
    nomor_st = ctx.get("Nomor ST", "[Nomor ST]")
    tgl_st = ctx.get("Tanggal ST", "[Tanggal ST]")
    dasar = ctx.get("Dasar Penugasan", "[Dasar Penugasan]")
    penerima = ctx.get("Penerima Laporan", "[Penerima Laporan]")
    tahun_anggaran = None
    m = re.search(r"TA\s*(\d{4})|Tahun\s+Anggaran\s+(\d{4})", obyek)
    if m:
        tahun_anggaran = m.group(1) or m.group(2)
    if not tahun_anggaran:
        tahun_anggaran = "[TA]"

    fields = [
        ("Kepada Yth.", penerima),
        ("Dari", "Inspektur II"),
        ("Hal", f"Catatan Hasil Reviu / Laporan Hasil Reviu Rencana Kerja dan Anggaran {unit_diawasi} Tahun Anggaran {tahun_anggaran} - {obyek}"),
        ("Klasifikasi", "Terbatas"),
        ("Lampiran", "Satu laporan"),
        ("Tanggal", "[DIISI AUDITOR]"),
    ]
    t = doc.add_table(rows=len(fields), cols=3)
    t.autofit = False
    for row in t.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(13.0)
    for i, (k, v) in enumerate(fields):
        for j, val in enumerate([k, ":", v]):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(11)
            if j == 0:
                r.bold = True
    doc.add_paragraph()

    _add_para(doc,
              f"Menindaklanjuti {dasar}, kami telah menerbitkan Surat Tugas Nomor {nomor_st} tanggal {tgl_st} untuk "
              f"melakukan reviu atas dokumen Rencana Kerja (Renja) dan Anggaran {unit_diawasi} TA {tahun_anggaran} "
              f"pada {tahap}, khususnya untuk {obyek}. "
              f"Bersama ini kami sampaikan Laporan Hasil Reviu sebagaimana terlampir pada Nota Dinas ini.",
              justify=True)
    _add_para(doc,
              "Kami mohon kesediaan Saudara/i untuk mengisi survei kepuasan auditan melalui link: "
              "https://simwas.komdigi.go.id/api/survey/formulir/[DIISI AUDITOR] (minimal lima pegawai).",
              justify=True)
    _add_para(doc,
              "Terima kasih telah membantu kami dalam menjaga integritas. Demikian Nota Dinas ini kami sampaikan. "
              "Atas perhatian dan kerja sama Bapak/Ibu kami ucapkan terima kasih.",
              justify=True)
    doc.add_paragraph()
    _add_para(doc, "[TTD - DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, italic=True)
    _add_para(doc, "Muhammad Arief", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    _add_para(doc, "Inspektur II", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    doc.add_paragraph()
    _add_para(doc, "Tembusan:")
    _add_para(doc, "Kepala Biro Perencanaan", indent_cm=0.5)

    doc.add_page_break()

    for line in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI", "INSPEKTORAT JENDERAL"]:
        _add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    for line in ["LAPORAN HASIL REVIU / CATATAN HASIL REVIU", "RENCANA KERJA DAN ANGGARAN"]:
        _add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph()
    _add_para(doc, unit_diawasi.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _add_para(doc, "KEMENTERIAN KOMUNIKASI DAN DIGITAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _add_para(doc, f"TAHUN ANGGARAN {tahun_anggaran}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_paragraph()
    _add_para(doc, "Rincian Output yang Direviu:", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    _add_para(doc, tor_meta.get("ro") or obyek, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    if tor_meta.get("program_kode"):
        _add_para(doc,
                  f"(Program {tor_meta['program_kode']} - Kegiatan {tor_meta.get('kegiatan_kode', '-')} - KRO {tor_meta.get('kro', '-')})",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    if rab_meta.get("total_pagu"):
        _add_para(doc, f"Pagu: Rp {rab_meta['total_pagu']:,}",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()
    _add_para(doc, "Nomor : B-[DIISI AUDITOR]/IJ.3/PW.04.04/[bulan]/[tahun]",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _add_para(doc, "Tanggal : [DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_page_break()

    _add_heading(doc, "1. RINGKASAN EKSEKUTIF", size=12)
    anomalies = ano_data.get("anomalies", [])
    anomalies = [a for a in anomalies if a.get("severity") != "ERROR"]
    n_total = len(anomalies)
    by_aspek = ano_data.get("summary_by_aspek", {})

    _add_para(doc,
              f"Dalam rangka penyusunan RKA-K/L Tahun Anggaran {tahun_anggaran}, Aparat Pengawasan Intern "
              f"Pemerintah (APIP) melaksanakan reviu atas Rencana Kerja (Renja) dan data dukungnya sesuai "
              f"Pasal 61 PMK 107/2024. Inspektorat II melaksanakan reviu terhadap {obyek} pada {unit_diawasi}.",
              justify=True)
    if tor_meta.get("biaya_total"):
        _add_para(doc,
                  f"Rincian Output ini dialokasikan Rp {tor_meta['biaya_total']:,} dengan target "
                  f"{tor_meta.get('volume') or '[volume]'} {tor_meta.get('satuan') or 'Orang'} "
                  f"di {len(tor_meta.get('lokasi') or [])} lokasi "
                  f"({', '.join((tor_meta.get('lokasi') or [])[:4]) or '[lokasi]'}).",
                  justify=True)
    _add_para(doc,
              f"Berdasarkan hasil reviu secara terbatas atas TOR dan RAB yang disampaikan, diidentifikasi "
              f"{n_total} catatan yang perlu ditindaklanjuti oleh unit penyusun. Catatan dikelompokkan per "
              f"aspek Pasal 61 ayat (2) PMK 107/2024:",
              justify=True)
    for aspek_code in ["A", "B", "C", "D", "E", "F"]:
        cnt = by_aspek.get(aspek_code, 0)
        if cnt > 0:
            _add_para(doc,
                      f"  {aspek_code}. {ASPEK_TITLES.get(aspek_code, '')} - {cnt} catatan.",
                      justify=True, indent_cm=0.5)
    _add_para(doc,
              f"Berdasarkan keseluruhan hasil reviu, disimpulkan bahwa sepanjang {n_total} catatan dalam "
              f"laporan ini ditindaklanjuti, tidak terdapat hal-hal yang membuat kami yakin bahwa Renja dan "
              f"data dukung {unit_diawasi} TA {tahun_anggaran} tidak disusun sesuai Peraturan Menteri "
              f"Keuangan Nomor 107 Tahun 2024.",
              justify=True)

    doc.add_paragraph()
    _add_heading(doc, "2. DASAR HUKUM", size=12)
    dasar_hukum_items = [
        "a. Peraturan Pemerintah Nomor 60 Tahun 2008 tentang Sistem Pengendalian Intern Pemerintah;",
        "b. Peraturan Menteri Keuangan Nomor 107 Tahun 2024 tentang Perubahan Atas PMK 62/2023, "
        "khususnya Pasal 61;",
        "c. Peraturan Menteri Keuangan Nomor 102/PMK.02/2018 jo. PMK 187/PMK.02/2019 tentang "
        "Klasifikasi Anggaran;",
        "d. Kriteria Substansi TOR/KAK Inspektorat II (Kriteria IR2);",
        f"e. {dasar};",
        f"f. Surat Tugas Inspektur II Nomor {nomor_st} tanggal {tgl_st}.",
    ]
    for item in dasar_hukum_items:
        _add_para(doc, item, indent_cm=0.5, justify=True)

    for hd, body in [
        ("3. TUJUAN REVIU",
         "Tujuan pelaksanaan reviu adalah memberikan keyakinan terbatas (limited assurance) "
         "bahwa Renja dan data dukung telah disusun sesuai kaidah penganggaran."),
        ("4. RUANG LINGKUP REVIU",
         f"Ruang lingkup reviu terbatas pada penelaahan dokumen TOR dan RAB untuk {obyek} "
         f"TA {tahun_anggaran}."),
        ("5. METODOLOGI REVIU",
         "Reviu dilaksanakan dengan desk review + pipeline pre-digest dan cross-check "
         "deterministik (21 rules + 18 alt-rules) + verifikasi manual substantif."),
    ]:
        doc.add_paragraph()
        _add_heading(doc, hd, size=12)
        _add_para(doc, body, justify=True)

    doc.add_paragraph()
    _add_heading(doc, "7. URAIAN HASIL REVIU", size=12)
    per_aspek = {}
    for a in anomalies:
        per_aspek.setdefault(a.get("aspek", "?"), []).append(a)

    for aspek_code in ["A", "B", "C", "D", "E", "F"]:
        if aspek_code not in per_aspek:
            continue
        catatan_list = per_aspek[aspek_code]
        _add_heading(doc,
                     f"7.{aspek_code.lower()}. {ASPEK_TITLES.get(aspek_code, '')}",
                     size=11)
        for a in catatan_list:
            draft = a.get("draft_catatan") or {}
            rid = a.get("rule_id", "?")
            judul = a.get("judul", "?")
            p = doc.add_paragraph()
            r = p.add_run(f"{rid}. {judul}")
            r.bold = True
            r.font.size = Pt(11)
            for label, key in [("Kondisi", "kondisi"), ("Kriteria", "kriteria"),
                               ("Akibat", "akibat"), ("Rekomendasi", "rekomendasi")]:
                val = draft.get(key)
                if not val:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                r = p.add_run(f"{label}: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(val)
                r2.font.size = Pt(10)

    doc.add_paragraph()
    _add_heading(doc, "8. SIMPULAN", size=12)
    _add_para(doc,
              f"Berdasarkan hasil reviu, ditemukan {n_total} catatan yang perlu ditindaklanjuti.",
              justify=True)
    doc.add_paragraph()
    _add_heading(doc, "9. APRESIASI", size=12)
    _add_para(doc,
              "Inspektorat II menyampaikan terima kasih atas kerja sama unit yang direviu. "
              "Terima kasih telah membantu kami dalam menjaga integritas.",
              justify=True)

    doc.add_paragraph()
    _add_para(doc, "[TTD - DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, italic=True)
    _add_para(doc, "Inspektur II,", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    _add_para(doc, "Muhammad Arief", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    _add_para(doc, "NIP. [DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path, {"mode": "single", "n_anomalies": n_total}


# ============================================================
# Public entry point
# ============================================================
def render_lhr(anomalies_path, context_path,
               output_path,
               tor_json_path=None,
               rab_json_path=None,
               tor_dir=None,
               rab_dir=None,
               judul=None,
               nomor=None,
               tanggal=None,
               penerima=None):
    """Main render function - auto-detect single vs multi-RO mode."""
    ano = json.loads(Path(anomalies_path).read_text(encoding="utf-8"))
    ctx = _parse_context_md(Path(context_path)) if context_path else {}

    mode = _detect_mode(ano)
    if mode == "multi":
        return _render_multi(ano, ctx, Path(output_path),
                             tor_dir, rab_dir,
                             judul, nomor, tanggal, penerima)
    return _render_single(ano, ctx, Path(output_path),
                          tor_json_path, rab_json_path)


def _self_check_ast():
    import ast
    try:
        ast.parse(open(__file__, "r", encoding="utf-8").read())
    except SyntaxError as e:
        print(f"Self-check AST gagal di {__file__}: {e}", file=sys.stderr)
        sys.exit(2)


def main(argv=None):
    _self_check_ast()
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("anomalies", help="Path anomalies.json (single) atau anomalies-master.json (multi-RO)")
    ap.add_argument("context", nargs="?", default=None,
                    help="Path context.md / KP-Reviu.md (opsional)")
    ap.add_argument("-o", "--output", default="LHR-DRAFT.docx",
                    help="Output path docx (default: LHR-DRAFT.docx)")
    ap.add_argument("--tor", default=None,
                    help="(single-RO) Path tor.json untuk metadata cover")
    ap.add_argument("--rab", default=None,
                    help="(single-RO) Path rab.json untuk metadata pagu")
    ap.add_argument("--tor-dir", default=None,
                    help="(multi-RO) Direktori berisi tor-1.json...tor-N.json")
    ap.add_argument("--rab-dir", default=None,
                    help="(multi-RO) Direktori berisi rab-1.json...rab-N.json")
    ap.add_argument("--judul", default=None,
                    help='Judul laporan (default: "Laporan Hasil Reviu RKA-K/L Direktorat ...")')
    ap.add_argument("--nomor", default=None,
                    help='Nomor surat (default: placeholder DIISI AUDITOR)')
    ap.add_argument("--tanggal", default=None,
                    help='Tanggal surat "DD Bulan YYYY" (default: placeholder)')
    ap.add_argument("--penerima", default=None,
                    help='Penerima ("Direktur ...") (default: placeholder)')
    args = ap.parse_args(argv)

    out_path, info = render_lhr(
        Path(args.anomalies),
        Path(args.context) if args.context else None,
        Path(args.output),
        Path(args.tor) if args.tor else None,
        Path(args.rab) if args.rab else None,
        Path(args.tor_dir) if args.tor_dir else None,
        Path(args.rab_dir) if args.rab_dir else None,
        args.judul, args.nomor, args.tanggal, args.penerima,
    )
    print(f"OK - written: {out_path}")
    print(f"   Mode: {info.get('mode')}")
    if info.get("mode") == "multi":
        print(f"   RO direviu: {info.get('n_ro')}")
        print(f"   Catatan: {info.get('n_anomalies')} "
              f"(PERINGATAN={info.get('n_peringatan')}, INFO={info.get('n_info')})")
        print(f"   A.3 dedup: {info.get('a3_dedup_ro_count')} RO digabung jadi 1 catatan")
        print(f"   Top rekomendasi: {info.get('n_top_recs')}")
    else:
        print(f"   Catatan ter-render: {info.get('n_anomalies')}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
