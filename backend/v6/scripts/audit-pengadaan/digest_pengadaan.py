"""
digest_pengadaan.py — Parser konsolidasi dokumen pengadaan → JSON terstruktur.

Scan folder penugasan, identifikasi dokumen berdasarkan pola nama file, lalu
parse setiap dokumen ke schema yang sesuai. Output menyertakan:
  - KAK (Kerangka Acuan Kerja)
  - HPS (Harga Perkiraan Sendiri)
  - Kontrak/SPK/SSUKSSKK
  - SPPBJ (Surat Penunjukan Penyedia)
  - BAST / BA Rekonsiliasi
  - Dokumen Pembayaran (LS, SPTB)

Usage:
    python digest_pengadaan.py <folder-penugasan> [-o output.json]

Pendekatan: regex + filename pattern matching. Cocok untuk audit-pengadaan
dan reviu-pengadaan (sharing schema).
"""

from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    sys.exit("ERROR: pdfplumber not installed. Run: pip install pdfplumber")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# ============================================================
# DOCUMENT TYPE CLASSIFICATION (by filename heuristic)
# ============================================================

FILENAME_PATTERNS = {
    "kak": [r"\bKAK\b", r"Kerangka\s+Acuan", r"TOR\b"],
    "hps": [r"\bHPS\b", r"Harga\s+Perkiraan"],
    "hps_detail": [r"Tabel\s+Penyusun.+HPS", r"Rekap\s+Komponen.+HPS"],
    "identifikasi_pengadaan": [r"Identifikasi\s+Pengadaan"],
    "rfi": [r"\bRFI\b", r"Request\s+For\s+Information"],
    "kontrak": [r"SSUKSSKK", r"Salinan\s+Jasa\s+Lainnya", r"Kontrak", r"\bSPK\b"],
    "sppbj": [r"SPPBJ"],
    "perjanjian_kerahasiaan": [r"Perjanjian\s+Kerahasiaan", r"NDA"],
    "permohonan_jaminan": [r"Permohonan\s+.+Jaminan"],
    "pembayaran_ls": [r"\bLS\b.+(Jan|Feb|Mar|Apr|Mei|Jun|Jul|Agu|Sep|Okt|Nov|Des)",
                      r"SPM.+LS", r"Pembayaran\s+LS"],
    "sptb": [r"SPTB", r"Surat\s+Pernyataan\s+Tanggungjawab\s+Belanja"],
    "ba_rekonsiliasi": [r"BA\s+Rekonsiliasi", r"Berita\s+Acara.+SLA"],
    # Dokumen PEMERIKSAAN/penerimaan hasil pekerjaan oleh PPK/PPHP/PjPHP/tim teknis
    # — INI pivot audit output-vs-kontrak (verifikasi kuantitas/spek barang diterima),
    # berbeda dari BAST (serah terima, sering formalitas tanda tangan).
    # Pola nama LONGGAR (nama dokumen beda tiap direktorat) — klasifikasi utama
    # untuk dokumen pemeriksaan ada di classify_content() berbasis ISI di bawah.
    "ba_pemeriksaan": [
        r"pemeriksa", r"penerimaan\s+hasil", r"PPHP", r"PjPHP",
        r"BA\s*P?HP\b", r"cek\s*fisik", r"uji\s*terima", r"serah\s*terima\s+barang",
    ],
    "laporan_bulanan": [r"Laporan\s+Bulanan"],
}

# ============================================================
# KLASIFIKASI BERBASIS ISI (general lintas direktorat)
# Dipakai HANYA bila klasifikasi nama file gagal — banyak dokumen (mis.
# pemeriksaan/penerimaan hasil pekerjaan) dinamai beda-beda tiap direktorat,
# jadi nama file tak bisa diandalkan. Deteksi dari FUNGSI dokumen di teksnya.
# ============================================================

_CONTENT_SIGNALS = {
    # Dokumen pemeriksaan/penerimaan hasil pekerjaan: butuh sinyal "aksi memeriksa"
    # + "konteks objek/kesesuaian" agar tak salah tangkap kontrak/laporan biasa.
    "ba_pemeriksaan": {
        "any": [r"pemeriksa", r"memeriksa", r"diperiksa", r"uji\s*terima",
                r"pemeriksaan\s+hasil", r"penerimaan\s+hasil", r"\bPPHP\b", r"\bPjPHP\b",
                r"berita\s+acara\s+(?:pemeriksaan|penerimaan|serah\s*terima)"],
        "context": [r"hasil\s+pekerjaan", r"barang", r"diterima", r"kesesuaian",
                    r"sesuai\s+(?:dengan\s+)?(?:kontrak|spesifikasi|spek|kak)",
                    r"kuantitas", r"volume", r"spesifikasi"],
    },
}


def classify_content(text: str) -> str | None:
    """Klasifikasi dokumen dari ISI (fallback bila nama file tak dikenal).

    General lintas direktorat: deteksi berdasarkan fungsi dokumen, bukan namanya.
    Butuh sinyal `any` (aksi) DAN `context` (objek) agar tidak salah klasifikasi.
    """
    if not text:
        return None
    head = text[:6000]
    for doc_type, sig in _CONTENT_SIGNALS.items():
        if any(re.search(p, head, re.I) for p in sig["any"]) and \
           any(re.search(p, head, re.I) for p in sig["context"]):
            return doc_type
    return None


def classify_file(filename: str) -> str | None:
    """Kembalikan jenis dokumen berdasarkan pola nama file."""
    for doc_type, patterns in FILENAME_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, filename, re.I):
                return doc_type
    return None


# ============================================================
# EXTRACTORS (per file type)
# ============================================================

def _extract_pdf_text(path: Path) -> list[str]:
    pages = []
    try:
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                pages.append(p.extract_text() or "")
    except Exception as e:
        return [f"[ERROR: {e}]"]
    return pages


def _extract_docx_text(path: Path) -> list[str]:
    """Ekstrak teks dari docx; kembalikan single-item list untuk kompatibilitas."""
    if DocxDocument is None:
        return ["[docx not supported — install python-docx]"]
    try:
        doc = DocxDocument(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # tambah teks dari tabel
        for t in doc.tables:
            for row in t.rows:
                text += "\n" + " | ".join(c.text for c in row.cells)
        return [text]
    except Exception as e:
        return [f"[ERROR: {e}]"]


def _extract_text(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx_text(path)
    return []


def _rupiah_to_int(s: str) -> int | None:
    if not s:
        return None
    m = re.search(r"([\d\.\,]+)", s)
    if not m:
        return None
    num = m.group(1).replace(".", "").replace(",", "")
    try:
        return int(num)
    except ValueError:
        return None


def _extract_periode(text: str) -> str | None:
    """Ekstrak periode pelaksanaan dari teks dokumen pengadaan.

    Aturan ketat untuk hindari false positive:
      1. Prefer pola dengan keyword konteks: "jangka waktu / selama / durasi /
         periode pelaksanaan / masa kontrak / masa pelaksanaan / waktu pelaksanaan".
      2. Bila tidak ada konteks, terima "(\\d{1,3})\\s*(bulan)" dengan angka 1-60.
         JANGAN terima "X tahun" tanpa konteks (terlalu rawan ketabrak nomor PP/UU
         seperti "PP No. 45 Tahun 2013").
      3. Reject angka >99 (jelas bukan durasi proyek).
      4. Reject jika angka berupa tahun (>=1900) atau berdekatan dengan kata
         "Tahun Anggaran".
    """
    # 1) Pola dengan konteks
    ctx_pattern = (
        r"(?:jangka\s+waktu|selama|durasi|periode\s+pelaksanaan|"
        r"masa\s+(?:kontrak|pelaksanaan)|waktu\s+pelaksanaan)"
        r"[^\n]{0,40}?(\d{1,2})\s*(?:\(\w+\))?\s*(bulan|tahun)"
    )
    for m in re.finditer(ctx_pattern, text, re.I):
        num = int(m.group(1))
        unit = m.group(2).lower()
        if unit == "bulan" and 1 <= num <= 60:
            return f"{num} bulan"
        if unit == "tahun" and 1 <= num <= 5:
            return f"{num} tahun"

    # 2) Fallback: hanya pola "X bulan" (1..60), bukan "X tahun"
    for m in re.finditer(r"(?<!\d)(\d{1,2})\s*(?:\(\w+\))?\s*bulan\b", text, re.I):
        num = int(m.group(1))
        if 1 <= num <= 60:
            # Cek 30-char di kiri: jangan kalau didahului "Pasal", "No.", dsb
            left = text[max(0, m.start()-30):m.start()].lower()
            if re.search(r"pasal|no\.|nomor|tahun\s+anggaran|t\.?a\.?\b", left):
                continue
            return f"{num} bulan"

    return None


# Komponen ruang lingkup yang lazim disebut di KAK namun kerap terlupa
# dialokasikan di HPS. Dipakai rule P.4 versi UMUM (lintas jenis pengadaan:
# barang/konstruksi/jasa/konsultansi) — bukan lagi spesifik migrasi/TI.
_LINGKUP_KOMPONEN = {
    "migrasi": r"\bmigrasi\b|migration",
    "instalasi": r"instalasi|pemasangan|\binstall",
    "pelatihan": r"pelatihan|training|bimbingan\s+teknis|bimtek",
    "pemeliharaan": r"pemeliharaan|maintenance|perawatan|masa\s+pemeliharaan",
    "garansi": r"garansi|warranty|jaminan\s+purna\s*jual",
    "pengujian": r"uji\s*coba|commissioning|pengujian|\bUAT\b|\bSIT\b",
    "lisensi": r"lisensi|license|langganan|subscription",
}

# Label manusiawi untuk narasi temuan (dipakai cross_check P.4).
_LINGKUP_LABEL = {
    "migrasi": "migrasi data/sistem", "instalasi": "instalasi/pemasangan",
    "pelatihan": "pelatihan/alih pengetahuan", "pemeliharaan": "pemeliharaan/perawatan",
    "garansi": "garansi/masa pemeliharaan", "pengujian": "pengujian/commissioning",
    "lisensi": "lisensi/langganan",
}


def _detect_lingkup_komponen(text: str) -> dict:
    """Deteksi presence-only komponen ruang lingkup di teks dokumen (heuristik keyword)."""
    return {k: bool(re.search(pat, text, re.I)) for k, pat in _LINGKUP_KOMPONEN.items()}


def parse_kak(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nama_pekerjaan": None,
        "nilai_hps": None, "periode": None, "sla_disebut": False,
        "sla_value": None, "sla_all_values": [], "migrasi_disebut": False,
        "lingkup_komponen": {}, "kapasitas_disebut": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    m = re.search(r"(\d{1,2}\s+\w+\s+\d{4})", text[:1500])
    if m:
        out["tanggal"] = m.group(1)
    # nama pekerjaan
    m = re.search(r"(?:Pengadaan|Penyediaan|Jasa)\s+([A-Z][^\n.]{10,200})", text[:2000])
    if m:
        out["nama_pekerjaan"] = m.group(0)[:200]
    # periode — STRICT: harus ada konteks ("jangka waktu/selama/durasi/periode/masa")
    # ATAU "X bulan" murni dengan X<=60. Reject angka >99 atau >1900 (tahun rujukan).
    out["periode"] = _extract_periode(text)
    # SLA - scan SELURUH text, tangkap semua angka SLA berbeda (untuk RP.11)
    if re.search(r"\bSLA\b|uptime|service\s+level", text, re.I):
        out["sla_disebut"] = True
        sm = re.search(r"(?:SLA|uptime|service\s+level).{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
        # Kumpulkan SEMUA angka SLA % yang muncul di seluruh dokumen
        all_sla = re.findall(r"(?:SLA|uptime|service\s+level)[^\n]{0,100}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        # Normalize titik → koma untuk konsistensi pembanding
        unique_sla = sorted(set(s.replace(".", ",") for s in all_sla))
        out["sla_all_values"] = [f"{v}%" for v in unique_sla]
    # migrasi (dipertahankan utk kompatibilitas) + komponen ruang lingkup umum
    if re.search(r"\bmigrasi\b|migration", text, re.I):
        out["migrasi_disebut"] = True
    out["lingkup_komponen"] = _detect_lingkup_komponen(text)
    # kapasitas
    km = re.search(r"(\d+(?:[,.]\d+)?)\s*(Tbps|Gbps|Mbps|TB|GB)", text, re.I)
    if km:
        out["kapasitas_disebut"] = km.group(0)
    # HPS jika ada
    hm = re.search(r"Rp\s*([\d\.,]+)", text)
    if hm:
        out["nilai_hps"] = _rupiah_to_int(hm.group(1))
    # Kelengkapan 5 elemen justifikasi/dokumen persiapan (deteksi keyword full-text).
    # Dipakai rule kelengkapan justifikasi (reviu/audit pengadaan). Heuristik
    # presence-only — bisa false-negative bila frasa tak lazim → rule berseverity
    # PERINGATAN + minta konfirmasi manual.
    out["elemen_justifikasi"] = {
        "kebutuhan": bool(re.search(
            r"latar\s+belakang|identifikasi\s+kebutuhan|analisis\s+kebutuhan|maksud\s+dan\s+tujuan",
            text, re.I)),
        "spesifikasi_teknis": bool(re.search(
            r"spesifikasi\s+teknis|persyaratan\s+teknis|spesifikasi\s+barang|spek\s+teknis|"
            r"spesifikasi\s+fungsi|kebutuhan\s+fungsional",
            text, re.I)),
        "metode_pengadaan": bool(re.search(
            r"metode\s+(?:pengadaan|pemilihan)|cara\s+pengadaan|tender\b|seleksi\b|"
            r"penunjukan\s+langsung|pengadaan\s+langsung|e-?purchasing|e-?katalog|e-?katalogue",
            text, re.I)),
        # waktu: reuse hasil _extract_periode atau frasa jadwal/waktu penyelesaian
        "waktu_penyelesaian": bool(out["periode"]) or bool(re.search(
            r"jangka\s+waktu|waktu\s+pelaksanaan|jadwal\s+pelaksanaan|masa\s+pelaksanaan|"
            r"waktu\s+penyelesaian|jadwal\s+kegiatan",
            text, re.I)),
        "output": bool(re.search(
            r"keluaran|output|deliverable|hasil\s+yang\s+diharapkan|hasil\s+pekerjaan|"
            r"produk\s+yang\s+dihasilkan|sasaran\s+keluaran",
            text, re.I)),
    }
    # Analisis kebutuhan (untuk rule reviu RP.13). Inti: yang penting ADA
    # IDENTIFIKASI KEBUTUHAN yang mendasari pengadaan — bukan asal sebut angka.
    # GENERIK lintas konteks (barang/jasa/konstruksi), bukan hanya komputer.
    # Heuristik presence-only — reviewer wajib konfirmasi.
    #
    # (a) kuantitas/volume disebut? — satuan UMUM + frasa "sebanyak/sejumlah N",
    #     bukan jenis barang spesifik (berlaku barang/jasa/konstruksi):
    out["kuantitas_pengadaan_disebut"] = bool(re.search(
        r"\b(?:sebanyak|sejumlah)\s+\d{1,6}\b|"
        r"\b\d{1,6}(?:[.,]\d{3})*\s*(?:unit|buah|set|pcs|paket|rim|dus|box|kotak|"
        r"lisensi|user|node|titik|lembar|orang[-\s]?bulan|\bOB\b|man[-\s]?month|"
        r"meter|\bm2\b|\bm3\b|\bkm\b|kg|ton|liter|kavling|bidang)\b",
        text, re.I))
    # (b) identifikasi kebutuhan yang MENDASARI pengadaan ada? — lebih ketat dari
    #     sekadar "latar belakang" naratif (itu sudah dicek RP.12). Butuh identifikasi/
    #     analisis/perhitungan kebutuhan ATAU dasar kuantitatif (pegawai/ABK/aset/standar).
    #     Catatan: "orang" telanjang TIDAK dipakai (tabrakan dgn satuan jasa "orang-bulan").
    out["identifikasi_kebutuhan"] = bool(re.search(
        r"identifikasi\s+kebutuhan|analisis\s+kebutuhan|kajian\s+kebutuhan|"
        r"perhitungan\s+kebutuhan|rincian\s+kebutuhan|dasar\s+(?:perhitungan|kebutuhan)|"
        r"justifikasi\s+(?:jumlah|kuantitas|volume|kebutuhan)|kebutuhan\s+riil|"
        r"jumlah\s+pegawai|sebanyak\s+\d+\s+(?:pegawai|orang(?![-\s]?bulan)|staf|personil|pejabat)|"
        r"\d+\s+(?:pegawai|staf|personil|pejabat)\b|analisis\s+beban\s+kerja|\bABK\b|beban\s+kerja|"
        r"unit\s+kerja|satuan\s+kerja|kondisi\s+(?:existing|saat\s+ini|eksisting)|"
        r"aset\s+(?:existing|eksisting|yang\s+(?:ada|dimiliki))|telah\s+dimiliki|barang\s+lama|"
        r"standar\s+(?:barang|kebutuhan|harga)|rasio\s+kebutuhan|kapasitas\s+terpasang",
        text, re.I))
    return out


def parse_hps(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "total": None,
        "komponen_count": 0, "sla_disebut": False, "sla_value": None,
        "sla_all_values": [],
        "periode": None, "migrasi_disebut": False, "lingkup_komponen": {},
        "ada_dokumen_pembentuk_harga": False, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # total HPS — biasanya ada "Total" atau "Jumlah" besar
    tm = re.search(r"(?:Total|Jumlah|Grand\s+Total)\s+(?:Rp\s*)?([\d\.,]{10,})", text)
    if tm:
        out["total"] = _rupiah_to_int(tm.group(1))
    # komponen count — baris dengan "Rp" di akhir
    rupiah_lines = re.findall(r"Rp\s*[\d\.,]{6,}", text)
    out["komponen_count"] = len(rupiah_lines)
    # SLA
    if re.search(r"\bSLA\b|uptime|service\s+level", text, re.I):
        out["sla_disebut"] = True
        sm = re.search(r"(?:SLA|uptime|service\s+level).{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
        all_sla = re.findall(r"(?:SLA|uptime|service\s+level)[^\n]{0,100}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        unique_sla = sorted(set(s.replace(".", ",") for s in all_sla))
        out["sla_all_values"] = [f"{v}%" for v in unique_sla]
    # periode — STRICT (lihat _extract_periode)
    out["periode"] = _extract_periode(text)
    # migrasi (kompatibilitas) + komponen ruang lingkup umum (untuk P.4)
    if re.search(r"\bmigrasi\b|migration", text, re.I):
        out["migrasi_disebut"] = True
    out["lingkup_komponen"] = _detect_lingkup_komponen(text)
    # dokumen pembentuk harga (quotation vendor, market research)
    if re.search(r"penawaran\s+vendor|quotation|market\s+research|RFI", text, re.I):
        out["ada_dokumen_pembentuk_harga"] = True
    return out


def parse_kontrak(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nilai_kontrak": None,
        "penyedia": None, "periode_mulai": None, "periode_selesai": None,
        "sla_clause": False, "sla_value": None, "jaminan_pelaksanaan": None,
        "metode_pembayaran": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # nilai kontrak
    nm = re.search(r"(?:Nilai\s+Kontrak|Total\s+Kontrak|sebesar|senilai)[^\n]{0,50}?Rp\s*([\d\.\,]+)", text, re.I)
    if nm:
        out["nilai_kontrak"] = _rupiah_to_int(nm.group(1))
    # penyedia
    pm = re.search(r"(?:Penyedia|Kontraktor)\s*:?\s*([A-Z][A-Za-z\s\.\,&]{5,80})", text)
    if pm:
        out["penyedia"] = pm.group(1).strip()[:80]
    # SLA
    if re.search(r"\bSLA\b|Service\s+Level\s+Agreement", text, re.I):
        out["sla_clause"] = True
        sm = re.search(r"SLA.{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
        if sm:
            out["sla_value"] = sm.group(1) + "%"
    # jaminan pelaksanaan (biasanya 5%)
    jm = re.search(r"[Jj]aminan\s+[Pp]elaksanaan.{0,100}?(\d+)\s*%", text)
    if jm:
        out["jaminan_pelaksanaan"] = f"{jm.group(1)}%"
    # metode pembayaran
    if re.search(r"\bLS\b|Langsung|Lump.?sum", text):
        out["metode_pembayaran"] = "LS/Lump-sum"
    elif re.search(r"[Bb]ulanan|termin", text):
        out["metode_pembayaran"] = "Termin/Bulanan"
    return out


def parse_bast(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "periode": None,
        "nilai_diterima": None, "sla_reported": None,
        "sla_passed": None, "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # periode reconciled
    m = re.search(r"(Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s*(\d{4})?", text, re.I)
    if m:
        out["periode"] = m.group(0)
    # SLA reported
    sm = re.search(r"SLA.{0,80}?(\d{2,3}(?:[,.]\d+)?)\s*%", text, re.I)
    if sm:
        out["sla_reported"] = sm.group(1) + "%"
    return out


def parse_pembayaran(pages: list[str]) -> dict:
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None, "nilai": None,
        "periode": None, "bukti_pendukung_lengkap": False,
        "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    nm = re.search(r"(?:sebesar|senilai)[^\n]{0,50}?Rp\s*([\d\.\,]+)", text, re.I)
    if nm:
        out["nilai"] = _rupiah_to_int(nm.group(1))
    # bukti pendukung
    if re.search(r"(?:BAST|Berita\s+Acara|BA\s+Rekonsiliasi|SPTB|Invoice|Kwitansi)", text, re.I):
        out["bukti_pendukung_lengkap"] = True
    return out


def parse_pemeriksaan(pages: list[str]) -> dict:
    """Dokumen pemeriksaan/penerimaan hasil pekerjaan (PPK/PPHP/PjPHP/tim teknis).

    Fokus audit output-vs-kontrak: apakah dokumen MERINCI kuantitas/spesifikasi
    per item (verifikasi nyata) atau hanya tanda tangan (formalitas), dan apakah
    mencatat ketidaksesuaian. Dipakai rule PL.2/PL.3 + analisis substantif #5.
    """
    text = "\n".join(pages)
    out = {
        "nomor": None, "tanggal": None,
        "rincian_per_item": False,          # ada daftar/tabel kuantitas-spesifikasi diperiksa
        "mencatat_ketidaksesuaian": False,  # dokumen mencatat kurang/cacat/tidak sesuai
        "halaman": len(pages),
    }
    m = re.search(r"Nomor\s*:?\s*(\S+)", text[:1500])
    if m:
        out["nomor"] = m.group(1)[:80]
    # rincian per item: indikasi kuantitas/spesifikasi yang diperiksa (bukan sekadar tanda tangan)
    qty_lines = re.findall(
        r"\b\d+\s*(?:unit|buah|set|pcs|lembar|paket|meter|m2|m3|kg|liter|titik|lisensi|user|node)\b",
        text, re.I)
    has_kolom = bool(re.search(
        r"kuantitas|kuantiti|volume|jumlah\s+diterima|spesifikasi\s+diterima|"
        r"hasil\s+pemeriksaan|sesuai\s*/\s*tidak\s+sesuai", text, re.I))
    out["rincian_per_item"] = len(qty_lines) >= 3 or (has_kolom and len(qty_lines) >= 1)
    # mencatat ketidaksesuaian (kurang/cacat/selisih)
    if re.search(r"\bkurang\b|tidak\s+sesuai|selisih|cacat|tidak\s+lengkap|kekurangan|reject",
                 text, re.I):
        out["mencatat_ketidaksesuaian"] = True
    return out


PARSERS = {
    "kak": parse_kak,
    "hps": parse_hps,
    "kontrak": parse_kontrak,
    "ba_rekonsiliasi": parse_bast,
    "ba_pemeriksaan": parse_pemeriksaan,
    "pembayaran_ls": parse_pembayaran,
    "sptb": parse_pembayaran,
}


# ============================================================
# RFI PARSER (untuk multi-source HPS validation)
# ============================================================

_RE_REFUSAL = re.compile(
    r"(belum\s+dapat\s+(?:berpartisipasi|ber\s*partisipasi|ikut|mengikuti)|"
    r"tidak\s+dapat\s+(?:berpartisipasi|memberikan|menyampaikan|ikut)|"
    r"menolak|"
    r"decline|not\s+interested|cannot\s+participate|unable\s+to\s+participate|"
    r"belum\s+sanggup|tidak\s+bersedia)",
    re.I,
)
_RE_PRICE = re.compile(r"(Rp\s*[\d.,]+|IDR\s*[\d.,]+|USD\s*[\d.,]+|harga\s+(?:total|estimasi|perkiraan))", re.I)


def parse_rfi(pages: list[str]) -> dict:
    """Parse RFI: deteksi apakah vendor memberikan harga atau menolak partisipasi."""
    text = "\n".join(pages) if pages else ""
    refusal = bool(_RE_REFUSAL.search(text))
    has_price = bool(_RE_PRICE.search(text))
    # Vendor: cari nama vendor di header atau signature
    vendor_match = re.search(r"(?:^|\n)\s*(?:PT\.?\s+|CV\.?\s+|VP\s+|Hormat\s+kami,?\s*\n+\s*)([A-Z][A-Z\s&.,-]{4,60})", text)
    vendor = vendor_match.group(1).strip() if vendor_match else None
    return {
        "vendor_terdeteksi": vendor,
        "memberikan_harga": has_price and not refusal,
        "menolak_partisipasi": refusal,
        "punya_indikasi_harga": has_price,
        "halaman": len(pages) if pages else 0,
    }


PARSERS["rfi"] = parse_rfi


# ============================================================
# FOLDER SCANNER
# ============================================================

def scan_folder(folder: Path) -> dict:
    """Scan folder, klasifikasi + parse setiap dokumen."""
    out = {
        "metadata": {
            "folder_source": str(folder),
            "parser_version": "v0.2",
        },
        "dokumen": {},
        "unclassified_files": [],
        "missing_types": [],
    }

    # walk
    all_files = list(folder.rglob("*"))
    for f in all_files:
        if not f.is_file():
            continue
        if f.suffix.lower() not in (".pdf", ".docx", ".doc"):
            continue
        if f.name == "desktop.ini" or f.name.startswith("~"):
            continue

        doc_type = classify_file(f.name)
        pages = None
        classified_by = "nama"
        if not doc_type:
            # Nama file tak dikenal → baca ISI lalu klasifikasi berbasis fungsi
            # (general lintas direktorat). Pages di-cache agar tak baca dua kali.
            pages = _extract_text(f)
            doc_type = classify_content("\n".join(pages) if pages else "")
            classified_by = "isi"
            if not doc_type:
                out["unclassified_files"].append(str(f.relative_to(folder)))
                continue

        parser = PARSERS.get(doc_type)
        entry = {
            "filename": f.name,
            "path": str(f.relative_to(folder)),
            "jenis_dokumen": doc_type,
            "classified_by": classified_by,
            "parsed": None,
        }
        if parser:
            try:
                if pages is None:
                    pages = _extract_text(f)
                entry["parsed"] = parser(pages)
                entry["parsed"]["_raw_first_chars"] = ("\n".join(pages))[:2500] if pages else ""
            except Exception as e:
                entry["parsed"] = {"_error": str(e)}

        out["dokumen"].setdefault(doc_type, []).append(entry)

    # Cek missing types yang penting
    penting = ["kak", "hps", "kontrak"]
    for t in penting:
        if t not in out["dokumen"]:
            out["missing_types"].append(t)

    return out


def _self_check_ast() -> None:
    """Preflight: pastikan script ini sendiri syntactically valid sebelum run."""
    import ast
    try:
        ast.parse(open(__file__, "r", encoding="utf-8").read())
    except SyntaxError as e:
        print(f"Self-check AST gagal di {__file__}: {e}", file=sys.stderr)
        print("   File mungkin korup. Lihat backup atau git restore.", file=sys.stderr)
        sys.exit(2)


def main(argv=None):
    _self_check_ast()
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input_dir", help="Folder berisi PDF/dokumen pengadaan")
    ap.add_argument("-o", "--output", default="pengadaan-digest.json")
    args = ap.parse_args(argv)

    in_dir = Path(args.input_dir)
    if not in_dir.exists():
        print(f"Input dir tidak ditemukan: {in_dir}", file=sys.stderr)
        return 1

    digest = scan_folder(in_dir)
    Path(args.output).write_text(json.dumps(digest, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"OK: {args.output}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
