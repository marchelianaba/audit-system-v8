"""
render_lha.py — Render anomalies + context.md → draft LHA (Laporan Hasil Audit).

Usage:
    python render_lha.py <anomalies.json> <context.md> [-o LHA-DRAFT.docx]

Output: draft LHA dengan struktur:
  - Nota Dinas Pengantar
  - Cover LHA
  - Ringkasan Eksekutif
  - Dasar Hukum, Tujuan, Ruang Lingkup, Metodologi
  - Uraian Hasil Audit (per aspek: Perencanaan/Kontrak/Pelaksanaan/Pembayaran)
  - Simpulan
  - Rekomendasi
  - Apresiasi
"""

from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")


ASPEK_TITLES = {
    "Perencanaan": "Aspek Perencanaan Pengadaan (KAK, HPS, Dokumen Pembentuk Harga)",
    "Kontrak": "Aspek Kontrak (SPPBJ, Klausul, Jaminan)",
    "Pelaksanaan": "Aspek Pelaksanaan Kontrak (BAST, SLA, Rekonsiliasi)",
    "Pembayaran": "Aspek Pembayaran (LS, SPTB, Kelengkapan Bukti)",
    "Dokumentasi": "Aspek Dokumentasi dan Ketertiban Arsip",
}


def _parse_context_md(path: Path) -> dict:
    if not path.exists():
        return {}
    text = path.read_text(encoding="utf-8")
    ctx = {}
    for m in re.finditer(r"^([A-Z][A-Za-z\s]+?)\s*:\s*(.+?)$", text, re.M):
        ctx[m.group(1).strip()] = m.group(2).strip()
    return ctx


def _add_h(doc, text, size=12, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _add_p(doc, text, bold=False, italic=False, size=11, align=None, indent=None, justify=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def render_lha(anomalies_path: Path, context_path: Path, output_path: Path):
    ano = json.loads(anomalies_path.read_text(encoding="utf-8"))
    ctx = _parse_context_md(context_path)
    anomalies = [a for a in ano.get("anomalies", []) if a.get("severity") != "ERROR"]

    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # === Nota Dinas ===
    for l in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI", "INSPEKTORAT JENDERAL", "INSPEKTORAT II"]:
        _add_p(doc, l, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _add_p(doc, "Jl. Medan Merdeka Barat No. 9, Jakarta 10110", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True)
    doc.add_paragraph()
    _add_h(doc, "NOTA DINAS", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_p(doc, "Nomor: [DIISI AUDITOR — dari SIMWAS]/IJ.3/PW.04.04/[bulan]/[tahun]",
           align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    obyek = ctx.get("Obyek", "[Obyek Pengadaan]")
    unit = ctx.get("Unit yang Diawasi", "[Unit]")
    nomor_st = ctx.get("Nomor ST", "[Nomor ST]")
    tgl_st = ctx.get("Tanggal ST", "[Tanggal ST]")
    penerima = ctx.get("Penerima Laporan", "[Penerima]")

    fields = [
        ("Kepada Yth.", penerima),
        ("Dari", "Inspektur II"),
        ("Hal", f"Laporan Hasil Audit Pengadaan — {obyek}"),
        ("Klasifikasi", "Terbatas"),
        ("Lampiran", "Satu laporan"),
        ("Tanggal", "[DIISI AUDITOR]"),
    ]
    t = doc.add_table(rows=len(fields), cols=3)
    t.autofit = False
    for row in t.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(0.5)
        row.cells[2].width = Cm(13.0)
    for i, (k, v) in enumerate(fields):
        for j, val in enumerate([k, ":", v]):
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(11)
            if j == 0:
                r.bold = True
    doc.add_paragraph()

    _add_p(doc,
        f"Menindaklanjuti Surat Tugas Inspektur II Nomor {nomor_st} tanggal {tgl_st} perihal audit "
        f"pengadaan atas {obyek} di {unit}, bersama ini kami sampaikan Laporan Hasil Audit (LHA) "
        f"sebagaimana terlampir.",
        justify=True)
    _add_p(doc, "Terima kasih telah membantu kami dalam menjaga integritas.", justify=True)
    doc.add_paragraph()
    _add_p(doc, "[TTD — DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, italic=True)
    _add_p(doc, "Muhammad Arief", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    _add_p(doc, "Inspektur II", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    doc.add_page_break()

    # === Cover LHA ===
    for l in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI", "INSPEKTORAT JENDERAL"]:
        _add_p(doc, l, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()
    _add_p(doc, "LAPORAN HASIL AUDIT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _add_p(doc, "AUDIT PENGADAAN BARANG/JASA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph()
    _add_p(doc, obyek.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _add_p(doc, unit.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()

    # === 1. Ringkasan Eksekutif ===
    _add_h(doc, "1. RINGKASAN EKSEKUTIF", size=12)
    n = len(anomalies)
    _add_p(doc,
        f"Inspektorat II telah melakukan audit pengadaan atas {obyek} di {unit} dengan pendekatan "
        f"keyakinan memadai (reasonable assurance) sesuai Standar Audit Intern Pemerintah. Audit "
        f"mencakup aspek Perencanaan (KAK, HPS), Pemilihan Penyedia, Kontrak, Pelaksanaan, dan "
        f"Pembayaran.",
        justify=True)
    _add_p(doc,
        f"Berdasarkan hasil audit, diidentifikasi {n} temuan yang perlu ditindaklanjuti:",
        justify=True)
    by_aspek = ano.get("summary_by_aspek", {})
    for aspek, cnt in by_aspek.items():
        _add_p(doc, f"  {aspek}: {cnt} temuan", indent=0.5)

    # === 2-5 ===
    doc.add_paragraph()
    _add_h(doc, "2. DASAR HUKUM", size=12)
    for d in [
        "a. Peraturan Pemerintah Nomor 60 Tahun 2008 tentang Sistem Pengendalian Intern Pemerintah;",
        "b. Peraturan Presiden Nomor 16 Tahun 2018 jo. Perpres 12/2021 dan Perpres 46/2025 tentang "
        "Pengadaan Barang/Jasa Pemerintah;",
        "c. Peraturan Menteri Keuangan Nomor 190/PMK.05/2012 tentang Tata Cara Pembayaran atas "
        "Beban APBN;",
        "d. Peraturan Lembaga LKPP Nomor 12 Tahun 2021 tentang Pedoman Pelaksanaan Pengadaan;",
        f"e. Surat Tugas Inspektur II Nomor {nomor_st} tanggal {tgl_st}.",
    ]:
        _add_p(doc, d, indent=0.5, justify=True)

    doc.add_paragraph()
    _add_h(doc, "3. TUJUAN AUDIT", size=12)
    _add_p(doc,
        "Tujuan audit adalah memberikan keyakinan memadai (reasonable assurance) bahwa pengadaan "
        "barang/jasa telah dilaksanakan sesuai dengan ketentuan peraturan perundang-undangan, "
        "prinsip pengadaan, dan kontrak yang berlaku, serta memberikan rekomendasi perbaikan "
        "atas kelemahan yang ditemukan.",
        justify=True)

    doc.add_paragraph()
    _add_h(doc, "4. RUANG LINGKUP", size=12)
    _add_p(doc,
        f"Audit mencakup seluruh tahapan pengadaan {obyek} — mulai perencanaan (KAK, HPS), pemilihan "
        f"penyedia (SPPBJ), penandatanganan kontrak, pelaksanaan (BAST, monitoring SLA), hingga "
        f"pembayaran (LS, SPTB, bukti pendukung). Pengujian dilakukan atas kecukupan dokumen, "
        f"kesesuaian prosedur, konsistensi nilai antar dokumen, dan kepatuhan terhadap regulasi.",
        justify=True)

    doc.add_paragraph()
    _add_h(doc, "5. METODOLOGI", size=12)
    _add_p(doc,
        "Audit dilaksanakan sesuai Standar Audit Intern Pemerintah (AAIPI). Metodologi mencakup: "
        "(a) pre-digest dokumen pengadaan menggunakan pipeline deterministik untuk klasifikasi "
        "dan ekstraksi terstruktur; (b) cross-check rule-based atas konsistensi antar dokumen "
        "(KAK ↔ HPS ↔ Kontrak ↔ Pembayaran); (c) verifikasi substantif atas temuan yang memerlukan "
        "judgment; (d) konfirmasi kepada auditan atas kondisi yang berisiko kerugian negara.",
        justify=True)

    # === 6. Uraian Hasil Audit (per aspek) ===
    doc.add_paragraph()
    _add_h(doc, "6. URAIAN HASIL AUDIT", size=12)

    per_aspek = {}
    for a in anomalies:
        per_aspek.setdefault(a.get("aspek", "?"), []).append(a)

    for aspek, title in ASPEK_TITLES.items():
        if aspek not in per_aspek:
            continue
        _add_h(doc, f"6.{aspek[0]}. {title}", size=11)
        for a in per_aspek[aspek]:
            draft = a.get("draft_catatan") or {}
            _add_p(doc, f"{a.get('rule_id','?')}. {a.get('judul','?')}", bold=True, size=11)
            for label, key in [("Kondisi", "kondisi"), ("Kriteria", "kriteria"),
                                ("Sebab", "sebab"), ("Akibat", "akibat")]:
                val = draft.get(key)
                if not val:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                r = p.add_run(f"{label}: ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(val)
                r2.font.size = Pt(10)

    # === 7. Simpulan ===
    doc.add_paragraph()
    _add_h(doc, "7. SIMPULAN", size=12)
    _add_p(doc,
        f"Berdasarkan hasil audit pengadaan atas {obyek}, ditemukan {n} temuan yang perlu "
        f"ditindaklanjuti oleh auditan. Sepanjang temuan tersebut ditindaklanjuti secara memadai, "
        f"tidak terdapat hal-hal material yang membuat kami yakin bahwa pengadaan tidak dilaksanakan "
        f"sesuai ketentuan peraturan perundang-undangan.",
        justify=True)

    doc.add_paragraph()
    _add_h(doc, "8. REKOMENDASI", size=12)
    _add_p(doc,
        "Rekomendasi atas temuan dirumuskan dalam butir-butir terlampir di bagian 6. "
        "[Rumusan rekomendasi final dilakukan setelah gate auditor — DIISI AUDITOR]",
        justify=True, italic=True)

    doc.add_paragraph()
    _add_h(doc, "9. APRESIASI", size=12)
    _add_p(doc,
        f"Inspektorat II menyampaikan terima kasih atas kerja sama seluruh pejabat/pegawai "
        f"di lingkungan {unit} selama pelaksanaan audit ini.",
        justify=True)

    doc.add_paragraph()
    _add_p(doc, "[TTD — DIISI AUDITOR]", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, italic=True)
    _add_p(doc, "Inspektur II,", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    _add_p(doc, "Muhammad Arief", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path, n


def _self_check_ast() -> None:
    """Preflight: pastikan script ini sendiri syntactically valid sebelum run.
    Mencegah eksekusi dengan file korup (mis. akibat OneDrive sync artifact)."""
    import ast
    try:
        ast.parse(open(__file__, "r", encoding="utf-8").read())
    except SyntaxError as e:
        print(f"Self-check AST gagal di {__file__}: {e}", file=sys.stderr)
        print("   File mungkin korup. Lihat backup atau git restore.", file=sys.stderr)
        sys.exit(2)


def main(argv=None):
    _self_check_ast()
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("anomalies")
    ap.add_argument("context")
    ap.add_argument("-o", "--output", default="LHA-DRAFT.docx")
    args = ap.parse_args(argv)
    out, n = render_lha(Path(args.anomalies), Path(args.context), Path(args.output))
    print(f"OK — written: {out}")
    print(f"   Temuan ter-render: {n}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
