#!/usr/bin/env python3
"""
render_kkp.py — Renderer KKP DOCX dari _KKP/temuan.json (v4.0.4).

Input  : penugasan/[nama]/_KKP/temuan.json + penugasan/[nama]/context.md
Output : penugasan/[nama]/_KKP/KKP-[Nama-Anggota].docx

Konsumsi sumber kebenaran (temuan.json) dan menghasilkan view per anggota
tim sesuai filter `anggota_tim.nama_lengkap`. Format kolom mengikuti
paradigma jenis pengawasan:
  - audit / reviu / evaluasi non-LKE: No, Sasaran, Judul, Kondisi, Kriteria, Sebab, Akibat, Sumber
  - pemantauan-*                   : No, Sasaran, Judul, Kondisi, Kriteria, Sebab, Sumber   (tanpa Akibat)
  - evaluasi ber-LKE (RB/SAKIP/SPIP): No, Sasaran, Judul, Kondisi, Kriteria, Akibat, Sumber  (TANPA Sebab — rezim LKE)
  Sejak 17 Jun 2026 kolom Sebab ada di semua jenis ber-KKSA (anti-mengarang: "Tidak ditemukan
  penyebab"/"Tidak cukup data" bila tak terbukti). KECUALI evaluasi ber-LKE (RB/SAKIP/SPIP) yang
  memakai instrumen LKE + AoI, bukan KKSA → tanpa kolom Sebab.
  - konsultasi-pengadaan           : No, Pertanyaan, Analisis, Dasar Hukum, Rekomendasi    (kolom khusus)

Contoh:
  python3 scripts/render_kkp.py --penugasan penugasan/[nama] --anggota "Sarah Aulia"
  python3 scripts/render_kkp.py --penugasan penugasan/[nama] --all-anggota
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx tidak terpasang. Jalankan: pip install python-docx\n")
    sys.exit(2)


JENIS_KOLOM = {
    "audit-pengadaan":          ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Akibat", "Sumber Dokumen"],
    "audit-kinerja":            ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Akibat", "Sumber Dokumen"],
    "reviu-pengadaan":          ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Akibat", "Sumber Dokumen"],
    "reviu-rka-kl":             ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Akibat", "Sumber Dokumen"],
    # Evaluasi ber-LKE (RB/SAKIP/SPIP): rezim LKE — instrumen + AoI, BUKAN KKSA → tanpa kolom Sebab
    "evaluasi-sakip":           ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Akibat", "Sumber Dokumen"],
    "evaluasi-spip":            ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Akibat", "Sumber Dokumen"],
    "evaluasi-reformasi-birokrasi": ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Akibat", "Sumber Dokumen"],
    "evaluasi-manajemen-risiko":   ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Akibat", "Sumber Dokumen"],
    "pemantauan-pengadaan":     ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Sumber Dokumen"],
    "pemantauan-tindak-lanjut": ["No", "Sasaran", "Judul Temuan", "Kondisi", "Kriteria", "Sebab", "Sumber Dokumen"],
}

WIDTHS = {
    8: [0.7, 1.2, 4.0, 5.5, 4.5, 3.5, 4.0, 4.0],   # 8 kolom (audit)
    7: [0.8, 1.2, 4.5, 6.5, 5.5, 4.5, 4.5],         # 7 kolom (reviu/evaluasi)
    6: [0.8, 1.2, 4.5, 7.5, 6.0, 5.5],              # 6 kolom (pemantauan)
}


def shade(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_p(doc, text, bold=False, italic=False, size=11, align=None, font="Arial"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    return p


def parse_context_meta(context_path: Path) -> dict:
    """Ambil identitas penugasan dari context.md (gunakan tabel Markdown)."""
    out = {"nomor_st": "", "tanggal_st": "", "obyek": "", "ketua_tim": "",
           "periode": "", "tahun_anggaran": ""}
    if not context_path.exists():
        return out
    text = context_path.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("|"):
            parts = [c.strip() for c in line.split("|") if c.strip()]
            if len(parts) >= 2:
                k = parts[0].lower()
                v = parts[1]
                if "nomor st" in k:
                    out["nomor_st"] = v
                elif "tanggal st" in k:
                    out["tanggal_st"] = v
                elif "objek" in k or "obyek" in k:
                    if not out["obyek"]:
                        out["obyek"] = v
                elif "periode pelaksanaan" in k:
                    out["periode"] = v
                elif "tahun anggaran" in k:
                    out["tahun_anggaran"] = v
        # Tabel tim — cari Ketua Tim
        if "Ketua Tim" in line and out["ketua_tim"] == "":
            parts = [c.strip() for c in line.split("|") if c.strip()]
            for i, p in enumerate(parts):
                if p == "Ketua Tim" and i > 0:
                    # nama biasanya 2 kolom sebelum
                    out["ketua_tim"] = parts[1] if len(parts) > 1 else ""
                    break
    return out


# Trio evaluasi ber-LKE: KKP memuat rekap skor/predikat (dari penilaian-lke-<skill>.json
# yang ditulis tool write_penilaian_lke) + tabel AoI. Sumber skor tunggal = file ini.
_LKE_SKILLS = {"evaluasi-sakip", "evaluasi-spip", "evaluasi-reformasi-birokrasi"}


def render_lke_recap(doc, penugasan_dir: Path, jenis: str) -> bool:
    """Render tabel 'Rekap Penilaian (LKE)' untuk skill ber-LKE bila
    _KKP/penilaian-lke-<jenis>.json ada. Return True bila ter-render."""
    if jenis not in _LKE_SKILLS:
        return False
    src = penugasan_dir / "_KKP" / f"penilaian-lke-{jenis}.json"
    if not src.exists():
        return False
    try:
        pen = json.loads(src.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return False
    komponen = pen.get("komponen") or []
    if not komponen:
        return False
    add_p(doc, "REKAP PENILAIAN (LKE)", bold=True, size=12)
    cols = ["No", "Komponen / Unsur", "Bobot", "Nilai PM (auditee)", "Nilai APIP", "Predikat"]
    t = doc.add_table(rows=len(komponen) + 1, cols=len(cols))
    t.style = "Table Grid"
    for j, h in enumerate(cols):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, "D9D9D9")
    for i, k in enumerate(komponen, start=1):
        vals = [str(i), str(k.get("nama", "")), str(k.get("bobot", "")),
                str(k.get("nilai_pm", "")), str(k.get("nilai_apip", "")), str(k.get("predikat", ""))]
        for j, v in enumerate(vals):
            cell = t.rows[i].cells[j]; cell.text = ""
            cell.paragraphs[0].add_run(v).font.size = Pt(9)
    total_apip = pen.get("total_apip") or pen.get("nilai_akhir") or ""
    total_pm = pen.get("total_pm") or ""
    predikat = pen.get("predikat_akhir") or pen.get("predikat") or ""
    ringkas = "Hasil Penilaian APIP"
    if total_apip != "":
        ringkas += f": {total_apip}"
    if predikat:
        ringkas += f" (Predikat {predikat})"
    if total_pm != "":
        ringkas += f" — Penilaian Mandiri auditee: {total_pm}"
    add_p(doc, ringkas, bold=True, size=10)
    doc.add_paragraph()
    return True


def render_kkp_for_anggota(penugasan_dir: Path, anggota_nama: str) -> Path:
    kkp_path = penugasan_dir / "_KKP" / "temuan.json"
    if not kkp_path.exists():
        raise FileNotFoundError(f"temuan.json tidak ditemukan: {kkp_path}")

    data = json.loads(kkp_path.read_text(encoding="utf-8"))
    jenis = data["penugasan"]["jenis_pengawasan"]
    kolom = JENIS_KOLOM.get(jenis, JENIS_KOLOM["reviu-pengadaan"])

    my_temuan = [t for t in data["temuan"] if t["anggota_tim"]["nama_lengkap"] == anggota_nama]
    ctx = parse_context_meta(penugasan_dir / "context.md")

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)
        s.orientation = 1  # landscape
        s.page_width, s.page_height = s.page_height, s.page_width

    # Header
    add_p(doc, "KERTAS KERJA PENGAWASAN — DRAFT",
          bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"{jenis.replace('-', ' ').title()} | {data['penugasan']['obyek']}",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    info_line = f"ST: {data['penugasan']['nomor_st']}"
    if ctx.get("periode"):
        info_line += f" | Periode: {ctx['periode']}"
    if ctx.get("ketua_tim"):
        info_line += f" | Ketua Tim: {ctx['ketua_tim']}"
    add_p(doc, info_line, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    sasaran_list = sorted({t["sasaran_id"] for t in my_temuan}) if my_temuan else []
    add_p(doc, f"Anggota Pengisi KKP: {anggota_nama} | Sasaran: {', '.join(sasaran_list) if sasaran_list else '(tidak ada temuan)'}",
          size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    # Skill ber-LKE: rekap skor/predikat dulu (kertas kerja mencerminkan instrumen), baru AoI
    render_lke_recap(doc, penugasan_dir, jenis)
    add_p(doc, "DAFTAR CATATAN / TEMUAN", bold=True, size=12)

    if not my_temuan:
        add_p(doc, "(Tidak ada temuan untuk anggota ini.)", italic=True, size=10)
    else:
        t = doc.add_table(rows=len(my_temuan) + 1, cols=len(kolom))
        t.style = "Table Grid"

        # Header row
        for j, h in enumerate(kolom):
            c = t.rows[0].cells[j]; c.text = ""
            run = c.paragraphs[0].add_run(h)
            run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade(c, "D9D9D9")

        widths = WIDTHS.get(len(kolom), [3.0] * len(kolom))
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)

        for i, temuan in enumerate(my_temuan, start=1):
            sumber_text = "\n\n".join(
                f"• {ds['file']} (hal. {ds['halaman']}):\n  \"{ds['kutipan']}\""
                for ds in temuan.get("dokumen_sumber", [])
            )
            row_vals = []
            for col in kolom:
                if col == "No":
                    row_vals.append(str(i))
                elif col == "Sasaran":
                    row_vals.append(temuan["sasaran_id"])
                elif col == "Judul Temuan":
                    judul = temuan["judul_temuan"]
                    # Kodefikasi temuan (SIM-HP) — tampilkan di bawah judul bila ada.
                    _kodes = []
                    if temuan.get("kode_kondisi"):
                        _kodes.append(f"Kondisi {temuan['kode_kondisi']}")
                    if temuan.get("kode_penyebab"):
                        _kodes.append(f"Penyebab {temuan['kode_penyebab']}")
                    if temuan.get("kode_rekomendasi"):
                        _kodes.append(f"Rekomendasi {temuan['kode_rekomendasi']}")
                    if _kodes:
                        judul += "\n[Kode: " + " · ".join(_kodes) + "]"
                    row_vals.append(judul)
                elif col == "Kondisi":
                    row_vals.append(temuan["kondisi"])
                elif col == "Kriteria":
                    row_vals.append(temuan["kriteria"])
                elif col == "Sebab":
                    row_vals.append(temuan.get("sebab") or "—")
                elif col == "Akibat":
                    row_vals.append(temuan["akibat"])
                elif col == "Sumber Dokumen":
                    row_vals.append(sumber_text)
                else:
                    row_vals.append("")
            for j, v in enumerate(row_vals):
                c = t.rows[i].cells[j]; c.text = ""
                for line in v.split("\n"):
                    p = c.add_paragraph()
                    p.add_run(line).font.size = Pt(9)
                p0 = c.paragraphs[0]
                if p0.text == "":
                    p0._element.getparent().remove(p0._element)

    doc.add_paragraph()
    if "Sebab" in kolom:
        catatan_sebab = ("Catatan: kolom Sebab diisi berdasarkan bukti; bila penyebab tidak ditemukan "
                         "atau bukti tidak cukup, ditulis \"Tidak ditemukan penyebab\"/\"Tidak cukup data\" "
                         "(tidak dikarang). Rekomendasi final dirumuskan saat penyusunan Laporan Hasil bersama Ketua Tim.")
    else:
        catatan_sebab = ("Catatan: penilaian memakai instrumen Lembar Kerja Evaluasi (LKE) — skor/predikat per "
                         "kriteria/unsur + Area of Improvement (AoI); tidak memuat unsur Sebab (bukan format KKSA). "
                         "Rekomendasi final dirumuskan saat penyusunan Laporan Hasil bersama Ketua Tim.")
    add_p(doc, catatan_sebab, italic=True, size=9)
    doc.add_paragraph()
    add_p(doc, "⚠ KKP INI ADALAH DRAFT — Mohon berikan feedback ke Ketua Tim sebelum LHP dibuat.",
          bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # File name: KKP-[Nama-Tanpa-Spasi].docx
    safe_name = anggota_nama.replace(",", "").replace(".", "").replace("  ", " ")
    safe_name = "-".join(safe_name.split())
    out = penugasan_dir / "_KKP" / f"KKP-{safe_name}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main() -> int:
    ap = argparse.ArgumentParser(description="Render KKP DOCX dari _KKP/temuan.json")
    ap.add_argument("--penugasan", required=True, help="Folder penugasan (mis. test/uji-coba-XXX)")
    grp = ap.add_mutually_exclusive_group(required=True)
    grp.add_argument("--anggota", help="Nama lengkap anggota tim (sesuai temuan.json)")
    grp.add_argument("--all-anggota", action="store_true", help="Render KKP untuk semua anggota yang punya temuan")
    args = ap.parse_args()

    pen_dir = Path(args.penugasan)
    if not pen_dir.exists():
        sys.stderr.write(f"Folder tidak ada: {pen_dir}\n")
        return 1

    kkp_json = pen_dir / "_KKP" / "temuan.json"
    if not kkp_json.exists():
        sys.stderr.write(f"temuan.json belum ada di {kkp_json} — jalankan Task 03 dulu\n")
        return 1

    if args.all_anggota:
        data = json.loads(kkp_json.read_text(encoding="utf-8"))
        anggota_list = sorted({t["anggota_tim"]["nama_lengkap"] for t in data["temuan"]})
        if not anggota_list:
            print("(Tidak ada temuan dalam temuan.json)")
            return 0
        for nama in anggota_list:
            out = render_kkp_for_anggota(pen_dir, nama)
            print(f"OK: {out}")
    else:
        out = render_kkp_for_anggota(pen_dir, args.anggota)
        print(f"OK: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
