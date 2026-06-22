"""Renderer DETERMINISTIK untuk Laporan Survei Pendahuluan (.docx).

Khusus skill audit (audit-umum, audit-pengadaan, audit-kinerja). Survei
pendahuluan adalah tahap A0 — orientasi obyek, pemetaan risiko, inventarisasi
dokumen, dan hipotesis area pengujian SEBELUM analisis substantif.

File ini di-generate BERSAMAAN dengan context.md (sumber data sama):
  - context.md  → markdown, dibaca agen
  - Survei .docx → laporan Word, bisa di-download auditor

Dibangun DARI NOL dengan python-docx — TIDAK memakai/menyentuh template .docx
apa pun. Format paragraf: Arial 12pt, spasi 1,5 (selaras laporan lain).

Sumber data:
  - context.md  (Identitas, Tujuan, Ruang Lingkup, Gambaran Umum, Ringkasan Obyek)
  - _INGESTED/*.json digest (inventarisasi dokumen + petunjuk risiko)
  - _PKP/sasaran-assignment.json (tim + sasaran → hipotesis awal)
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Skill audit yang WAJIB punya survei pendahuluan.
AUDIT_SKILLS = {"audit-umum", "audit-pengadaan", "audit-kinerja"}

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def is_audit_skill(skill: str | None) -> bool:
    return (skill or "").strip().lower() in AUDIT_SKILLS


# ---------------------------------------------------------------------------
# Parsing context.md → dict of sections
# ---------------------------------------------------------------------------

def _clean(s: str) -> str:
    """Buang penanda bold/backtick/spasi di tepi (untuk key & value identitas)."""
    return re.sub(r"^[*\s]+|[*\s]+$", "", s).strip()


def _parse_md_table(text: str) -> list[dict[str, str]]:
    """Parse tabel markdown (baris pertama = header) → list dict per baris.

    Abaikan baris pemisah '|---|---|'. Toleran terhadap teks lain di sekitar tabel.
    """
    rows: list[list[str]] = []
    for line in text.splitlines():
        ls = line.strip()
        if not ls.startswith("|"):
            continue
        cells = [c.strip() for c in ls.strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
            continue  # baris pemisah
        rows.append(cells)
    if len(rows) < 2:
        return []
    header = [_clean(h).lower() for h in rows[0]]
    out: list[dict[str, str]] = []
    for r in rows[1:]:
        d = {header[i]: _clean(r[i]) for i in range(min(len(header), len(r)))}
        out.append(d)
    return out


def _parse_context_md(folder: Path) -> dict[str, Any]:
    """Pecah context.md jadi map section (key lowercase) → teks isi.

    Juga ekstrak inline 'Tujuan:' / 'Ruang Lingkup:', identitas, serta tabel
    Tim & Sasaran. Toleran terhadap key ber-markdown bold (**Kode:**).
    """
    out: dict[str, Any] = {
        "sections": {}, "tujuan": "", "ruang_lingkup": "", "identitas": {},
        "tim": [], "sasaran": [],
    }
    path = folder / "context.md"
    if not path.is_file():
        return out
    text = path.read_text(encoding="utf-8")

    # Section per heading "## ..."
    cur_key = None
    buf: list[str] = []
    for line in text.splitlines():
        m = re.match(r"^##\s+(.+?)\s*$", line)
        if m:
            if cur_key is not None:
                out["sections"][cur_key] = "\n".join(buf).strip()
            cur_key = m.group(1).strip().lower()
            buf = []
            continue
        if cur_key is not None:
            buf.append(line)
        # inline Tujuan / Ruang Lingkup (boleh ber-bold: **Tujuan:**)
        mt = re.match(r"^\s*\**\s*Tujuan\s*\**\s*:\s*(.+)$", line, re.IGNORECASE)
        if mt:
            out["tujuan"] = _clean(mt.group(1))
        mr = re.match(r"^\s*\**\s*Ruang Lingkup\s*\**\s*:\s*(.+)$", line, re.IGNORECASE)
        if mr:
            out["ruang_lingkup"] = _clean(mr.group(1))
    if cur_key is not None:
        out["sections"][cur_key] = "\n".join(buf).strip()

    # Identitas: parse bullet "- **Key:** value" (strip bold di key & value)
    ident_raw = out["sections"].get("identitas penugasan", "")
    for line in ident_raw.splitlines():
        mi = re.match(r"^\s*[-*]\s*(.+?)\s*:\s*(.+)$", line)
        if mi:
            out["identitas"][_clean(mi.group(1)).lower()] = _clean(mi.group(2))

    # Tim & Sasaran dari tabel markdown di section masing-masing
    out["tim"] = _parse_md_table(out["sections"].get("tim", ""))
    out["sasaran"] = _parse_md_table(out["sections"].get("sasaran pengawasan", "")
                                     or out["sections"].get("sasaran", ""))
    return out


def _section(ctx: dict, *names: str) -> str:
    """Ambil isi section pertama yang cocok dari beberapa kandidat nama."""
    for n in names:
        v = ctx["sections"].get(n)
        if v:
            return v
    return ""


def _gambaran_narrative(text: str) -> str:
    """Ambil paragraf naratif gambaran umum SEBELUM baris 'Tujuan:'."""
    parts: list[str] = []
    for line in text.splitlines():
        if re.match(r"^\s*\**\s*(Tujuan|Ruang Lingkup)\s*\**\s*:", line, re.IGNORECASE):
            break
        parts.append(line)
    return "\n".join(parts).strip()


def _strip_md(text: str) -> str:
    """Buang penanda markdown ringan (**bold**, backtick) agar rapi di Word."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    return text.strip()


# ---------------------------------------------------------------------------
# python-docx helpers (Arial 12, spasi 1,5)
# ---------------------------------------------------------------------------

def _set_spacing_15(p) -> None:
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _para(doc, text: str, *, bold: bool = False, italic: bool = False,
          size: int = 12, align: str | None = None, center: bool = False):
    p = doc.add_paragraph()
    if center or align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing_15(p)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Arial"
    return p


def _heading(doc, text: str):
    return _para(doc, text, bold=True, size=12)


def _table_arial(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Arial"
    return table


# ---------------------------------------------------------------------------
# Pemetaan risiko per jenis audit
# ---------------------------------------------------------------------------

_RISIKO_3E = [
    ("Ekonomis", "Apakah input (SDM, barang, modal) diperoleh dengan biaya wajar?",
     "HPS kemahalan, spesifikasi berlebihan, biaya overhead tidak proporsional"),
    ("Efisien", "Apakah rasio input:output optimal?",
     "Produksi rendah per rupiah, proses berulang tidak otomatis, idle asset"),
    ("Efektif", "Apakah output mencapai outcome?",
     "IK outcome tidak ada, target tidak SMART, manfaat tidak sampai sasaran"),
]


def _build_pemetaan_risiko(doc, skill: str, sasaran_desc: list[str], aspek_audit: str) -> None:
    skill = (skill or "").lower()
    # Bila context.md sudah punya "Aspek Audit (dari Survei Pendahuluan)" — itu
    # hasil survei yang sesungguhnya, pakai langsung.
    if aspek_audit:
        for para in aspek_audit.split("\n"):
            para = _strip_md(para)
            if para:
                _para(doc, para, align="justify")
        return
    if skill == "audit-kinerja":
        _para(doc,
              "Pemetaan risiko menggunakan kerangka 3E (Ekonomis, Efisien, Efektif). "
              "Tiap dimensi menjadi acuan penajaman sasaran dan langkah pengujian.",
              align="justify")
        _table_arial(doc, ["Dimensi", "Pertanyaan Kunci", "Risiko Tipikal"],
                     [[d, q, r] for d, q, r in _RISIKO_3E])
        return
    # audit-umum / audit-pengadaan → area pengujian berbasis sasaran
    _para(doc,
          "Pemetaan risiko diarahkan untuk menentukan area pengujian yang paling "
          "berisiko, sehingga prosedur audit terfokus (bukan memeriksa seluruh hal "
          "secara merata). Area pengujian awal:",
          align="justify")
    if sasaran_desc:
        for i, desc in enumerate(sasaran_desc):
            _para(doc, f"{chr(ord('a') + i)}. {desc}", align="justify")
    else:
        _para(doc, "[DIISI AUDITOR — area pengujian diturunkan dari sasaran pada Kartu Penugasan]",
              italic=True)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _load_sasaran(folder: Path) -> list[dict]:
    p = folder / "_PKP" / "sasaran-assignment.json"
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    if isinstance(data, dict):
        s = data.get("sasaran")
        return s if isinstance(s, list) else []
    return []


def render_survey_pendahuluan(
    folder: str | Path,
    *,
    skill: str,
    obyek: str | None = None,
    judul: str | None = None,
) -> Path:
    """Render Laporan Survei Pendahuluan ke `_SURVEY/Survei-Pendahuluan.docx`.

    Return path file hasil. Raise ValueError bila skill bukan skill audit.
    """
    if not is_audit_skill(skill):
        raise ValueError(f"Survei pendahuluan hanya untuk skill audit, bukan '{skill}'")

    folder = Path(folder)
    ctx = _parse_context_md(folder)
    from app.context_template import (  # late import (hindari circular)
        _collect_digest_summaries,
        _format_rupiah,
        _load_tim,
    )
    summaries = _collect_digest_summaries(folder)
    # Tim: utamakan tabel di context.md (lengkap), fallback ke _PKP json.
    tim_rows = ctx["tim"] or [
        {"peran": m.get("peran") or m.get("role") or "Anggota",
         "nama lengkap": m.get("nama") or m.get("name") or "",
         "nip": m.get("nip") or "", "jabfung": m.get("jabfung") or m.get("jabatan") or ""}
        for m in _load_tim(folder)
    ]
    # Sasaran (deskripsi): utamakan tabel context.md, fallback _PKP json.
    sasaran_desc: list[str] = []
    for r in ctx["sasaran"]:
        d = _strip_md(r.get("sasaran") or r.get("deskripsi") or "")
        if d:
            sasaran_desc.append(d)
    if not sasaran_desc:
        for s in _load_sasaran(folder):
            d = _strip_md(s.get("deskripsi") or s.get("sasaran") or "")
            if d:
                sasaran_desc.append(d)

    ident = ctx["identitas"]
    obyek_txt = obyek or ident.get("obyek") or "[DIISI AUDITOR]"
    judul_txt = judul or obyek_txt
    tujuan = ctx["tujuan"] or "[DIISI AUDITOR — Tujuan dari Kartu Penugasan]"
    ruang_lingkup = ctx["ruang_lingkup"] or "[DIISI AUDITOR — Ruang Lingkup dari Kartu Penugasan]"
    gambaran = _strip_md(_gambaran_narrative(
        _section(ctx, "gambaran umum obyek", "gambaran umum")))
    ringkasan = _strip_md(_section(ctx, "ringkasan obyek"))
    aspek_audit = _strip_md(_section(ctx, "aspek audit (dari survei pendahuluan)",
                                     "aspek audit"))

    doc = Document()

    # ── Kepala ───────────────────────────────────────────────────────────
    _para(doc, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True, center=True)
    _para(doc, "INSPEKTORAT JENDERAL", bold=True, center=True)
    _para(doc, "", size=12)
    _para(doc, "LAPORAN SURVEI PENDAHULUAN", bold=True, center=True, size=14)
    _para(doc, judul_txt, bold=True, center=True)
    _para(doc, "", size=12)

    # ── I. PENDAHULUAN ─────────────────────────────────────────────────────
    _heading(doc, "I.  PENDAHULUAN")
    _para(doc,
          f"Survei pendahuluan ini merupakan tahap orientasi (Tahap A0) sebelum "
          f"pelaksanaan audit terhadap {obyek_txt}. Survei dilaksanakan untuk "
          f"memahami obyek, memetakan risiko, menginventarisasi dokumen, dan "
          f"menyusun hipotesis area pengujian sebagai dasar penyusunan Program "
          f"Kerja Audit (PKA) berbasis risiko.", align="justify")
    ident_rows = []
    for label, key in (("Kode Penugasan", "kode"), ("Obyek", "obyek"),
                       ("Jenis Pengawasan", "skill / jenis pengawasan"),
                       ("Nomor ST", "nomor st"), ("Tanggal ST", "tanggal st")):
        val = ident.get(key)
        if val:
            ident_rows.append([label, _strip_md(val)])
    if ident_rows:
        _table_arial(doc, ["Uraian", "Keterangan"], ident_rows)
    _para(doc, "", size=12)

    # ── II. GAMBARAN UMUM OBYEK ────────────────────────────────────────────
    _heading(doc, "II.  GAMBARAN UMUM OBYEK")
    if gambaran:
        _para(doc, gambaran, align="justify")
    if ringkasan:
        _para(doc, ringkasan, align="justify")
    if not gambaran and not ringkasan:
        _para(doc, "[DIISI AUDITOR — gambaran umum obyek 3–5 kalimat]", italic=True)
    _para(doc, "", size=12)

    # ── III. TUJUAN DAN RUANG LINGKUP ──────────────────────────────────────
    _heading(doc, "III.  TUJUAN DAN RUANG LINGKUP")
    _para(doc, f"1. Tujuan dari dilaksanakannya audit adalah {tujuan}", align="justify")
    _para(doc, f"2. Ruang lingkup audit meliputi {ruang_lingkup}", align="justify")
    _para(doc, "", size=12)

    # ── IV. PEMETAAN RISIKO / AREA PENGUJIAN ───────────────────────────────
    _heading(doc, "IV.  PEMETAAN RISIKO DAN AREA PENGUJIAN")
    _build_pemetaan_risiko(doc, skill, sasaran_desc, aspek_audit)
    _para(doc, "", size=12)

    # ── V. INVENTARISASI DOKUMEN ───────────────────────────────────────────
    _heading(doc, "V.  INVENTARISASI DOKUMEN")
    if summaries:
        rows = []
        for i, s in enumerate(summaries, start=1):
            jenis = s.get("jenis") or "—"
            fname = s.get("file") or "?"
            extras: list[str] = []
            for k in ("ro", "total_biaya", "total_pagu", "nilai_hps", "jumlah_komponen"):
                v = s.get(k)
                if v not in (None, "", 0):
                    if k in ("total_biaya", "total_pagu", "nilai_hps"):
                        v = _format_rupiah(v) or v
                    extras.append(f"{k}={v}")
            rows.append([str(i), jenis, fname, ", ".join(extras) or "—"])
        _table_arial(doc, ["No", "Jenis", "Dokumen", "Catatan"], rows)
    else:
        _para(doc, "[DIISI AUDITOR — belum ada dokumen ter-ingest]", italic=True)
    _para(doc, "", size=12)

    # ── VI. HIPOTESIS / FOKUS PENGUJIAN ────────────────────────────────────
    _heading(doc, "VI.  HIPOTESIS DAN FOKUS PENGUJIAN AWAL")
    if sasaran_desc:
        _para(doc, "Berdasarkan sasaran penugasan, fokus pengujian awal diarahkan pada:",
              align="justify")
        for i, desc in enumerate(sasaran_desc, start=1):
            _para(doc, f"{i}. {desc}", align="justify")
    else:
        _para(doc, "[DIISI AUDITOR — hipotesis audit awal yang akan diuji di KKP]", italic=True)
    _para(doc, "", size=12)

    # ── VII. TIM SURVEI ────────────────────────────────────────────────────
    _heading(doc, "VII.  TIM")
    if tim_rows:
        rows = []
        for m in tim_rows:
            rows.append([
                _strip_md(str(m.get("peran") or m.get("role") or "Anggota")),
                _strip_md(str(m.get("nama lengkap") or m.get("nama") or m.get("name") or "[DIISI]")),
                _strip_md(str(m.get("nip") or "[DIISI]")),
                _strip_md(str(m.get("jabfung") or m.get("jabatan") or "Auditor")),
            ])
        _table_arial(doc, ["Peran", "Nama", "NIP", "Jabfung"], rows)
    else:
        _para(doc, "[DIISI AUDITOR — komposisi tim]", italic=True)
    _para(doc, "", size=12)

    _para(doc,
          "Catatan: Survei pendahuluan bersifat orientasi dan penajaman — tidak "
          "menyimpulkan penyimpangan. Hipotesis di atas diverifikasi pada tahap "
          "pengujian substantif (KKP).", italic=True, align="justify")

    out_dir = folder / "_SURVEY"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Survei-Pendahuluan.docx"
    doc.save(str(out_path))
    return out_path
