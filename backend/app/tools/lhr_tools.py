"""Tools untuk Agen Ketua Tim: baca temuan, completeness check, render LHR, QC LHP sync.

Schema rekomendasi.json yang dipakai V6 render_lhp.py:

    {
        "T-001": "Rekomendasi tegas untuk perbaikan...",
        "T-002": "...",
        ...
    }

Note: Function `request_qc_lhp` lama (async-flag) DIGANTI dengan `run_qc_lhp`
sync — sama pola dengan `run_qc_kkp` di kkp_tools.py. Pola lama bermasalah:
agen tidak dapat hasil → improvisasi.
"""
import json
import re
from datetime import datetime
from pathlib import Path

from claude_agent_sdk import tool
from docx import Document

from app.config import get_settings
from app.tools.v6_bridge import qc_summary_counts, run_v6_script, safe_read_json

settings = get_settings()

# Template LHP placeholder-driven, dimiliki app (bukan V6) supaya backend/v6/
# tetap read-only. render_lhp.py V6 menerima override path lewat --template.
_APP_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"


def _slug(skill: str) -> str:
    return re.sub(r"[^a-z0-9\-]", "-", str(skill).strip().lower())


def resolve_lhp_template(skill: str) -> Path | None:
    """Pilih template LHP per jenis pengawasan.

    Prioritas:
      1. Skeleton per-skill di APP_TEMPLATES_PATH/_skeleton-lhp/template-lhp-[skill].docx
         (placeholder {{...}} V6, satu file per jenis laporan).
      2. Skeleton GENERIK (template-lhp-generic.docx, kata "Reviu"→"Pengawasan")
         untuk skill tanpa skeleton khusus (mis. *-umum, kepatuhan-saipi).
      3. Fallback terakhir: template app reviu-rka-kl (sudah teruji).
    Return path absolut atau None bila tidak ada satupun.
    """
    slug = _slug(skill)
    skel_dir = settings.templates_path / "_skeleton-lhp"
    for candidate in (skel_dir / f"template-lhp-{slug}.docx", skel_dir / "template-lhp-generic.docx"):
        if candidate.is_file():
            return candidate
    fallback = _APP_TEMPLATE_DIR / "template-lhp-reviu-rka-kl.docx"
    return fallback if fallback.is_file() else None


@tool(
    "write_sasaran_assignment",
    "Tulis (overwrite) _PKP/sasaran-assignment.json. PAKAI HANYA di mode 'Setup Penugasan' "
    "saat sasaran-assignment masih kosong/draft. Input `sasaran` adalah list of dict dengan "
    "field: sasaran_id (mis. 'S-PBJ-01'), deskripsi, assigned_to (list[str] nama anggota), "
    "langkah_kerja (list[str]), status (default 'AKTIF'). KT primary path tetap via UI form — "
    "tool ini fallback untuk agent-driven setup.",
    {"penugasan_folder": str, "sasaran": list},
)
async def write_sasaran_assignment(args: dict) -> dict:
    folder = Path(args["penugasan_folder"])
    path = folder / "_PKP" / "sasaran-assignment.json"
    path.parent.mkdir(parents=True, exist_ok=True)

    raw_sasaran = args.get("sasaran", [])
    if not isinstance(raw_sasaran, list):
        return {
            "content": [{"type": "text", "text": "FAILED|sasaran harus list of dict"}],
            "is_error": True,
        }

    # Normalize + validasi
    sasaran_clean: list[dict] = []
    seen_ids: set[str] = set()
    for s in raw_sasaran:
        if not isinstance(s, dict):
            continue
        sid = str(s.get("sasaran_id", "")).strip()
        if not sid:
            continue
        if sid in seen_ids:
            return {
                "content": [{"type": "text", "text": f"FAILED|sasaran_id duplikat: {sid}"}],
                "is_error": True,
            }
        seen_ids.add(sid)
        assigned = s.get("assigned_to", [])
        if isinstance(assigned, str):
            assigned = [assigned]
        langkah = s.get("langkah_kerja", [])
        if isinstance(langkah, str):
            langkah = [langkah]
        sasaran_clean.append({
            "sasaran_id": sid,
            "deskripsi": str(s.get("deskripsi", "")).strip(),
            "assigned_to": [str(x).strip() for x in assigned if str(x).strip()],
            "langkah_kerja": [str(x).strip() for x in langkah if str(x).strip()],
            "status": str(s.get("status", "AKTIF")).strip() or "AKTIF",
        })

    # Preserve existing envelope kalau file sudah ada, supaya penugasan_id/skill tidak hilang
    existing = safe_read_json(path) if path.exists() else {}
    data = {
        "penugasan_id": existing.get("penugasan_id", folder.name),
        "skill": existing.get("skill", ""),
        "schema_version": "v4.0.0",
        "tanggal_dibuat": datetime.utcnow().isoformat() + "Z",
        "sasaran": sasaran_clean,
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "content": [{
            "type": "text",
            "text": f"OK|total_sasaran={len(sasaran_clean)}|path={path.name}",
        }]
    }


@tool(
    "read_temuan_json",
    "Baca _KKP/temuan.json penugasan. Return JSON lengkap dengan envelope penugasan + array temuan.",
    {"penugasan_folder": str},
)
async def read_temuan_json(args: dict) -> dict:
    path = Path(args["penugasan_folder"]) / "_KKP" / "temuan.json"
    if not path.exists():
        return {
            "content": [{"type": "text", "text": "FAILED|temuan.json tidak ada"}],
            "is_error": True,
        }
    data = safe_read_json(path)
    return {"content": [{"type": "text", "text": json.dumps(data, ensure_ascii=False)}]}


@tool(
    "check_completeness",
    "Pastikan semua sasaran di sasaran-assignment.json sudah DISETUJUI_KT (sudah di-approve "
    "oleh Ketua Tim). Kalau ada yang masih AKTIF (belum ada temuan) atau SELESAI_KKP "
    "(sudah ada temuan tapi belum approve), STOP — minta KT approve dulu lewat UI Setup.",
    {"penugasan_folder": str},
)
async def check_completeness(args: dict) -> dict:
    folder = Path(args["penugasan_folder"])
    assignment = safe_read_json(folder / "_PKP" / "sasaran-assignment.json")
    sasaran_list = assignment.get("sasaran", []) if isinstance(assignment, dict) else []

    # Approved statuses yang siap LHR
    APPROVED = {"DISETUJUI_KT"}

    belum = [s for s in sasaran_list if s.get("status") not in APPROVED]
    if belum:
        text = "BELUM_LENGKAP|sasaran_belum=" + json.dumps(
            [
                {
                    "id": s.get("sasaran_id"),
                    "status_current": s.get("status"),
                    "assigned_to": s.get("assigned_to"),
                }
                for s in belum
            ],
            ensure_ascii=False,
        )
        return {"content": [{"type": "text", "text": text}], "is_error": False}
    return {
        "content": [{
            "type": "text",
            "text": f"OK|total_sasaran={len(sasaran_list)}|all_disetujui_kt=true"
        }]
    }


@tool(
    "write_rekomendasi_json",
    "Tulis _LHP/rekomendasi.json — mapping id_temuan ke teks rekomendasi.",
    {"penugasan_folder": str, "rekomendasi": dict},
)
async def write_rekomendasi_json(args: dict) -> dict:
    path = Path(args["penugasan_folder"]) / "_LHP" / "rekomendasi.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(args["rekomendasi"], ensure_ascii=False, indent=2), encoding="utf-8")
    return {"content": [{"type": "text", "text": f"OK|n_rekomendasi={len(args['rekomendasi'])}"}]}


# A1 — penyesuaian jenis (lapisan app, V6 read-only). render_lhp.py menghasilkan
# docx ber-judul "REVIU" generik + nama file LHP-SUBSTANSI; di sini judul/kata &
# nama file disesuaikan per jenis pengawasan: audit→LHA, evaluasi→LHE,
# pemantauan→LP, reviu→LHR.
_JENIS_LABEL = {
    "audit": ("AUDIT", "Audit", "LHA"),
    "kepatuhan": ("AUDIT", "Audit", "LHA"),
    "evaluasi": ("EVALUASI", "Evaluasi", "LHE"),
    "pemantauan": ("PEMANTAUAN", "Pemantauan", "LP"),
    "reviu": ("REVIU", "Reviu", "LHR"),
}


def _jenis_meta(skill: str) -> tuple[str, str, str]:
    s = _slug(skill)
    for key, val in _JENIS_LABEL.items():
        if s.startswith(key):
            return val
    return ("PENGAWASAN", "Pengawasan", "LHP")


def _para_replace(p, subs: list[tuple]) -> None:
    full = "".join(r.text for r in p.runs)
    if not full:
        return
    new = full
    for pat, repl in subs:
        new = pat.sub(repl, new)
    if new != full and p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""


def _family(skill: str) -> str:
    s = _slug(skill)
    for key in ("audit", "kepatuhan", "evaluasi", "pemantauan", "reviu"):
        if s.startswith(key):
            return "audit" if key == "kepatuhan" else key
    return "lain"


def _terbilang(n: int) -> str:
    d = ["nol", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "delapan",
         "sembilan", "sepuluh", "sebelas"]
    return d[n] if 0 <= n < len(d) else str(n)


def _counts(folder: Path) -> tuple[int, int]:
    n = m = 0
    tj = folder / "_KKP" / "temuan.json"
    if tj.exists():
        try:
            d = json.loads(tj.read_text(encoding="utf-8"))
            items = d if isinstance(d, list) else (d.get("temuan") or d.get("items") or [])
            n = len(items)
        except (json.JSONDecodeError, OSError):
            pass
    sj = folder / "_PKP" / "sasaran-assignment.json"
    if sj.exists():
        try:
            d = json.loads(sj.read_text(encoding="utf-8"))
            sas = d.get("sasaran") if isinstance(d, dict) else d
            m = len(sas or [])
        except (json.JSONDecodeError, OSError):
            pass
    return n, m


# A2 — paragraf Metodologi/Intro/Simpulan di-hardcode paradigma REVIU oleh V6
# (read-only). Di lapisan app diganti dengan paragraf sesuai jenis. Paragraf
# diidentifikasi lewat penanda stabil dari teks asli V6.
_SUBSTANCE = {
    "audit": {
        "desk review": "Audit dilaksanakan sesuai Standar Audit Intern Pemerintah Indonesia (SAIPI) melalui penelaahan dokumen, pengujian bukti secara memadai, klarifikasi/wawancara dengan pihak terkait, serta — bila relevan — pemeriksaan dan uji petik atas hasil pekerjaan. Tim juga memanfaatkan analisis pre-digest dan cross-check otomatis untuk mendeteksi anomali antar dokumen.",
        "dikelompokkan ke dalam": "Berdasarkan pengujian atas dokumen dan bukti audit, tim Inspektorat II memperoleh {n} ({nt}) temuan yang dikelompokkan ke dalam {m} ({mt}) aspek sesuai sasaran audit. Temuan dirumuskan dengan paradigma audit (Kondisi-Kriteria-Sebab-Akibat-Rekomendasi) dengan tingkat keyakinan memadai sebagaimana diatur dalam SAIPI.",
        "limited assurance": "Berdasarkan hasil audit yang kami lakukan dengan tingkat keyakinan memadai, terdapat {n} ({nt}) temuan yang perlu ditindaklanjuti auditi sesuai rekomendasi pada laporan ini. Simpulan ini didasarkan pada bukti yang cukup dan memadai; tindak lanjut atas rekomendasi menjadi tanggung jawab pimpinan auditi.",
    },
    "evaluasi": {
        "desk review": "Evaluasi dilaksanakan melalui penelaahan dokumen, analisis data kinerja/capaian, dan klarifikasi kepada unit terkait, dengan membandingkan kondisi yang ada terhadap kriteria evaluasi yang ditetapkan.",
        "dikelompokkan ke dalam": "Berdasarkan penelaahan, tim Inspektorat II memperoleh {n} ({nt}) catatan evaluasi yang dikelompokkan ke dalam {m} ({mt}) aspek sesuai sasaran evaluasi. Catatan dirumuskan dengan paradigma Kondisi-Kriteria-Akibat-Rekomendasi.",
        "limited assurance": "Berdasarkan hasil evaluasi dengan tingkat keyakinan terbatas, terdapat {n} ({nt}) catatan yang perlu ditindaklanjuti untuk meningkatkan kualitas pelaksanaan pada aspek yang dievaluasi.",
    },
    "pemantauan": {
        "desk review": "Pemantauan dilaksanakan melalui penelaahan laporan berkala dan data status pelaksanaan dari auditi/pengawas pekerjaan, dengan membandingkan realisasi terhadap target serta ketentuan kontrak/rencana.",
        "dikelompokkan ke dalam": "Berdasarkan pemantauan, tim Inspektorat II mencatat {n} ({nt}) isu/kondisi yang perlu perhatian, dikelompokkan ke dalam {m} ({mt}) aspek. Pemantauan bersifat pelaporan status dan tidak memberikan keyakinan atas kebenaran substansi.",
        "limited assurance": "Berdasarkan hasil pemantauan, terdapat {n} ({nt}) kondisi yang perlu perhatian sebagaimana diuraikan. Laporan ini bersifat pelaporan status — tidak memberikan keyakinan dan tidak menyimpulkan pelanggaran; tindak lanjut menjadi kewenangan pihak pengelola kegiatan.",
    },
}


def _set_para(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def _apply_jenis(docx_path: Path, family: str, up: str, title: str, n: int, m: int) -> None:
    """A2 substance (paragraf per jenis) + A1 wording (kata Reviu→jenis). Untuk non-reviu."""
    subst = _SUBSTANCE.get(family, {})
    fmt = {"n": n, "nt": _terbilang(n), "m": m, "mt": _terbilang(m)}
    subs = [
        (re.compile(r"\bREVIU\b"), up),
        (re.compile(r"\bReviu\b"), title),
        (re.compile(r"\breviu\b"), title.lower()),
    ]
    doc = Document(str(docx_path))

    def handle(p):
        for marker, tpl in subst.items():
            if marker in p.text:
                _set_para(p, tpl.format(**fmt))
                return
        _para_replace(p, subs)

    for p in doc.paragraphs:
        handle(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    handle(p)
    for sec in doc.sections:
        for area in (sec.header, sec.footer):
            for p in area.paragraphs:
                _para_replace(p, subs)
    doc.save(str(docx_path))


def _finalize_jenis(folder: Path, skill: str) -> str | None:
    """A1 nama-file + A2 substansi/kata per jenis. Return nama file final."""
    up, title, prefix = _jenis_meta(skill)
    family = _family(skill)
    outs = sorted((folder / "_LHP").glob("LHP-SUBSTANSI*.docx"), key=lambda p: p.stat().st_mtime)
    if not outs:
        return None
    produced = outs[-1]
    if family != "reviu":
        try:
            n, m = _counts(folder)
            _apply_jenis(produced, family, up, title, n, m)
        except Exception:
            pass
    final = produced.with_name(produced.name.replace("LHP-SUBSTANSI", prefix, 1))
    if final != produced:
        produced.replace(final)
    return final.name


async def _render_kksa(folder: Path, args: dict) -> dict:
    """Render LHP paradigma KKSA via render_lhp.py V6 (placeholder {{...}})."""
    rekomendasi = folder / "_LHP" / "rekomendasi.json"
    if not rekomendasi.exists():
        return {"content": [{"type": "text", "text": "FAILED|rekomendasi.json belum ada"}], "is_error": True}
    skill = args.get("skill") or ""
    # B — bab Gambaran Umum tidak boleh kosong/placeholder; paksa agen mengisinya.
    gu = (args.get("gambaran_umum") or "").strip()
    if not gu or gu.upper().startswith("[DIISI"):
        return {
            "content": [{
                "type": "text",
                "text": "FAILED|gambaran_umum kosong/placeholder. Susun 3–5 kalimat substantif "
                        "(obyek, nilai anggaran/HPS, mekanisme/periode) dari KP+PKP+digest+context, lalu render ulang.",
            }],
            "is_error": True,
        }
    template = resolve_lhp_template(skill)
    if template is None:
        return {
            "content": [{
                "type": "text",
                "text": f"FAILED|template LHP tidak ada untuk skill='{skill}' (cek APP_TEMPLATES_PATH/_skeleton-lhp/)",
            }],
            "is_error": True,
        }
    render_args = [
        "--penugasan", str(folder),
        "--rekomendasi-file", str(rekomendasi),
        "--template", str(template),
        "--judul", args["judul"],
        "--auditi", args["auditi"],
        "--dasar-permintaan", args["dasar_permintaan"],
        "--gambaran-umum", args["gambaran_umum"],
        "--tanggal-exit-meeting", args["tanggal_exit_meeting"],
    ]
    if args.get("tembusan"):
        render_args += ["--tembusan", args["tembusan"]]
    code, out, err = await run_v6_script(
        "scripts/render_lhp.py",
        render_args,
        timeout=120,
    )
    if code != 0:
        return {"content": [{"type": "text", "text": f"FAILED|exit={code}|err={err[:400]}"}], "is_error": True}
    # A1: sesuaikan judul/kata + nama file per jenis (LHA/LHR/LHE/LP).
    final_name = _finalize_jenis(folder, skill)
    tail = f"|file={final_name}" if final_name else ""
    return {"content": [{"type": "text", "text": f"OK|format=kksa|template={template.name}{tail}|{out[:120]}"}]}


@tool(
    "render_report",
    "Render laporan hasil sesuai PROFIL FORMAT skill (otomatis): 'kksa' (reviu/audit → "
    "LHP KKSA via V6), 'memo' (Konsultansi → Memo pendapat/saran, butuh _LHP/saran.json), "
    "'rb-4dim' (Eval RB → tabel 4-dimensi, butuh _LHP/penilaian-rb.json). Pakai tool ini "
    "sebagai jalur utama penyusunan laporan untuk SEMUA skill kecuali reviu-pengadaan "
    "(pakai render_lhr_pbj).",
    {
        "penugasan_folder": str, "skill": str, "judul": str, "auditi": str,
        "dasar_permintaan": str, "gambaran_umum": str, "tanggal_exit_meeting": str,
        "tembusan": str,
    },
)
async def render_report(args: dict) -> dict:
    from app.format_registry import format_profile

    folder = Path(args["penugasan_folder"])
    skill = args.get("skill") or ""
    profile = format_profile(skill)
    if profile == "memo":
        return _render_memo(folder, args)
    if profile == "rb-4dim":
        return _render_rb(folder, args)
    if profile == "pendampingan":
        return _render_pendampingan(folder, args)
    return await _render_kksa(folder, args)


@tool(
    "render_lhr_pbj",
    "Render LHR Reviu Pengadaan via scripts/reviu-pengadaan/run_batch.py V6 mode KT. "
    "Script baca context.md dan _LHP/rekomendasi.json dari folder penugasan.",
    {"penugasan_folder": str, "tembusan": str},
)
async def render_lhr_pbj(args: dict) -> dict:
    """Pipeline dua tahap untuk reviu-pengadaan:
    1. run_batch.py (digest + cross_check) — TANPA --render karena template path
       di run_batch.py hardcoded ke v6/templates/ yang tidak ada. Render dilakukan
       terpisah di tahap 2 dengan template path dari settings (APP_TEMPLATES_PATH).
    2. render_lhp.py langsung dengan --template dari resolve_lhp_template().
    """
    folder = Path(args["penugasan_folder"])

    # Tahap 1: digest pengadaan + cross_check (tanpa render)
    code, out, err = await run_v6_script(
        "scripts/reviu-pengadaan/run_batch.py",
        [
            "--penugasan", args["penugasan_folder"],
            "--role", "KT",
            # Tidak pakai --render: template path-nya hardcoded ke v6/templates/
            # yang tidak ada. Render dikerjakan di tahap 2.
        ],
        timeout=180,
    )
    if code != 0:
        return {
            "content": [{"type": "text", "text": f"FAILED (digest/cross_check)|exit={code}|err={err[:400]}"}],
            "is_error": True,
        }

    # Tahap 2: render LHP dengan template path yang benar dari settings
    template = resolve_lhp_template("reviu-pengadaan")
    if template is None:
        return {
            "content": [{
                "type": "text",
                "text": (
                    "FAILED (render)|template LHP reviu-pengadaan tidak ditemukan "
                    "(cek APP_TEMPLATES_PATH/_skeleton-lhp/template-lhp-reviu-pengadaan.docx)"
                ),
            }],
            "is_error": True,
        }

    render_cmd = ["--penugasan", args["penugasan_folder"], "--template", str(template)]
    rekomendasi = folder / "_LHP" / "rekomendasi.json"
    if rekomendasi.exists():
        render_cmd += ["--rekomendasi-file", str(rekomendasi)]
    if args.get("tembusan"):
        render_cmd += ["--tembusan", args["tembusan"]]

    code2, out2, err2 = await run_v6_script(
        "scripts/render_lhp.py",
        render_cmd,
        timeout=120,
    )
    if code2 != 0:
        return {
            "content": [{"type": "text", "text": f"FAILED (render_lhp)|exit={code2}|err={err2[:400]}"}],
            "is_error": True,
        }
    return {"content": [{"type": "text", "text": f"OK|digest+crosscheck+render|template={template.name}|{out2[:160]}"}]}


@tool(
    "run_qc_lhp",
    "Jalankan QC SAIPI stage LHP secara SYNCHRONOUS. Memanggil scripts/qc_saipi.py "
    "V6 dengan --stage lhp lalu return status + breakdown severity + excerpt laporan. "
    "Pakai SETELAH render_lhr selesai untuk gate kepatuhan SAIPI tahap pelaporan.",
    {"penugasan_folder": str},
)
async def run_qc_lhp(args: dict) -> dict:
    """Sync version dari QC LHP — ganti pola async marker-flag yang lama
    (`request_qc_lhp` writer flag). Pola lama bermasalah: agen yang memanggilnya
    tidak dapat hasil → improvisasi sendiri.
    """
    folder = Path(args["penugasan_folder"])
    if not folder.exists():
        return {
            "content": [{
                "type": "text",
                "text": f"FAILED|folder penugasan tidak ada: {folder} — cek path (typo?), jangan anggap PASS",
            }],
            "is_error": True,
        }
    code, out, err = await run_v6_script(
        "scripts/qc_saipi.py",
        ["--penugasan", str(folder), "--stage", "lhp"],
        timeout=120,
    )

    checklist = safe_read_json(folder / "_QA-SAIPI" / "checklist-lhp.json")
    total_kritis, total_peringatan, total_needs_review, total_ok = qc_summary_counts(checklist)

    if total_kritis > 0:
        status_label = "BLOCKED_KRITIS"
    elif total_peringatan > 0 or total_needs_review > 0:
        status_label = "PASS_WITH_WARNINGS"
    else:
        status_label = "PASS"

    laporan_path = folder / "_QA-SAIPI" / "laporan-qa-lhp.md"
    laporan_excerpt = ""
    if laporan_path.exists():
        laporan_excerpt = laporan_path.read_text(encoding="utf-8")[:4000]

    return {
        "content": [
            {
                "type": "text",
                "text": (
                    f"stage=lhp|status={status_label}|exit_code={code}|"
                    f"kritis={total_kritis}|peringatan={total_peringatan}|"
                    f"needs_review={total_needs_review}|ok={total_ok}|"
                    f"laporan_path={laporan_path}\n\n"
                    f"=== LAPORAN QA (excerpt) ===\n{laporan_excerpt}"
                ),
            }
        ]
    }


# =============================================================================
# FORMAT NON-KKSA — Memo Konsultansi + Evaluasi RB (tabel 4-dimensi)
# Renderer MILIK APP (python-docx) — V6 tetap read-only.
# =============================================================================


def _ctx_lines(folder: Path) -> dict:
    """Ambil beberapa field identitas dari context.md (best-effort)."""
    out: dict = {}
    p = folder / "context.md"
    if p.is_file():
        for line in p.read_text(encoding="utf-8").splitlines():
            m = re.match(r"^\s*[-*]?\s*(Kode|Obyek|Nomor ST|Tanggal ST)\s*:\s*(.+)$", line, re.IGNORECASE)
            if m:
                out[m.group(1).strip().lower()] = m.group(2).strip()
    return out


def _render_memo(folder: Path, args: dict) -> dict:
    """Memo Konsultansi: dasar hukum + pertanyaan→pendapat/saran. Tanpa KKSA/keyakinan."""
    saran_path = folder / "_LHP" / "saran.json"
    if not saran_path.exists():
        return {"content": [{"type": "text", "text": "FAILED|_LHP/saran.json belum ada (pakai append_saran)"}], "is_error": True}
    try:
        items = json.loads(saran_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        return {"content": [{"type": "text", "text": f"FAILED|baca saran.json: {e}"}], "is_error": True}
    if not isinstance(items, list) or not items:
        return {"content": [{"type": "text", "text": "FAILED|saran.json kosong"}], "is_error": True}

    ctx = _ctx_lines(folder)
    # Mulai dari template ber-KOP (kop-only) bila tersedia, lalu tambahkan isi
    # memo di bawahnya. Fallback ke dokumen kosong bila template tak ada.
    kop_tpl = settings.templates_path / "_skeleton-lhp" / "template-lhp-konsultansi-umum.docx"
    doc = Document(str(kop_tpl)) if kop_tpl.exists() else Document()
    doc.add_heading("MEMO KONSULTANSI", level=0)
    doc.add_paragraph(args.get("judul") or "Memo Konsultansi")
    meta = doc.add_paragraph()
    meta.add_run(f"Auditan: {args.get('auditi') or ctx.get('obyek', '-')}\n")
    meta.add_run(f"Dasar: {args.get('dasar_permintaan') or ctx.get('nomor st', '-')}\n")
    meta.add_run(f"Kode penugasan: {ctx.get('kode', folder.name)}")
    doc.add_paragraph(
        "Catatan: dokumen ini berisi PENDAPAT/SARAN konsultansi berbasis dasar hukum, "
        "TIDAK memberikan keyakinan, dan TIDAK mengikat pejabat berwenang.",
    ).italic = True

    # Dasar hukum gabungan (unik)
    dh: list[str] = []
    for it in items:
        for d in (it.get("dasar_hukum") or []):
            if d and d not in dh:
                dh.append(d)
    if dh:
        doc.add_heading("Dasar Hukum", level=1)
        for d in dh:
            doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("Pendapat dan Saran", level=1)
    for i, it in enumerate(items, start=1):
        doc.add_heading(f"{i}. {it.get('pertanyaan', '(pertanyaan)')}", level=2)
        if it.get("pendapat"):
            p = doc.add_paragraph(); p.add_run("Pendapat: ").bold = True; p.add_run(str(it["pendapat"]))
        if it.get("saran"):
            p = doc.add_paragraph(); p.add_run("Saran: ").bold = True; p.add_run(str(it["saran"]))

    out_path = folder / "_LHP" / "LHP-SUBSTANSI-MEMO.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return {"content": [{"type": "text", "text": f"OK|format=memo|n_pertanyaan={len(items)}|{out_path.name}"}]}


_RB_DIM = [
    ("ketepatan", "Ketepatan Pelaksanaan"),
    ("ketercapaian", "Ketercapaian Output"),
    ("kualitas", "Kualitas Pelaksanaan"),
    ("kesesuaian", "Kesesuaian Waktu"),
]


def _render_rb(folder: Path, args: dict) -> dict:
    """Evaluasi RB: tabel komponen Renaksi × 4 dimensi (Sesuai/Tidak Sesuai)."""
    pen_path = folder / "_LHP" / "penilaian-rb.json"
    if not pen_path.exists():
        return {"content": [{"type": "text", "text": "FAILED|_LHP/penilaian-rb.json belum ada (pakai write_penilaian_rb)"}], "is_error": True}
    try:
        data = json.loads(pen_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        return {"content": [{"type": "text", "text": f"FAILED|baca penilaian-rb.json: {e}"}], "is_error": True}
    komponen = data.get("komponen") if isinstance(data, dict) else None
    if not isinstance(komponen, list) or not komponen:
        return {"content": [{"type": "text", "text": "FAILED|penilaian-rb.json: 'komponen' kosong"}], "is_error": True}

    ctx = _ctx_lines(folder)
    doc = Document()
    doc.add_heading("LAPORAN HASIL EVALUASI REFORMASI BIROKRASI", level=0)
    doc.add_paragraph(args.get("judul") or "Laporan Hasil Evaluasi Reformasi Birokrasi")
    meta = doc.add_paragraph()
    meta.add_run(f"Auditan: {args.get('auditi') or ctx.get('obyek', '-')}\n")
    meta.add_run(f"Dasar: {args.get('dasar_permintaan') or ctx.get('nomor st', '-')}\n")
    meta.add_run(f"Kode penugasan: {ctx.get('kode', folder.name)}")

    doc.add_heading("Penilaian per Komponen Rencana Aksi (4 Dimensi)", level=1)
    table = doc.add_table(rows=1, cols=2 + len(_RB_DIM))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Komponen Renaksi"
    for j, (_, label) in enumerate(_RB_DIM, start=1):
        hdr[j].text = label
    hdr[1 + len(_RB_DIM)].text = "Catatan"
    for k in komponen:
        row = table.add_row().cells
        row[0].text = str(k.get("nama", "-"))
        for j, (key, _) in enumerate(_RB_DIM, start=1):
            row[j].text = str(k.get(key, "-"))
        row[1 + len(_RB_DIM)].text = str(k.get("catatan", ""))

    if data.get("analisis_dampak"):
        doc.add_heading("Analisis Dampak", level=1)
        doc.add_paragraph(str(data["analisis_dampak"]))
    aoi = data.get("aoi") or []
    if isinstance(aoi, list) and aoi:
        doc.add_heading("Area of Improvement (AoI)", level=1)
        for a in aoi:
            doc.add_paragraph(str(a), style="List Bullet")

    out_path = folder / "_LHP" / "LHP-SUBSTANSI-RB.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return {"content": [{"type": "text", "text": f"OK|format=rb-4dim|n_komponen={len(komponen)}|{out_path.name}"}]}


def _render_pendampingan(folder: Path, args: dict) -> dict:
    """Laporan Hasil Pendampingan (profil baru utk konsultasi-pengadaan).

    Beda dengan Memo: bukan jawab pertanyaan, melainkan LOG kegiatan
    pendampingan yang sudah diselesaikan. Sesuai pola Inspektorat II Komdigi
    yang sering pendampingan berkelanjutan (hadir rapat, reviu bertahap,
    klarifikasi proses) bukan konsultasi sekali jadi.

    Input: `_LHP/kegiatan-pendampingan.json` — list of:
        {tanggal, jenis_kegiatan, deskripsi, hasil,
         pihak_didampingi?, dokumen_pendukung?[], tindak_lanjut?}

    Output: `_LHP/LHP-PENDAMPINGAN.docx` — BAB I Kegiatan Pendampingan +
    BAB II Tindak Lanjut + Kesimpulan.
    """
    keg_path = folder / "_LHP" / "kegiatan-pendampingan.json"
    if not keg_path.exists():
        return {
            "content": [{
                "type": "text",
                "text": (
                    "FAILED|_LHP/kegiatan-pendampingan.json belum ada "
                    "(pakai append_kegiatan_pendampingan)"
                ),
            }],
            "is_error": True,
        }
    try:
        items = json.loads(keg_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        return {
            "content": [{"type": "text", "text": f"FAILED|baca kegiatan: {e}"}],
            "is_error": True,
        }
    if not isinstance(items, list) or not items:
        return {
            "content": [{"type": "text", "text": "FAILED|kegiatan-pendampingan.json kosong"}],
            "is_error": True,
        }

    ctx = _ctx_lines(folder)
    doc = Document()
    doc.add_heading("LAPORAN HASIL PENDAMPINGAN PENGADAAN", level=0)
    doc.add_paragraph(args.get("judul") or "Laporan Hasil Pendampingan Pengadaan")

    meta = doc.add_paragraph()
    meta.add_run(f"Auditan: {args.get('auditi') or ctx.get('obyek', '-')}\n")
    meta.add_run(f"Dasar Penugasan: {args.get('dasar_permintaan') or ctx.get('nomor st', '-')}\n")
    meta.add_run(f"Kode penugasan: {ctx.get('kode', folder.name)}\n")
    # Periode pendampingan: tanggal min-max dari kegiatan
    tgls = sorted([str(it.get("tanggal", "")) for it in items if it.get("tanggal")])
    if tgls:
        periode = f"{tgls[0]} s.d. {tgls[-1]}" if tgls[0] != tgls[-1] else tgls[0]
        meta.add_run(f"Periode Pendampingan: {periode}\n")

    p = doc.add_paragraph()
    p.add_run(
        "Catatan: Laporan ini berisi rangkaian KEGIATAN PENDAMPINGAN yang "
        "telah diselesaikan tim Inspektorat II atas permintaan unit kerja. "
        "Pendampingan bersifat advisory dan preventif — tidak memberikan "
        "keyakinan dan tidak mengikat pejabat berwenang."
    ).italic = True

    # Gambaran umum (opsional)
    if args.get("gambaran_umum"):
        doc.add_heading("Gambaran Umum", level=1)
        doc.add_paragraph(str(args["gambaran_umum"]))

    # BAB I — Kegiatan Pendampingan
    doc.add_heading(f"I. Kegiatan Pendampingan yang Telah Diselesaikan ({len(items)})", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["No", "Tanggal", "Jenis Kegiatan", "Pihak Didampingi", "Deskripsi", "Hasil"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for i, it in enumerate(items, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(it.get("tanggal", "-"))
        row[2].text = str(it.get("jenis_kegiatan", "-"))
        row[3].text = str(it.get("pihak_didampingi", "-"))
        row[4].text = str(it.get("deskripsi", "-"))
        row[5].text = str(it.get("hasil", "-"))

    # Dokumen pendukung (per kegiatan, kalau ada)
    has_dok = any(it.get("dokumen_pendukung") for it in items)
    if has_dok:
        doc.add_heading("Dokumen Pendukung per Kegiatan", level=2)
        for i, it in enumerate(items, start=1):
            doks = it.get("dokumen_pendukung") or []
            if not isinstance(doks, list) or not doks:
                continue
            doc.add_paragraph(f"Kegiatan #{i} ({it.get('tanggal', '-')})").bold = True
            for d in doks:
                doc.add_paragraph(str(d), style="List Bullet")

    # BAB II — Tindak Lanjut
    tindak_lanjut = [it for it in items if it.get("tindak_lanjut")]
    if tindak_lanjut:
        doc.add_heading(f"II. Hal yang Masih Memerlukan Tindak Lanjut ({len(tindak_lanjut)})", level=1)
        for i, it in enumerate(tindak_lanjut, start=1):
            p2 = doc.add_paragraph()
            p2.add_run(f"{i}. {it.get('jenis_kegiatan', '-')} ({it.get('tanggal', '-')}): ").bold = True
            p2.add_run(str(it.get("tindak_lanjut", "")))

    # Kesimpulan (template singkat — auditor bisa edit DOCX)
    doc.add_heading("III. Kesimpulan", level=1)
    doc.add_paragraph(
        args.get("kesimpulan")
        or (
            f"Tim Inspektorat II telah menyelesaikan {len(items)} kegiatan pendampingan "
            f"pengadaan pada {args.get('auditi') or ctx.get('obyek', 'unit kerja')}. "
            f"Seluruh kegiatan diarahkan untuk mencegah penyimpangan prosedur "
            f"pengadaan dan memberikan masukan teknis berbasis Perpres 16/2018 "
            f"jo. 12/2021. Pendampingan ini tidak menggantikan kewenangan "
            f"PPK/PA/KPA atas keputusan pengadaan."
        )
    )

    out_path = folder / "_LHP" / "LHP-PENDAMPINGAN.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return {
        "content": [{
            "type": "text",
            "text": (
                f"OK|format=pendampingan|n_kegiatan={len(items)}|"
                f"n_tindak_lanjut={len(tindak_lanjut)}|{out_path.name}"
            ),
        }]
    }


@tool(
    "append_kegiatan_pendampingan",
    "Tambah satu kegiatan pendampingan ke `_LHP/kegiatan-pendampingan.json` "
    "(skill konsultasi-pengadaan, profil 'pendampingan'). Schema: "
    "{tanggal: 'YYYY-MM-DD', jenis_kegiatan: 'Rapat|Reviu|Klarifikasi|Pendampingan langsung|...', "
    "deskripsi: 'apa yg auditor lakukan', hasil: 'apa yg berhasil diselesaikan', "
    "pihak_didampingi?: 'PPK/PA/Pokja Pemilihan/dst', "
    "dokumen_pendukung?: ['Notulen rapat tanggal X', 'Draft KAK rev-2', ...], "
    "tindak_lanjut?: 'hal yg masih perlu diselesaikan auditi'}. "
    "Konsultasi-pengadaan TIDAK pakai temuan KKSA dan TIDAK pakai memo — pakai "
    "log kegiatan ini, lalu render_report.",
    {"penugasan_folder": str, "kegiatan": dict},
)
async def append_kegiatan_pendampingan(args: dict) -> dict:
    path = Path(args["penugasan_folder"]) / "_LHP" / "kegiatan-pendampingan.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    items: list = []
    if path.exists():
        try:
            items = json.loads(path.read_text(encoding="utf-8"))
            if not isinstance(items, list):
                items = []
        except (json.JSONDecodeError, OSError):
            items = []
    new = args.get("kegiatan")
    if isinstance(new, list):
        items.extend(new)
    elif isinstance(new, dict):
        items.append(new)
    else:
        return {
            "content": [{"type": "text", "text": "FAILED|param 'kegiatan' harus dict atau list of dict"}],
            "is_error": True,
        }
    path.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "content": [{
            "type": "text",
            "text": f"OK|total_kegiatan={len(items)} @ {path}",
        }]
    }


@tool(
    "append_saran",
    "Tambah butir Memo Konsultansi ke _LHP/saran.json (untuk skill konsultansi). "
    "`saran` = dict/list of {pertanyaan, dasar_hukum[], pendapat, saran}. Konsultansi "
    "tidak memakai temuan KKSA — pakai ini lalu render_report.",
    {"penugasan_folder": str, "saran": dict},
)
async def append_saran(args: dict) -> dict:
    path = Path(args["penugasan_folder"]) / "_LHP" / "saran.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    items = []
    if path.exists():
        try:
            items = json.loads(path.read_text(encoding="utf-8"))
            if not isinstance(items, list):
                items = []
        except (json.JSONDecodeError, OSError):
            items = []
    new = args.get("saran")
    items.extend(new if isinstance(new, list) else [new])
    path.write_text(json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"content": [{"type": "text", "text": f"OK|total_saran={len(items)}"}]}


@tool(
    "write_penilaian_rb",
    "Tulis (overwrite) _LHP/penilaian-rb.json untuk Evaluasi RB. Struktur: "
    "{komponen:[{nama, ketepatan, ketercapaian, kualitas, kesesuaian, catatan}], "
    "analisis_dampak, aoi:[...]}. Nilai dimensi 'Sesuai'/'Tidak Sesuai'. Sumber: hasil "
    "gate evaluasi RB. Lalu render_report.",
    {"penugasan_folder": str, "penilaian": dict},
)
async def write_penilaian_rb(args: dict) -> dict:
    path = Path(args["penugasan_folder"]) / "_LHP" / "penilaian-rb.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(args["penilaian"], ensure_ascii=False, indent=2), encoding="utf-8")
    n = len((args["penilaian"] or {}).get("komponen", []))
    return {"content": [{"type": "text", "text": f"OK|n_komponen={n}"}]}


LHR_TOOLS = [
    write_sasaran_assignment,  # Setup Penugasan mode
    read_temuan_json,
    check_completeness,
    write_rekomendasi_json,
    render_report,     # dispatcher per-profil (kksa/memo/rb-4dim/pendampingan) — jalur utama LHP
    render_lhr_pbj,    # pipeline khusus reviu-pengadaan
    append_saran,      # Memo Konsultansi (konsultansi-umum)
    append_kegiatan_pendampingan,  # Laporan Pendampingan (konsultasi-pengadaan)
    write_penilaian_rb,  # Evaluasi RB 4-dimensi
    run_qc_lhp,
]
