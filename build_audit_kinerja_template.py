"""
Build template-lhp-audit-kinerja.docx dari awal mengikuti format SKILL.md:

Nota Dinas
───────────────────────────────
LHA Kinerja
  Bab 1: PENDAHULUAN
         1.1 Latar Belakang
         1.2 Dasar Penugasan       → {{A_DASAR}}
         1.3 Tujuan Audit          → {{B_TUJUAN}}
         1.4 Pertanyaan Audit      → {{C_PERTANYAAN_AUDIT}}
         1.5 Ruang Lingkup & Met.  → {{D_RUANG_LINGKUP}} / {{D_METODOLOGI}}
         1.6 Batasan Audit         → [teks tetap]
         1.7 Komposisi Tim         → {{E_KOMPOSISI_TIM}}
  Bab 2: GAMBARAN UMUM PROGRAM    → {{GAMBARAN_UMUM}}
  Bab 3: METODOLOGI AUDIT KINERJA → {{METODOLOGI_REVIU}}
  Bab 4: TEMUAN DAN ANALISIS
         4.1 Efektivitas           → {{HASIL_EFEKTIVITAS}}
         4.2 Efisiensi             → {{HASIL_EFISIENSI}}
  Bab 5: SIMPULAN                 → {{SIMPULAN_REVIU}}
  Bab 6: REKOMENDASI              → [tabel rekomendasi] {{REKOMENDASI_LOOP}}
TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-kinerja.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False, size_pt=None):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    effective_pt = size_pt if size_pt else 12
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", str(effective_pt * 2))
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", str(effective_pt * 2))
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body):
    body.append(mkpara(""))


def section(body, num, judul, bold=True):
    """Tambah heading Bab / sub-bab."""
    body.append(mkpara(f"{num}  {judul}", bold=bold))


def placeholder(body, key):
    """Paragraf berisi {{KEY}} sebagai expand point."""
    body.append(mkpara(f"{{{{{key}}}}}"))


def add(body, txt, **kwargs):
    body.append(mkpara(txt, **kwargs))


# ── Baca template lama (ambil tabel kop kalau ada, lalu hapus semua para) ──
doc = Document(str(TPL))
body = doc.element.body

# Hapus semua paragraf — pertahankan tabel (jika ada kop tabel)
for el in list(body):
    body.remove(el)

# ═══════════════════════════════════════════════════════════
# NOTA DINAS
# ═══════════════════════════════════════════════════════════
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
    "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Audit Kinerja "
    "terhadap {{NAMA_AUDITI}} Tahun Anggaran {{TAHUN_ANGGARAN}}.")
empty(body)
add(body,
    "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Audit (LHA) Kinerja "
    "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
empty(body)
add(body,
    "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)
add(body, "Inspektur II,")
empty(body)
empty(body)
empty(body)
empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
placeholder(body, "TEMBUSAN_LIST")

# ═══════════════════════════════════════════════════════════
# HALAMAN COVER LHA
# ═══════════════════════════════════════════════════════════
empty(body)
add(body, "─" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL AUDIT KINERJA", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_3}}")
empty(body)
add(body, "Nomor LHA : {{NOMOR_LHR}}")
add(body, "Surat Tugas: {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode Audit: {{PERIODE_PELAKSANAAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══════════════════════════════════════════════════════════
# TUBUH LHA — BAB 1 s.d. 6
# ═══════════════════════════════════════════════════════════
empty(body)
add(body, "─" * 60)
empty(body)

# Pembuka LHA
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Audit Kinerja "
    "terhadap {{NAMA_AUDITI}} Tahun Anggaran {{TAHUN_ANGGARAN}} berdasarkan Surat "
    "Tugas Inspektur Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
    "Pelaksanaan audit berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
empty(body)

# ── Bab 1 ──────────────────────────────────────────────────
section(body, "BAB 1", "PENDAHULUAN", bold=True)
empty(body)

section(body, "1.1", "Latar Belakang")
add(body,
    "Dalam kerangka pengawasan intern berbasis risiko, Inspektorat Jenderal "
    "Kementerian Komunikasi dan Digital melaksanakan Audit Kinerja terhadap "
    "{{NAMA_AUDITI}} Tahun Anggaran {{TAHUN_ANGGARAN}}. Audit ini dilaksanakan "
    "untuk memberikan keyakinan memadai atas efektivitas dan efisiensi pelaksanaan "
    "program/kegiatan serta mengidentifikasi area perbaikan yang diperlukan.")
empty(body)

section(body, "1.2", "Dasar Penugasan")
placeholder(body, "A_DASAR")
empty(body)

section(body, "1.3", "Tujuan Audit")
placeholder(body, "B_TUJUAN")
empty(body)

section(body, "1.4", "Pertanyaan Audit")
placeholder(body, "C_PERTANYAAN_AUDIT")
empty(body)

section(body, "1.5", "Ruang Lingkup dan Metodologi")
add(body, "Ruang Lingkup:")
placeholder(body, "D_RUANG_LINGKUP")
add(body, "Metodologi:")
placeholder(body, "D_METODOLOGI")
empty(body)

section(body, "1.6", "Batasan Audit")
add(body,
    "Audit ini dibatasi pada aspek efektivitas dan efisiensi pelaksanaan "
    "program/kegiatan {{NAMA_AUDITI}} tahun anggaran {{TAHUN_ANGGARAN}} sesuai "
    "sasaran yang ditetapkan dalam Surat Tugas dan Kerangka Penugasan. Penilaian "
    "atas ekonomisitas (kewajaran harga pengadaan) di luar lingkup audit ini. "
    "Apabila ditemukan indikasi permasalahan di luar lingkup, akan disampaikan "
    "sebagai catatan terpisah atau eskalasi kepada pimpinan.")
empty(body)

section(body, "1.7", "Komposisi Tim dan Jangka Waktu")
placeholder(body, "E_KOMPOSISI_TIM")
empty(body)

# ── Bab 2 ──────────────────────────────────────────────────
section(body, "BAB 2", "GAMBARAN UMUM PROGRAM", bold=True)
empty(body)

section(body, "2.1", "Tujuan dan Desain Program")
placeholder(body, "BAB2_1_TUJUAN_DESAIN")
empty(body)

section(body, "2.2", "Logika Intervensi")
placeholder(body, "BAB2_2_LOGIKA")
empty(body)

section(body, "2.3", "Anggaran dan Sumber Daya")
placeholder(body, "BAB2_3_ANGGARAN")
empty(body)

section(body, "2.4", "Pelaksana dan Mekanisme")
placeholder(body, "BAB2_4_PELAKSANA")
empty(body)

# ── Bab 3 ──────────────────────────────────────────────────
section(body, "BAB 3", "METODOLOGI AUDIT KINERJA", bold=True)
empty(body)
placeholder(body, "METODOLOGI_REVIU")
empty(body)

# ── Bab 4 ──────────────────────────────────────────────────
section(body, "BAB 4", "TEMUAN DAN ANALISIS", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil audit kinerja yang dilaksanakan, tim Inspektorat II "
    "menemukan hal-hal berikut yang dikelompokkan menurut dimensi efektivitas "
    "dan efisiensi pelaksanaan program:")
empty(body)

section(body, "4.1", "Efektivitas Pencapaian Target")
add(body,
    "Temuan terkait efektivitas menilai sejauh mana program/kegiatan mencapai "
    "tujuan dan target yang telah ditetapkan:")
placeholder(body, "HASIL_EFEKTIVITAS")
empty(body)

section(body, "4.2", "Efisiensi Penggunaan Sumber Daya")
add(body,
    "Temuan terkait efisiensi menilai penggunaan sumber daya (anggaran, SDM, "
    "waktu) dalam pelaksanaan program/kegiatan:")
placeholder(body, "HASIL_EFISIENSI")
empty(body)

# ── Bab 5 ──────────────────────────────────────────────────
section(body, "BAB 5", "SIMPULAN", bold=True)
empty(body)
placeholder(body, "SIMPULAN_REVIU")
empty(body)

# ── Bab 6 ──────────────────────────────────────────────────
section(body, "BAB 6", "REKOMENDASI", bold=True)
empty(body)
add(body,
    "Berdasarkan temuan dan analisis dalam Bab 4, tim Inspektorat II merekomendasikan "
    "hal-hal berikut kepada {{NAMA_AUDITI}}:")
placeholder(body, "REKOMENDASI_LOOP")
empty(body)

add(body,
    "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama {{NAMA_AUDITI}}, "
    "kami ucapkan terima kasih.")
empty(body)

add(body, "Inspektur II,")
empty(body)
empty(body)
empty(body)
empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
placeholder(body, "TEMBUSAN_LIST")

# ── Simpan ─────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci — perubahan tersimpan di: {TMP}")
    print("    Tutup Word terlebih dahulu lalu rename manual.")
    exit()

# ── Verifikasi ─────────────────────────────────────────────
doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt:
        print(f"{i:3}: {txt[:80]}")
