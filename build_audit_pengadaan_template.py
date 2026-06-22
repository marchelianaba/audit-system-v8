"""
Rebuild template-lhp-audit-pengadaan.docx sesuai SKILL.md Format LHP:

  Bab 1: Pendahuluan (1.1 Dasar, 1.2 Tujuan, 1.3 Ruang Lingkup)
  Bab 2: Gambaran Umum Obyek Audit
  Bab 3: Metodologi Audit          ← bab tersendiri, bukan sub-bab 1
  Bab 4: Hasil Audit               (ringkasan per area)
  Bab 5: Temuan dan Rekomendasi    (detail CCSAA)
  Bab 6: Kesimpulan
  Lampiran 1: Daftar Dokumen
  Lampiran 2: Matriks Temuan (CCSAA)

Tidak ada "Ringkasan Eksekutif" — tidak disebut di SKILL.md.
"""
from docx import Document
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-pengadaan.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def E(body): body.append(mkpara(""))
def A(body, txt, **kw): body.append(mkpara(txt, **kw))
def PH(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def S(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


from docx.oxml.ns import qn as _qn

# Buka dari audit-umum (memiliki sectPr valid) sebagai base
BASE_TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")
doc = Document(str(BASE_TPL if BASE_TPL.exists() else TPL))
body = doc.element.body
# Simpan sectPr agar doc.sections tetap valid (dibutuhkan doc.add_table)
sect_pr = body.find(_qn("w:sectPr"))
for el in list(body):
    body.remove(el)
# Restore atau buat sectPr minimal di akhir body
if sect_pr is None:
    sect_pr = OxmlElement("w:sectPr")
body.append(sect_pr)

# ═══ NOTA DINAS ═══
A(body, "NOTA DINAS", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
   "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Audit "
   "Kepatuhan Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
E(body)
A(body, "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Audit (LHA) "
   "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
E(body)
A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")

# ═══ COVER LHA ═══
E(body); A(body, "=" * 60); E(body)
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL"); E(body)
A(body, "LAPORAN HASIL AUDIT", bold=True)
A(body, "KEPATUHAN PENGADAAN BARANG/JASA", bold=True)
A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
E(body)
A(body, "Nomor LHA     : {{NOMOR_LHR}}")
A(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
A(body, "Periode Audit : {{PERIODE_PELAKSANAAN}}")
A(body, "Tahun Anggaran: {{TAHUN_ANGGARAN}}")
E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")

# ═══ TUBUH LHA ═══
E(body); A(body, "=" * 60); E(body)
A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Audit "
   "Kepatuhan Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} "
   "berdasarkan Surat Tugas Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
   "Pelaksanaan audit berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
E(body)

# ── Bab 1 ────────────────────────────────────────────────────────────────────
S(body, "BAB 1", "PENDAHULUAN", bold=True); E(body)
S(body, "1.1", "Dasar Penugasan"); PH(body, "BAB1_DASAR"); E(body)
S(body, "1.2", "Tujuan Audit"); PH(body, "BAB1_TUJUAN"); E(body)
S(body, "1.3", "Ruang Lingkup"); PH(body, "BAB1_RUANG_LINGKUP"); E(body)

# ── Bab 2 ────────────────────────────────────────────────────────────────────
S(body, "BAB 2", "GAMBARAN UMUM OBYEK AUDIT", bold=True)
A(body, "Berikut gambaran umum obyek audit yang menjadi objek pemeriksaan:")
PH(body, "BAB2_GAMBARAN_UMUM"); E(body)

# ── Bab 3 ────────────────────────────────────────────────────────────────────
S(body, "BAB 3", "METODOLOGI AUDIT", bold=True); PH(body, "BAB3_METODOLOGI"); E(body)

# ── Bab 4 ────────────────────────────────────────────────────────────────────
S(body, "BAB 4", "HASIL AUDIT", bold=True)
A(body, "Berikut ringkasan hasil audit per area pemeriksaan:")
PH(body, "BAB4_HASIL_AUDIT"); E(body)

# ── Bab 5 ────────────────────────────────────────────────────────────────────
S(body, "BAB 5", "TEMUAN DAN REKOMENDASI", bold=True)
A(body, "Berikut uraian temuan audit beserta rekomendasi perbaikan "
   "dalam format Kondisi-Kriteria-Sebab-Akibat-Rekomendasi (CCSAA):")
PH(body, "BAB5_TEMUAN_REKOMENDASI"); E(body)

# ── Bab 6 ────────────────────────────────────────────────────────────────────
S(body, "BAB 6", "KESIMPULAN", bold=True); PH(body, "BAB6_KESIMPULAN"); E(body)

# ── TTD ──────────────────────────────────────────────────────────────────────
A(body, "Atas kerja sama {{NAMA_AUDITI}} selama pelaksanaan audit ini, "
   "Inspektorat II menyampaikan apresiasi dan terima kasih.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")

# ═══ LAMPIRAN ═══
E(body); A(body, "=" * 60); E(body)
A(body, "LAMPIRAN 1: DAFTAR DOKUMEN SUMBER", bold=True)
A(body, "Dokumen yang digunakan sebagai bukti dan referensi dalam pelaksanaan audit:")
PH(body, "LAMPIRAN_1_DAFTAR_DOKUMEN"); E(body)

A(body, "LAMPIRAN 2: MATRIKS TEMUAN (CCSAA)", bold=True)
A(body, "Rekapitulasi temuan audit dalam format matriks kondisi-kriteria-sebab-akibat-rekomendasi:")
PH(body, "LAMPIRAN_2_MATRIKS_TEMUAN"); E(body)

# ── Simpan ────────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists(): os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    d2 = Document(str(TPL))
    print(f"Template berhasil diperbarui: {TPL}")
    print(f"paragraphs: {len(d2.paragraphs)}")
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
