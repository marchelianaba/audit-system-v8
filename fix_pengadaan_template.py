"""
Fix template-lhp-audit-pengadaan.docx:
  1. Ganti ${ttd_pengirim} → {{TTD_INSPEKTUR}}
  2. Ganti 'Muhammad Arief' → {{NAMA_INSPEKTUR}}; tambah NIP. {{NIP_INSPEKTUR}} setelahnya
     (berlaku untuk semua occurrence — ada di Nota Dinas & Penutup LHA)
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-pengadaan.docx")
doc = Document(str(TPL))
XML_NS = "http://www.w3.org/XML/1998/namespace"


def gtext(p):
    return "".join(t.text or "" for t in p._element.iter(qn("w:t")))


def stext(p_elem, txt):
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p_elem.append(r)


def mkpara(txt):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


for p in list(doc.paragraphs):
    txt = gtext(p).strip()

    if "${ttd_pengirim}" in txt:
        stext(p._element, "{{TTD_INSPEKTUR}}")

    elif txt in ("Muhammad Arief", "Muhammad Arief."):
        stext(p._element, "{{NAMA_INSPEKTUR}}")
        nxt = p._element.getnext()
        nxt_txt = ("".join(t.text or "" for t in nxt.iter(qn("w:t")))).strip() if nxt is not None else ""
        if "NIP" not in nxt_txt:
            nip_elem = mkpara("NIP. {{NIP_INSPEKTUR}}")
            p._element.addnext(nip_elem)

TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template diperbarui:", TPL)
except PermissionError:
    print(f"[!] Terkunci — perubahan di: {TMP}")
    print("    Tutup Word lalu jalankan ulang.")
    exit()

doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
print("\n=== Nota Dinas (para 0-18) ===")
for i, p in enumerate(doc2.paragraphs):
    if i < 19:
        txt = p.text.strip()
        if txt:
            print(f"{i:3}: {txt[:90]}")
print("\n=== Semua TTD/NAMA/NIP ===")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if any(k in txt for k in ("TTD_INSPEKTUR", "NAMA_INSPEKTUR", "NIP_INSPEKTUR", "NIP.")):
        print(f"{i:3}: {txt[:90]}")
