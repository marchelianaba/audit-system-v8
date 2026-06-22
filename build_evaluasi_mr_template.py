"""
Build template-lhp-evaluasi-manajemen-risiko.docx sesuai SKILL.md:

Nota Dinas → [Paragraf pembuka]
A. Dasar Pelaksanaan Evaluasi   → {{A_DASAR}}
B. Tujuan Evaluasi              → {{B_TUJUAN}}
C. Ruang Lingkup Evaluasi       → {{C_RUANG_LINGKUP}}
D. Metodologi Evaluasi          → {{D_METODOLOGI}}
E. Gambaran Umum                → {{E_GAMBARAN_UMUM}}
F. Hasil Evaluasi               → {{F_HASIL_EVALUASI}}
G. Rekomendasi                  → {{G_REKOMENDASI}}
H. Apresiasi                    (teks tetap)
TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-evaluasi-manajemen-risiko.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body):
    body.append(mkpara(""))


def add(body, txt, **kw):
    body.append(mkpara(txt, **kw))


def ph(body, key):
    body.append(mkpara(f"{{{{{key}}}}}"))


def sec(body, lbl, judul, bold=True):
    body.append(mkpara(f"{lbl}  {judul}", bold=bold))


# ── buka template lama, hapus semua elemen ───────────────────────────────────
doc = Document(str(TPL))
body = doc.element.body
for el in list(body):
    body.remove(el)

# ═══════════════════════════════════════════════════════════════════════════════
# NOTA DINAS PENGANTAR
# ═══════════════════════════════════════════════════════════════════════════════
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
    "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Evaluasi "
    "Efektivitas Manajemen Risiko terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
empty(body)
add(body,
    "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Evaluasi (LHE) Manajemen Risiko "
    "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
empty(body)
add(body,
    "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)
add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER LHE MANAJEMEN RISIKO
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL EVALUASI", bold=True)
add(body, "EFEKTIVITAS MANAJEMEN RISIKO", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}")
empty(body)
add(body, "Nomor LHE     : {{NOMOR_LHR}}")
add(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode Eval. : {{PERIODE_PELAKSANAAN}}")
add(body, "Tahun         : {{TAHUN_ANGGARAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══════════════════════════════════════════════════════════════════════════════
# TUBUH LHE
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Evaluasi Efektivitas "
    "Manajemen Risiko terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan "
    "Surat Tugas Inspektur Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
    "Pelaksanaan evaluasi berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
empty(body)

# ── A. Dasar Pelaksanaan ─────────────────────────────────────────────────────
sec(body, "A.", "Dasar Pelaksanaan Evaluasi", bold=True)
ph(body, "A_DASAR")
empty(body)

# ── B. Tujuan ────────────────────────────────────────────────────────────────
sec(body, "B.", "Tujuan Evaluasi", bold=True)
ph(body, "B_TUJUAN")
empty(body)

# ── C. Ruang Lingkup ─────────────────────────────────────────────────────────
sec(body, "C.", "Ruang Lingkup Evaluasi", bold=True)
ph(body, "C_RUANG_LINGKUP")
empty(body)

# ── D. Metodologi ────────────────────────────────────────────────────────────
sec(body, "D.", "Metodologi Evaluasi", bold=True)
ph(body, "D_METODOLOGI")
empty(body)

# ── E. Gambaran Umum ─────────────────────────────────────────────────────────
sec(body, "E.", "Gambaran Umum", bold=True)
ph(body, "E_GAMBARAN_UMUM")
empty(body)

# ── F. Hasil Evaluasi ────────────────────────────────────────────────────────
sec(body, "F.", "Hasil Evaluasi", bold=True)
empty(body)
add(body,
    "Berdasarkan evaluasi yang dilaksanakan, terdapat hal-hal yang perlu mendapat "
    "perhatian, sebagai berikut:")
empty(body)
ph(body, "F_HASIL_EVALUASI")
empty(body)

# ── G. Rekomendasi ───────────────────────────────────────────────────────────
sec(body, "G.", "Rekomendasi", bold=True)
empty(body)
add(body,
    "Berdasarkan kondisi-kondisi tersebut, Inspektorat II merekomendasikan agar:")
ph(body, "G_REKOMENDASI")
empty(body)

# ── H. Apresiasi ─────────────────────────────────────────────────────────────
sec(body, "H.", "Apresiasi", bold=True)
empty(body)
add(body,
    "Inspektorat II Inspektorat Jenderal Kementerian Komunikasi dan Digital "
    "menyampaikan penghargaan dan terima kasih atas kerja sama {{NAMA_AUDITI}} "
    "selama pelaksanaan evaluasi manajemen risiko ini.")
add(body,
    "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)

add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
    exit()

doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
