"""
Generator untuk 3 template workflow Reviu RKA-K/L.
Outputs:
  1. PKP-Reviu-RKA-KL-template.xlsx
  2. KKR-Reviu-RKA-KL-template.xlsx
  3. LHR-Reviu-RKA-KL-skeleton.docx
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.comments import Comment
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

HEADER_FILL = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
THIN_BORDER = Border(
    left=Side(style="thin", color="CCCCCC"),
    right=Side(style="thin", color="CCCCCC"),
    top=Side(style="thin", color="CCCCCC"),
    bottom=Side(style="thin", color="CCCCCC"),
)
WRAP_CENTER = Alignment(wrap_text=True, vertical="center", horizontal="left")
HEADER_ALIGN = Alignment(wrap_text=True, vertical="center", horizontal="center")


def style_header_row(ws, row_idx, n_cols):
    for col in range(1, n_cols + 1):
        c = ws.cell(row=row_idx, column=col)
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = HEADER_ALIGN
        c.border = THIN_BORDER


def apply_borders(ws, min_row, max_row, min_col, max_col):
    for row in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
        for cell in row:
            cell.border = THIN_BORDER
            if not cell.alignment.wrap_text:
                cell.alignment = WRAP_CENTER


def build_pkp():
    wb = Workbook()
    ws = wb.active
    ws.title = "Cover"
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 70

    cover_rows = [
        ("Item", "Value"),
        ("Penugasan", "{{ISI: Reviu RKA-K/L ... TA YYYY}}"),
        ("Surat Tugas", "{{ISI: nomor + tanggal}}"),
        ("ND Permintaan", "{{ISI: nomor + tanggal — kosongkan jika berbasis PKPT}}"),
        ("Periode Reviu", "{{ISI: range tanggal}}"),
        ("Skill", "reviu-rka-kl v3.0"),
        ("Total RO", "{{ISI: jumlah RO}}"),
        ("Total Aspek", "6 (A-F)"),
        ("Tim", "{{ISI: nama tim}}"),
    ]
    for r_idx, (k, v) in enumerate(cover_rows, start=1):
        ws.cell(row=r_idx, column=1, value=k)
        ws.cell(row=r_idx, column=2, value=v)
    style_header_row(ws, 1, 2)
    for r_idx in range(2, len(cover_rows) + 1):
        ws.cell(row=r_idx, column=1).font = Font(bold=True)
        ws.cell(row=r_idx, column=1).alignment = WRAP_CENTER
        ws.cell(row=r_idx, column=2).alignment = WRAP_CENTER
    apply_borders(ws, 1, len(cover_rows), 1, 2)
    ws.freeze_panes = "A2"

    ws2 = wb.create_sheet("Matriks PKP")
    headers = ["No", "Aspek", "Sub-aspek", "Pertanyaan Pengujian",
               "Bukti yang Dicari", "Rules Pipeline", "RO #1"]
    for c_idx, h in enumerate(headers, start=1):
        ws2.cell(row=1, column=c_idx, value=h)
    style_header_row(ws2, 1, len(headers))

    rows = [
        (1, "A. Kelayakan vs SBM", "Honorarium", "Apakah honorarium per OJ/OB sesuai PMK SBM TA berjalan?", "TOR, RAB, PMK SBM", "A1, A1-alt"),
        (2, "A. Kelayakan vs SBM", "Perjalanan Dinas", "Apakah biaya perjadin sesuai SBM (transport, uang harian, penginapan)?", "RAB, ToR, SBM Perjadin", "A2, A2-alt"),
        (3, "A. Kelayakan vs SBM", "Konsumsi/Rapat", "Apakah uang makan/snack dan paket meeting sesuai SBM?", "RAB, kuitansi indikatif", "A3, A3-alt"),
        (4, "A. Kelayakan vs SBM", "Sewa & Bahan", "Apakah sewa kendaraan/ruang/peralatan & ATK wajar?", "RAB, harga pasar", "A4, A4-alt"),
        (5, "B. Kepatuhan Kaidah Penganggaran", "Akun Belanja", "Apakah akun belanja (51/52/53) sesuai jenis output?", "RAB, BAS Kemenkeu", "B1, B1-alt"),
        (6, "B. Kepatuhan Kaidah Penganggaran", "Klasifikasi Komponen", "Apakah komponen utama vs pendukung dipisah benar?", "TOR, RAB", "B2, B2-alt"),
        (7, "B. Kepatuhan Kaidah Penganggaran", "Volume vs Satuan", "Apakah volume × harga satuan menghasilkan total yang konsisten?", "RAB perhitungan", "B3, B3-alt"),
        (8, "B. Kepatuhan Kaidah Penganggaran", "Duplikasi", "Apakah ada duplikasi komponen antar RO?", "RAB lintas RO", "B4, B4-alt"),
        (9, "C. Penandaan Anggaran", "Tagging Tematik", "Apakah penandaan PN/Direktif/Gender/SDGs sudah benar?", "Form penandaan, KRISNA", "C1, C1-alt"),
        (10, "C. Penandaan Anggaran", "Konsistensi Tag", "Apakah tagging konsisten antar dokumen (RKA, RAB, KRISNA)?", "RKA, RAB, screenshot KRISNA", "C2, C2-alt"),
        (11, "D. Kelengkapan Dokumen", "TOR Lengkap", "Apakah TOR memuat 5W1H, output, dan rincian biaya?", "Dokumen TOR", "D1, D1-alt"),
        (12, "D. Kelengkapan Dokumen", "RAB Lengkap", "Apakah RAB memuat komponen, volume, harga satuan, total?", "Dokumen RAB", "D2, D2-alt"),
        (13, "D. Kelengkapan Dokumen", "Dukungan Dokumen", "Apakah ada dokumen pendukung (KAK, RKAS, MoU jika perlu)?", "Folder dokumen RO", "D3, D3-alt"),
        (14, "E. Kelayakan Rincian Baru", "Justifikasi Baru", "Apakah komponen baru (vs TA sebelumnya) memiliki justifikasi?", "TOR, RAB tahun lalu", "E1, E1-alt"),
        (15, "E. Kelayakan Rincian Baru", "Benchmarking", "Apakah harga komponen baru di-benchmark ke pasar/instansi lain?", "Survey harga, RAB", "E2, E2-alt"),
        (16, "F. Pengalokasian Tematik", "Prioritas Nasional", "Apakah RO mendukung PN sesuai Renja-K/L?", "Renja, Matriks PN", "F1, F1-alt"),
        (17, "F. Pengalokasian Tematik", "Direktif Pimpinan", "Apakah alokasi direktif Presiden/Menteri tercermin?", "Memo direktif, RKA", "F2, F2-alt"),
        (18, "F. Pengalokasian Tematik", "Mandatory Spending", "Apakah mandatory spending (gender, perubahan iklim) terpenuhi?", "RKA tagging", "F3, F3-alt"),
    ]
    for i, (no, aspek, sub, q, bukti, rules) in enumerate(rows, start=2):
        ws2.cell(row=i, column=1, value=no)
        ws2.cell(row=i, column=2, value=aspek)
        ws2.cell(row=i, column=3, value=sub)
        ws2.cell(row=i, column=4, value=q)
        ws2.cell(row=i, column=5, value=bukti)
        ws2.cell(row=i, column=6, value=rules)
        ws2.cell(row=i, column=7, value="—")

    widths = {"A": 5, "B": 28, "C": 22, "D": 50, "E": 32, "F": 18, "G": 14}
    for col, w in widths.items():
        ws2.column_dimensions[col].width = w
    apply_borders(ws2, 1, len(rows) + 1, 1, len(headers))
    ws2.freeze_panes = "A2"
    ws2.cell(row=1, column=7).comment = Comment(
        "Duplikasi kolom ini sesuai jumlah RO yang direviu.\n"
        "Sel dapat diisi: '—' (tidak ada catatan), ID rule yang TRIGGERED (mis. A1, B3-alt), atau 'PERINGATAN'/'INFO'.",
        "audit-system-v4"
    )

    ws3 = wb.create_sheet("Status Per RO")
    s3_headers = ["RO ID", "RO Nama", "Total Anomali", "A", "B", "C", "D", "E", "F", "Status"]
    for c_idx, h in enumerate(s3_headers, start=1):
        ws3.cell(row=1, column=c_idx, value=h)
    style_header_row(ws3, 1, len(s3_headers))
    ws3.cell(row=2, column=1, value="{{RO_ID_1}}")
    ws3.cell(row=2, column=2, value="{{RO_NAMA_1}}")
    for col in range(3, 10):
        ws3.cell(row=2, column=col, value=0)
    ws3.cell(row=2, column=10, value="{{Status: TERPENUHI / TERPENUHI DENGAN CATATAN / TIDAK TERPENUHI}}")

    dv_status = DataValidation(type="list",
        formula1='"TERPENUHI,TERPENUHI DENGAN CATATAN,TIDAK TERPENUHI"',
        allow_blank=True)
    dv_status.add("J2:J100")
    ws3.add_data_validation(dv_status)

    widths3 = {"A": 14, "B": 38, "C": 14, "D": 6, "E": 6, "F": 6, "G": 6, "H": 6, "I": 6, "J": 32}
    for col, w in widths3.items():
        ws3.column_dimensions[col].width = w
    apply_borders(ws3, 1, 2, 1, len(s3_headers))
    ws3.freeze_panes = "A2"

    out = os.path.join(OUT_DIR, "PKP-Reviu-RKA-KL-template.xlsx")
    wb.save(out)
    return out


def build_kkr():
    wb = Workbook()
    ws = wb.active
    ws.title = "Cover"
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 60

    rows = [
        ("Item", "Value"),
        ("Penugasan", "{{ISI: Reviu RKA-K/L ... TA YYYY}}"),
        ("Surat Tugas", "{{ISI: nomor + tanggal}}"),
        ("ND Permintaan", "{{ISI: nomor + tanggal — kosongkan jika PKPT}}"),
        ("Periode Reviu", "{{ISI: range tanggal}}"),
        ("Skill", "reviu-rka-kl v3.0"),
        ("Total RO Direviu", "{{ISI: jumlah RO}}"),
        ("Tim", "{{ISI: nama tim}}"),
        ("", ""),
        ("RINGKASAN AGREGAT", ""),
        ("Total Catatan Reviu", "{{ISI: integer}}"),
        ("PERINGATAN (severity tinggi)", "{{ISI: integer}}"),
        ("INFO (severity rendah)", "{{ISI: integer}}"),
        ("", ""),
        ("Aspek A — Kelayakan vs SBM", "{{ISI: jumlah catatan}}"),
        ("Aspek B — Kepatuhan Kaidah Penganggaran", "{{ISI: jumlah catatan}}"),
        ("Aspek C — Penandaan Anggaran", "{{ISI: jumlah catatan}}"),
        ("Aspek D — Kelengkapan Dokumen", "{{ISI: jumlah catatan}}"),
        ("Aspek E — Kelayakan Rincian Baru", "{{ISI: jumlah catatan}}"),
        ("Aspek F — Pengalokasian Tematik", "{{ISI: jumlah catatan}}"),
    ]
    for r_idx, (k, v) in enumerate(rows, start=1):
        ws.cell(row=r_idx, column=1, value=k)
        ws.cell(row=r_idx, column=2, value=v)
    style_header_row(ws, 1, 2)
    style_header_row(ws, 10, 2)
    for r_idx in range(2, len(rows) + 1):
        if rows[r_idx - 1][0] in ("", "RINGKASAN AGREGAT"):
            continue
        ws.cell(row=r_idx, column=1).font = Font(bold=True)
        ws.cell(row=r_idx, column=1).alignment = WRAP_CENTER
        ws.cell(row=r_idx, column=2).alignment = WRAP_CENTER
    apply_borders(ws, 1, len(rows), 1, 2)
    ws.freeze_panes = "A2"

    ws2 = wb.create_sheet("Catatan Reviu")
    headers = ["No", "RO ID", "RO Nama", "Aspek", "Rule ID", "Severity",
               "Kondisi (Judul)", "Kriteria", "Catatan/Akibat",
               "Rekomendasi", "Status", "Bukti"]
    for c_idx, h in enumerate(headers, start=1):
        ws2.cell(row=1, column=c_idx, value=h)
    style_header_row(ws2, 1, len(headers))

    dummy = [1, "{{RO_ID}}", "{{RO_NAMA}}", "A. Kelayakan vs SBM", "A1", "PERINGATAN",
             "{{Judul singkat kondisi: misal 'Honorarium Narasumber Eselon I melebihi SBM'}}",
             "{{Kriteria: PMK xxx/202x ttg SBM TA YYYY}}",
             "{{Akibat/dampak — boros, tidak efisien, dll}}",
             "{{Rekomendasi — sesuaikan dengan SBM, kurangi alokasi, dll}}",
             "TIDAK TERPENUHI",
             "{{Path/lampiran bukti, mis. RAB-RO1.xlsx baris 12}}"]
    for c_idx, v in enumerate(dummy, start=1):
        ws2.cell(row=2, column=c_idx, value=v)

    dv_sev = DataValidation(type="list", formula1='"PERINGATAN,INFO"', allow_blank=True)
    dv_sev.add("F2:F500")
    ws2.add_data_validation(dv_sev)
    dv_status = DataValidation(type="list",
        formula1='"TERPENUHI,TERPENUHI DENGAN CATATAN,TIDAK TERPENUHI"',
        allow_blank=True)
    dv_status.add("K2:K500")
    ws2.add_data_validation(dv_status)

    widths = {"A": 5, "B": 12, "C": 28, "D": 24, "E": 10, "F": 13,
              "G": 38, "H": 32, "I": 38, "J": 38, "K": 22, "L": 28}
    for col, w in widths.items():
        ws2.column_dimensions[col].width = w
    apply_borders(ws2, 1, 2, 1, len(headers))
    ws2.freeze_panes = "A2"

    ws3 = wb.create_sheet("Rekap per RO")
    h3 = ["RO ID", "RO Nama", "Total Anomali", "A", "B", "C", "D", "E", "F",
          "PERINGATAN", "INFO", "Simpulan"]
    for c_idx, h in enumerate(h3, start=1):
        ws3.cell(row=1, column=c_idx, value=h)
    style_header_row(ws3, 1, len(h3))
    ws3.cell(row=2, column=1, value="{{RO_ID_1}}")
    ws3.cell(row=2, column=2, value="{{RO_NAMA_1}}")
    for col in range(3, 12):
        ws3.cell(row=2, column=col, value=0)
    ws3.cell(row=2, column=12, value="{{TERPENUHI / TERPENUHI DENGAN CATATAN / TIDAK TERPENUHI}}")

    dv_simp = DataValidation(type="list",
        formula1='"TERPENUHI,TERPENUHI DENGAN CATATAN,TIDAK TERPENUHI"',
        allow_blank=True)
    dv_simp.add("L2:L200")
    ws3.add_data_validation(dv_simp)

    widths3 = {"A": 14, "B": 38, "C": 14, "D": 5, "E": 5, "F": 5, "G": 5, "H": 5, "I": 5,
               "J": 12, "K": 8, "L": 36}
    for col, w in widths3.items():
        ws3.column_dimensions[col].width = w
    apply_borders(ws3, 1, 2, 1, len(h3))
    ws3.freeze_panes = "A2"

    ws4 = wb.create_sheet("Rekap per Aspek")
    h4 = ["Aspek", "Deskripsi", "Total Catatan", "RO Terdampak", "Top Rule"]
    for c_idx, h in enumerate(h4, start=1):
        ws4.cell(row=1, column=c_idx, value=h)
    style_header_row(ws4, 1, len(h4))

    aspek_rows = [
        ("A", "Kelayakan vs SBM", 0, "{{daftar RO}}", "{{rule_id}}"),
        ("B", "Kepatuhan Kaidah Penganggaran", 0, "{{daftar RO}}", "{{rule_id}}"),
        ("C", "Penandaan Anggaran", 0, "{{daftar RO}}", "{{rule_id}}"),
        ("D", "Kelengkapan Dokumen", 0, "{{daftar RO}}", "{{rule_id}}"),
        ("E", "Kelayakan Rincian Baru", 0, "{{daftar RO}}", "{{rule_id}}"),
        ("F", "Pengalokasian Tematik", 0, "{{daftar RO}}", "{{rule_id}}"),
    ]
    for i, r in enumerate(aspek_rows, start=2):
        for c_idx, v in enumerate(r, start=1):
            ws4.cell(row=i, column=c_idx, value=v)

    widths4 = {"A": 8, "B": 36, "C": 14, "D": 32, "E": 14}
    for col, w in widths4.items():
        ws4.column_dimensions[col].width = w
    apply_borders(ws4, 1, 7, 1, len(h4))
    ws4.freeze_panes = "A2"

    out = os.path.join(OUT_DIR, "KKR-Reviu-RKA-KL-template.xlsx")
    wb.save(out)
    return out


def build_lhr():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    for line in ["KEMENTERIAN KOMUNIKASI DAN DIGITAL RI",
                 "INSPEKTORAT JENDERAL", "INSPEKTORAT II"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Jl. Medan Merdeka Barat No. 9, Jakarta 10110").font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Telp./Fax. (021) 3859627").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("_" * 90).font.size = Pt(8)

    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = True
    header_data = [
        ("Nomor    : {{NOMOR_LHR}}", "Jakarta, {{TANGGAL_LHR}}"),
        ("Lampiran : Satu berkas", ""),
        ("Hal      : {{JUDUL_LAPORAN}}", ""),
    ]
    for i, (left, right) in enumerate(header_data):
        tbl.rows[i].cells[0].text = left
        tbl.rows[i].cells[1].text = right
        tbl.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph("Kepada Yth.")
    doc.add_paragraph("{{PENERIMA}}")
    doc.add_paragraph("di Jakarta")
    doc.add_paragraph()
    doc.add_paragraph("{{PARAGRAF_PEMBUKA}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("A. Dasar").bold = True
    doc.add_paragraph("    1. {{DASAR_1 — biasanya ST}}")
    doc.add_paragraph("    2. {{DASAR_2 — ND permintaan jika ada}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("B. Tujuan").bold = True
    doc.add_paragraph("    {{TUJUAN — narasi 2-3 kalimat ttg keyakinan terbatas atas RKA-K/L TA YYYY untuk 6 aspek (A-F sesuai PMK 107/2024 Pasal 61).}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("C. Ruang Lingkup").bold = True
    doc.add_paragraph("    Reviu mencakup {{N}} Rincian Output (RO) Direktorat {{NAMA_DIREKTORAT}} TA {{TAHUN}}, yaitu:")
    doc.add_paragraph("    {{LIST_RO — 1 baris per RO}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("D. Metodologi").bold = True
    doc.add_paragraph("    Reviu dilakukan melalui:")
    doc.add_paragraph("    1. Telaah dokumen TOR dan RAB tiap RO menggunakan pipeline 21 rules deterministik + 18 alt-rule")
    doc.add_paragraph("    2. Konfirmasi dengan auditan bila diperlukan klarifikasi")
    doc.add_paragraph("    3. Penilaian berdasarkan PMK 107/2024 dan Standar Biaya Masukan (PMK terkait)")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("E. Hasil Reviu").bold = True
    aspek_blocks = [
        ("E.1 Aspek A — Kelayakan vs SBM", "{{HASIL_A}}"),
        ("E.2 Aspek B — Kepatuhan Kaidah Penganggaran", "{{HASIL_B}}"),
        ("E.3 Aspek C — Penandaan Anggaran", "{{HASIL_C}}"),
        ("E.4 Aspek D — Kelengkapan Dokumen", "{{HASIL_D}}"),
        ("E.5 Aspek E — Kelayakan Rincian Baru", "{{HASIL_E}}"),
        ("E.6 Aspek F — Pengalokasian Tematik", "{{HASIL_F}}"),
    ]
    for title, placeholder in aspek_blocks:
        p = doc.add_paragraph(); p.add_run("    " + title).bold = True
        doc.add_paragraph("    " + placeholder)
        doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("F. Catatan & Rekomendasi Prioritas").bold = True
    doc.add_paragraph("    {{REKOMENDASI_LIST — top 5 PERINGATAN sebagai rekomendasi singkat}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("G. Simpulan").bold = True
    doc.add_paragraph("    {{SIMPULAN — bahasa keyakinan terbatas; pilih kalimat sesuai jumlah PERINGATAN: 0 PERINGATAN = \"Berdasarkan hasil reviu, tidak terdapat hal-hal...\"; ada PERINGATAN = \"Berdasarkan hasil reviu, masih ditemukan beberapa catatan...\". Lihat skill panduan-format-umum/PANDUAN.md}}")
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run("H. Apresiasi & Penutup").bold = True
    doc.add_paragraph("    Terima kasih telah membantu kami dalam menjaga integritas. Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara kami ucapkan terima kasih.")
    doc.add_paragraph()
    doc.add_paragraph()

    sign_tbl = doc.add_table(rows=4, cols=2)
    sign_tbl.rows[0].cells[1].text = "Inspektur II,"
    sign_tbl.rows[1].cells[1].text = ""
    sign_tbl.rows[2].cells[1].text = "{{NAMA_INSPEKTUR}}"
    sign_tbl.rows[3].cells[1].text = "NIP. {{NIP_INSPEKTUR}}"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[Catatan BSrE digital signature di footer setiap halaman]").italic = True

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAMPIRAN 1 — MATRIKS CATATAN REVIU (KRITERIA PERINGATAN)")
    run.bold = True; run.font.size = Pt(12)
    doc.add_paragraph()

    lamp_tbl = doc.add_table(rows=2, cols=6)
    lamp_tbl.style = "Light Grid Accent 1"
    headers = ["No", "RO", "Aspek", "Rule ID", "Kondisi", "Rekomendasi"]
    for c_idx, h in enumerate(headers):
        cell = lamp_tbl.rows[0].cells[c_idx]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    dummy_lamp = ["1", "{{RO_ID}}", "A", "A1", "{{Judul kondisi}}", "{{Rekomendasi}}"]
    for c_idx, v in enumerate(dummy_lamp):
        lamp_tbl.rows[1].cells[c_idx].text = v

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("(Tabel di atas adalah baris dummy untuk demo format. Tambah baris sesuai jumlah catatan PERINGATAN aktual.)").italic = True

    out = os.path.join(OUT_DIR, "LHR-Reviu-RKA-KL-skeleton.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    p1 = build_pkp()
    p2 = build_kkr()
    p3 = build_lhr()
    for p in (p1, p2, p3):
        size = os.path.getsize(p)
        print(f"OK  {os.path.basename(p):50s}  {size:>8} B")
