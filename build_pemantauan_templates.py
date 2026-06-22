"""
Build 3 template pemantauan sekaligus:
- template-lhp-pemantauan-pengadaan.docx  (A–H: pendahuluan 1-6, profil, status, isu, perub kontrak, TL sebelumnya, simpulan+rek, apresiasi)
- template-lhp-pemantauan-tindak-lanjut.docx  (1–5: ringkasan, statistik, aging, kritis, rekomendasi)
- template-lhp-pemantauan-umum.docx  (A–H: dasar, tujuan, periode, metodologi, ringkasan status, hasil per item, rekomendasi, apresiasi)
"""
from docx import Document
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def rebuild(tpl_path):
    doc = Document(str(tpl_path))
    body = doc.element.body
    for el in list(body): body.remove(el)
    return doc, body


def save(doc, body, tpl_path):
    TMP = tpl_path.with_suffix(".tmp.docx")
    doc.save(str(TMP))
    try:
        if tpl_path.exists(): os.remove(str(tpl_path))
        shutil.move(str(TMP), str(tpl_path))
        d2 = Document(str(tpl_path))
        print(f"{tpl_path.name}: paragraphs={len(d2.paragraphs)}")
    except PermissionError:
        print(f"[!] File terkunci: {TMP}")


def E(body): body.append(mkpara(""))
def A(body, txt, **kw): body.append(mkpara(txt, **kw))
def PH(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def S(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


BASE = Path("knowledge/templates/_skeleton-lhp")

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PEMANTAUAN-PENGADAAN
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-pemantauan-pengadaan.docx"
doc, body = rebuild(tpl)

A(body, "NOTA DINAS", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
   "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Pemantauan "
   "Pelaksanaan Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
E(body)
A(body, "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Pemantauan Pengadaan "
   "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
E(body)
A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
E(body); A(body, "=" * 60); E(body)
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL"); E(body)
A(body, "LAPORAN HASIL PEMANTAUAN PENGADAAN BARANG/JASA", bold=True)
A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
E(body)
A(body, "Nomor LHP     : {{NOMOR_LHR}}")
A(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
A(body, "Periode       : {{PERIODE_PELAKSANAAN}}")
E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")
E(body); A(body, "=" * 60); E(body)
A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pemantauan "
   "Pelaksanaan Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} "
   "berdasarkan Surat Tugas Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
   "Pelaksanaan pemantauan berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
E(body)
S(body, "A.", "Pendahuluan", bold=True); E(body)
S(body, "1.", "Latar Belakang"); PH(body, "A1_LATAR_BELAKANG"); E(body)
S(body, "2.", "Dasar Pelaksanaan"); PH(body, "A2_DASAR"); E(body)
S(body, "3.", "Tujuan dan Ruang Lingkup"); PH(body, "A3_TUJUAN_RUANG_LINGKUP"); E(body)
S(body, "4.", "Metodologi"); PH(body, "A4_METODOLOGI"); E(body)
S(body, "5.", "Periode Pemantauan"); PH(body, "A5_PERIODE"); E(body)
S(body, "6.", "Komposisi Tim"); PH(body, "A6_KOMPOSISI_TIM"); E(body)
S(body, "B.", "Profil Pekerjaan", bold=True); PH(body, "B_PROFIL_PEKERJAAN"); E(body)
S(body, "C.", "Status Pelaksanaan", bold=True)
A(body, "Berikut gambaran progres pelaksanaan pekerjaan pada tanggal pemantauan:")
PH(body, "C_STATUS_PELAKSANAAN"); E(body)
S(body, "D.", "Isu dan Permasalahan", bold=True); PH(body, "D_ISU_PERMASALAHAN"); E(body)
S(body, "E.", "Perubahan Kontrak", bold=True); PH(body, "E_PERUBAHAN_KONTRAK"); E(body)
S(body, "F.", "Tindak Lanjut Pemantauan Sebelumnya", bold=True); PH(body, "F_TINDAK_LANJUT_SEBELUMNYA"); E(body)
S(body, "G.", "Simpulan dan Rekomendasi", bold=True); PH(body, "G_SIMPULAN_REKOMENDASI"); E(body)
S(body, "H.", "Apresiasi", bold=True); E(body)
A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
   "selama pelaksanaan pemantauan pengadaan ini.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
save(doc, body, tpl)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. PEMANTAUAN-TINDAK-LANJUT
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-pemantauan-tindak-lanjut.docx"
doc, body = rebuild(tpl)

A(body, "NOTA DINAS", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pemantauan "
   "Tindak Lanjut Hasil Pengawasan (TLHP) terhadap {{NAMA_AUDITI}} berdasarkan "
   "Surat Tugas Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}.")
E(body)
A(body, "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Pemantauan TLHP "
   "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
E(body)
A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
E(body); A(body, "=" * 60); E(body)
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL"); E(body)
A(body, "LAPORAN HASIL PEMANTAUAN TINDAK LANJUT HASIL PENGAWASAN", bold=True)
A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
E(body)
A(body, "Nomor LHP     : {{NOMOR_LHR}}")
A(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
A(body, "Cut-Off Date  : {{PERIODE_PELAKSANAAN}}")
E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")
E(body); A(body, "=" * 60); E(body)
A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pemantauan "
   "Tindak Lanjut Hasil Pengawasan (TLHP) terhadap {{NAMA_AUDITI}} berdasarkan "
   "Surat Tugas Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}.")
E(body)
S(body, "1.", "Ringkasan Eksekutif", bold=True); PH(body, "RINGKASAN_EKSEKUTIF_TLHP"); E(body)
S(body, "2.", "Statistik Umum", bold=True)
A(body, "Berikut statistik tindak lanjut rekomendasi hasil pengawasan:")
PH(body, "STATISTIK_TLHP"); E(body)
S(body, "3.", "Aging Rekomendasi per PIC", bold=True); PH(body, "AGING_PER_PIC"); E(body)
S(body, "4.", "Daftar Rekomendasi Kritis (> 365 Hari)", bold=True); PH(body, "REKOMENDASI_KRITIS"); E(body)
S(body, "5.", "Rekomendasi Percepatan Tindak Lanjut", bold=True); PH(body, "REKOMENDASI_PERCEPATAN"); E(body)
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
save(doc, body, tpl)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. PEMANTAUAN-UMUM
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-pemantauan-umum.docx"
doc, body = rebuild(tpl)

A(body, "NOTA DINAS", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
   "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Pemantauan "
   "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
E(body)
A(body, "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Pemantauan "
   "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
E(body)
A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
E(body); A(body, "=" * 60); E(body)
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL"); E(body)
A(body, "LAPORAN HASIL PEMANTAUAN", bold=True)
A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
E(body)
A(body, "Nomor LHP     : {{NOMOR_LHR}}")
A(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
A(body, "Periode       : {{PERIODE_PELAKSANAAN}}")
E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")
E(body); A(body, "=" * 60); E(body)
A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pemantauan "
   "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan Surat Tugas "
   "Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
   "Pelaksanaan pemantauan berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
E(body)
S(body, "A.", "Dasar Pemantauan", bold=True); PH(body, "A_DASAR"); E(body)
S(body, "B.", "Tujuan dan Ruang Lingkup", bold=True); PH(body, "B_TUJUAN"); E(body)
S(body, "C.", "Periode Pemantauan dan Cut-Off Date", bold=True); PH(body, "C_PERIODE"); E(body)
S(body, "D.", "Metodologi", bold=True); PH(body, "D_METODOLOGI"); E(body)
S(body, "E.", "Ringkasan Status", bold=True)
A(body, "Berikut ringkasan status capaian per item yang dipantau:")
PH(body, "E_RINGKASAN_STATUS"); E(body)
S(body, "F.", "Hasil Pemantauan per Item", bold=True)
A(body, "Narasi rinci item yang berstatus KUNING atau MERAH:")
PH(body, "F_HASIL_PER_ITEM"); E(body)
S(body, "G.", "Rekomendasi dan Tindakan Percepatan", bold=True); PH(body, "G_REKOMENDASI"); E(body)
S(body, "H.", "Apresiasi", bold=True); E(body)
A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
   "selama pelaksanaan pemantauan ini.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")
save(doc, body, tpl)
