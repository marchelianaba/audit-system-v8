"""
Build template-lhp-evaluasi-spip.docx mengikuti SKILL.md evaluasi-spip
(Penjaminan Kualitas Maturitas SPIP Terintegrasi — PerBPKP 5/2021).

Struktur LHE PK SPIP:
  Nota Dinas Pengantar
  ─────────────────────────────────────────────────────────────
  LHE PK SPIP
    Ringkasan Eksekutif          → {{RINGKASAN_EKSEKUTIF}}
    I.  Gambaran Umum
        A. Dasar Hukum           → {{I_A_DASAR_HUKUM}}
        B. Tujuan PK             → {{I_B_TUJUAN}}
        C. Ruang Lingkup         → {{I_C_RUANG_LINGKUP}}
        D. Metodologi            → {{I_D_METODOLOGI}}
    II. Tindak Lanjut            → {{II_TINDAK_LANJUT}}
    III. Hasil PK Maturitas SPIP
        Tabel rekapitulasi nilai → {{TABEL_NILAI_SPIP}}
        A. Penetapan Tujuan      → {{III_A_PENETAPAN_TUJUAN}}
        B. Struktur dan Proses   → {{III_B_STRUKTUR_PROSES}}
        C. Pencapaian Tujuan     → {{III_C_PENCAPAIAN_TUJUAN}}
    IV. Area of Improvement      → {{IV_AOI}}
    V.  Rekomendasi              → {{V_REKOMENDASI}}
    VI. Penutup                  (teks tetap)
  TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-evaluasi-spip.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body):
    body.append(mkpara(""))


def add(body, txt, **kw):
    body.append(mkpara(txt, **kw))


def ph(body, key):
    body.append(mkpara(f"{{{{{key}}}}}"))


def sec(body, num, judul, bold=True):
    body.append(mkpara(f"{num}  {judul}", bold=bold))


# ── buka template lama, hapus semua elemen ───────────────────────────────────
doc = Document(str(TPL))
body = doc.element.body
for el in list(body):
    body.remove(el)

# ═══════════════════════════════════════════════════════════════════════════════
# NOTA DINAS PENGANTAR
# ═══════════════════════════════════════════════════════════════════════════════
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
    "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Penjaminan "
    "Kualitas (PK) atas Penilaian Mandiri Maturitas Penyelenggaraan SPIP Terintegrasi "
    "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
empty(body)
add(body,
    "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Evaluasi (LHE) Penjaminan "
    "Kualitas Maturitas SPIP Terintegrasi {{JUDUL_LHR_INLINE}} sebagai "
    "pertanggungjawaban pelaksanaan penugasan dimaksud.")
empty(body)
add(body,
    "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)
add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER LHE PK SPIP
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL EVALUASI", bold=True)
add(body, "PENJAMINAN KUALITAS MATURITAS PENYELENGGARAAN", bold=True)
add(body, "SISTEM PENGENDALIAN INTERN PEMERINTAH (SPIP) TERINTEGRASI", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}")
empty(body)
add(body, "Nomor LHE     : {{NOMOR_LHR}}")
add(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode PK    : {{PERIODE_PELAKSANAAN}}")
add(body, "Tahun         : {{TAHUN_ANGGARAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══════════════════════════════════════════════════════════════════════════════
# TUBUH LHE PK SPIP
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Penjaminan Kualitas "
    "(PK) atas Penilaian Mandiri (PM) Maturitas Penyelenggaraan SPIP Terintegrasi "
    "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan Surat Tugas Nomor "
    "{{NOMOR_ST}} tanggal {{TANGGAL_ST}}. Pelaksanaan PK berlangsung mulai "
    "{{PERIODE_PELAKSANAAN}}.")
empty(body)

# ── Ringkasan Eksekutif ───────────────────────────────────────────────────────
add(body, "RINGKASAN EKSEKUTIF", bold=True)
ph(body, "RINGKASAN_EKSEKUTIF")
empty(body)

# ── I. GAMBARAN UMUM ─────────────────────────────────────────────────────────
sec(body, "I.", "GAMBARAN UMUM", bold=True)
empty(body)

add(body, "A.  Dasar Hukum", bold=True)
add(body,
    "Penjaminan Kualitas Maturitas Penyelenggaraan SPIP Terintegrasi terhadap "
    "{{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} dilaksanakan berdasarkan:")
ph(body, "I_A_DASAR_HUKUM")
empty(body)

add(body, "B.  Tujuan", bold=True)
ph(body, "I_B_TUJUAN")
empty(body)

add(body, "C.  Ruang Lingkup", bold=True)
ph(body, "I_C_RUANG_LINGKUP")
empty(body)

add(body, "D.  Metodologi", bold=True)
ph(body, "I_D_METODOLOGI")
empty(body)

# ── II. TINDAK LANJUT ────────────────────────────────────────────────────────
sec(body, "II.", "TINDAK LANJUT ATAS HASIL PK TAHUN SEBELUMNYA", bold=True)
empty(body)
ph(body, "II_TINDAK_LANJUT")
empty(body)

# ── III. HASIL PK MATURITAS SPIP ─────────────────────────────────────────────
sec(body, "III.", "HASIL PENJAMINAN KUALITAS MATURITAS SPIP TERINTEGRASI", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil PK atas PM Maturitas SPIP Terintegrasi terhadap "
    "{{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}, diperoleh nilai maturitas SPIP "
    "sebagai berikut:")
ph(body, "TABEL_NILAI_SPIP")
empty(body)

add(body, "A.  Penetapan Tujuan", bold=True)
ph(body, "III_A_PENETAPAN_TUJUAN")
empty(body)

add(body, "B.  Struktur dan Proses", bold=True)
ph(body, "III_B_STRUKTUR_PROSES")
empty(body)

add(body, "C.  Pencapaian Tujuan", bold=True)
ph(body, "III_C_PENCAPAIAN_TUJUAN")
empty(body)

# ── IV. AREA OF IMPROVEMENT ──────────────────────────────────────────────────
sec(body, "IV.", "AREA OF IMPROVEMENT (AoI)", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil PK, terdapat area yang perlu diperbaiki untuk meningkatkan "
    "tingkat maturitas SPIP Terintegrasi {{NAMA_AUDITI}}:")
ph(body, "IV_AOI")
empty(body)

# ── V. REKOMENDASI ───────────────────────────────────────────────────────────
sec(body, "V.", "REKOMENDASI", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil PK di atas, direkomendasikan kepada "
    "{{PENERIMA_LHP}} agar:")
ph(body, "V_REKOMENDASI")
empty(body)

# ── VI. PENUTUP ──────────────────────────────────────────────────────────────
sec(body, "VI.", "PENUTUP", bold=True)
empty(body)
add(body,
    "Demikian laporan hasil Penjaminan Kualitas Maturitas Penyelenggaraan SPIP "
    "Terintegrasi ini kami sampaikan. Hasil PK ini digunakan sebagai bahan masukan "
    "dalam penyempurnaan Laporan Hasil Penilaian Mandiri (LHPM) dan peningkatan "
    "efektivitas sistem pengendalian intern {{NAMA_AUDITI}}. Atas perhatian dan "
    "kerja sama Saudara, kami ucapkan terima kasih.")
empty(body)

add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
    exit()

# ── Verifikasi ───────────────────────────────────────────────────────────────
doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt:
        try:
            print(f"{i:3}: {txt[:80]}")
        except UnicodeEncodeError:
            print(f"{i:3}: [non-printable]")
