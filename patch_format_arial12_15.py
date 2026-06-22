"""
Patch one-time: set font Arial 12pt + line spacing 1,5 di semua template.

Yang diubah:
  1. mkpara() di semua build_*_template*.py  — paragraf statis template
  2. make_p()  di render_lhp.py              — paragraf dinamis saat render

Tidak menyentuh logika/workflow apapun, hanya formatting.
"""
import re
from pathlib import Path

ROOT = Path(__file__).parent

# ── 1. Ganti mkpara di semua build scripts ──────────────────────────────────

NEW_MKPARA = '''\
def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p\
'''

MKPARA_PAT = re.compile(
    r'def mkpara\(txt, bold=False, italic=False\):.*?return p',
    re.DOTALL
)

build_scripts = list(ROOT.glob("build_*template*.py")) + list(ROOT.glob("build_reviu_templates.py")) + list(ROOT.glob("build_pemantauan_templates.py"))
# dedup
build_scripts = list({p.resolve(): p for p in build_scripts}.values())

for script in sorted(build_scripts):
    txt = script.read_text(encoding="utf-8")
    if "def mkpara" not in txt:
        print(f"  SKIP (no mkpara): {script.name}")
        continue
    new_txt = MKPARA_PAT.sub(NEW_MKPARA, txt)
    if new_txt == txt:
        print(f"  UNCHANGED: {script.name}")
    else:
        script.write_text(new_txt, encoding="utf-8")
        print(f"  PATCHED: {script.name}")

# ── 2. Patch make_p() di render_lhp.py ─────────────────────────────────────

RENDER = ROOT / "backend" / "v6" / "scripts" / "render_lhp.py"

OLD_MAKE_P = '''\
    def make_p(text, fmt=None):
        fmt = fmt or {}
        new = doc.add_paragraph()
        if fmt.get("align") == "justify":
            new.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif fmt.get("align") == "center":
            new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif fmt.get("align") == "right":
            new.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if fmt.get("indent"):
            new.paragraph_format.left_indent = Cm(fmt["indent"])
        r = new.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        r.font.size = Pt(fmt.get("size", 11))
        r.font.name = "Arial"
        return new'''

NEW_MAKE_P = '''\
    def make_p(text, fmt=None):
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        fmt = fmt or {}
        new = doc.add_paragraph()
        if fmt.get("align") == "justify":
            new.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif fmt.get("align") == "center":
            new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif fmt.get("align") == "right":
            new.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if fmt.get("indent"):
            new.paragraph_format.left_indent = Cm(fmt["indent"])
        # Spasi 1,5 via XML (paling reliable lintas versi python-docx)
        pPr = new._element.get_or_add_pPr()
        sp_el = _OE("w:spacing")
        sp_el.set(_qn("w:line"), "360")
        sp_el.set(_qn("w:lineRule"), "auto")
        pPr.append(sp_el)
        r = new.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        r.font.size = Pt(fmt.get("size", 12))
        r.font.name = "Arial"
        return new'''

render_txt = RENDER.read_text(encoding="utf-8")
if OLD_MAKE_P in render_txt:
    render_txt = render_txt.replace(OLD_MAKE_P, NEW_MAKE_P)
    RENDER.write_text(render_txt, encoding="utf-8")
    print(f"\n  PATCHED: {RENDER.name}")
else:
    print(f"\n  WARNING: make_p() di render_lhp.py tidak ditemukan — cek manual!")

print("\nDone. Jalankan rebuild_all_templates.py untuk rebuild semua template.")
