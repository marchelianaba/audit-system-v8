"""
Script sekali-pakai: restruktur template-lhp-audit-umum.docx
agar sesuai struktur bab LHA audit-umum (A–H + 2 Lampiran).

Jalankan dari root project:
  python build_audit_umum_template.py
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pathlib import Path

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def get_text(p_elem):
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def set_para_text(p_elem, text, bold=False, size=12):
    """Hapus semua run di paragraph dan isi ulang dengan teks baru."""
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    # Spasi 1,5
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_elem.append(r)


def new_para_elem(text, bold=False, size=12):
    """Buat elemen paragraf baru."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), "360")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def new_page_break():
    """Buat paragraf page break."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def main():
    doc = Document(str(TPL))
    body = doc.element.body

    # ----------------------------------------------------------------
    # Peta transformasi placeholder di zona konten LHA:
    # old_placeholder → (heading_text, new_placeholder)
    # ----------------------------------------------------------------
    REMAP = {
        "{{DASAR_HUKUM_LIST}}": ("A.  DASAR", "{{A_DASAR}}"),
        "{{SASARAN_LIST}}":      ("B.  TUJUAN", "{{B_TUJUAN}}"),
        "{{RUANG_LINGKUP}}":     ("C.  RUANG LINGKUP", "{{C_RUANG_LINGKUP}}"),
        "{{GAMBARAN_UMUM}}":     ("E.  GAMBARAN UMUM OBJEK", "{{E_GAMBARAN_UMUM}}"),
        "{{HASIL_REVIU_LOOP}}":  ("F.  HASIL AUDIT", "{{F_HASIL_AUDIT}}"),
    }
    # Paragraf {{METODOLOGI_REVIU}} perlu penanganan khusus:
    # teksnya "{{METODOLOGI_REVIU}} Reviu dilaksanakan..." — kita ganti seluruhnya
    METODOLOGI_OLD = "{{METODOLOGI_REVIU}}"
    HASIL_INTRO    = "{{HASIL_REVIU_INTRO}}"   # hapus
    SIMPULAN_OLD   = "{{SIMPULAN_REVIU}}"      # hapus, sisipkan G + H di sini

    # ----------------------------------------------------------------
    # Pass 1: identifikasi semua paragraf yang perlu diproses
    # ----------------------------------------------------------------
    # Cari batas konten LHA: mulai dari {{PENERIMA_LHP}} sampai TTD ke-2
    # (yang pertama ada di Nota Dinas, yang kedua di isi LHA)
    tdd_count = 0
    anchor_tdd = None          # TTD ke-2 (awal blok tanda tangan isi LHA)
    tembusan_elem = None       # {{TEMBUSAN_LIST}} di isi LHA

    for p in body.iter(qn("w:p")):
        txt = get_text(p)
        if "{{TTD_INSPEKTUR}}" in txt:
            tdd_count += 1
            if tdd_count == 2:
                anchor_tdd = p
        if "{{TEMBUSAN_LIST}}" in txt and tdd_count == 2:
            tembusan_elem = p

    # ----------------------------------------------------------------
    # Pass 2: transformasi paragraf zona konten
    # ----------------------------------------------------------------
    simpulan_parent = None
    simpulan_ref    = None     # simpan referensi setelah simpulan dihapus

    for p in list(body.iter(qn("w:p"))):
        txt = get_text(p)

        # --- Hapus HASIL_REVIU_INTRO ---
        if HASIL_INTRO in txt:
            p.getparent().remove(p)
            continue

        # --- Ganti SIMPULAN dengan G + H (sisipkan di posisinya, lalu hapus) ---
        if SIMPULAN_OLD in txt:
            parent = p.getparent()
            # Sisipkan H dulu (addprevious → urut terbalik)
            h_ph   = new_para_elem("{{H_APRESIASI}}")
            h_head = new_para_elem("H.  APRESIASI DAN PENUTUP", bold=True)
            g_ph   = new_para_elem("{{G_REKOMENDASI}}")
            g_head = new_para_elem("G.  REKOMENDASI", bold=True)
            p.addprevious(g_head)
            p.addprevious(g_ph)
            p.addprevious(h_head)
            p.addprevious(h_ph)
            parent.remove(p)
            continue

        # --- Ganti METODOLOGI ---
        if METODOLOGI_OLD in txt:
            # Sisipkan heading sebelum ini
            p.addprevious(new_para_elem("D.  METODOLOGI", bold=True))
            set_para_text(p, "{{D_METODOLOGI}}")
            continue

        # --- Remap placeholder generik → audit-umum ---
        for old_ph, (heading, new_ph) in REMAP.items():
            if old_ph in txt:
                p.addprevious(new_para_elem(heading, bold=True))
                set_para_text(p, new_ph)
                break

    # ----------------------------------------------------------------
    # Pass 3: tambahkan Lampiran setelah {{TEMBUSAN_LIST}}
    # ----------------------------------------------------------------
    if tembusan_elem is not None:
        # Sisipkan setelah tembusan_elem (addnext = setelah)
        lamp2_ph   = new_para_elem("{{LAMPIRAN_2_DOKUMEN_SUMBER}}")
        lamp2_head = new_para_elem("LAMPIRAN 2: DAFTAR DOKUMEN SUMBER", bold=True)
        pb2        = new_page_break()
        lamp1_ph   = new_para_elem("{{LAMPIRAN_1_MATRIKS_TEMUAN}}")
        lamp1_head = new_para_elem("LAMPIRAN 1: MATRIKS TEMUAN (CCSAA)", bold=True)
        pb1        = new_page_break()

        # Urutan addnext (setiap addnext menyisipkan tepat setelah tembusan):
        # Akhirnya urutan di dokumen: pb1 → lamp1_head → lamp1_ph → pb2 → lamp2_head → lamp2_ph
        tembusan_elem.addnext(lamp2_ph)
        tembusan_elem.addnext(lamp2_head)
        tembusan_elem.addnext(pb2)
        tembusan_elem.addnext(lamp1_ph)
        tembusan_elem.addnext(lamp1_head)
        tembusan_elem.addnext(pb1)

    doc.save(str(TPL))
    print(f"Template disimpan: {TPL}")

    # Verifikasi: tampilkan daftar placeholder yang ada
    print("\nPlaceholder di template baru:")
    doc2 = Document(str(TPL))
    for p in doc2.paragraphs:
        if "{{" in p.text:
            print(f"  {p.text[:80]}")


if __name__ == "__main__":
    main()
