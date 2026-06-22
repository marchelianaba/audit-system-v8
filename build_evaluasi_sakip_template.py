"""
Build template-lhp-evaluasi-sakip.docx mengikuti 02-template-lhe.md (SKILL.md evaluasi-sakip).

Struktur:
  Nota Dinas Pengantar
  ───────────────────────────────────────────────────────
  LHE AKIP
    Ringkasan Eksekutif          → {{RINGKASAN_EKSEKUTIF}}
    I.  Gambaran Umum
        A. Dasar Hukum           → {{I_A_DASAR_HUKUM}}
        B. Tujuan                → {{I_B_TUJUAN}}
        C. Ruang Lingkup         (teks tetap — 4 komponen)
        D. Metodologi            → {{I_D_METODOLOGI}}
    II. Tindak Lanjut            → {{II_TINDAK_LANJUT}}
    III. Hasil Evaluasi
        Tabel rekapitulasi nilai → {{TABEL_NILAI_LKE}}
        A. Perencanaan Kinerja   → {{III_A_PERENCANAAN}}
        B. Pengukuran Kinerja    → {{III_B_PENGUKURAN}}
        C. Pelaporan Kinerja     → {{III_C_PELAPORAN}}
        D. Evaluasi Internal     → {{III_D_EVALUASI_INTERNAL}}
    IV. Rekomendasi              → {{IV_REKOMENDASI}}
    V.  Penutup                  (teks tetap)
  TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-evaluasi-sakip.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body):
    body.append(mkpara(""))


def add(body, txt, **kw):
    body.append(mkpara(txt, **kw))


def ph(body, key):
    body.append(mkpara(f"{{{{{key}}}}}"))


def sec(body, num, judul, bold=True):
    body.append(mkpara(f"{num}  {judul}", bold=bold))


# ── buka template lama, hapus semua elemen ───────────────────────────────────
doc = Document(str(TPL))
body = doc.element.body
for el in list(body):
    body.remove(el)

# ═══════════════════════════════════════════════════════════════════════════════
# NOTA DINAS PENGANTAR
# ═══════════════════════════════════════════════════════════════════════════════
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
    "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Evaluasi "
    "Akuntabilitas Kinerja Instansi Pemerintah (AKIP) terhadap {{NAMA_AUDITI}} "
    "Tahun Anggaran {{TAHUN_ANGGARAN}}.")
empty(body)
add(body,
    "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Evaluasi (LHE) AKIP "
    "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
empty(body)
add(body,
    "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)
add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER LHE
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL EVALUASI", bold=True)
add(body, "AKUNTABILITAS KINERJA INSTANSI PEMERINTAH (AKIP)", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}")
empty(body)
add(body, "Nomor LHE     : {{NOMOR_LHR}}")
add(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode Eval. : {{PERIODE_PELAKSANAAN}}")
add(body, "Tahun Evaluasi: {{TAHUN_ANGGARAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══════════════════════════════════════════════════════════════════════════════
# TUBUH LHE
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Evaluasi AKIP "
    "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan Surat Tugas "
    "Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. Pelaksanaan evaluasi berlangsung "
    "mulai {{PERIODE_PELAKSANAAN}}.")
empty(body)

# ── Ringkasan Eksekutif ───────────────────────────────────────────────────────
add(body, "RINGKASAN EKSEKUTIF", bold=True)
ph(body, "RINGKASAN_EKSEKUTIF")
empty(body)

# ── I. GAMBARAN UMUM ─────────────────────────────────────────────────────────
sec(body, "I.", "GAMBARAN UMUM", bold=True)
empty(body)

add(body, "A.  Dasar Hukum", bold=True)
add(body,
    "Evaluasi AKIP {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} dilaksanakan berdasarkan:")
ph(body, "I_A_DASAR_HUKUM")
empty(body)

add(body, "B.  Tujuan", bold=True)
ph(body, "I_B_TUJUAN")
empty(body)

add(body, "C.  Ruang Lingkup", bold=True)
add(body,
    "Evaluasi dilaksanakan terhadap empat komponen penilaian AKIP sesuai "
    "PermenPAN-RB Nomor 88 Tahun 2021, yaitu:")
add(body, "a.  Perencanaan Kinerja;")
add(body, "b.  Pengukuran Kinerja;")
add(body, "c.  Pelaporan Kinerja; dan")
add(body, "d.  Evaluasi Akuntabilitas Kinerja Internal.")
empty(body)

add(body, "D.  Metodologi", bold=True)
ph(body, "I_D_METODOLOGI")
empty(body)

# ── II. TINDAK LANJUT ────────────────────────────────────────────────────────
sec(body, "II.", "TINDAK LANJUT ATAS HASIL EVALUASI TAHUN SEBELUMNYA", bold=True)
empty(body)
ph(body, "II_TINDAK_LANJUT")
empty(body)

# ── III. HASIL EVALUASI ──────────────────────────────────────────────────────
sec(body, "III.", "HASIL EVALUASI", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil evaluasi AKIP terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}, "
    "diperoleh nilai dan predikat sebagai berikut:")
ph(body, "TABEL_NILAI_LKE")
empty(body)

add(body, "A.  Perencanaan Kinerja", bold=True)
ph(body, "III_A_PERENCANAAN")
empty(body)

add(body, "B.  Pengukuran Kinerja", bold=True)
ph(body, "III_B_PENGUKURAN")
empty(body)

add(body, "C.  Pelaporan Kinerja", bold=True)
ph(body, "III_C_PELAPORAN")
empty(body)

add(body, "D.  Evaluasi Akuntabilitas Kinerja Internal", bold=True)
ph(body, "III_D_EVALUASI_INTERNAL")
empty(body)

# ── IV. REKOMENDASI ──────────────────────────────────────────────────────────
sec(body, "IV.", "REKOMENDASI", bold=True)
empty(body)
add(body,
    "Berdasarkan hasil evaluasi di atas, direkomendasikan kepada "
    "{{PENERIMA_LHP}} agar:")
ph(body, "IV_REKOMENDASI")
empty(body)

# ── V. PENUTUP ───────────────────────────────────────────────────────────────
sec(body, "V.", "PENUTUP", bold=True)
empty(body)
add(body,
    "Terima kasih telah membantu kami dalam menjaga integritas dan akuntabilitas "
    "kinerja instansi pemerintah. Demikian laporan hasil evaluasi ini disampaikan "
    "sebagai bahan pertimbangan dalam meningkatkan akuntabilitas kinerja "
    "{{NAMA_AUDITI}}. Atas perhatian dan kerja sama Saudara, kami ucapkan "
    "terima kasih.")
empty(body)

add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
    exit()

# ── Verifikasi ───────────────────────────────────────────────────────────────
doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt:
        try:
            print(f"{i:3}: {txt[:80]}")
        except UnicodeEncodeError:
            print(f"{i:3}: [non-printable]")
