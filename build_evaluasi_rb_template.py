"""
Build template-lhp-evaluasi-reformasi-birokrasi.docx sesuai SKILL.md:

Nota Dinas → [Paragraf pembuka RB]
[Tujuan evaluasi]
[Metode evaluasi]
Hasil Evaluasi:
1. Gambaran Umum Pelaksanaan RB       → {{SEKSI_1_GAMBARAN_UMUM}}
2. Analisis Dampak RB Tematik         → {{SEKSI_2_DAMPAK}}
3. Hasil Evaluasi Pelaksanaan Renaksi → {{SEKSI_3_RENAKSI}}
4. Penutup                            (teks tetap)
TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-evaluasi-reformasi-birokrasi.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body):
    body.append(mkpara(""))


def add(body, txt, **kw):
    body.append(mkpara(txt, **kw))


def ph(body, key):
    body.append(mkpara(f"{{{{{key}}}}}"))


def sec(body, lbl, judul, bold=True):
    body.append(mkpara(f"{lbl}  {judul}", bold=bold))


# ── buka template lama, hapus semua elemen ───────────────────────────────────
doc = Document(str(TPL))
body = doc.element.body
for el in list(body):
    body.remove(el)

# ═══════════════════════════════════════════════════════════════════════════════
# NOTA DINAS PENGANTAR
# ═══════════════════════════════════════════════════════════════════════════════
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Sesuai dengan Peraturan Menteri Pendayagunaan Aparatur Negara dan Reformasi "
    "Birokrasi Nomor 9 Tahun 2023 tentang Evaluasi Reformasi Birokrasi, Keputusan "
    "Menteri PAN-RB Nomor 182 Tahun 2024 tentang Juknis Evaluasi RB Tahun 2024, "
    "dan Surat Edaran Menteri PAN-RB Nomor 6 Tahun 2025 tentang Pelaksanaan "
    "Reformasi Birokrasi pada Periode Transisi, Evaluator Internal Kementerian "
    "Komunikasi dan Digital telah melakukan evaluasi Reformasi Birokrasi "
    "{{JUDUL_LHR_INLINE}} dengan menerbitkan Surat Tugas Nomor {{NOMOR_ST}} "
    "tanggal {{TANGGAL_ST}}. Bersama ini kami sampaikan laporan hasil evaluasi "
    "terkait hal tersebut.")
empty(body)
add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER LHEI
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL EVALUASI INTERNAL (LHEI)", bold=True)
add(body, "PELAKSANAAN REFORMASI BIROKRASI", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}")
empty(body)
add(body, "Nomor LHEI    : {{NOMOR_LHR}}")
add(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode Eval. : {{PERIODE_PELAKSANAAN}}")
add(body, "Tahun         : {{TAHUN_ANGGARAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══════════════════════════════════════════════════════════════════════════════
# TUBUH LHEI
# ═══════════════════════════════════════════════════════════════════════════════
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)

# Paragraf pembuka
add(body,
    "Sesuai dengan {{DASAR_PERMINTAAN}}, kami telah melaksanakan Evaluasi "
    "Pelaksanaan Reformasi Birokrasi di Lingkungan {{NAMA_AUDITI}} Tahun "
    "{{TAHUN_ANGGARAN}} berdasarkan Surat Tugas Inspektur Jenderal Nomor "
    "{{NOMOR_ST}} tanggal {{TANGGAL_ST}}. Pelaksanaan evaluasi berlangsung "
    "mulai {{PERIODE_PELAKSANAAN}}.")
empty(body)

# Tujuan evaluasi
ph(body, "TUJUAN_EVALUASI_RB")
empty(body)

# Metode evaluasi
ph(body, "METODE_EVALUASI_RB")
empty(body)

# Judul blok hasil
add(body, "Hasil Evaluasi:", bold=True)
empty(body)

# ── 1. Gambaran Umum ─────────────────────────────────────────────────────────
sec(body, "1.", "Gambaran Umum Pelaksanaan Reformasi Birokrasi", bold=True)
empty(body)
ph(body, "SEKSI_1_GAMBARAN_UMUM")
empty(body)

# ── 2. Analisis Dampak ───────────────────────────────────────────────────────
sec(body, "2.", "Analisis Dampak Reformasi Birokrasi Tematik", bold=True)
empty(body)
ph(body, "SEKSI_2_DAMPAK")
empty(body)

# ── 3. Hasil Evaluasi Renaksi ─────────────────────────────────────────────────
sec(body, "3.", "Hasil Evaluasi Pelaksanaan Rencana Aksi", bold=True)
empty(body)
ph(body, "SEKSI_3_RENAKSI")
empty(body)

# ── 4. Penutup ────────────────────────────────────────────────────────────────
sec(body, "4.", "Penutup", bold=True)
empty(body)
add(body,
    "Tim evaluasi mengucapkan terima kasih atas kerja sama {{NAMA_AUDITI}} "
    "dalam pelaksanaan evaluasi Reformasi Birokrasi ini.")
add(body,
    "Laporan ini diharapkan dapat memberikan informasi yang memadai sebagai "
    "dasar perbaikan pelaksanaan Reformasi Birokrasi di lingkungan {{NAMA_AUDITI}}.")
add(body,
    "Kegiatan evaluasi ini telah dilaksanakan sesuai dengan Standar Audit Intern "
    "Pemerintah Indonesia (SAIPI).")
empty(body)

add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
    exit()

doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
