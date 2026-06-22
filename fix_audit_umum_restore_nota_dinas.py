"""
Restore paragraf penutup Nota Dinas yang terhapus:
  "Terima kasih telah membantu kami dalam menjaga integritas.
   Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama
   Saudara, kami mengucapkan terima kasih."
Sisipkan SEBELUM {{TTD_INSPEKTUR}} di section Nota Dinas.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")
doc = Document(str(TPL))

CLOSING = (
    "Terima kasih telah membantu kami dalam menjaga integritas. "
    "Demikian Nota Dinas ini kami sampaikan. "
    "Atas perhatian dan kerja sama Saudara, kami mengucapkan terima kasih."
)

XML_NS = "http://www.w3.org/XML/1998/namespace"


def gtext(p):
    return "".join(t.text or "" for t in p._element.iter(qn("w:t")))


def mkpara_full(txt):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


# Cek apakah 'Demikian Nota Dinas' sudah ada — kalau ada, skip
already = any("Demikian Nota Dinas" in gtext(p) for p in doc.paragraphs)
if already:
    print("'Demikian Nota Dinas' sudah ada — tidak perlu diperbaiki.")
else:
    # Cari paragraf {{TTD_INSPEKTUR}} PERTAMA (yang ada di Nota Dinas)
    ttd_para = None
    for p in doc.paragraphs:
        if "{{TTD_INSPEKTUR}}" in gtext(p):
            ttd_para = p
            break

    if ttd_para is None:
        print("ERROR: {{TTD_INSPEKTUR}} tidak ditemukan.")
    else:
        # Sisipkan closing + 1 empty para sebelum TTD
        empty_elem = mkpara_full("")
        closing_elem = mkpara_full(CLOSING)
        ttd_para._element.addprevious(empty_elem)
        ttd_para._element.addprevious(closing_elem)
        print("Paragraf penutup Nota Dinas ditambahkan sebelum {{TTD_INSPEKTUR}}.")

    TMP = TPL.with_suffix(".tmp.docx")
    doc.save(str(TMP))
    try:
        if TPL.exists():
            os.remove(str(TPL))
        shutil.move(str(TMP), str(TPL))
        print("Template diperbarui:", TPL)
    except PermissionError:
        print(f"[!] File terkunci — perubahan tersimpan ke: {TMP}")
        print("    Tutup file di Word lalu jalankan ulang.")

# Verifikasi
doc2 = Document(str(TPL if TPL.exists() else TPL.with_suffix(".tmp.docx")))
print(f"\nparagraphs: {len(doc2.paragraphs)}")
for i, p in enumerate(doc2.paragraphs):
    if i < 16:
        txt = p.text.strip()
        if txt:
            print(f"{i:3}: {txt[:90]}")
