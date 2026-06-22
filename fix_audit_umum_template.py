"""
Fix template-lhp-audit-umum.docx:
  1. Hapus paragraf survei hardcoded di Nota Dinas
  2. Ganti ${ttd_pengirim} → {{TTD_INSPEKTUR}}
  3. Ganti 'Muhammad Arief' → {{NAMA_INSPEKTUR}}; tambah NIP. {{NIP_INSPEKTUR}} setelahnya
  4. Hapus teks residual lama di section H_APRESIASI ("Terima kasih... Demikian laporan")
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")
TMP_PREV = TPL.with_suffix(".tmp.docx")
# Gunakan asli (jika bisa dibaca), bukan tmp dari run sebelumnya
doc = Document(str(TPL))
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def gtext(p):
    return "".join(t.text or "" for t in p._element.iter(qn("w:t")))


def stext(p_elem, txt):
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = txt
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_elem.append(r)


def delp(p):
    parent = p._element.getparent()
    if parent is not None:
        parent.remove(p._element)


def mkpara(txt):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = txt
    r.append(t)
    p.append(r)
    return p


nama_para = None   # referensi untuk insert NIP setelahnya

for p in list(doc.paragraphs):
    txt = gtext(p).strip()

    # 1. Hapus paragraf survei hardcoded di Nota Dinas
    if "survei perihal kepuasan auditan" in txt or "mengisi survei" in txt.lower():
        delp(p)

    # 2. Ganti ${ttd_pengirim} → {{TTD_INSPEKTUR}}
    elif "${ttd_pengirim}" in txt:
        stext(p._element, "{{TTD_INSPEKTUR}}")

    # 3. Ganti nama hardcoded → {{NAMA_INSPEKTUR}}; tambah NIP setelahnya jika belum ada
    elif txt in ("Muhammad Arief", "Muhammad Arief."):
        stext(p._element, "{{NAMA_INSPEKTUR}}")
        # Cek sibling berikutnya (lxml element) — tambah NIP jika belum ada
        nxt = p._element.getnext()
        nxt_txt = ("".join(t.text or "" for t in nxt.iter(qn("w:t")))).strip() if nxt is not None else ""
        if "NIP" not in nxt_txt:
            nip_elem = mkpara("NIP. {{NIP_INSPEKTUR}}")
            p._element.addnext(nip_elem)

    # 4. Hapus teks penutup residual di section H_APRESIASI
    # Bedakan dari closing Nota Dinas dengan "Demikian Nota Dinas" vs "Demikian laporan"
    elif "membantu kami dalam menjaga integritas" in txt and "Demikian laporan" in txt:
        delp(p)

TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
import shutil, os
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    renamed = True
except PermissionError:
    renamed = False
if renamed:
    print("Template diperbarui:", TPL)
else:
    print(f"[!] File terkunci — perubahan tersimpan ke: {TMP}")
    print("    Tutup file di Word lalu jalankan ulang script ini.")

# Verifikasi
doc2 = Document(str(TPL))
print("\n=== STRUKTUR SETELAH FIX ===")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt and i < 20:
        print(f"{i:3}: {txt[:100]}")
print("  ...")
for i, p in enumerate(doc2.paragraphs):
    txt = p.text.strip()
    if txt and any(k in txt for k in ("APRESIASI", "{{H_", "integritas", "Terima kasih yang")):
        print(f"{i:3}: {txt[:100]}")
