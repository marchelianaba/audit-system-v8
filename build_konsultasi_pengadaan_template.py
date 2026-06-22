"""
Rebuild template-lhp-konsultasi-pengadaan.docx sesuai SKILL.md v3.0:

  [Header + Cover Laporan Hasil Pendampingan]
  [Paragraf pembuka — catatan advisory]

  I.  Kegiatan Pendampingan yang Telah Diselesaikan → {{BAB1_KEGIATAN}}
      (Dokumen Pendukung per Kegiatan)              → {{BAB1_DOKUMEN_PENDUKUNG}}
  II. Hal yang Masih Memerlukan Tindak Lanjut       → {{BAB2_TINDAK_LANJUT}}
  III. Kesimpulan                                   → {{BAB3_KESIMPULAN}}
  TTD + Tembusan
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn
from pathlib import Path
import shutil, os

TPL    = Path("knowledge/templates/_skeleton-lhp/template-lhp-konsultasi-pengadaan.docx")
BASE   = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def E(body): body.append(mkpara(""))
def A(body, txt, **kw): body.append(mkpara(txt, **kw))
def PH(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def S(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


# Buka base yang punya sectPr valid
src = BASE if BASE.exists() else TPL
doc = Document(str(src))
body = doc.element.body
sect_pr = body.find(_qn("w:sectPr"))
for el in list(body):
    body.remove(el)
if sect_pr is None:
    sect_pr = OxmlElement("w:sectPr")
body.append(sect_pr)

# ═══ NOTA DINAS PENGANTAR ═══
A(body, "NOTA DINAS", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pendampingan "
   "Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} berdasarkan Surat Tugas Nomor "
   "{{NOMOR_ST}} tanggal {{TANGGAL_ST}}.")
E(body)
A(body, "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Pendampingan Pengadaan "
   "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
E(body)
A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")

# ═══ COVER LAPORAN ═══
E(body); A(body, "=" * 60); E(body)
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL"); E(body)
A(body, "LAPORAN HASIL PENDAMPINGAN PENGADAAN BARANG/JASA", bold=True)
A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
E(body)
A(body, "Auditan          : {{NAMA_AUDITI}}")
A(body, "Dasar Penugasan  : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
A(body, "Periode          : {{PERIODE_PELAKSANAAN}}")
E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")

# ═══ TUBUH LAPORAN ═══
E(body); A(body, "=" * 60); E(body)
A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
A(body, "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Pendampingan "
   "Pengadaan Barang/Jasa terhadap {{NAMA_AUDITI}} berdasarkan Surat Tugas Nomor "
   "{{NOMOR_ST}} tanggal {{TANGGAL_ST}}.")
E(body)
A(body, "Catatan: Laporan ini berisi rangkaian kegiatan pendampingan yang telah "
   "diselesaikan tim Inspektorat II atas permintaan unit kerja. Pendampingan bersifat "
   "advisory dan preventif — tidak memberikan keyakinan dan tidak mengikat keputusan "
   "pejabat yang berwenang.", italic=True)
E(body)

# ── Bab I ────────────────────────────────────────────────────────────────────
S(body, "I.", "KEGIATAN PENDAMPINGAN YANG TELAH DISELESAIKAN", bold=True)
A(body, "Berikut kegiatan pendampingan yang telah diselesaikan tim Inspektorat II:")
PH(body, "BAB1_KEGIATAN"); E(body)
A(body, "Dokumen Pendukung per Kegiatan:", bold=True)
PH(body, "BAB1_DOKUMEN_PENDUKUNG"); E(body)

# ── Bab II ───────────────────────────────────────────────────────────────────
S(body, "II.", "HAL YANG MASIH MEMERLUKAN TINDAK LANJUT", bold=True)
PH(body, "BAB2_TINDAK_LANJUT"); E(body)

# ── Bab III ──────────────────────────────────────────────────────────────────
S(body, "III.", "KESIMPULAN", bold=True)
PH(body, "BAB3_KESIMPULAN"); E(body)

# ── TTD ──────────────────────────────────────────────────────────────────────
A(body, "Pendampingan ini dilaksanakan sesuai Standar Audit Intern Pemerintah Indonesia "
   "(SAIPI). Atas kerja sama {{NAMA_AUDITI}}, kami ucapkan terima kasih.")
A(body, "Demikian laporan ini kami sampaikan.")
E(body)
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists(): os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    d2 = Document(str(TPL))
    print(f"Template berhasil diperbarui: {TPL}")
    print(f"paragraphs: {len(d2.paragraphs)}")
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
