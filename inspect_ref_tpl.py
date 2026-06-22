from docx import Document
doc = Document("knowledge/templates/Laporan Hasil Audit.docx")
print(f"Paragraf total: {len(doc.paragraphs)}, Tabel: {len(doc.tables)}")
print("\n=== SEMUA PARAGRAF ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name if p.style else "-"
    if txt:
        print(f"{i:3} [{style[:20]}] {txt[:100]}")
    else:
        print(f"{i:3} [{style[:20]}] <empty>")

print("\n=== ISI TABEL ===")
for ti, tbl in enumerate(doc.tables):
    print(f"\nTabel {ti}:")
    for row in tbl.rows:
        cells = [c.text.strip()[:50] for c in row.cells]
        print(f"  {cells}")
