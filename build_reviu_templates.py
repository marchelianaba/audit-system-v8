"""
Build 3 template reviu sekaligus:
- template-lhp-reviu-pengadaan.docx  (A–F: Pendahuluan 1-7, Gambaran Umum, Hasil [C.1 Perencanaan / C.2 Pemilihan], Simpulan, Rekomendasi, Apresiasi)
- template-lhp-reviu-rka-kl.docx     (A–F: Pendahuluan 1-5, Gambaran Umum, Hasil [C.1–C.6], Simpulan, Rekomendasi, Apresiasi)
- template-lhp-reviu-umum.docx       (A–G: Dasar, Tujuan, Metodologi, Hasil Reviu, Catatan+Rek, Simpulan, Apresiasi)
"""
from docx import Document
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def rebuild(tpl_path):
    doc = Document(str(tpl_path))
    body = doc.element.body
    for el in list(body): body.remove(el)
    return doc, body


def save(doc, body, tpl_path):
    TMP = tpl_path.with_suffix(".tmp.docx")
    doc.save(str(TMP))
    try:
        if tpl_path.exists(): os.remove(str(tpl_path))
        shutil.move(str(TMP), str(tpl_path))
        d2 = Document(str(tpl_path))
        print(f"{tpl_path.name}: paragraphs={len(d2.paragraphs)}")
    except PermissionError:
        print(f"[!] File terkunci: {TMP}")


def E(body): body.append(mkpara(""))
def A(body, txt, **kw): body.append(mkpara(txt, **kw))
def PH(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def S(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


BASE = Path("knowledge/templates/_skeleton-lhp")


def nota_dinas_reviu(body, jenis_label, judul_prefix):
    A(body, "NOTA DINAS", bold=True)
    A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
    E(body)
    A(body, "Kepada   : {{PENERIMA_LHP}}")
    A(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
    A(body, "Hal       : {{HAL_LHR}}")
    A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
    E(body)
    A(body, f"Menindaklanjuti {{{{DASAR_PERMINTAAN}}}}, kami telah menerbitkan Surat Tugas "
       f"Inspektur Jenderal Nomor {{{{NOMOR_ST}}}} tanggal {{{{TANGGAL_ST}}}} untuk "
       f"melaksanakan {jenis_label} terhadap {{{{NAMA_AUDITI}}}} Tahun {{{{TAHUN_ANGGARAN}}}}.")
    E(body)
    A(body, f"Bersama Nota Dinas ini kami sampaikan {judul_prefix} {{{{JUDUL_LHR_INLINE}}}} "
       "sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
    E(body)
    A(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
       "kami ucapkan terima kasih.")
    E(body)
    A(body, "Inspektur II,")
    E(body); E(body); E(body); E(body)
    A(body, "{{TTD_INSPEKTUR}}"); E(body)
    A(body, "{{NAMA_INSPEKTUR}}")
    A(body, "NIP. {{NIP_INSPEKTUR}}")
    E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")


def cover_lhr(body, judul_laporan):
    E(body); A(body, "=" * 60); E(body)
    A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
    A(body, "INSPEKTORAT JENDERAL"); E(body)
    A(body, judul_laporan, bold=True)
    A(body, "{{JUDUL_LHR_LINE_1}}", bold=True); A(body, "{{JUDUL_LHR_LINE_2}}")
    E(body)
    A(body, "Nomor LHR     : {{NOMOR_LHR}}")
    A(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
    A(body, "Periode Reviu : {{PERIODE_PELAKSANAAN}}")
    A(body, "Tahun         : {{TAHUN_ANGGARAN}}")
    E(body); A(body, "INSPEKTORAT II"); A(body, "{{BULAN_TAHUN}}")


def ttd_tembusan(body):
    E(body)
    A(body, "Inspektur II,")
    E(body); E(body); E(body); E(body)
    A(body, "{{TTD_INSPEKTUR}}"); E(body)
    A(body, "{{NAMA_INSPEKTUR}}")
    A(body, "NIP. {{NIP_INSPEKTUR}}")
    E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")


def pembuka_tubuh(body, jenis_label):
    E(body); A(body, "=" * 60); E(body)
    A(body, "Kepada Yth."); A(body, "{{PENERIMA_LHP}}"); A(body, "di Jakarta"); E(body)
    A(body, f"Menindaklanjuti {{{{DASAR_PERMINTAAN}}}}, kami telah melaksanakan {jenis_label} "
       "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan Surat Tugas "
       "Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. "
       "Pelaksanaan reviu berlangsung mulai {{PERIODE_PELAKSANAAN}}.")
    E(body)


def apresiasi_reviu(body):
    S(body, "F." if True else "G.", "Apresiasi", bold=True); E(body)
    A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
       "selama pelaksanaan reviu ini.")
    A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
       "kami ucapkan terima kasih.")
    E(body)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. REVIU-PENGADAAN (A–F)
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-reviu-pengadaan.docx"
doc, body = rebuild(tpl)
nota_dinas_reviu(body, "Reviu Pengadaan Barang/Jasa", "Laporan Hasil Reviu (LHR) Pengadaan")
cover_lhr(body, "LAPORAN HASIL REVIU PENGADAAN BARANG/JASA")
pembuka_tubuh(body, "Reviu Pengadaan Barang/Jasa")

S(body, "A.", "Pendahuluan", bold=True); E(body)
S(body, "1.", "Latar Belakang"); PH(body, "A1_LATAR_BELAKANG"); E(body)
S(body, "2.", "Dasar Pelaksanaan"); PH(body, "A2_DASAR"); E(body)
S(body, "3.", "Tujuan dan Sasaran"); PH(body, "A3_TUJUAN"); E(body)
S(body, "4.", "Ruang Lingkup"); PH(body, "A4_RUANG_LINGKUP"); E(body)
S(body, "5.", "Metodologi"); PH(body, "A5_METODOLOGI"); E(body)
S(body, "6.", "Jangka Waktu"); PH(body, "A6_JANGKA_WAKTU"); E(body)
S(body, "7.", "Komposisi Tim"); PH(body, "A7_KOMPOSISI_TIM"); E(body)

S(body, "B.", "Gambaran Umum Paket Pengadaan", bold=True); PH(body, "B_GAMBARAN_UMUM"); E(body)

S(body, "C.", "Hasil Reviu", bold=True); E(body)
S(body, "C.1", "Hasil Reviu Perencanaan Pengadaan"); PH(body, "C1_PERENCANAAN"); E(body)
S(body, "C.2", "Hasil Reviu Pemilihan Penyedia"); PH(body, "C2_PEMILIHAN"); E(body)

S(body, "D.", "Simpulan", bold=True); PH(body, "D_SIMPULAN"); E(body)
S(body, "E.", "Rekomendasi", bold=True)
A(body, "Berdasarkan hasil reviu, Inspektorat II merekomendasikan agar:")
PH(body, "E_REKOMENDASI"); E(body)

S(body, "F.", "Apresiasi", bold=True); E(body)
A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
   "selama pelaksanaan reviu pengadaan ini.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
ttd_tembusan(body)
save(doc, body, tpl)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. REVIU-RKA-KL (A–F, 6 aspek reviu)
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-reviu-rka-kl.docx"
doc, body = rebuild(tpl)
nota_dinas_reviu(body, "Reviu Rencana Kerja dan Anggaran (RKA-K/L)", "Laporan Hasil Reviu (LHR) RKA-K/L")
cover_lhr(body, "LAPORAN HASIL REVIU RENCANA KERJA DAN ANGGARAN (RKA-K/L)")
pembuka_tubuh(body, "Reviu Rencana Kerja dan Anggaran (RKA-K/L)")

S(body, "A.", "Pendahuluan", bold=True); E(body)
S(body, "1.", "Latar Belakang dan Dasar Hukum"); PH(body, "A1_LATAR_DASAR"); E(body)
S(body, "2.", "Tujuan"); PH(body, "A2_TUJUAN"); E(body)
S(body, "3.", "Ruang Lingkup"); PH(body, "A3_RUANG_LINGKUP"); E(body)
S(body, "4.", "Metodologi"); PH(body, "A4_METODOLOGI"); E(body)
S(body, "5.", "Jangka Waktu dan Komposisi Tim"); PH(body, "A5_JANGKA_TIM"); E(body)

S(body, "B.", "Gambaran Umum RKA-K/L", bold=True); PH(body, "B_GAMBARAN_UMUM"); E(body)

S(body, "C.", "Hasil Reviu", bold=True); E(body)
S(body, "C.1", "Kelayakan SBM/SBK"); PH(body, "C1_SBM_SBK"); E(body)
S(body, "C.2", "Kepatuhan Kaidah Penganggaran"); PH(body, "C2_KAIDAH"); E(body)
S(body, "C.3", "Penandaan Tematik"); PH(body, "C3_PENANDAAN"); E(body)
S(body, "C.4", "Kelengkapan Dokumen Pendukung"); PH(body, "C4_KELENGKAPAN"); E(body)
S(body, "C.5", "Kelayakan Rincian Baru"); PH(body, "C5_RINCIAN_BARU"); E(body)
S(body, "C.6", "Pengalokasian Tematik"); PH(body, "C6_ALOKASI_TEMATIK"); E(body)

S(body, "D.", "Simpulan", bold=True); PH(body, "D_SIMPULAN"); E(body)
S(body, "E.", "Rekomendasi", bold=True)
A(body, "Berdasarkan hasil reviu, Inspektorat II merekomendasikan agar:")
PH(body, "E_REKOMENDASI"); E(body)

S(body, "F.", "Apresiasi", bold=True); E(body)
A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
   "selama pelaksanaan reviu RKA-K/L ini.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
ttd_tembusan(body)
save(doc, body, tpl)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. REVIU-UMUM (A–G: status per aspek checkmark)
# ═══════════════════════════════════════════════════════════════════════════════
tpl = BASE / "template-lhp-reviu-umum.docx"
doc, body = rebuild(tpl)
nota_dinas_reviu(body, "Reviu", "Laporan Hasil Reviu (LHR)")
cover_lhr(body, "LAPORAN HASIL REVIU")
pembuka_tubuh(body, "Reviu")

S(body, "A.", "Dasar", bold=True); PH(body, "A_DASAR"); E(body)
S(body, "B.", "Tujuan dan Ruang Lingkup", bold=True); PH(body, "B_TUJUAN"); E(body)
S(body, "C.", "Metodologi", bold=True); PH(body, "C_METODOLOGI"); E(body)

S(body, "D.", "Hasil Reviu", bold=True); E(body)
A(body, "Status per aspek yang direviu (TERPENUHI / TERPENUHI DENGAN CATATAN / TIDAK TERPENUHI):")
PH(body, "D_HASIL_REVIU"); E(body)

S(body, "E.", "Catatan dan Rekomendasi", bold=True)
A(body, "Berdasarkan hasil reviu, catatan dan rekomendasi yang perlu ditindaklanjuti:")
PH(body, "E_CATATAN_REKOMENDASI"); E(body)

S(body, "F.", "Simpulan", bold=True); PH(body, "F_SIMPULAN"); E(body)

S(body, "G.", "Apresiasi", bold=True); E(body)
A(body, "Inspektorat II menyampaikan terima kasih atas kerja sama {{NAMA_AUDITI}} "
   "selama pelaksanaan reviu ini.")
A(body, "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
   "kami ucapkan terima kasih.")
E(body)
ttd_tembusan(body)
save(doc, body, tpl)
