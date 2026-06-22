"""
Cleanup template-lhp-audit-umum.docx:
- Hapus heading lama (duplikat dari template generik)
- Hapus paragraf konten generik yang sudah diganti
- Hapus paragraf ST yang bermasalah (ke-2 dari Nota Dinas)
"""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")


def get_text(p_elem):
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


# Paragraf yang dihapus berdasarkan pencocokan teks (substring)
REMOVE_IF_CONTAINS = [
    # Heading lama dari template generik
    "A. Dasar Pelaksanaan Pengawasan",
    "B. Tujuan dan Sasaran Pengawasan",
    "Tujuan: {{TUJUAN_REVIU}}",
    "Sasaran pengawasan adalah memastikan",
    "C. Ruang Lingkup Pengawasan",
    "D. Metodologi Pengawasan",
    "E. Gambaran Umum Pengadaan",
    "F. Hasil Pengawasan",
    "G. Simpulan Pengawasan",
    "H. Tanggapan Auditi",
    "I. Apresiasi",
    # Konten tanggapan lama (bukan untuk audit-umum)
    "Atas seluruh catatan reviu yang disampaikan",
    "[DIISI AUDITI",
    "Inspektorat II Inspektorat Jenderal Kementerian Komunikasi dan Digital menyampaikan terima kasih",
    "Demikian laporan ini kami sampaikan. Atas perhatian",
]

doc = Document(str(TPL))
body = doc.element.body

# Hapus paragraf yang mengandung teks-teks di atas
# Khusus "Menindaklanjuti {{DASAR_PERMINTAAN}}" — hapus HANYA yang ke-2
dasar_permintaan_count = 0
removed = []

for p in list(body.iter(qn("w:p"))):
    txt = get_text(p)

    # Tangani khusus paragraf Menindaklanjuti ST (yang ke-2 = di isi LHA, bukan Nota Dinas)
    if "Menindaklanjuti {{DASAR_PERMINTAAN}}" in txt:
        dasar_permintaan_count += 1
        if dasar_permintaan_count == 2:
            p.getparent().remove(p)
            removed.append(f"[ke-2] {txt[:70]}")
        continue

    # Hapus berdasarkan daftar
    for marker in REMOVE_IF_CONTAINS:
        if marker in txt:
            p.getparent().remove(p)
            removed.append(txt[:70])
            break

doc.save(str(TPL))
print(f"Dibersihkan {len(removed)} paragraf:")
for r in removed:
    print(f"  - {r}")

print("\n=== HASIL AKHIR (paragraf berisi teks) ===")
doc2 = Document(str(TPL))
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip():
        print(f"{i:3}: {p.text[:100]}")
