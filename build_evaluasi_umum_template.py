"""
Build template-lhp-evaluasi-umum.docx sesuai SKILL.md:

Nota Dinas →
A. Dasar                       → {{A_DASAR}}
B. Tujuan & Ruang Lingkup      → {{B_TUJUAN}}
C. Metodologi                  → {{C_METODOLOGI}}
D. Gambaran Umum Objek         → {{D_GAMBARAN_UMUM}}
E. Hasil Evaluasi
   E.1 Skor per Dimensi        → {{E1_SKOR_DIMENSI}}
   E.2 Predikat & Posisi       → {{E2_PREDIKAT}}
   E.3 Analisis per Dimensi    → {{E3_ANALISIS}}
F. Temuan & Catatan            → {{F_TEMUAN}}
G. Rekomendasi                 → {{G_REKOMENDASI}}
H. Simpulan                    → {{H_SIMPULAN}}
I. Apresiasi                   (teks tetap)
TTD + Tembusan
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil, os

TPL = Path("knowledge/templates/_skeleton-lhp/template-lhp-evaluasi-umum.docx")
XML_NS = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def empty(body): body.append(mkpara(""))
def add(body, txt, **kw): body.append(mkpara(txt, **kw))
def ph(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def sec(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


doc = Document(str(TPL))
body = doc.element.body
for el in list(body):
    body.remove(el)

# ═══ NOTA DINAS ═══
add(body, "NOTA DINAS", bold=True)
add(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
empty(body)
add(body, "Kepada   : {{PENERIMA_LHP}}")
add(body, "Dari      : Inspektur II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
add(body, "Hal       : {{HAL_LHR}}")
add(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah menerbitkan Surat Tugas Inspektur "
    "Jenderal Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}} untuk melaksanakan Evaluasi "
    "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}}.")
empty(body)
add(body,
    "Bersama Nota Dinas ini kami sampaikan Laporan Hasil Evaluasi (LHE) "
    "{{JUDUL_LHR_INLINE}} sebagai pertanggungjawaban pelaksanaan penugasan dimaksud.")
empty(body)
add(body, "Demikian Nota Dinas ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)
add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ═══ COVER LHE ═══
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
add(body, "INSPEKTORAT JENDERAL")
empty(body)
add(body, "LAPORAN HASIL EVALUASI", bold=True)
add(body, "{{JUDUL_LHR_LINE_1}}", bold=True)
add(body, "{{JUDUL_LHR_LINE_2}}")
empty(body)
add(body, "Nomor LHE     : {{NOMOR_LHR}}")
add(body, "Surat Tugas   : {{NOMOR_ST}} tanggal {{TANGGAL_ST}}")
add(body, "Periode Eval. : {{PERIODE_PELAKSANAAN}}")
add(body, "Tahun         : {{TAHUN_ANGGARAN}}")
empty(body)
add(body, "INSPEKTORAT II")
add(body, "{{BULAN_TAHUN}}")

# ═══ TUBUH LHE ═══
empty(body)
add(body, "=" * 60)
empty(body)
add(body, "Kepada Yth.")
add(body, "{{PENERIMA_LHP}}")
add(body, "di Jakarta")
empty(body)
add(body,
    "Menindaklanjuti {{DASAR_PERMINTAAN}}, kami telah melaksanakan Evaluasi "
    "terhadap {{NAMA_AUDITI}} Tahun {{TAHUN_ANGGARAN}} berdasarkan Surat Tugas "
    "Nomor {{NOMOR_ST}} tanggal {{TANGGAL_ST}}. Pelaksanaan evaluasi berlangsung "
    "mulai {{PERIODE_PELAKSANAAN}}.")
empty(body)

sec(body, "A.", "Dasar", bold=True)
ph(body, "A_DASAR")
empty(body)

sec(body, "B.", "Tujuan dan Ruang Lingkup", bold=True)
ph(body, "B_TUJUAN")
empty(body)

sec(body, "C.", "Metodologi", bold=True)
ph(body, "C_METODOLOGI")
empty(body)

sec(body, "D.", "Gambaran Umum Objek Evaluasi", bold=True)
ph(body, "D_GAMBARAN_UMUM")
empty(body)

sec(body, "E.", "Hasil Evaluasi", bold=True)
empty(body)
sec(body, "E.1", "Skor per Dimensi")
ph(body, "E1_SKOR_DIMENSI")
empty(body)
sec(body, "E.2", "Predikat dan Posisi")
ph(body, "E2_PREDIKAT")
empty(body)
sec(body, "E.3", "Analisis per Dimensi")
ph(body, "E3_ANALISIS")
empty(body)

sec(body, "F.", "Temuan dan Catatan", bold=True)
empty(body)
add(body,
    "Berdasarkan evaluasi yang dilaksanakan, terdapat hal-hal yang perlu "
    "mendapat perhatian:")
empty(body)
ph(body, "F_TEMUAN")
empty(body)

sec(body, "G.", "Rekomendasi", bold=True)
empty(body)
add(body, "Berdasarkan kondisi-kondisi tersebut, Inspektorat II merekomendasikan agar:")
ph(body, "G_REKOMENDASI")
empty(body)

sec(body, "H.", "Simpulan", bold=True)
ph(body, "H_SIMPULAN")
empty(body)

sec(body, "I.", "Apresiasi", bold=True)
empty(body)
add(body,
    "Inspektorat II menyampaikan penghargaan dan terima kasih atas kerja sama "
    "{{NAMA_AUDITI}} selama pelaksanaan evaluasi ini.")
add(body,
    "Demikian laporan ini kami sampaikan. Atas perhatian dan kerja sama Saudara, "
    "kami ucapkan terima kasih.")
empty(body)

add(body, "Inspektur II,")
empty(body); empty(body); empty(body); empty(body)
add(body, "{{TTD_INSPEKTUR}}")
empty(body)
add(body, "{{NAMA_INSPEKTUR}}")
add(body, "NIP. {{NIP_INSPEKTUR}}")
empty(body)
add(body, "Tembusan:")
ph(body, "TEMBUSAN_LIST")

# ── Simpan ───────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists():
        os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    print("Template berhasil diperbarui:", TPL)
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
    exit()
doc2 = Document(str(TPL))
print(f"paragraphs: {len(doc2.paragraphs)}")
