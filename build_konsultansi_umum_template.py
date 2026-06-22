"""
Rebuild template-lhp-konsultansi-umum.docx sesuai SKILL.md Format Memo Konsultasi:

  A. Dasar            — ND permintaan + ST       → {{A_DASAR}}
  B. Pertanyaan       — daftar pertanyaan        → {{B_PERTANYAAN}}
  C. Dasar Hukum      — kompilasi referensi      → {{C_DASAR_HUKUM}}
  D. Telaah / Analisis — narasi per pertanyaan   → {{D_TELAAH}}
  E. Pendapat / Saran  — jawaban ringkas          → {{E_PENDAPAT}}
  F. Asumsi & Batasan  — batas dan catatan        → {{F_ASUMSI_BATASAN}}
  G. Penutup           — teks tetap
  TTD + Tembusan
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn
from pathlib import Path
import shutil, os

TPL     = Path("knowledge/templates/_skeleton-lhp/template-lhp-konsultansi-umum.docx")
BASE    = Path("knowledge/templates/_skeleton-lhp/template-lhp-audit-umum.docx")
XML_NS  = "http://www.w3.org/XML/1998/namespace"


def mkpara(txt, bold=False, italic=False):
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(f"{{{_W}}}line", "360")
    spacing.set(f"{{{_W}}}lineRule", "auto")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(f"{{{_W}}}ascii", "Arial")
    rFonts.set(f"{{{_W}}}hAnsi", "Arial")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(f"{{{_W}}}val", "24")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(f"{{{_W}}}val", "24")
    rpr.append(szCs)
    if bold: rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = txt
    t.set(f"{{{XML_NS}}}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def E(body): body.append(mkpara(""))
def A(body, txt, **kw): body.append(mkpara(txt, **kw))
def PH(body, key): body.append(mkpara(f"{{{{{key}}}}}"))
def S(body, lbl, judul, bold=True): body.append(mkpara(f"{lbl}  {judul}", bold=bold))


# Buka base template yang punya sectPr valid
src = BASE if BASE.exists() else TPL
doc = Document(str(src))
body = doc.element.body
sect_pr = body.find(_qn("w:sectPr"))
for el in list(body):
    body.remove(el)
if sect_pr is None:
    sect_pr = OxmlElement("w:sectPr")
body.append(sect_pr)

# ═══ HEADER MEMO KONSULTASI ═══
A(body, "KEMENTERIAN KOMUNIKASI DAN DIGITAL REPUBLIK INDONESIA", bold=True)
A(body, "INSPEKTORAT JENDERAL", bold=True)
E(body)
A(body, "MEMO KONSULTASI", bold=True)
A(body, "Nomor: {{NOMOR_NOTA_DINAS}}")
E(body)
A(body, "Kepada   : {{PENERIMA_LHP}}")
A(body, "Dari      : Inspektorat II, Inspektorat Jenderal Kementerian Komunikasi dan Digital")
A(body, "Hal       : {{HAL_LHR}}")
A(body, "Tanggal  : {{TANGGAL_NOTA_DINAS}}")
E(body)

# ═══ BADAN MEMO ═══
S(body, "A.", "Dasar", bold=True)
PH(body, "A_DASAR"); E(body)

S(body, "B.", "Pertanyaan", bold=True)
A(body, "Pertanyaan yang dijawab dalam Memo Konsultasi ini adalah sebagai berikut:")
PH(body, "B_PERTANYAAN"); E(body)

S(body, "C.", "Dasar Hukum", bold=True)
A(body, "Dasar hukum dan referensi yang digunakan dalam pemberian pendapat:")
PH(body, "C_DASAR_HUKUM"); E(body)

S(body, "D.", "Telaah / Analisis", bold=True)
A(body, "Berikut telaah dan analisis atas setiap pertanyaan:")
PH(body, "D_TELAAH"); E(body)

S(body, "E.", "Pendapat / Saran", bold=True)
A(body, "Berdasarkan telaah di atas, kami menyampaikan pendapat sebagai berikut:")
PH(body, "E_PENDAPAT"); E(body)

S(body, "F.", "Asumsi dan Batasan", bold=True)
PH(body, "F_ASUMSI_BATASAN"); E(body)

S(body, "G.", "Penutup", bold=True); E(body)
A(body, "Memo Konsultasi ini bersifat advisory dan tidak mengikat secara hukum. "
   "Pelaksanaan keputusan tetap menjadi kewenangan dan tanggung jawab pejabat yang berwenang "
   "di {{NAMA_AUDITI}}. Pendapat ini diberikan berdasarkan informasi yang disampaikan dalam "
   "{{DASAR_PERMINTAAN}} dan peraturan perundang-undangan yang berlaku pada saat memo ini "
   "diterbitkan.")
A(body, "Apabila terdapat perubahan regulasi atau fakta yang belum disampaikan, "
   "pendapat ini dapat berubah. Untuk konfirmasi atau pertanyaan lebih lanjut, "
   "dapat menghubungi Inspektorat II.")
A(body, "Demikian Memo Konsultasi ini kami sampaikan. Atas perhatian dan kerja sama "
   "Saudara, kami ucapkan terima kasih.")
E(body)

# ═══ TTD ═══
A(body, "Inspektur II,")
E(body); E(body); E(body); E(body)
A(body, "{{TTD_INSPEKTUR}}"); E(body)
A(body, "{{NAMA_INSPEKTUR}}")
A(body, "NIP. {{NIP_INSPEKTUR}}")
E(body); A(body, "Tembusan:"); PH(body, "TEMBUSAN_LIST")

# ── Simpan ──────────────────────────────────────────────────────────────────
TMP = TPL.with_suffix(".tmp.docx")
doc.save(str(TMP))
try:
    if TPL.exists(): os.remove(str(TPL))
    shutil.move(str(TMP), str(TPL))
    d2 = Document(str(TPL))
    print(f"Template berhasil diperbarui: {TPL}")
    print(f"paragraphs: {len(d2.paragraphs)}")
except PermissionError:
    print(f"[!] File terkunci: {TMP}")
